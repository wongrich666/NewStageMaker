from __future__ import annotations

import io
import json
import re
import sqlite3
import threading
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable

from docx import Document
from pypdf import PdfReader

from .deepseek_agent import DeepSeekAgentError, deepseek_agent_client, deepseek_agent_status
from .runtime_paths import get_runtime_data_dir


SKILL_MODULE_SPECS = {
    "genre_profile": {
        "file": "genre-profile.md",
        "label": "题材画像与情绪承诺",
        "stages": ("showrunner",),
    },
    "story_architecture": {
        "file": "story-architecture.md",
        "label": "主线支线与因果架构",
        "stages": ("story_architect",),
    },
    "hook_craft": {
        "file": "hook-craft.md",
        "label": "垂类钩子与追剧问题",
        "stages": ("episode_continuity", "script_writer", "final_editor"),
    },
    "character_emotion": {
        "file": "character-emotion.md",
        "label": "人物关系债与情感共鸣",
        "stages": ("character_emotion", "script_writer", "final_editor"),
    },
    "continuity": {
        "file": "continuity.md",
        "label": "分集因果与场景连续",
        "stages": ("episode_continuity", "script_writer", "state_recorder", "final_editor"),
    },
    "dialogue_voice": {
        "file": "dialogue-voice.md",
        "label": "对白声音与去AI味",
        "stages": ("character_emotion", "script_writer", "final_editor"),
    },
    "adversity_payoff": {
        "file": "adversity-payoff.md",
        "label": "逆风压抑与情绪兑现",
        "stages": ("story_architect", "episode_continuity", "script_writer", "final_editor"),
    },
    "anti_patterns": {
        "file": "anti-patterns.md",
        "label": "题材反模式与失效边界",
        "stages": ("showrunner", "story_architect", "character_emotion", "episode_continuity", "script_writer", "final_editor"),
    },
    "quality_gate": {
        "file": "quality-gate.md",
        "label": "终审质量门",
        "stages": ("final_editor",),
    },
}
SKILL_MODULE_KEYS = tuple(SKILL_MODULE_SPECS)
# Backward-compatible import name for older callers. Values now represent new-workflow modules.
STAGE_PROMPT_KEYS = SKILL_MODULE_KEYS
EVIDENCE_SCHEMA_VERSION = "script-team-evidence/v3"
SKILL_SCHEMA_VERSION = "script-team-skill/v1"

SURFACE_ELEMENT_KEYS = (
    "character_names",
    "relationship_gimmicks",
    "identity_jobs",
    "props_and_evidence",
    "locations_and_world_rules",
    "concrete_incidents",
    "medical_or_biological_elements",
)

RUN_STAGES = (
    ("parse", "素材解析"),
    ("evidence", "单篇证据提取"),
    ("synthesis", "跨样本证据验证"),
    ("compile", "新工作流 Skill 编译"),
    ("evaluate", "发布质量门"),
)

ALLOWED_SUFFIXES = {".docx", ".pdf", ".txt", ".md"}
MAX_SOURCE_BYTES = 30 * 1024 * 1024


def _now() -> str:
    return datetime.now(timezone.utc).astimezone().isoformat()


def _json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))


def _loads(value: Any, default: Any) -> Any:
    try:
        parsed = json.loads(str(value or ""))
    except (TypeError, ValueError, json.JSONDecodeError):
        return default
    return parsed


def _safe_name(value: str) -> str:
    name = Path(str(value or "source")).name
    name = re.sub(r"[^\w\-.()（）\u4e00-\u9fff]+", "_", name, flags=re.UNICODE)
    return name[:160] or "source.txt"


def _read_source(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        document = Document(str(path))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return "\n".join(blocks)
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    return path.read_text(encoding="utf-8", errors="ignore")


def _version_number(existing: list[str]) -> str:
    numbers: list[int] = []
    for value in existing:
        match = re.fullmatch(r"v1\.(\d+)", str(value or ""))
        if match:
            numbers.append(int(match.group(1)))
    return f"v1.{max(numbers, default=-1) + 1}"


def _distillation_sample(text: str, max_chars: int = 32000) -> str:
    """Keep opening, distributed middle evidence, and ending without sending a whole long script."""
    value = str(text or "").strip()
    if len(value) <= max_chars:
        return value
    head_size = int(max_chars * 0.42)
    tail_size = int(max_chars * 0.23)
    middle_budget = max_chars - head_size - tail_size
    middle_start = head_size
    middle_end = len(value) - tail_size
    window_count = 4
    window_size = max(800, middle_budget // window_count)
    span = max(1, middle_end - middle_start - window_size)
    windows: list[str] = []
    for index in range(window_count):
        offset = middle_start + int(span * index / max(1, window_count - 1))
        windows.append(value[offset : offset + window_size])
    return (
        "[开篇样本]\n"
        + value[:head_size]
        + "\n\n[中段分布样本]\n"
        + "\n\n---\n\n".join(windows)
        + "\n\n[结尾样本]\n"
        + value[-tail_size:]
    )


def _surface_terms(evidence: list[dict[str, Any]]) -> list[str]:
    terms: set[str] = set()
    for card in evidence:
        if str(card.get("schema_version") or "") != EVIDENCE_SCHEMA_VERSION:
            continue
        surface = card.get("surface_elements")
        if not isinstance(surface, dict):
            continue
        for key in SURFACE_ELEMENT_KEYS:
            values = surface.get(key)
            if not isinstance(values, list):
                continue
            for value in values:
                term = re.sub(r"\s+", "", str(value or "").strip(" `*#：:，,。.;；"))
                if 2 <= len(term) <= 80:
                    terms.add(term)
    return sorted(terms)


def _surface_leaks(
    modules: dict[str, Any],
    evidence: list[dict[str, Any]],
    *,
    skill_md: str = "",
) -> dict[str, list[str]]:
    terms = _surface_terms(evidence)
    contents = {key: str(value or "") for key, value in modules.items()}
    if skill_md:
        contents["skill_md"] = str(skill_md)
    leaks: dict[str, list[str]] = {}
    for key, content in contents.items():
        matched = [term for term in terms if term in re.sub(r"\s+", "", content)]
        if matched:
            leaks[str(key)] = matched
    return leaks


def _evidence_abstraction_errors(card: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    if not isinstance(card.get("structure_map"), dict):
        errors.append("structure_map")
    surface = card.get("surface_elements")
    if not isinstance(surface, dict):
        return errors + ["surface_elements"]
    for key in SURFACE_ELEMENT_KEYS:
        if not isinstance(surface.get(key), list):
            errors.append(f"surface_elements.{key}")
    return errors


def _skill_manifest(project: sqlite3.Row | dict[str, Any], version: str) -> dict[str, Any]:
    return {
        "schema_version": SKILL_SCHEMA_VERSION,
        "runtime": "codebuddy_npc_script_team",
        "skill_id": str(project["id"]),
        "name": str(project["name"]),
        "genre": str(project["genre"] or ""),
        "market": str(project["market"] or ""),
        "audience": str(project["audience"] or ""),
        "version": version,
        "loading_policy": "selected_skill_only",
        "modules": [
            {
                "key": key,
                "file": spec["file"],
                "label": spec["label"],
                "stages": list(spec["stages"]),
            }
            for key, spec in SKILL_MODULE_SPECS.items()
        ],
    }


def _complete_json_with_repair(
    prompt: str,
    *,
    system_prompt: str,
    max_tokens: int,
    timeout_seconds: int,
) -> dict[str, Any]:
    try:
        return deepseek_agent_client.complete_json(
            prompt,
            system_prompt=system_prompt,
            max_tokens=max_tokens,
            timeout_seconds=timeout_seconds,
        )["structured_output"]
    except DeepSeekAgentError as exc:
        if "合法JSON" not in str(exc) and "数据格式" not in str(exc):
            raise
        repair_system = (
            system_prompt
            + " 严格只返回一个可被json.loads解析的JSON对象；不得使用Markdown代码围栏、注释、尾随逗号或正文说明。"
        )
        return deepseek_agent_client.complete_json(
            prompt,
            system_prompt=repair_system,
            max_tokens=max_tokens,
            timeout_seconds=timeout_seconds,
        )["structured_output"]


class DistillationLabStore:
    def __init__(self, base_dir: str | Path | None = None) -> None:
        self.base_dir = Path(base_dir).resolve() if base_dir else get_runtime_data_dir() / "distillation_lab"
        self.upload_dir = self.base_dir / "uploads"
        self.export_dir = self.base_dir / "exports"
        self.db_path = self.base_dir / "distillation_lab.sqlite3"
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.export_dir.mkdir(parents=True, exist_ok=True)
        self._write_lock = threading.RLock()
        self._threads: dict[str, threading.Thread] = {}
        self._init_schema()

    def _connect(self) -> sqlite3.Connection:
        connection = sqlite3.connect(self.db_path, timeout=30, check_same_thread=False)
        connection.row_factory = sqlite3.Row
        connection.execute("PRAGMA journal_mode=WAL")
        connection.execute("PRAGMA foreign_keys=ON")
        return connection

    def _init_schema(self) -> None:
        with self._connect() as db:
            db.executescript(
                """
                CREATE TABLE IF NOT EXISTS projects (
                    id TEXT PRIMARY KEY,
                    user_id INTEGER NOT NULL,
                    name TEXT NOT NULL,
                    genre TEXT NOT NULL DEFAULT '',
                    market TEXT NOT NULL DEFAULT '',
                    audience TEXT NOT NULL DEFAULT '',
                    description TEXT NOT NULL DEFAULT '',
                    status TEXT NOT NULL DEFAULT 'draft',
                    active_version_id TEXT NOT NULL DEFAULT '',
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                );
                CREATE INDEX IF NOT EXISTS idx_distill_projects_user
                    ON projects(user_id, updated_at DESC);

                CREATE TABLE IF NOT EXISTS sources (
                    id TEXT PRIMARY KEY,
                    project_id TEXT NOT NULL,
                    user_id INTEGER NOT NULL,
                    original_name TEXT NOT NULL,
                    stored_path TEXT NOT NULL,
                    suffix TEXT NOT NULL,
                    polarity TEXT NOT NULL DEFAULT 'positive',
                    weight REAL NOT NULL DEFAULT 1.0,
                    status TEXT NOT NULL DEFAULT 'uploaded',
                    char_count INTEGER NOT NULL DEFAULT 0,
                    extracted_path TEXT NOT NULL DEFAULT '',
                    analysis_json TEXT NOT NULL DEFAULT '{}',
                    error TEXT NOT NULL DEFAULT '',
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL,
                    FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
                );
                CREATE INDEX IF NOT EXISTS idx_distill_sources_project
                    ON sources(project_id, created_at);

                CREATE TABLE IF NOT EXISTS runs (
                    id TEXT PRIMARY KEY,
                    project_id TEXT NOT NULL,
                    user_id INTEGER NOT NULL,
                    status TEXT NOT NULL DEFAULT 'queued',
                    current_stage TEXT NOT NULL DEFAULT '',
                    progress INTEGER NOT NULL DEFAULT 0,
                    error TEXT NOT NULL DEFAULT '',
                    result_json TEXT NOT NULL DEFAULT '{}',
                    created_at TEXT NOT NULL,
                    started_at TEXT NOT NULL DEFAULT '',
                    completed_at TEXT NOT NULL DEFAULT '',
                    FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
                );
                CREATE INDEX IF NOT EXISTS idx_distill_runs_project
                    ON runs(project_id, created_at DESC);

                CREATE TABLE IF NOT EXISTS run_stages (
                    run_id TEXT NOT NULL,
                    stage_key TEXT NOT NULL,
                    stage_name TEXT NOT NULL,
                    status TEXT NOT NULL DEFAULT 'pending',
                    output_json TEXT NOT NULL DEFAULT '{}',
                    error TEXT NOT NULL DEFAULT '',
                    started_at TEXT NOT NULL DEFAULT '',
                    completed_at TEXT NOT NULL DEFAULT '',
                    PRIMARY KEY(run_id, stage_key),
                    FOREIGN KEY(run_id) REFERENCES runs(id) ON DELETE CASCADE
                );

                CREATE TABLE IF NOT EXISTS skill_versions (
                    id TEXT PRIMARY KEY,
                    project_id TEXT NOT NULL,
                    user_id INTEGER NOT NULL,
                    version TEXT NOT NULL,
                    status TEXT NOT NULL DEFAULT 'candidate',
                    skill_md TEXT NOT NULL DEFAULT '',
                    stage_prompts_json TEXT NOT NULL DEFAULT '{}',
                    assets_json TEXT NOT NULL DEFAULT '{}',
                    evidence_json TEXT NOT NULL DEFAULT '[]',
                    score_json TEXT NOT NULL DEFAULT '{}',
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL,
                    published_at TEXT NOT NULL DEFAULT '',
                    UNIQUE(project_id, version),
                    FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
                );
                CREATE INDEX IF NOT EXISTS idx_distill_versions_project
                    ON skill_versions(project_id, created_at DESC);
                """
            )

    @staticmethod
    def _project(row: sqlite3.Row | None) -> dict[str, Any] | None:
        return dict(row) if row else None

    @staticmethod
    def _source(row: sqlite3.Row) -> dict[str, Any]:
        item = dict(row)
        item["analysis"] = _loads(item.pop("analysis_json", "{}"), {})
        item.pop("stored_path", None)
        item.pop("extracted_path", None)
        return item

    @staticmethod
    def _version(row: sqlite3.Row) -> dict[str, Any]:
        item = dict(row)
        modules = _loads(item.pop("stage_prompts_json", "{}"), {})
        item["modules"] = modules
        item["stage_prompts"] = modules
        item["assets"] = _loads(item.pop("assets_json", "{}"), {})
        item["evidence"] = _loads(item.pop("evidence_json", "[]"), [])
        item["score"] = _loads(item.pop("score_json", "{}"), {})
        return item

    def list_projects(self, user_id: int) -> list[dict[str, Any]]:
        with self._connect() as db:
            rows = db.execute(
                """
                SELECT p.*,
                    (SELECT COUNT(*) FROM sources s WHERE s.project_id=p.id) AS source_count,
                    (SELECT COUNT(*) FROM skill_versions v WHERE v.project_id=p.id) AS version_count,
                    (SELECT status FROM runs r WHERE r.project_id=p.id ORDER BY created_at DESC LIMIT 1) AS latest_run_status
                FROM projects p WHERE p.user_id=? ORDER BY p.updated_at DESC
                """,
                (user_id,),
            ).fetchall()
        return [dict(row) for row in rows]

    def create_project(self, user_id: int, payload: dict[str, Any]) -> dict[str, Any]:
        name = str(payload.get("name") or "").strip()
        if not name:
            raise ValueError("请输入蒸馏项目名称。")
        project_id = f"dst-{uuid.uuid4().hex[:16]}"
        now = _now()
        values = (
            project_id,
            user_id,
            name[:100],
            str(payload.get("genre") or "").strip()[:80],
            str(payload.get("market") or "").strip()[:80],
            str(payload.get("audience") or "").strip()[:120],
            str(payload.get("description") or "").strip()[:2000],
            now,
            now,
        )
        with self._write_lock, self._connect() as db:
            db.execute(
                """INSERT INTO projects
                (id,user_id,name,genre,market,audience,description,created_at,updated_at)
                VALUES (?,?,?,?,?,?,?,?,?)""",
                values,
            )
        return self.get_project(user_id, project_id)

    def get_project(self, user_id: int, project_id: str) -> dict[str, Any]:
        with self._connect() as db:
            row = db.execute(
                "SELECT * FROM projects WHERE id=? AND user_id=?", (project_id, user_id)
            ).fetchone()
            if not row:
                raise KeyError("蒸馏项目不存在。")
            project = dict(row)
            project["sources"] = [
                self._source(item)
                for item in db.execute(
                    "SELECT * FROM sources WHERE project_id=? ORDER BY created_at", (project_id,)
                ).fetchall()
            ]
            project["runs"] = [
                self._run(db, item)
                for item in db.execute(
                    "SELECT * FROM runs WHERE project_id=? ORDER BY created_at DESC LIMIT 20", (project_id,)
                ).fetchall()
            ]
            project["versions"] = [
                self._version(item)
                for item in db.execute(
                    "SELECT * FROM skill_versions WHERE project_id=? ORDER BY created_at DESC", (project_id,)
                ).fetchall()
            ]
        return project

    def update_project(self, user_id: int, project_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        current = self.get_project(user_id, project_id)
        values = {
            "name": str(payload.get("name", current["name"]) or "").strip()[:100],
            "genre": str(payload.get("genre", current["genre"]) or "").strip()[:80],
            "market": str(payload.get("market", current["market"]) or "").strip()[:80],
            "audience": str(payload.get("audience", current["audience"]) or "").strip()[:120],
            "description": str(payload.get("description", current["description"]) or "").strip()[:2000],
        }
        if not values["name"]:
            raise ValueError("项目名称不能为空。")
        with self._write_lock, self._connect() as db:
            db.execute(
                """UPDATE projects SET name=?,genre=?,market=?,audience=?,description=?,updated_at=?
                WHERE id=? AND user_id=?""",
                (*values.values(), _now(), project_id, user_id),
            )
        return self.get_project(user_id, project_id)

    def delete_project(self, user_id: int, project_id: str) -> None:
        project = self.get_project(user_id, project_id)
        for source in project["sources"]:
            source_dir = self.upload_dir / str(user_id) / project_id / str(source["id"])
            if source_dir.exists():
                for path in source_dir.iterdir():
                    path.unlink(missing_ok=True)
                source_dir.rmdir()
        with self._write_lock, self._connect() as db:
            db.execute("DELETE FROM projects WHERE id=? AND user_id=?", (project_id, user_id))

    def add_source(
        self,
        user_id: int,
        project_id: str,
        filename: str,
        content: bytes,
        *,
        polarity: str = "positive",
        weight: float = 1.0,
    ) -> dict[str, Any]:
        self.get_project(user_id, project_id)
        suffix = Path(filename).suffix.lower()
        if suffix not in ALLOWED_SUFFIXES:
            raise ValueError("仅支持 Word、PDF、TXT 和 Markdown 文件。")
        if not content:
            raise ValueError("上传文件为空。")
        if len(content) > MAX_SOURCE_BYTES:
            raise ValueError("单个素材不能超过30MB。")
        source_id = f"src-{uuid.uuid4().hex[:16]}"
        source_dir = self.upload_dir / str(user_id) / project_id / source_id
        source_dir.mkdir(parents=True, exist_ok=True)
        stored_path = source_dir / _safe_name(filename)
        stored_path.write_bytes(content)
        now = _now()
        with self._write_lock, self._connect() as db:
            db.execute(
                """INSERT INTO sources
                (id,project_id,user_id,original_name,stored_path,suffix,polarity,weight,created_at,updated_at)
                VALUES (?,?,?,?,?,?,?,?,?,?)""",
                (
                    source_id,
                    project_id,
                    user_id,
                    _safe_name(filename),
                    str(stored_path),
                    suffix,
                    "negative" if polarity == "negative" else "positive",
                    min(3.0, max(0.1, float(weight))),
                    now,
                    now,
                ),
            )
            db.execute("UPDATE projects SET updated_at=? WHERE id=?", (now, project_id))
            row = db.execute("SELECT * FROM sources WHERE id=?", (source_id,)).fetchone()
        return self._source(row)

    def delete_source(self, user_id: int, project_id: str, source_id: str) -> None:
        with self._connect() as db:
            row = db.execute(
                "SELECT * FROM sources WHERE id=? AND project_id=? AND user_id=?",
                (source_id, project_id, user_id),
            ).fetchone()
        if not row:
            raise KeyError("素材不存在。")
        for key in ("stored_path", "extracted_path"):
            value = str(row[key] or "")
            if value:
                Path(value).unlink(missing_ok=True)
        with self._write_lock, self._connect() as db:
            db.execute("DELETE FROM sources WHERE id=?", (source_id,))

    def _run(self, db: sqlite3.Connection, row: sqlite3.Row) -> dict[str, Any]:
        item = dict(row)
        item["result"] = _loads(item.pop("result_json", "{}"), {})
        item["stages"] = [
            {**dict(stage), "output": _loads(stage["output_json"], {})}
            for stage in db.execute(
                "SELECT * FROM run_stages WHERE run_id=? ORDER BY rowid", (item["id"],)
            ).fetchall()
        ]
        for stage in item["stages"]:
            stage.pop("output_json", None)
        return item

    def get_run(self, user_id: int, run_id: str) -> dict[str, Any]:
        with self._connect() as db:
            row = db.execute("SELECT * FROM runs WHERE id=? AND user_id=?", (run_id, user_id)).fetchone()
            if not row:
                raise KeyError("蒸馏任务不存在。")
            return self._run(db, row)

    def start_run(self, user_id: int, project_id: str) -> dict[str, Any]:
        project = self.get_project(user_id, project_id)
        if not project["sources"]:
            raise ValueError("请先上传至少一份剧本或文章素材。")
        with self._connect() as db:
            active = db.execute(
                "SELECT id FROM runs WHERE project_id=? AND status IN ('queued','running') LIMIT 1",
                (project_id,),
            ).fetchone()
        if active:
            return self.get_run(user_id, str(active["id"]))
        run_id = f"run-{uuid.uuid4().hex[:16]}"
        now = _now()
        with self._write_lock, self._connect() as db:
            db.execute(
                "INSERT INTO runs (id,project_id,user_id,status,created_at) VALUES (?,?,?,?,?)",
                (run_id, project_id, user_id, "queued", now),
            )
            db.executemany(
                "INSERT INTO run_stages (run_id,stage_key,stage_name) VALUES (?,?,?)",
                [(run_id, key, name) for key, name in RUN_STAGES],
            )
            db.execute("UPDATE projects SET status='distilling',updated_at=? WHERE id=?", (now, project_id))
        thread = threading.Thread(
            target=self._execute_run,
            args=(user_id, project_id, run_id),
            daemon=True,
            name=f"distillation-{run_id[-8:]}",
        )
        self._threads[run_id] = thread
        thread.start()
        return self.get_run(user_id, run_id)

    def _stage(self, run_id: str, key: str, status: str, *, output: Any = None, error: str = "") -> None:
        now = _now()
        with self._write_lock, self._connect() as db:
            current = db.execute(
                "SELECT started_at FROM run_stages WHERE run_id=? AND stage_key=?", (run_id, key)
            ).fetchone()
            started = str(current["started_at"] or "") if current else ""
            if status == "running" and not started:
                started = now
            completed = now if status in {"completed", "failed"} else ""
            db.execute(
                """UPDATE run_stages SET status=?,output_json=?,error=?,started_at=?,completed_at=?
                WHERE run_id=? AND stage_key=?""",
                (status, _json(output or {}), error[:2000], started, completed, run_id, key),
            )
            completed_count = db.execute(
                "SELECT COUNT(*) AS count FROM run_stages WHERE run_id=? AND status='completed'", (run_id,)
            ).fetchone()["count"]
            db.execute(
                "UPDATE runs SET current_stage=?,progress=? WHERE id=?",
                (key, int(completed_count * 100 / len(RUN_STAGES)), run_id),
            )

    def _execute_run(self, user_id: int, project_id: str, run_id: str) -> None:
        try:
            with self._write_lock, self._connect() as db:
                db.execute(
                    "UPDATE runs SET status='running',started_at=? WHERE id=?", (_now(), run_id)
                )
            self._stage(run_id, "parse", "running")
            corpus = self._parse_sources(user_id, project_id)
            self._stage(
                run_id,
                "parse",
                "completed",
                output={"source_count": len(corpus), "total_characters": sum(x["char_count"] for x in corpus)},
            )

            self._stage(run_id, "evidence", "running")
            evidence = self._extract_evidence(user_id, project_id, corpus)
            self._stage(run_id, "evidence", "completed", output={"items": evidence})

            self._stage(run_id, "synthesis", "running")
            synthesis = self._synthesize(project_id, evidence)
            self._stage(run_id, "synthesis", "completed", output=synthesis)

            self._stage(run_id, "compile", "running")
            version = self._compile_version(user_id, project_id, synthesis, evidence)
            self._stage(
                run_id,
                "compile",
                "completed",
                output={"version_id": version["id"], "version": version["version"]},
            )

            self._stage(run_id, "evaluate", "running")
            score = self._evaluate(version, evidence)
            with self._write_lock, self._connect() as db:
                db.execute(
                    "UPDATE skill_versions SET score_json=?,updated_at=? WHERE id=?",
                    (_json(score), _now(), version["id"]),
                )
                db.execute(
                    """UPDATE runs SET status='completed',progress=100,completed_at=?,result_json=?
                    WHERE id=?""",
                    (_now(), _json({"version_id": version["id"], "score": score}), run_id),
                )
                db.execute(
                    "UPDATE projects SET status='candidate',updated_at=? WHERE id=?", (_now(), project_id)
                )
            self._stage(run_id, "evaluate", "completed", output=score)
        except Exception as exc:
            message = str(exc) or exc.__class__.__name__
            with self._write_lock, self._connect() as db:
                db.execute(
                    "UPDATE runs SET status='failed',error=?,completed_at=? WHERE id=?",
                    (message[:2000], _now(), run_id),
                )
                db.execute(
                    "UPDATE projects SET status='failed',updated_at=? WHERE id=?", (_now(), project_id)
                )
                current = db.execute("SELECT current_stage FROM runs WHERE id=?", (run_id,)).fetchone()
            if current and current["current_stage"]:
                self._stage(run_id, str(current["current_stage"]), "failed", error=message)
        finally:
            self._threads.pop(run_id, None)

    def _parse_sources(self, user_id: int, project_id: str) -> list[dict[str, Any]]:
        with self._connect() as db:
            rows = db.execute(
                "SELECT * FROM sources WHERE project_id=? AND user_id=? ORDER BY created_at",
                (project_id, user_id),
            ).fetchall()
        corpus: list[dict[str, Any]] = []
        for row in rows:
            source = dict(row)
            try:
                text = _read_source(Path(source["stored_path"])).strip()
                if len(text) < 80:
                    raise ValueError("未提取到足够的有效文本。")
                extracted_path = Path(source["stored_path"]).parent / "extracted.txt"
                extracted_path.write_text(text, encoding="utf-8")
                with self._write_lock, self._connect() as db:
                    db.execute(
                        """UPDATE sources SET status='parsed',char_count=?,extracted_path=?,error='',updated_at=?
                        WHERE id=?""",
                        (len(text), str(extracted_path), _now(), source["id"]),
                    )
                corpus.append(
                    {
                        "id": source["id"],
                        "name": source["original_name"],
                        "polarity": source["polarity"],
                        "weight": source["weight"],
                        "char_count": len(text),
                        "text": text,
                    }
                )
            except Exception as exc:
                with self._write_lock, self._connect() as db:
                    db.execute(
                        "UPDATE sources SET status='failed',error=?,updated_at=? WHERE id=?",
                        (str(exc)[:1000], _now(), source["id"]),
                    )
        if not corpus:
            raise ValueError("全部素材解析失败，请检查文件是否包含可复制文本。")
        return corpus

    def _extract_evidence(
        self, user_id: int, project_id: str, corpus: list[dict[str, Any]]
    ) -> list[dict[str, Any]]:
        evidence: list[dict[str, Any]] = []
        for item in corpus:
            cached = self._source_analysis(item["id"])
            if cached:
                evidence.append(cached)
                continue
            sample = _distillation_sample(item["text"])
            prompt = f"""分析以下剧本或文章抽样，生成可追溯的“写法结构证据卡”。不要续写，不要评价作者。
素材名：{item['name']}
样本属性：{'反面样本' if item['polarity'] == 'negative' else '正面样本'}
权重：{item['weight']}

输出JSON字段：
summary、genre_signals、audience_emotions、structure_map、story_architecture、character_patterns、hook_patterns、emotion_curve、continuity_patterns、dialogue_style、adversity_payoff、effective_patterns、failure_patterns、surface_elements。
structure_map必须描述写法机制：主线驱动函数、升级阶梯、支线功能、钩子信息结构、情绪释放顺序、场景交接方式和描写技法。
所有patterns必须是数组；每项包含rule、evidence（概述并标注集数或位置）、conditions、failure_conditions、abstract_slots、transfer_test。
rule只能写可迁移的叙事功能、信息顺序、节拍关系、人物选择机制和描写手法，不能写原作人物、身份、职业、地点、道具、证据方式、疾病、具体事件或固定男女角色分工。
evidence只负责证明rule，可以提到原作内容；abstract_slots说明新故事需要自行填充的功能槽位；transfer_test说明换掉人物、时代、题材和道具后规则是否仍成立。
surface_elements必须是对象，完整包含：character_names、relationship_gimmicks、identity_jobs、props_and_evidence、locations_and_world_rules、concrete_incidents、medical_or_biological_elements；每项都是数组，只登记本素材的表层内容，作为“禁止迁移清单”，不得把它们写进rule。
只提取该素材真实证明的内容；普通编剧常识不作为题材证据。不要大段复制原文，不要把剧情摘要冒充结构。

素材正文：
{sample}"""
            result = _complete_json_with_repair(
                prompt,
                system_prompt="你是剧本结构研究员。只输出合法JSON；严格分离可迁移写法与不可迁移的故事表层。",
                max_tokens=7000,
                timeout_seconds=900,
            )
            abstraction_errors = _evidence_abstraction_errors(result)
            if abstraction_errors:
                result = _complete_json_with_repair(
                    f"""修复以下证据卡的结构蒸馏字段，保留已有证据，不得续写剧情。
缺失或错误字段：{_json(abstraction_errors)}
structure_map必须是写法机制对象；surface_elements必须完整包含{', '.join(SURFACE_ELEMENT_KEYS)}七个数组，并只登记不可迁移的样本表层元素。
所有pattern的rule只能写可迁移结构，原作内容只能留在evidence或surface_elements。
原JSON：
{_json(result)}""",
                    system_prompt="你是证据卡结构修复器。只输出一个合法JSON对象。",
                    max_tokens=7000,
                    timeout_seconds=900,
                )
                abstraction_errors = _evidence_abstraction_errors(result)
            if abstraction_errors:
                raise ValueError("证据卡未完成结构/表层分离：" + "、".join(abstraction_errors))
            result.update(
                {
                    "schema_version": EVIDENCE_SCHEMA_VERSION,
                    "source_id": item["id"],
                    "source_name": item["name"],
                    "polarity": item["polarity"],
                    "weight": item["weight"],
                }
            )
            with self._write_lock, self._connect() as db:
                db.execute(
                    "UPDATE sources SET status='analyzed',analysis_json=?,updated_at=? WHERE id=?",
                    (_json(result), _now(), item["id"]),
                )
            evidence.append(result)
        return evidence

    def _source_analysis(self, source_id: str) -> dict[str, Any] | None:
        with self._connect() as db:
            row = db.execute("SELECT analysis_json FROM sources WHERE id=?", (source_id,)).fetchone()
        data = _loads(row["analysis_json"], {}) if row else {}
        return (
            data
            if isinstance(data, dict)
            and data.get("source_id")
            and data.get("schema_version") == EVIDENCE_SCHEMA_VERSION
            else None
        )

    def _synthesize(self, project_id: str, evidence: list[dict[str, Any]]) -> dict[str, Any]:
        with self._connect() as db:
            project = db.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
        positive = [item for item in evidence if item.get("polarity") != "negative"]
        negative = [item for item in evidence if item.get("polarity") == "negative"]
        source_ids = [str(item.get("source_id") or "") for item in evidence if item.get("source_id")]
        return {
            "method": "nuwa-inspired-evidence-gate",
            "model_calls": "one cached extraction per new source plus one compile call per version",
            "project": {
                "name": project["name"],
                "genre": project["genre"],
                "market": project["market"],
                "audience": project["audience"],
            },
            "source_count": len(evidence),
            "positive_count": len(positive),
            "negative_count": len(negative),
            "source_ids": source_ids,
            "verification_rules": [
                "重复性：三份及以上样本时，稳定规律原则上需要至少两个独立来源；单来源只能标为候选假设。",
                "生成性：规律必须能转写为条件、创作动作和失败边界，能指导未见过的新剧情。",
                "区分度：删除任何题材都成立的空泛常识，保留与目标受众和情绪承诺有关的差异规律。",
                "反证：反面样本只用于识别失败模式，不与正面样本做简单多数表决。",
            ],
        }

    def _compile_version(
        self,
        user_id: int,
        project_id: str,
        synthesis: dict[str, Any],
        evidence: list[dict[str, Any]],
    ) -> dict[str, Any]:
        with self._connect() as db:
            project = db.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
            versions = [
                row["version"]
                for row in db.execute(
                    "SELECT version FROM skill_versions WHERE project_id=?", (project_id,)
                ).fetchall()
            ]
        version = _version_number(versions)
        evidence_payload = _json(evidence)
        if len(evidence_payload) > 120000:
            evidence_payload = evidence_payload[:120000]
        manifest = _skill_manifest(project, version)
        prompt = f"""把证据卡编译成一个供新剧本团队选择加载的“结构与写法Skill”。不要写具体剧本。
项目：{project['name']}；题材：{project['genre']}；版本：{version}
市场：{project['market']}；受众：{project['audience']}；边界：{project['description']}

输出JSON必须包含：
 skill_md：完整Markdown，含YAML frontmatter(name/description)；name必须是64字符内的小写英文连字符格式，description要写清题材、受众及何时触发；正文控制在500行内；
modules：对象，且必须完整包含这些新工作流模块：{', '.join(SKILL_MODULE_KEYS)}；
verified_rules：通过验证的规则数组；hypotheses：证据不足但值得保留的假设数组；
source_conflicts：样本之间互相冲突的规律；confidence_notes：能力边界。

证据门：
1. 三份及以上正面样本时，稳定规律原则上至少有两个独立source_id；单来源只能进入hypotheses。
2. 每条verified_rule必须具有rule、conditions、failure_conditions、source_ids、confidence。
3. 规则必须能指导新剧情，并具有题材或受众区分度；删除“加强冲突”等空泛常识。
4. 反面样本提炼anti_patterns，不可当成正向范式。
5. 模块只放对应节点真正需要的规则，禁止九份模块重复同一内容。
6. 只迁移结构，不迁移故事：规则写“叙事功能+触发条件+变量槽位+节拍关系+失败边界”。
7. surface_elements是隔离区，其中的人物名、固定关系套路、身份职业、场景世界观、道具、证据手段、疾病和具体事件不得出现在skill_md或modules；即使多篇样本重复出现，也只能说明样本同质，不能升级为结构规律。
8. 不得规定“男主必须如何、女主必须如何”等固定性别分工；除非项目边界明确要求，否则统一改写为主角、关系对手、权力方、受压方、盟友等功能角色。
9. 关键物品只能抽象为“承载某段关系/信息/代价且会递进变化的媒介”，不能沿用样本的物品类别；支线只能规定功能、进入条件、回撞主线方式和退出条件，不能沿用样本支线事件。
10. 每条模块规则必须通过迁移测试：把原作人物、时代、关系、地点、道具全部替换后仍能指导一个不同故事，否则删除或降为hypotheses。
11. skill_md只负责触发条件、使用顺序、模块导航和边界，详细规则放modules，节约运行上下文。

新工作流路由清单：
{_json(manifest)}

证据验证配置：
{_json(synthesis)}

逐篇证据卡：
{evidence_payload}"""
        compiled = _complete_json_with_repair(
            prompt,
            system_prompt="你是Agent Skill编译器，只输出合法JSON。",
            max_tokens=16000,
            timeout_seconds=1200,
        )
        first_modules = (
            compiled.get("modules") if isinstance(compiled.get("modules"), dict) else {}
        )
        first_leaks = _surface_leaks(
            first_modules,
            evidence,
            skill_md=str(compiled.get("skill_md") or ""),
        )
        if first_leaks:
            repair_prompt = f"""上一版Skill误把样本表层内容写进了创作规则，必须做结构抽象修复。
检测到的污染位置与词项：{_json(first_leaks)}

请在不减少模块、不改变JSON字段的前提下重写上一版：
1. 删除这些样本专属名词及其同义改写，不得简单换成另一种具体人物、道具或事件。
2. 将其改为叙事功能、变量槽位、信息顺序、节拍关系和失败边界。
3. 保留证据支持的写法机制；不续写故事，不新增剧情模板。
4. 输出仍须包含skill_md、modules、verified_rules、hypotheses、source_conflicts、confidence_notes。

上一版JSON：
{_json(compiled)}"""
            compiled = _complete_json_with_repair(
                repair_prompt,
                system_prompt="你是Skill结构纯化器。只输出合法JSON，只迁移写法，不迁移故事表层。",
                max_tokens=16000,
                timeout_seconds=1200,
            )
        raw_modules = compiled.get("modules") if isinstance(compiled.get("modules"), dict) else {}
        modules = {key: str(raw_modules.get(key) or "").strip() for key in SKILL_MODULE_KEYS}
        skill_md = str(compiled.get("skill_md") or "").strip()
        if not skill_md:
            raise ValueError("模型没有生成SKILL.md。")
        version_id = f"ver-{uuid.uuid4().hex[:16]}"
        now = _now()
        with self._write_lock, self._connect() as db:
            db.execute(
                """INSERT INTO skill_versions
                (id,project_id,user_id,version,status,skill_md,stage_prompts_json,assets_json,evidence_json,created_at,updated_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
                (
                    version_id,
                    project_id,
                    user_id,
                    version,
                    "candidate",
                    skill_md,
                    _json(modules),
                    _json(
                        {
                            "manifest": manifest,
                            "methodology": synthesis,
                            "verified_rules": compiled.get("verified_rules") or [],
                            "hypotheses": compiled.get("hypotheses") or [],
                            "source_conflicts": compiled.get("source_conflicts") or [],
                            "confidence_notes": compiled.get("confidence_notes") or [],
                        }
                    ),
                    _json(evidence),
                    now,
                    now,
                ),
            )
            row = db.execute("SELECT * FROM skill_versions WHERE id=?", (version_id,)).fetchone()
        return self._version(row)

    @staticmethod
    def _evaluate(version: dict[str, Any], evidence: list[dict[str, Any]]) -> dict[str, Any]:
        modules = version.get("modules") or version.get("stage_prompts") or {}
        filled = sum(1 for key in SKILL_MODULE_KEYS if str(modules.get(key) or "").strip())
        source_count = len({str(item.get("source_id") or "") for item in evidence})
        assets = version.get("assets") if isinstance(version.get("assets"), dict) else {}
        verified_rules = assets.get("verified_rules") if isinstance(assets.get("verified_rules"), list) else []
        valid_source_ids = {
            str(item.get("source_id") or "") for item in evidence if str(item.get("source_id") or "")
        }
        required_sources = 2 if source_count >= 2 else 1
        validated_rules = 0
        for rule in verified_rules:
            cited = {
                str(source_id)
                for source_id in (rule.get("source_ids") if isinstance(rule, dict) else []) or []
                if str(source_id)
            }
            if len(cited) >= required_sources and cited.issubset(valid_source_ids):
                validated_rules += 1
        validation_ratio = validated_rules / len(verified_rules) if verified_rules else 0
        abstracted_cards = sum(
            1
            for item in evidence
            if str(item.get("schema_version") or "") == EVIDENCE_SCHEMA_VERSION
            and not _evidence_abstraction_errors(item)
        )
        abstraction_ratio = abstracted_cards / source_count if source_count else 0
        surface_leaks = _surface_leaks(
            modules,
            evidence,
            skill_md=str(version.get("skill_md") or ""),
        )
        checks = {
            "skill_structure": 100 if "---" in version.get("skill_md", "") else 60,
            "module_coverage": round(filled / len(SKILL_MODULE_KEYS) * 100),
            "evidence_coverage": min(100, 45 + source_count * 9),
            "evidence_traceability": round(validation_ratio * 100),
            "boundary_clarity": 100 if any(word in version.get("skill_md", "") for word in ("边界", "失效", "不适用")) else 65,
            "new_workflow_manifest": 100 if assets.get("manifest", {}).get("schema_version") == SKILL_SCHEMA_VERSION else 40,
            "structural_purity": 100 if not surface_leaks else 0,
            "structure_surface_separation": round(abstraction_ratio * 100),
        }
        total = round(sum(checks.values()) / len(checks))
        return {
            "total": total,
            "grade": "A" if total >= 85 else "B" if total >= 75 else "C",
            "checks": checks,
            "source_count": source_count,
            "ready_to_publish": (
                total >= 78
                and filled == len(SKILL_MODULE_KEYS)
                and bool(verified_rules)
                and validated_rules == len(verified_rules)
                and not surface_leaks
                and abstracted_cards == source_count
            ),
            "validation_mode": "deterministic_evidence_gate",
            "deep_blind_test": "not_run",
            "verified_rule_count": len(verified_rules),
            "validated_rule_count": validated_rules,
            "surface_leaks": surface_leaks,
            "abstracted_source_count": abstracted_cards,
        }

    def update_version(
        self, user_id: int, project_id: str, version_id: str, payload: dict[str, Any]
    ) -> dict[str, Any]:
        version = self.get_version(user_id, project_id, version_id)
        if version["status"] != "candidate":
            raise ValueError("只有候选版本可以修改，请重新蒸馏生成新候选版本。")
        skill_md = str(payload.get("skill_md", version["skill_md"]) or "").strip()
        modules = payload.get("modules", payload.get("stage_prompts", version.get("modules") or {}))
        if not isinstance(modules, dict):
            raise ValueError("modules必须是对象。")
        normalized = {key: str(modules.get(key) or "").strip() for key in SKILL_MODULE_KEYS}
        candidate = {
            **version,
            "skill_md": skill_md,
            "modules": normalized,
            "stage_prompts": normalized,
        }
        score = self._evaluate(candidate, version.get("evidence") or [])
        with self._write_lock, self._connect() as db:
            db.execute(
                """UPDATE skill_versions
                SET skill_md=?,stage_prompts_json=?,score_json=?,updated_at=? WHERE id=?""",
                (skill_md, _json(normalized), _json(score), _now(), version_id),
            )
        return self.get_version(user_id, project_id, version_id)

    def get_version(self, user_id: int, project_id: str, version_id: str) -> dict[str, Any]:
        with self._connect() as db:
            row = db.execute(
                "SELECT * FROM skill_versions WHERE id=? AND project_id=? AND user_id=?",
                (version_id, project_id, user_id),
            ).fetchone()
        if not row:
            raise KeyError("Skill版本不存在。")
        return self._version(row)

    def publish_version(self, user_id: int, project_id: str, version_id: str) -> dict[str, Any]:
        version = self.get_version(user_id, project_id, version_id)
        score = self._evaluate(version, version.get("evidence") or [])
        if not score.get("ready_to_publish"):
            raise ValueError("该版本尚未通过完整性评测，不能发布。")
        now = _now()
        with self._write_lock, self._connect() as db:
            db.execute(
                "UPDATE skill_versions SET status='retired',updated_at=? WHERE project_id=? AND status='published'",
                (now, project_id),
            )
            db.execute(
                "UPDATE skill_versions SET status='published',score_json=?,published_at=?,updated_at=? WHERE id=?",
                (_json(score), now, now, version_id),
            )
            db.execute(
                "UPDATE projects SET status='published',active_version_id=?,updated_at=? WHERE id=? AND user_id=?",
                (version_id, now, project_id, user_id),
            )
        return self.get_version(user_id, project_id, version_id)

    def unpublish_version(self, user_id: int, project_id: str, version_id: str) -> dict[str, Any]:
        version = self.get_version(user_id, project_id, version_id)
        if version.get("status") != "published":
            raise ValueError("只有已发布版本可以取消发布。")
        now = _now()
        with self._write_lock, self._connect() as db:
            db.execute(
                """UPDATE skill_versions
                SET status='candidate',published_at='',updated_at=?
                WHERE id=? AND project_id=? AND user_id=?""",
                (now, version_id, project_id, user_id),
            )
            db.execute(
                """UPDATE projects
                SET status='candidate',active_version_id='',updated_at=?
                WHERE id=? AND user_id=? AND active_version_id=?""",
                (now, project_id, user_id, version_id),
            )
        return self.get_version(user_id, project_id, version_id)

    def list_published_skills(self, user_id: int) -> list[dict[str, Any]]:
        """Return lightweight cards for Skills that can be attached to a script job."""
        with self._connect() as db:
            rows = db.execute(
                """
                SELECT p.id AS project_id,p.name,p.genre,p.market,p.audience,p.description,
                       p.active_version_id,v.version,v.skill_md,v.stage_prompts_json,
                       v.assets_json,v.evidence_json,v.score_json,v.updated_at,v.published_at
                FROM projects p
                JOIN skill_versions v ON v.id=p.active_version_id
                WHERE p.user_id=? AND p.status='published' AND v.status='published'
                ORDER BY COALESCE(v.published_at,v.updated_at) DESC
                """,
                (user_id,),
            ).fetchall()
        cards: list[dict[str, Any]] = []
        for row in rows:
            item = dict(row)
            item.pop("score_json", None)
            version = {
                "skill_md": str(item.pop("skill_md", "") or ""),
                "modules": _loads(item.pop("stage_prompts_json", "{}"), {}),
                "assets": _loads(item.pop("assets_json", "{}"), {}),
            }
            evidence = _loads(item.pop("evidence_json", "[]"), [])
            score = self._evaluate(version, evidence)
            if not score.get("ready_to_publish"):
                continue
            identity = f"{item.get('name', '')} {item.get('genre', '')}"
            cover_name = (
                "romance-angst.png"
                if any(word in identity for word in ("虐恋", "追妻", "言情", "爱情"))
                else "contemporary-emotion.png"
            )
            cards.append(
                {
                    **item,
                    "skill_id": item["project_id"],
                    "version_id": item["active_version_id"],
                    "score": int(score.get("total") or 0),
                    "grade": str(score.get("grade") or ""),
                    "cover_url": f"/static/distillation_skill_covers/{cover_name}",
                    "module_count": len(SKILL_MODULE_KEYS),
                    "schema_version": SKILL_SCHEMA_VERSION,
                }
            )
        return cards

    def resolve_runtime_skill(
        self, user_id: int, skill_id: str, version_id: str = ""
    ) -> dict[str, Any]:
        """Resolve a published Skill into an immutable job-time runtime snapshot."""
        project = self.get_project(user_id, skill_id)
        selected_version_id = str(version_id or project.get("active_version_id") or "").strip()
        if not selected_version_id:
            raise ValueError("所选蒸馏 Skill 尚未发布可用版本。")
        version = self.get_version(user_id, skill_id, selected_version_id)
        if version.get("status") != "published":
            raise ValueError("只能关联已发布的蒸馏 Skill 版本。")
        assets = version.get("assets") if isinstance(version.get("assets"), dict) else {}
        manifest = assets.get("manifest") if isinstance(assets.get("manifest"), dict) else {}
        if manifest.get("schema_version") != SKILL_SCHEMA_VERSION:
            raise ValueError("所选 Skill 与当前新工作流不兼容，请重新蒸馏并发布。")
        modules = version.get("modules") if isinstance(version.get("modules"), dict) else {}
        missing = [key for key in SKILL_MODULE_KEYS if not str(modules.get(key) or "").strip()]
        if missing:
            raise ValueError("所选 Skill 缺少运行模块：" + "、".join(missing))
        runtime_score = self._evaluate(version, version.get("evidence") or [])
        if not runtime_score.get("ready_to_publish"):
            raise ValueError("所选 Skill 属于旧版剧情复刻型蒸馏，请重新蒸馏并发布结构版。")
        return {
            "schema_version": SKILL_SCHEMA_VERSION,
            "skill_id": skill_id,
            "project_id": skill_id,
            "name": str(project.get("name") or ""),
            "genre": str(project.get("genre") or ""),
            "market": str(project.get("market") or ""),
            "audience": str(project.get("audience") or ""),
            "version_id": selected_version_id,
            "version": str(version.get("version") or ""),
            "score": int(runtime_score.get("total") or 0),
            "skill_md": str(version.get("skill_md") or ""),
            "manifest": manifest,
            "modules": {key: str(modules.get(key) or "") for key in SKILL_MODULE_KEYS},
            "published_at": str(version.get("published_at") or ""),
        }

    def export_version(self, user_id: int, project_id: str, version_id: str) -> tuple[io.BytesIO, str]:
        project = self.get_project(user_id, project_id)
        version = self.get_version(user_id, project_id, version_id)
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("SKILL.md", version["skill_md"])
            modules = version.get("modules") or version.get("stage_prompts") or {}
            for key, spec in SKILL_MODULE_SPECS.items():
                archive.writestr(f"references/{spec['file']}", str(modules.get(key) or ""))
            assets = version.get("assets") if isinstance(version.get("assets"), dict) else {}
            archive.writestr(
                "manifest.json",
                json.dumps(assets.get("manifest") or {}, ensure_ascii=False, indent=2),
            )
            archive.writestr(
                "references/methodology.json",
                json.dumps(assets.get("methodology") or {}, ensure_ascii=False, indent=2),
            )
            archive.writestr(
                "references/verified_rules.json",
                json.dumps(assets.get("verified_rules") or [], ensure_ascii=False, indent=2),
            )
            archive.writestr(
                "references/hypotheses.json",
                json.dumps(assets.get("hypotheses") or [], ensure_ascii=False, indent=2),
            )
            archive.writestr(
                "references/source_conflicts.json",
                json.dumps(assets.get("source_conflicts") or [], ensure_ascii=False, indent=2),
            )
            archive.writestr(
                "references/evidence_index.json",
                json.dumps(version["evidence"], ensure_ascii=False, indent=2),
            )
            archive.writestr(
                "references/evaluation.json",
                json.dumps(version["score"], ensure_ascii=False, indent=2),
            )
        buffer.seek(0)
        filename = f"{_safe_name(project['name'])}_{version['version']}_skill.zip"
        return buffer, filename

    def overview(self, user_id: int) -> dict[str, Any]:
        projects = self.list_projects(user_id)
        return {
            "projects": projects,
            "counts": {
                "projects": len(projects),
                "sources": sum(int(item.get("source_count") or 0) for item in projects),
                "versions": sum(int(item.get("version_count") or 0) for item in projects),
                "running": sum(1 for item in projects if item.get("latest_run_status") in {"queued", "running"}),
            },
            "model": deepseek_agent_status(),
            "workflow_link_enabled": True,
            "skill_schema_version": SKILL_SCHEMA_VERSION,
            "module_keys": list(SKILL_MODULE_KEYS),
            "cost_policy": {
                "source_analysis": "cached_by_source_and_schema",
                "source_sample_char_limit": 32000,
                "compile_calls_per_version": 1,
                "blind_test": "optional_not_run",
            },
        }


distillation_lab_store = DistillationLabStore()
