from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, BinaryIO

from docx import Document


MAX_REPLACEMENT_OPERATIONS = 40
MAX_REPLACEMENT_CHARS = 6000


def _open_document(source: bytes | str | Path | BinaryIO):
    if isinstance(source, bytes):
        return Document(BytesIO(source))
    return Document(source)


def _editable_paragraphs(document) -> list[Any]:
    paragraphs: list[Any] = []
    seen: set[int] = set()

    def append(paragraph: Any) -> None:
        marker = id(paragraph._p)
        if marker in seen or not str(paragraph.text or "").strip():
            return
        seen.add(marker)
        paragraphs.append(paragraph)

    for paragraph in document.paragraphs:
        append(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    append(paragraph)
    return paragraphs


def list_docx_paragraphs(source: bytes | str | Path | BinaryIO) -> list[dict[str, str]]:
    document = _open_document(source)
    return [
        {
            "paragraph_id": f"P{index:04d}",
            "text": str(paragraph.text or ""),
        }
        for index, paragraph in enumerate(_editable_paragraphs(document), start=1)
    ]


def indexed_docx_text(source: bytes | str | Path | BinaryIO) -> str:
    return "\n".join(
        f"[{item['paragraph_id']}] {item['text']}"
        for item in list_docx_paragraphs(source)
    )


def apply_docx_replacements(
    source: bytes | str | Path | BinaryIO,
    operations: list[dict[str, Any]],
) -> tuple[bytes, list[dict[str, str]], list[dict[str, str]]]:
    document = _open_document(source)
    paragraphs = _editable_paragraphs(document)
    paragraph_map = {
        f"P{index:04d}": paragraph
        for index, paragraph in enumerate(paragraphs, start=1)
    }
    applied: list[dict[str, str]] = []
    skipped: list[dict[str, str]] = []
    seen: set[str] = set()

    for raw_operation in list(operations or [])[:MAX_REPLACEMENT_OPERATIONS]:
        operation = raw_operation if isinstance(raw_operation, dict) else {}
        paragraph_id = str(operation.get("paragraph_id") or "").strip().upper()
        original_text = str(operation.get("original_text") or "")
        replacement_text = str(operation.get("replacement_text") or "")
        reason = str(operation.get("reason") or "").strip()

        if not paragraph_id or paragraph_id in seen:
            skipped.append({"paragraph_id": paragraph_id, "reason": "段落编号为空或重复。"})
            continue
        seen.add(paragraph_id)
        paragraph = paragraph_map.get(paragraph_id)
        if paragraph is None:
            skipped.append({"paragraph_id": paragraph_id, "reason": "原 Word 中不存在该段落。"})
            continue
        current_text = str(paragraph.text or "")
        if current_text.strip() != original_text.strip():
            skipped.append({"paragraph_id": paragraph_id, "reason": "AI 返回的原文与 Word 段落不完全匹配。"})
            continue
        if not replacement_text.strip() or replacement_text.strip() == current_text.strip():
            skipped.append({"paragraph_id": paragraph_id, "reason": "优化后文本为空或没有变化。"})
            continue
        if len(replacement_text) > MAX_REPLACEMENT_CHARS:
            skipped.append({"paragraph_id": paragraph_id, "reason": "单段优化文本超过安全长度。"})
            continue

        runs = list(paragraph.runs)
        if runs:
            runs[0].text = replacement_text
            for run in runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(replacement_text)
        applied.append(
            {
                "paragraph_id": paragraph_id,
                "original_text": current_text,
                "replacement_text": replacement_text,
                "reason": reason,
            }
        )

    output = BytesIO()
    document.save(output)
    return output.getvalue(), applied, skipped
