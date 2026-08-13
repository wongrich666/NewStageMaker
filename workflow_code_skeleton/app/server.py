from __future__ import annotations

import json
import copy
import re
import threading
import tempfile
import time
import traceback
import uuid
import requests
from io import BytesIO
from functools import wraps
import os
from pathlib import Path
from datetime import datetime, timezone
from typing import Any
from urllib.parse import quote

from flask import (
    Flask,
    Response,
    jsonify,
    g,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from werkzeug.middleware.proxy_fix import ProxyFix

from .models.inputs import derive_script_title_content
from .services.auth_store import auth_store
from .services.admin_store import admin_store
from .services.agent_conversation_store import agent_conversation_store
from .services.platform_agent import platform_conversation_agent
from .services.deepseek_agent import DeepSeekAgentError
from .services.compliance_review import catalog as compliance_catalog
from .services.compliance_review import compliance_review_store
from .services.compliance_review import review_script as run_compliance_review
from .services.script_rating import rate_script as run_script_rating
from .services.script_rating import repair_script as run_script_rating_repair
from .services.script_rating import script_rating_store
from .services.distillation_lab import distillation_lab_store
from .services.codebuddy_npc import (
    STAGE_NAMES,
    STAGE_ORDER,
    CodeBuddyNpcClient,
    CodeBuddyNpcConfig,
    CodeBuddyNpcError,
    CodeBuddyNpcJobStore,
    start_stage_timing,
    public_job,
)
from .services.codebuddy_npc_stage_runner import CodeBuddyNpcStageRunner
from .services.fastgpt_client import FastGPTTransientError
from .services.framework_planner_service import (
    FRAMEWORK_PLANNER_STORAGE_KEY,
    FrameworkPlannerStageError,
    framework_planner_backend_ready,
    framework_planner_fastgpt_diagnostics,
    list_framework_stage_history,
    load_framework_stage_history,
    run_framework_planner_score,
    run_framework_planner_stage,
    save_framework_stage_history,
    write_framework_frontend_debug_event,
    write_framework_stage_exception_log,
)
from .services.simple_fastgpt_tools import ToolExecutionError, list_simple_tools, run_simple_tool
from .services.workbuddy_doctor import (
    build_doctor_optimization_prompt,
    build_doctor_prompt,
    clear_workbuddy_history,
    delete_workbuddy_history,
    doctor_timeout_seconds,
    extract_doctor_report,
    format_doctor_exception,
    list_workbuddy_history,
    load_workbuddy_optimized_document,
    load_workbuddy_source,
    load_workbuddy_history,
    parse_doctor_optimization_result,
    run_codebuddy_doctor,
    save_workbuddy_optimized_document,
    save_workbuddy_source,
    save_workbuddy_history,
    validate_doctor_payload,
    workbuddy_config_status,
)
from .services.workbuddy_docx_optimizer import (
    apply_docx_replacements,
    indexed_docx_text,
)
from .services.script_audit_ecg_parser import (
    COMPACT_SCHEMA_VERSION as SCRIPT_AUDIT_COMPACT_SCHEMA_VERSION,
    SCHEMA_VERSION as SCRIPT_AUDIT_ECG_SCHEMA_VERSION,
    SCHEMA_VERSION_V3 as SCRIPT_AUDIT_ECG_SCHEMA_VERSION_V3,
    build_audit_visualization_payload,
    build_script_audit_view_model,
    fallback_audit_from_text,
    normalize_compact_audit_payload,
    normalize_script_audit_ecg,
    parse_compact_audit_json,
    parse_model_json_loose,
    validate_compact_audit_schema,
)
from .services.task_manager import task_manager
from .services.user_knowledge_store import user_knowledge_store
from .services.workflow_preference_keys import (
    FRAMEWORK_TO_SCRIPT_STAGE_PREFS,
    inject_stage_preference,
    preference_keys_for,
    stage_prompt_key_for,
)
from .services.workflow_line import (
    PRODUCTION_WORKFLOW_LINE,
    TEST_WORKFLOW_LINE,
    WORKFLOW_LINE_HEADER,
    is_test_workflow_line,
    reset_workflow_line,
    set_workflow_line,
    test_workflow_status,
)
from .utils.logger import get_logger
from .utils.readable_labels import readable_label, readable_scalar, readable_text


logger = get_logger("server")


def default_workflow_spec_path() -> str:
    return str(Path.home() / "Downloads" / "剧本生成_0401_loops.json")


def create_app(*, workflow_spec_path: str | None = None) -> Flask:
    template_dir = Path(__file__).resolve().parent / "web" / "templates"
    static_dir = Path(__file__).resolve().parent / "web" / "static"
    app = Flask(
        __name__,
        template_folder=str(template_dir),
        static_folder=str(static_dir),
    )
    if str(os.getenv("SCRIPTMAKER_TRUST_REVERSE_PROXY", "false")).strip().lower() in {
        "1",
        "true",
        "yes",
        "on",
    }:
        app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)
    codebuddy_npc_config = CodeBuddyNpcConfig.from_env()
    codebuddy_npc_jobs = CodeBuddyNpcJobStore(codebuddy_npc_config)
    codebuddy_npc_client = CodeBuddyNpcClient(codebuddy_npc_config)
    codebuddy_npc_stage_runner = CodeBuddyNpcStageRunner(codebuddy_npc_jobs)
    codebuddy_npc_refresh_locks: dict[str, threading.RLock] = {}
    codebuddy_npc_refresh_locks_guard = threading.Lock()
    try:
        codebuddy_npc_remote_retry_limit = int(
            os.getenv("CODEBUDDY_NPC_REMOTE_STAGE_RETRIES", "2")
        )
    except (TypeError, ValueError):
        codebuddy_npc_remote_retry_limit = 2
    codebuddy_npc_remote_retry_limit = min(
        5, max(0, codebuddy_npc_remote_retry_limit)
    )
    try:
        codebuddy_npc_result_pending_timeout = int(
            os.getenv("CODEBUDDY_NPC_RESULT_PENDING_TIMEOUT", "120")
        )
    except (TypeError, ValueError):
        codebuddy_npc_result_pending_timeout = 120
    codebuddy_npc_result_pending_timeout = max(
        30, min(600, codebuddy_npc_result_pending_timeout)
    )

    def _submit_remote_npc_stage(
        job: dict[str, Any],
        *,
        stage: str,
        feedback: str = "",
        continue_after: bool = False,
        retry_count: int = 0,
    ) -> dict[str, Any]:
        trigger_result = codebuddy_npc_client.trigger_stage(
            job,
            stage=stage,
            feedback=feedback,
            continue_after=continue_after,
        )
        job["build"] = {
            "sn": trigger_result.get("sn"),
            "build_log_url": trigger_result.get("buildLogUrl"),
            "message": trigger_result.get("message"),
        }
        job["execution_target"] = "remote_cnb"
        job["remote_kind"] = "stage"
        job["remote_stage"] = stage
        job["remote_continue_after"] = bool(continue_after)
        job["remote_retry_count"] = max(0, int(retry_count))
        job["remote_retry_limit"] = codebuddy_npc_remote_retry_limit
        job["active_stage"] = stage
        job["cancel_requested"] = False
        start_stage_timing(
            job,
            stage,
            reset=retry_count == 0,
            execution_target="remote_cnb",
        )
        job.pop("fallback_reason", None)
        job.pop("result_pending_since", None)
        job.pop("result_pending_polls", None)
        job["status"] = "running"
        retry_suffix = (
            f"（自动重试 {retry_count}/{codebuddy_npc_remote_retry_limit}）"
            if retry_count
            else ""
        )
        job["status_text"] = (
            f"已提交 CNB，{STAGE_NAMES[stage]}正在远程运行{retry_suffix}"
        )
        job["error"] = ""
        job["progress"] = round(STAGE_ORDER.index(stage) / len(STAGE_ORDER) * 100)
        return codebuddy_npc_jobs.save(job)

    def _complete_code_state_recorder(
        job: dict[str, Any],
        *,
        continue_after: bool,
    ) -> dict[str, Any]:
        completed = codebuddy_npc_stage_runner.complete_state_recorder(
            job_id=str(job["job_id"]),
            user_id=int(job["user_id"]),
        )
        if not continue_after:
            return completed
        return _submit_remote_npc_stage(
            completed,
            stage="final_editor",
            continue_after=True,
            retry_count=0,
        )

    def _state_recorder_uses_code(job: dict[str, Any]) -> bool:
        request_data = job.get("request") if isinstance(job.get("request"), dict) else {}
        try:
            episodes = max(1, int(request_data.get("episodes") or 1))
        except (TypeError, ValueError):
            episodes = 1
        return episodes < 40

    def _refresh_remote_npc_job_unlocked(job: dict[str, Any]) -> dict[str, Any]:
        job_id = str(job.get("job_id") or "")
        if bool(job.get("cancel_requested")):
            return job
        if (
            str(job.get("execution_target") or "") != "remote_cnb"
            or str(job.get("status") or "") not in {"running", "result_pending"}
            or codebuddy_npc_stage_runner.is_running(job_id)
        ):
            return job
        try:
            requested_build_sn = str((job.get("build") or {}).get("sn") or "")
            refreshed_job = codebuddy_npc_client.refresh(job)
            latest_job = codebuddy_npc_jobs.load(
                job_id,
                user_id=int(job["user_id"]),
            )
            latest_build_sn = str(
                ((latest_job or {}).get("build") or {}).get("sn") or ""
            )
            if (
                requested_build_sn
                and latest_build_sn
                and latest_build_sn != requested_build_sn
            ):
                return latest_job
            job = refreshed_job
            if str(job.get("status") or "") == "result_pending":
                pending_since = str(job.get("result_pending_since") or "").strip()
                if not pending_since:
                    pending_since = _now_iso()
                    job["result_pending_since"] = pending_since
                job["result_pending_polls"] = (
                    max(0, int(job.get("result_pending_polls") or 0)) + 1
                )
                try:
                    pending_at = datetime.fromisoformat(
                        pending_since.replace("Z", "+00:00")
                    )
                    if pending_at.tzinfo is None:
                        pending_at = pending_at.replace(tzinfo=timezone.utc)
                    pending_seconds = max(
                        0,
                        (datetime.now(timezone.utc) - pending_at).total_seconds(),
                    )
                except (TypeError, ValueError):
                    pending_seconds = 0
                if pending_seconds >= codebuddy_npc_result_pending_timeout:
                    job["status"] = "failed"
                    job["status_text"] = (
                        f"{STAGE_NAMES.get(str(job.get('remote_stage') or ''), '远程节点')}"
                        "云端已结束，但产物读取超时"
                    )
                    job["error"] = (
                        "CNB 构建已成功，但平台未能在限定时间内读取节点产物。"
                    )
                    job["active_stage"] = ""
            else:
                job.pop("result_pending_since", None)
                job.pop("result_pending_polls", None)
            job = codebuddy_npc_jobs.save(job)
            if (
                str(job.get("status") or "") == "failed"
                and str(job.get("remote_kind") or "") == "stage"
                and str(job.get("remote_stage") or "") in STAGE_ORDER
            ):
                failed_stage = str(job["remote_stage"])
                remote_error = str(job.get("error") or "CNB远程节点执行失败")
                if failed_stage == "state_recorder" and _state_recorder_uses_code(job):
                    return _complete_code_state_recorder(
                        job,
                        continue_after=bool(job.get("remote_continue_after")),
                    )
                retry_count = max(0, int(job.get("remote_retry_count") or 0))
                retry_limit = max(
                    0,
                    int(
                        job.get("remote_retry_limit")
                        if job.get("remote_retry_limit") is not None
                        else codebuddy_npc_remote_retry_limit
                    ),
                )
                retryable_remote_failure = not any(
                    marker in remote_error.lower()
                    for marker in ("cancel", "取消", "集数不完整", "实际集号")
                )
                if retryable_remote_failure and retry_count < retry_limit:
                    try:
                        job = _submit_remote_npc_stage(
                            job,
                            stage=failed_stage,
                            continue_after=bool(job.get("remote_continue_after")),
                            retry_count=retry_count + 1,
                        )
                        job["last_remote_error"] = remote_error
                        return codebuddy_npc_jobs.save(job)
                    except CodeBuddyNpcError as retry_exc:
                        remote_error = (
                            f"{remote_error}；自动重试提交失败：{retry_exc}"
                        )
                job = codebuddy_npc_stage_runner.start(
                    job_id=job_id,
                    user_id=int(job["user_id"]),
                    stage=failed_stage,
                    continue_after=bool(job.get("remote_continue_after")),
                )
                job["fallback_reason"] = remote_error
                job["status_text"] = (
                    f"CNB远程节点失败，已从断点切换本地兜底运行{STAGE_NAMES[failed_stage]}"
                )
                return codebuddy_npc_jobs.save(job)
            if (
                str(job.get("status") or "") == "stage_ready"
                and bool(job.get("remote_continue_after"))
                and str(job.get("remote_stage") or "") in STAGE_ORDER
            ):
                current_index = STAGE_ORDER.index(str(job["remote_stage"]))
                if current_index + 1 < len(STAGE_ORDER):
                    next_stage = STAGE_ORDER[current_index + 1]
                    try:
                        if next_stage == "state_recorder" and _state_recorder_uses_code(job):
                            job = _complete_code_state_recorder(
                                job,
                                continue_after=True,
                            )
                        else:
                            job = _submit_remote_npc_stage(
                                job,
                                stage=next_stage,
                                continue_after=True,
                                retry_count=0,
                            )
                    except CodeBuddyNpcError as remote_exc:
                        job = codebuddy_npc_stage_runner.start(
                            job_id=job_id,
                            user_id=int(job["user_id"]),
                            stage=next_stage,
                            continue_after=True,
                        )
                        job["fallback_reason"] = str(remote_exc)
                        job["status_text"] = (
                            f"CNB 暂不可用，已切换本地兜底运行{STAGE_NAMES[next_stage]}"
                        )
                        job = codebuddy_npc_jobs.save(job)
        except CodeBuddyNpcError as exc:
            job["status_text"] = "暂时无法读取 CNB 进度"
            job["poll_warning"] = str(exc)
            job = codebuddy_npc_jobs.save(job)
        except Exception as exc:
            logger.exception("codebuddy npc refresh failed: job_id=%s", job_id)
            job["poll_warning"] = str(exc)
            job = codebuddy_npc_jobs.save(job)
        return job

    def _refresh_remote_npc_job(job: dict[str, Any]) -> dict[str, Any]:
        job_id = str(job.get("job_id") or "")
        if not job_id:
            return job
        with codebuddy_npc_refresh_locks_guard:
            refresh_lock = codebuddy_npc_refresh_locks.setdefault(
                job_id,
                threading.RLock(),
            )
        with refresh_lock:
            latest_job = codebuddy_npc_jobs.load(
                job_id,
                user_id=int(job["user_id"]),
            )
            return _refresh_remote_npc_job_unlocked(latest_job or job)
    app.config["WORKFLOW_SPEC_PATH"] = workflow_spec_path or default_workflow_spec_path()
    app.config["SECRET_KEY"] = os.getenv("SECRET_KEY") or os.getenv("FLASK_SECRET_KEY") or "scriptmaker-dev-secret"
    force_secure_cookies = str(
        os.getenv("SCRIPTMAKER_SECURE_COOKIES", "false")
    ).strip().lower() in {"1", "true", "yes", "on"}
    app.config.update(
        SESSION_COOKIE_SECURE=force_secure_cookies,
        SESSION_COOKIE_HTTPONLY=True,
        SESSION_COOKIE_SAMESITE="Lax",
    )
    app.config["REGISTRATION_ENABLED"] = str(
        os.getenv("SCRIPTMAKER_REGISTRATION_ENABLED", "true")
    ).strip().lower() in {"1", "true", "yes", "on"}
    app.config["ADMIN_ONLY_LOGIN"] = str(
        os.getenv("SCRIPTMAKER_ADMIN_ONLY_LOGIN", "false")
    ).strip().lower() in {"1", "true", "yes", "on"}
    app.config["LOGIN_ALLOWED_USERNAMES"] = {
        username.strip()
        for username in str(os.getenv("SCRIPTMAKER_LOGIN_ALLOWED_USERNAMES", "")).split(",")
        if username.strip()
    }
    app.config["LOGIN_ALLOWED_USER_IDS"] = {
        int(user_id.strip())
        for user_id in str(os.getenv("SCRIPTMAKER_LOGIN_ALLOWED_USER_IDS", "")).split(",")
        if user_id.strip().isdigit()
    }

    @app.before_request
    def _select_workflow_line():
        g.workflow_line_token = set_workflow_line(
            request.headers.get(WORKFLOW_LINE_HEADER, PRODUCTION_WORKFLOW_LINE)
        )

    @app.after_request
    def _reset_selected_workflow_line(response):
        token = getattr(g, "workflow_line_token", None)
        if token is not None:
            reset_workflow_line(token)
            g.workflow_line_token = None
        return response

    @app.after_request
    def _apply_security_headers(response):
        response.headers.setdefault("X-Content-Type-Options", "nosniff")
        response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
        response.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
        response.headers.setdefault("Permissions-Policy", "camera=(), microphone=(), geolocation=()")
        if request.is_secure:
            response.headers.setdefault(
                "Strict-Transport-Security",
                "max-age=31536000; includeSubDomains",
            )
        if request.method == "GET" and response.mimetype == "text/html":
            response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
        return response

    def _json_ok(**payload):
        return jsonify({"success": True, **payload})

    def _sanitize_error_message(
        message: str,
        *,
        status: int,
        fallback: str | None = None,
    ) -> str:
        text = str(message or "").strip()
        if status >= 500:
            return fallback or "服务暂时不可用，请稍后重试。"
        if not text:
            return fallback or "请求未能完成，请稍后重试。"
        technical_markers = (
            "fastgpt",
            "traceback",
            "http ",
            "response.text",
            "url:",
            "url：",
            "requests.",
            "exception",
            "typeerror",
            "keyerror",
            "bad gateway",
            "failed to fetch",
            "json",
            "校验失败",
            "无法转换",
            ".py",
        )
        lowered = text.lower()
        if len(text) > 120 or any(marker in lowered for marker in technical_markers):
            return fallback or "请求未能完成，请稍后重试。"
        return text

    def _json_error(message: str, status: int = 400, *, fallback: str | None = None, detail: dict | None = None):
        public_message = _sanitize_error_message(message, status=status, fallback=fallback)
        payload = {"success": False, "message": public_message}
        if isinstance(detail, dict) and detail:
            payload["detail"] = _strip_raw_fastgpt_fields(copy.deepcopy(detail))
        return jsonify(payload), status


    def _script_audit_ecg_extract_text(value, *, limit: int = 200000) -> str:
        texts = []
        priority_keys = (
            "answer_text",
            "answerText",
            "text",
            "textOutput",
            "output",
            "response",
            "content",
            "result",
            "data",
            "raw",
            "message",
        )

        def visit(item, depth: int = 0):
            if depth > 5:
                return
            if isinstance(item, str):
                if item.strip():
                    texts.append(item)
                return
            if isinstance(item, dict):
                for key in priority_keys:
                    if key in item:
                        visit(item.get(key), depth + 1)
                for key, child in item.items():
                    if key not in priority_keys:
                        visit(child, depth + 1)
                return
            if isinstance(item, list):
                for child in item[:50]:
                    visit(child, depth + 1)

        visit(value)
        if not texts:
            return ""
        return max(texts, key=len)[:limit]

    def _script_audit_ecg_should_try(tool_key: str, parsed: dict | None, raw_text: str) -> bool:
        if isinstance(parsed, dict) and parsed.get("schema_version") in {
            SCRIPT_AUDIT_ECG_SCHEMA_VERSION,
            SCRIPT_AUDIT_ECG_SCHEMA_VERSION_V3,
            SCRIPT_AUDIT_COMPACT_SCHEMA_VERSION,
        }:
            return True
        key_text = str(tool_key or "").lower()
        keywords = (
            "hot_review",
            "爆款文审核",
            "剧本审核",
            "心电图",
            "script_audit",
            "audit_ecg",
            "script_audit_ecg",
        )
        return (
            any(keyword.lower() in key_text for keyword in keywords)
            or SCRIPT_AUDIT_ECG_SCHEMA_VERSION in raw_text
            or SCRIPT_AUDIT_ECG_SCHEMA_VERSION_V3 in raw_text
            or SCRIPT_AUDIT_COMPACT_SCHEMA_VERSION in raw_text
        )


    def _maybe_enrich_script_audit_ecg_tool_result(
        tool_key: str,
        result: dict,
        flattened: dict,
    ) -> tuple[dict, dict]:
        if not isinstance(result, dict):
            return result, flattened

        raw_text = _script_audit_ecg_extract_text(result)
        parsed, parse_warnings = parse_model_json_loose(result)

        if not _script_audit_ecg_should_try(
            tool_key,
            parsed,
            raw_text,
        ):
            return result, flattened

        enriched_result = dict(result)
        enriched_flattened = dict(flattened)

        enriched_result.setdefault(
            "answer_text",
            raw_text,
        )
        enriched_flattened.setdefault(
            "answer_text",
            raw_text,
        )

        compact_candidate = None
        compact_parse_warnings = []
        if (
            isinstance(parsed, dict)
            and parsed.get("schema_version") == SCRIPT_AUDIT_COMPACT_SCHEMA_VERSION
        ):
            compact_candidate = parsed
        else:
            compact_sources = []
            if raw_text:
                compact_sources.append(raw_text)
            try:
                compact_sources.append(json.dumps(result, ensure_ascii=False, default=str))
            except Exception:
                pass
            for compact_source in compact_sources:
                try:
                    compact_candidate = parse_compact_audit_json(compact_source)
                    break
                except Exception as exc:
                    compact_parse_warnings.append(str(exc))

        if (
            isinstance(compact_candidate, dict)
            and (
                compact_candidate.get("schema_version") == SCRIPT_AUDIT_COMPACT_SCHEMA_VERSION
                or "global_dimensions" in compact_candidate
                or "episodes" in compact_candidate
            )
        ):
            audit, normalize_warnings = normalize_compact_audit_payload(compact_candidate)
            validation_warnings = validate_compact_audit_schema(audit)
            visualization = build_audit_visualization_payload(audit)
            view = build_script_audit_view_model(audit)
            warnings = [
                *(parse_warnings or []),
                *compact_parse_warnings[:1],
                *normalize_warnings,
                *validation_warnings,
            ]
            for target in (
                enriched_result,
                enriched_flattened,
            ):
                target["result_type"] = "script_audit_ecg"
                target["resultType"] = "script_audit_ecg"
                target["schema_version"] = SCRIPT_AUDIT_COMPACT_SCHEMA_VERSION
                target["parsed"] = True
                target["audit"] = audit
                target["visualization"] = visualization
                target["view"] = view
                target["warnings"] = warnings
                target["parse_warnings"] = warnings
                target["answer_text"] = raw_text

            return enriched_result, enriched_flattened

        def looks_like_audit_payload(value) -> bool:
            if not isinstance(value, dict):
                return False

            if (
                value.get("schema_version")
                in {
                    SCRIPT_AUDIT_ECG_SCHEMA_VERSION,
                    SCRIPT_AUDIT_ECG_SCHEMA_VERSION_V3,
                    SCRIPT_AUDIT_COMPACT_SCHEMA_VERSION,
                }
            ):
                return True

            overall = value.get("overall")
            if not isinstance(overall, dict):
                return False

            structural_keys = (
                "dimension_scores",
                "segments",
                "ecg",
                "global_review",
                "episode_reviews",
                "cross_episode_analysis",
                "episode_summaries",
                "satisfying_points",
                "key_issues",
                "risk_scan",
                "rewrite_plan",
                "visualization_config",
            )

            return any(
                key in value
                for key in structural_keys
            )

        candidate = parsed

        # 兼容模型在标准结果外面再包 audit/data/result/output。
        if isinstance(candidate, dict):
            for nested_key in (
                "audit",
                "data",
                "result",
                "output",
            ):
                nested = candidate.get(nested_key)
                if looks_like_audit_payload(nested):
                    candidate = nested
                    break

        if not looks_like_audit_payload(candidate):
            warnings = (
                parse_warnings
                or ["未解析到可识别的剧本心电图 JSON"]
            )
            fallback_audit, normalize_warnings = normalize_script_audit_ecg(
                fallback_audit_from_text(raw_text, warnings=warnings),
                raw_answer_text=raw_text,
            )
            fallback_view = build_script_audit_view_model(fallback_audit)
            for target in (
                enriched_result,
                enriched_flattened,
            ):
                target["result_type"] = "script_audit_ecg"
                target["resultType"] = "script_audit_ecg"
                target["parsed"] = False
                target["audit"] = fallback_audit
                target["view"] = fallback_view
                target["parse_warnings"] = [*warnings, *normalize_warnings]
                target["answer_text"] = raw_text

            return enriched_result, enriched_flattened

        compatibility_warnings = list(
            parse_warnings or []
        )

        if (
            candidate.get("schema_version")
            not in {SCRIPT_AUDIT_ECG_SCHEMA_VERSION, SCRIPT_AUDIT_ECG_SCHEMA_VERSION_V3}
        ):
            compatibility_warnings.append(
                "模型输出缺少或改变了 schema_version，"
                "已根据审核结果结构兼容解析。"
            )

        audit, normalize_warnings = (
            normalize_script_audit_ecg(
                candidate,
                raw_answer_text=raw_text,
            )
        )

        view = build_script_audit_view_model(audit)

        warnings = [
            *compatibility_warnings,
            *normalize_warnings,
        ]

        for target in (
            enriched_result,
            enriched_flattened,
        ):
            target["result_type"] = "script_audit_ecg"
            target["resultType"] = "script_audit_ecg"
            target["parsed"] = True
            target["audit"] = audit
            target["view"] = view
            target["parse_warnings"] = warnings
            target["answer_text"] = raw_text

        return enriched_result, enriched_flattened

    raw_fastgpt_response_keys = {
        "responseData",
        "choices",
        "reasoningText",
        "historyPreview",
        "newVariables",
        "updateVarResult",
        "raw_stage_responses",
        "raw_output",
        "raw",
        "answerText",
        "usage",
        "debug",
        "logs",
        "cache",
    }
    framework_stage_runs: dict[tuple[int, str, str], float] = {}
    framework_stage_runs_lock = threading.Lock()
    FRAMEWORK_STAGE_RUN_TTL_SECONDS = 4 * 60 * 60

    def _now_iso() -> str:
        return datetime.now(timezone.utc).astimezone().isoformat()

    def _stage12_debug_safe_name(value) -> str:
        text = str(value or "").strip()
        invalid = '<>:"/\\|?*'
        text = "".join("_" if (char in invalid or ord(char) < 32) else char for char in text)
        text = "_".join(text.split()).strip("._ ")
        return text[:80] or "未命名项目"

    def _stage12_debug_project_title(data: dict, framework_asset: dict | None) -> str:
        candidates = [
            data.get("project_title"),
            data.get("project_name"),
            data.get("title"),
            data.get("source_title"),
            (framework_asset or {}).get("project_title") if isinstance(framework_asset, dict) else "",
            (framework_asset or {}).get("title") if isinstance(framework_asset, dict) else "",
            (framework_asset or {}).get("source_title") if isinstance(framework_asset, dict) else "",
        ]
        package = data.get("framework_plan_package") if isinstance(data.get("framework_plan_package"), dict) else {}
        basic_config = package.get("basic_config") if isinstance(package.get("basic_config"), dict) else {}
        candidates.extend(
            [
                package.get("project_title"),
                package.get("title"),
                basic_config.get("project_title"),
                basic_config.get("title"),
            ]
        )
        for item in candidates:
            text = str(item or "").strip()
            if text:
                return text
        project_id = str(data.get("project_id") or data.get("framework_asset_id") or "").strip()
        return f"project_{project_id}" if project_id else "未命名项目"

    def _stage12_debug_preview(value, *, limit: int = 240):
        if value is None:
            return ""
        if isinstance(value, str):
            text = value
        else:
            try:
                text = json.dumps(value, ensure_ascii=False, default=str, separators=(",", ":"))
            except Exception:
                text = str(value)
        text = " ".join(str(text).split())
        if len(text) > limit:
            return text[:limit] + "..."
        return text

    def _stage12_debug_summary(value, *, preview_limit: int = 240):
        if value is None:
            return {"type": "null", "exists": False, "length": 0, "preview": ""}
        if isinstance(value, str):
            return {
                "type": "str",
                "exists": bool(value),
                "length": len(value),
                "preview": _stage12_debug_preview(value, limit=preview_limit),
            }
        if isinstance(value, dict):
            return {
                "type": "dict",
                "exists": bool(value),
                "length": len(value),
                "keys": sorted(str(key) for key in value.keys())[:40],
                "json_length": len(json.dumps(value, ensure_ascii=False, default=str)),
                "preview": _stage12_debug_preview(value, limit=preview_limit),
            }
        if isinstance(value, list):
            return {
                "type": "list",
                "exists": bool(value),
                "length": len(value),
                "json_length": len(json.dumps(value, ensure_ascii=False, default=str)),
                "preview": _stage12_debug_preview(value[:2], limit=preview_limit),
            }
        text = str(value)
        return {
            "type": type(value).__name__,
            "exists": bool(text),
            "length": len(text),
            "preview": _stage12_debug_preview(text, limit=preview_limit),
        }

    def _stage12_debug_dir(data: dict, framework_asset: dict | None) -> Path:
        title = _stage12_debug_project_title(data, framework_asset)
        path = Path(__file__).resolve().parents[2] / "cache" / _stage12_debug_safe_name(title) / "framework_to_script" / "stage12"
        path.mkdir(parents=True, exist_ok=True)
        return path

    def _write_stage12_debug_file(
        debug_record: dict,
        *,
        data: dict,
        framework_asset: dict | None,
        success: bool = False,
    ) -> str:
        try:
            start_episode = debug_record.get("batch_start_episode") or "unknown"
            end_episode = debug_record.get("batch_end_episode") or start_episode
            if not debug_record.get("_debug_timestamp"):
                debug_record["_debug_timestamp"] = datetime.now(timezone.utc).astimezone().strftime("%Y%m%d-%H%M%S-%f")[:-3]
            filename = (
                f"stage12_episodes_{start_episode}_{end_episode}_success.json"
                if success
                else f"stage12_episodes_{start_episode}_{end_episode}_{debug_record['_debug_timestamp']}.json"
            )
            path = _stage12_debug_dir(data, framework_asset) / filename
            public_record = {key: value for key, value in debug_record.items() if not str(key).startswith("_")}
            public_record["debug_path"] = str(path)
            path.write_text(json.dumps(public_record, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
            return str(path)
        except Exception:
            logger.exception("framework-to-script stage12 debug file write failed")
            return ""

    def _stage12_fastgpt_debug_summary(client, stage_name: str) -> dict:
        try:
            info = client.get_last_stage_debug_info(stage_name)
        except Exception:
            info = {}
        if not isinstance(info, dict):
            return {}
        return {
            "status": info.get("status"),
            "http_status": info.get("http_status"),
            "http_reason": info.get("http_reason"),
            "payload_stats": info.get("payload_stats"),
            "payload_stats_before_compact": info.get("payload_stats_before_compact"),
            "payload_compaction": info.get("payload_compaction"),
            "raw_output_source": info.get("raw_output_source"),
            "matched_fields": info.get("matched_fields"),
            "matched_aliases": info.get("matched_aliases"),
            "missing_fields": info.get("missing_fields"),
            "candidate_sources": info.get("candidate_sources"),
            "probable_truncated_json": info.get("probable_truncated_json"),
            "response_preview": _stage12_debug_preview(info.get("response_preview"), limit=600),
            "answer_text_preview": _stage12_debug_preview(info.get("answer_text_preview"), limit=600),
            "output_keys": info.get("output_keys"),
            "last_failure_reason": info.get("last_failure_reason"),
        }

    def _strip_raw_fastgpt_fields(value):
        if isinstance(value, list):
            return [_strip_raw_fastgpt_fields(item) for item in value]
        if not isinstance(value, dict):
            return value
        result = {}
        for key, item in value.items():
            if str(key) in raw_fastgpt_response_keys:
                continue
            result[key] = _strip_raw_fastgpt_fields(item)
        return result

    def _framework_state_from_project(project: dict) -> dict:
        if not isinstance(project, dict):
            return {}
        artifacts = project.get("artifacts") if isinstance(project.get("artifacts"), dict) else {}
        input_payload = project.get("input_payload") if isinstance(project.get("input_payload"), dict) else {}
        state = (
            project.get("framework_planner_state")
            or artifacts.get("framework_planner_state")
            or input_payload.get("framework_planner_state")
            or {}
        )
        state = copy.deepcopy(state) if isinstance(state, dict) else {}
        for key in (
            "basic_config",
            "source_brief",
            "worldview_plan",
            "character_plan",
            "beat_checkpoint_timeline",
            "checkpoint_explanation",
            "character_storylines",
            "storyline_decisions",
            "adaptation_guide",
            "framework_plan_package",
            "stage_outputs",
            "stageOutputs",
        ):
            if key not in state and key in artifacts:
                state[key] = copy.deepcopy(artifacts.get(key))
            if key not in state and key in project:
                state[key] = copy.deepcopy(project.get(key))
            if key not in state and key in input_payload:
                state[key] = copy.deepcopy(input_payload.get(key))
        return state

    def _framework_stage_outputs(framework_state: dict) -> dict:
        package = framework_state.get("framework_plan_package") if isinstance(framework_state.get("framework_plan_package"), dict) else {}
        direct_outputs = framework_state.get("stage_outputs") if isinstance(framework_state.get("stage_outputs"), dict) else {}
        camel_outputs = framework_state.get("stageOutputs") if isinstance(framework_state.get("stageOutputs"), dict) else {}
        outputs = {
            "source_brief": framework_state.get("source_brief") or package.get("source_brief") or direct_outputs.get("source_brief") or camel_outputs.get("source_brief") or {},
            "worldview_plan": framework_state.get("worldview_plan") or package.get("worldview_plan") or direct_outputs.get("worldview_plan") or camel_outputs.get("worldview_plan") or {},
            "character_plan": framework_state.get("character_plan") or package.get("character_plan") or direct_outputs.get("character_plan") or camel_outputs.get("character_plan") or {},
            "beat_checkpoint_timeline": framework_state.get("beat_checkpoint_timeline") or package.get("beat_checkpoint_timeline") or direct_outputs.get("beat_checkpoint_timeline") or camel_outputs.get("beat_checkpoint_timeline") or [],
            "checkpoint_explanation": framework_state.get("checkpoint_explanation") or package.get("checkpoint_explanation") or direct_outputs.get("checkpoint_explanation") or camel_outputs.get("checkpoint_explanation") or {},
            "character_storylines": framework_state.get("character_storylines") or package.get("character_storylines") or direct_outputs.get("character_storylines") or camel_outputs.get("character_storylines") or [],
            "storyline_decisions": framework_state.get("storyline_decisions") or package.get("storyline_decisions") or direct_outputs.get("storyline_decisions") or camel_outputs.get("storyline_decisions") or [],
            "adaptation_guide": framework_state.get("adaptation_guide") or package.get("adaptation_guide") or direct_outputs.get("adaptation_guide") or camel_outputs.get("adaptation_guide") or {},
        }
        for key, value in {**direct_outputs, **camel_outputs}.items():
            if _framework_value_present(value):
                outputs.setdefault(key, value)
        return _strip_raw_fastgpt_fields(outputs)

    def _hydrate_framework_planner_stage_payload(data: dict, user_id: int) -> dict:
        """Recover saved upstream fields when a stale browser cache sends an incomplete stage payload."""
        if not isinstance(data, dict):
            return {}
        hydrated = dict(data)
        try:
            project_id = int(str(hydrated.get("project_id") or hydrated.get("asset_id") or "").strip())
        except Exception:
            project_id = 0
        if project_id <= 0:
            return hydrated

        snapshot = task_manager.get_project_snapshot(project_id, user_id=user_id, public_view=False)
        if not isinstance(snapshot, dict):
            return hydrated
        framework_state = _framework_state_from_project(snapshot)
        stage_outputs = _framework_stage_outputs(framework_state)
        saved_basic = framework_state.get("basic_config") if isinstance(framework_state.get("basic_config"), dict) else {}
        saved_fields = {
            "basic_config": saved_basic,
            "locked_basic_config": saved_basic,
            "source_brief": stage_outputs.get("source_brief"),
            "worldview_plan": stage_outputs.get("worldview_plan"),
            "character_plan": stage_outputs.get("character_plan"),
            "beat_checkpoint_timeline": stage_outputs.get("beat_checkpoint_timeline"),
            "checkpoint_explanation": stage_outputs.get("checkpoint_explanation"),
            "character_storylines": stage_outputs.get("character_storylines"),
            "storyline_decisions": stage_outputs.get("storyline_decisions"),
            "adaptation_guide": stage_outputs.get("adaptation_guide"),
        }
        recovered: list[str] = []
        for key, value in saved_fields.items():
            if not _framework_value_present(hydrated.get(key)) and _framework_value_present(value):
                hydrated[key] = copy.deepcopy(value)
                recovered.append(key)
        if recovered:
            logger.warning(
                "framework planner request recovered saved upstream fields: project_id=%s fields=%s",
                project_id,
                recovered,
            )
        return hydrated

    def _framework_project_history_name(project_id: object, requested_name: object = "") -> str:
        requested = str(requested_name or "").strip()
        if requested and requested not in {"未命名项目", "unsaved", "draft", "null", "undefined"}:
            return requested
        try:
            numeric_project_id = int(str(project_id or "").strip())
        except Exception:
            numeric_project_id = 0
        if numeric_project_id <= 0:
            return requested
        snapshot = task_manager.get_project_snapshot(numeric_project_id, user_id=_require_user_id(), public_view=False)
        if not isinstance(snapshot, dict):
            return requested
        framework_state = _framework_state_from_project(snapshot)
        basic = framework_state.get("basic_config") if isinstance(framework_state.get("basic_config"), dict) else {}
        input_payload = snapshot.get("input_payload") if isinstance(snapshot.get("input_payload"), dict) else {}
        for value in (
            framework_state.get("project_title"),
            snapshot.get("title"),
            basic.get("project_title"),
            basic.get("source_title"),
            input_payload.get("title"),
        ):
            text = str(value or "").strip()
            if text and text != "未命名项目":
                return text
        return requested

    def _framework_import_package(framework_state: dict, stage_outputs: dict | None = None) -> dict:
        package = framework_state.get("framework_plan_package") if isinstance(framework_state.get("framework_plan_package"), dict) else {}
        if package:
            return copy.deepcopy(package)
        outputs = stage_outputs if isinstance(stage_outputs, dict) else _framework_stage_outputs(framework_state)
        synthesized = {
            "source_brief": outputs.get("source_brief") or framework_state.get("source_brief") or {},
            "worldview_plan": outputs.get("worldview_plan") or framework_state.get("worldview_plan") or {},
            "character_plan": outputs.get("character_plan") or framework_state.get("character_plan") or {},
            "beat_checkpoint_timeline": outputs.get("beat_checkpoint_timeline") or framework_state.get("beat_checkpoint_timeline") or [],
            "checkpoint_explanation": outputs.get("checkpoint_explanation") or framework_state.get("checkpoint_explanation") or {},
            "character_storylines": outputs.get("character_storylines") or framework_state.get("character_storylines") or [],
            "storyline_decisions": outputs.get("storyline_decisions") or framework_state.get("storyline_decisions") or [],
            "adaptation_guide": outputs.get("adaptation_guide") or framework_state.get("adaptation_guide") or {},
            "basic_config": framework_state.get("basic_config") if isinstance(framework_state.get("basic_config"), dict) else {},
        }
        meaningful_keys = (
            "source_brief",
            "worldview_plan",
            "character_plan",
            "beat_checkpoint_timeline",
            "checkpoint_explanation",
            "character_storylines",
            "storyline_decisions",
            "adaptation_guide",
        )
        if any(_framework_value_present(synthesized.get(key)) for key in meaningful_keys):
            synthesized["package_id"] = framework_state.get("project_id") or "synthesized_framework_asset"
            synthesized["import_package_synthesized"] = True
            return _strip_raw_fastgpt_fields(synthesized)
        return {}

    def _framework_value_present(value) -> bool:
        if value is None:
            return False
        if isinstance(value, str):
            return bool(value.strip())
        if isinstance(value, (list, tuple, set, dict)):
            return bool(value)
        return True

    def _framework_basic_config_from_stage_payload(data: dict, existing_state: dict) -> dict:
        basic = data.get("basic_config") if isinstance(data.get("basic_config"), dict) else {}
        if basic:
            return copy.deepcopy(basic)
        existing_basic = existing_state.get("basic_config") if isinstance(existing_state.get("basic_config"), dict) else {}
        merged = copy.deepcopy(existing_basic)
        for key in (
            "project_title",
            "mode",
            "source_text",
            "source_title",
            "target_format",
            "season_count",
            "episodes_per_season",
            "episode_word_count",
            "adaptation_direction",
            "user_constraints",
            "user_requirements",
        ):
            if key in data and _framework_value_present(data.get(key)):
                merged[key] = copy.deepcopy(data.get(key))
        return merged

    def _framework_stage_output_fields(stage_no: str) -> tuple[str, ...]:
        return {
            "01": ("source_brief",),
            "02": ("worldview_plan",),
            "03": ("character_plan",),
            "04": ("beat_checkpoint_timeline", "checkpoint_explanation"),
            "05": ("character_storylines", "storyline_decisions"),
            "06": ("adaptation_guide",),
            "07": ("framework_plan_package", "validation_report"),
        }.get(str(stage_no or "").zfill(2), ())

    def _merge_framework_stage_state(existing_state: dict, stage_no: str, stage_output: dict) -> dict:
        stage_key = _framework_planner_stage_key(stage_no)
        stage_state = copy.deepcopy(existing_state.get("stage_state") if isinstance(existing_state.get("stage_state"), dict) else {})
        if not stage_key:
            return stage_state
        has_output = any(_framework_value_present(stage_output.get(field)) for field in _framework_stage_output_fields(stage_no))
        if not has_output:
            return stage_state
        current = stage_state.get(stage_key) if isinstance(stage_state.get(stage_key), dict) else {}
        stage_state[stage_key] = {
            **current,
            "locked": False,
            "confirmed": False,
            "status": "generated",
            "stageCommitted": True,
            "stageDraftDirty": False,
        }
        sequence = ("basic", "worldview", "character", "beat", "storylines", "guide", "package")
        if stage_key in sequence:
            index = sequence.index(stage_key)
            if index + 1 < len(sequence):
                next_key = sequence[index + 1]
                next_stage = stage_state.get(next_key) if isinstance(stage_state.get(next_key), dict) else {}
                if not next_stage.get("confirmed"):
                    stage_state[next_key] = {**next_stage, "locked": False, "status": "idle"}
        return stage_state

    def _autosave_framework_planner_stage(
        *,
        user_id: int,
        stage: str,
        request_payload: dict,
        stage_payload: dict,
    ) -> dict | None:
        if not isinstance(request_payload, dict) or not isinstance(stage_payload, dict):
            return None
        stage_no = str(stage or "").zfill(2)
        stage_output = stage_payload.get("data") if isinstance(stage_payload.get("data"), dict) else {}
        if not stage_output:
            return None

        raw_project_id = (
            request_payload.get("project_id")
            or request_payload.get("asset_id")
            or request_payload.get("source_framework_project_id")
        )
        existing_state: dict = {}
        if raw_project_id:
            try:
                snapshot = task_manager.get_project_snapshot(int(raw_project_id), user_id=user_id, public_view=False)
            except Exception:
                snapshot = None
            if snapshot and str(snapshot.get("asset_kind") or "").strip() == "framework_planner":
                existing_state = _framework_state_from_project(snapshot)

        save_payload: dict = copy.deepcopy(existing_state)
        if raw_project_id:
            save_payload["project_id"] = raw_project_id
        save_payload["basic_config"] = _framework_basic_config_from_stage_payload(request_payload, existing_state)
        save_payload["project_title"] = (
            request_payload.get("project_title")
            or save_payload.get("project_title")
            or save_payload.get("title")
            or save_payload.get("basic_config", {}).get("project_title")
            or save_payload.get("basic_config", {}).get("source_title")
            or "未命名框架策划"
        )
        save_payload["title"] = save_payload["project_title"]

        for field in (
            "source_brief",
            "worldview_plan",
            "character_plan",
            "beat_checkpoint_timeline",
            "checkpoint_explanation",
            "character_storylines",
            "storyline_decisions",
            "adaptation_guide",
            "framework_plan_package",
            "validation_report",
        ):
            if field in request_payload and _framework_value_present(request_payload.get(field)):
                save_payload[field] = copy.deepcopy(request_payload.get(field))
            elif field not in save_payload:
                save_payload[field] = [] if field in {"beat_checkpoint_timeline", "character_storylines", "storyline_decisions"} else {}
            if field in stage_output and _framework_value_present(stage_output.get(field)):
                save_payload[field] = copy.deepcopy(stage_output.get(field))

        display_texts = save_payload.get("display_texts") if isinstance(save_payload.get("display_texts"), dict) else {}
        if isinstance(request_payload.get("display_texts"), dict):
            display_texts.update(copy.deepcopy(request_payload["display_texts"]))
        if _framework_value_present(stage_payload.get("display_text")):
            display_texts[stage_no] = stage_payload.get("display_text")
        save_payload["display_texts"] = display_texts

        for field in (
            "prompt_preferences",
            "preference_snapshot",
            "stage_prompts",
            "user_knowledge_stage_prompts",
            "selected_preference_tag_ids",
            "selected_preference_tags",
            "user_edit_history",
        ):
            if field in request_payload and _framework_value_present(request_payload.get(field)):
                save_payload[field] = copy.deepcopy(request_payload.get(field))
            elif field not in save_payload:
                save_payload[field] = [] if field in {"selected_preference_tag_ids", "selected_preference_tags", "user_edit_history"} else {}

        merged_stage_prompts = _merge_stage_prompts_non_empty(
            save_payload.get("stage_prompts"),
            save_payload.get("user_knowledge_stage_prompts"),
            (save_payload.get("prompt_preferences") or {}).get("stage_prompts") if isinstance(save_payload.get("prompt_preferences"), dict) else {},
        )
        save_payload["stage_prompts"] = merged_stage_prompts
        save_payload["user_knowledge_stage_prompts"] = merged_stage_prompts
        prompt_preferences = save_payload.get("prompt_preferences") if isinstance(save_payload.get("prompt_preferences"), dict) else {}
        prompt_preferences = dict(prompt_preferences)
        prompt_preferences["stage_prompts"] = merged_stage_prompts
        save_payload["prompt_preferences"] = prompt_preferences
        if isinstance(save_payload.get("framework_plan_package"), dict):
            save_payload["framework_plan_package"]["stage_prompts"] = copy.deepcopy(merged_stage_prompts)
            save_payload["framework_plan_package"]["prompt_preferences"] = {
                **(save_payload["framework_plan_package"].get("prompt_preferences") if isinstance(save_payload["framework_plan_package"].get("prompt_preferences"), dict) else {}),
                "stage_prompts": copy.deepcopy(merged_stage_prompts),
            }

        save_payload["stage_state"] = _merge_framework_stage_state(save_payload, stage_no, stage_output)
        current_view = {
            "01": "basic",
            "02": "worldview",
            "03": "character",
            "04": "beat_timeline",
            "05": "storylines",
            "06": "guide",
            "07": "package",
        }.get(stage_no)
        save_payload["current_view"] = current_view or save_payload.get("current_view") or "basic"
        asset_state = save_payload.get("asset_state") if isinstance(save_payload.get("asset_state"), dict) else {}
        save_payload["asset_state"] = {
            **asset_state,
            "asset_kind": "framework_planner",
            "asset_type": "framework",
            "project_id": raw_project_id or asset_state.get("project_id"),
            "asset_id": raw_project_id or asset_state.get("asset_id"),
            "status": "completed" if stage_no == "07" and _framework_value_present(save_payload.get("framework_plan_package")) else "in_progress",
            "current_stage": _framework_planner_stage_key(stage_no) or asset_state.get("current_stage") or "framework_planner",
            "last_action": f"autosave:stage{stage_no}",
        }

        asset = task_manager.save_framework_planner_asset(user_id=user_id, payload=save_payload)
        logger.info(
            "framework planner stage autosaved: user_id=%s project_id=%s stage=%s output_fields=%s",
            user_id,
            asset.get("project_id") if isinstance(asset, dict) else raw_project_id,
            stage_no,
            sorted(stage_output.keys()),
        )
        return asset

    def _save_new_workflow_test_stage_output(
        *,
        user_id: int,
        project_id: Any,
        stage: str,
        output: Any,
    ) -> None:
        try:
            numeric_project_id = int(str(project_id or "").strip())
        except Exception:
            return
        if numeric_project_id <= 0 or output in (None, "", [], {}):
            return
        snapshot = task_manager.get_project_snapshot(
            numeric_project_id,
            user_id=user_id,
            public_view=False,
        )
        if not snapshot:
            return
        artifacts = snapshot.setdefault("artifacts", {})
        if not isinstance(artifacts, dict):
            artifacts = {}
            snapshot["artifacts"] = artifacts
        test_state = artifacts.get("new_workflow_test_state")
        if not isinstance(test_state, dict):
            test_state = {}
        outputs = test_state.get("outputs")
        if not isinstance(outputs, dict):
            outputs = {}
        outputs[str(stage or "").zfill(2)] = _strip_raw_fastgpt_fields(copy.deepcopy(output))
        test_state["outputs"] = outputs
        test_state["updated_at"] = _now_iso()
        artifacts["new_workflow_test_state"] = test_state
        snapshot["updated_at"] = test_state["updated_at"]
        record = task_manager._projects.get(numeric_project_id)
        if record:
            with record.lock:
                record.snapshot = snapshot
            task_manager._persist_snapshot(record)
        else:
            task_manager._persist_project_snapshot_data(numeric_project_id, snapshot)

    def _framework_asset_payload(project: dict, *, include_detail: bool) -> dict:
        def _safe_positive_int(value, default=0):
            try:
                number = int(value)
                return number if number > 0 else default
            except Exception:
                return default

        framework_state = _framework_state_from_project(project)
        basic = framework_state.get("basic_config") if isinstance(framework_state.get("basic_config"), dict) else {}
        stage_outputs = _framework_stage_outputs(framework_state)
        package = _framework_import_package(framework_state, stage_outputs)
        artifacts = project.get("artifacts") if isinstance(project.get("artifacts"), dict) else {}
        workspace_state = artifacts.get("framework_to_script_state") if isinstance(artifacts.get("framework_to_script_state"), dict) else {}
        script_stages = workspace_state.get("scriptStages") if isinstance(workspace_state.get("scriptStages"), dict) else {}
        workspace_stage_outputs = workspace_state.get("stageOutputs") if isinstance(workspace_state.get("stageOutputs"), dict) else {}
        stage_drafts = workspace_state.get("stageDrafts") if isinstance(workspace_state.get("stageDrafts"), dict) else {}

        def _script_stage_has_content(stage_key: str) -> bool:
            value = script_stages.get(stage_key)
            return isinstance(value, dict) and _framework_value_present(value)

        completed_script_stage_numbers = [
            key.replace("stage", "")
            for key in ("stage08", "stage09", "stage10", "stage11", "stage12")
            if _script_stage_has_content(key)
        ]
        has_framework_to_script_state = bool(script_stages or workspace_stage_outputs or workspace_state.get("runningStage") or stage_drafts)
        latest_script_stage = completed_script_stage_numbers[-1] if completed_script_stage_numbers else ""
        input_payload = project.get("input_payload") if isinstance(project.get("input_payload"), dict) else {}
        metadata = project.get("metadata") if isinstance(project.get("metadata"), dict) else {}
        preference_snapshot = (
            framework_state.get("preference_snapshot")
            or project.get("preference_snapshot")
            or metadata.get("preference_snapshot")
            or artifacts.get("preference_snapshot")
            or input_payload.get("preference_snapshot")
            or {}
        )
        if not isinstance(preference_snapshot, dict):
            preference_snapshot = {}
        stage_prompts = _merge_stage_prompts_non_empty(
            framework_state.get("stage_prompts") if isinstance(framework_state.get("stage_prompts"), dict) else {},
            framework_state.get("user_knowledge_stage_prompts") if isinstance(framework_state.get("user_knowledge_stage_prompts"), dict) else {},
            (framework_state.get("prompt_preferences") or {}).get("stage_prompts") if isinstance(framework_state.get("prompt_preferences"), dict) else {},
            _stage_prompts_from_snapshot(preference_snapshot),
        )
        title = str(
            framework_state.get("project_title")
            or project.get("title")
            or basic.get("project_title")
            or basic.get("source_title")
            or "未命名框架资产"
        ).strip()
        source_title = str(basic.get("source_title") or title).strip()
        summary_source = (
            package.get("summary")
            or package.get("logline")
            or package.get("core_summary")
            or basic.get("adaptation_direction")
            or basic.get("source_text")
            or project.get("summary")
            or ""
        )
        summary = str(summary_source or "").replace("\r", "\n").strip()
        if len(summary) > 220:
            summary = summary[:220].rstrip() + "..."
        package_ready = bool(package)
        stage_output_keys = [key for key, value in (stage_outputs or {}).items() if _framework_value_present(value)]
        import_disabled_reason = ""
        import_readiness = "可导入：已找到 07 最终策划包。"
        if not package_ready:
            import_readiness = "不可导入：没有找到 07 最终策划包，也没有足够的阶段输出可合成框架包。"
            if stage_output_keys:
                import_readiness = "不可导入：找到部分阶段输出，但不足以恢复框架转剧本所需框架包。"
            import_disabled_reason = import_readiness
        asset = {
            "asset_id": str(project.get("project_id") or framework_state.get("project_id") or ""),
            "project_id": project.get("project_id") or framework_state.get("project_id"),
            "title": title,
            "source_title": source_title,
            "target_format": str(basic.get("target_format") or project.get("target_format") or "短剧"),
            "episodes_per_season": _safe_positive_int(basic.get("episodes_per_season") or project.get("total_episodes"), 0),
            "episode_word_count": _safe_positive_int(basic.get("episode_word_count"), 0),
            "season_count": _safe_positive_int(basic.get("season_count"), 1),
            "created_at": project.get("created_at") or framework_state.get("created_at"),
            "updated_at": project.get("updated_at") or framework_state.get("updated_at"),
            "summary": summary or "已保存的框架资产，可导入后继续 框架到剧本链路。",
            "can_import": package_ready,
            "import_readiness": import_readiness,
            "import_disabled_reason": import_disabled_reason,
            "framework_package_source": "framework_plan_package" if package_ready and not package.get("import_package_synthesized") else ("synthesized_stage_outputs" if package_ready else ""),
            "available_stage_output_keys": stage_output_keys,
            "has_framework_to_script_state": has_framework_to_script_state,
            "framework_to_script_progress": {
                "has_state": has_framework_to_script_state,
                "latest_stage": latest_script_stage,
                "completed_stages": completed_script_stage_numbers,
                "running_stage": str(workspace_state.get("runningStage") or ""),
                "stage_count": len(completed_script_stage_numbers),
            },
            "stage_prompts": _strip_raw_fastgpt_fields(copy.deepcopy(stage_prompts)),
            "preference_snapshot": _strip_raw_fastgpt_fields(copy.deepcopy(preference_snapshot)),
        }
        if include_detail:
            asset.update(
                {
                    "framework_plan_package": _strip_raw_fastgpt_fields(copy.deepcopy(package)),
                    "stage_outputs": _strip_raw_fastgpt_fields({**copy.deepcopy(stage_outputs), **copy.deepcopy(workspace_stage_outputs)}),
                    "framework_to_script_state": _strip_raw_fastgpt_fields(copy.deepcopy(workspace_state)),
                    "scriptStages": _strip_raw_fastgpt_fields(copy.deepcopy(workspace_state.get("scriptStages") or {})),
                    "stageDrafts": _strip_raw_fastgpt_fields(copy.deepcopy(workspace_state.get("stageDrafts") or {})),
                    "stage_prompts": _strip_raw_fastgpt_fields(copy.deepcopy(stage_prompts)),
                    "preference_snapshot": _strip_raw_fastgpt_fields(copy.deepcopy(preference_snapshot)),
                }
            )
        return asset

    def _load_framework_asset_for_user(asset_id: str, user_id: int) -> dict | None:
        project_id = 0
        try:
            project_id = int(str(asset_id or "").strip())
        except Exception:
            project_id = 0
        if project_id <= 0:
            return None
        project = task_manager.get_project_snapshot(project_id, user_id=user_id, public_view=False)
        if not project or str(project.get("asset_kind") or "").strip() != "framework_planner":
            return None
        asset = _framework_asset_payload(project, include_detail=True)
        return asset

    def _try_begin_framework_stage(user_id: int, asset_id: str, stage: str) -> bool:
        key = (int(user_id), str(asset_id or "").strip(), str(stage or "").strip())
        if not key[1] or not key[2]:
            return True
        now = time.monotonic()
        with framework_stage_runs_lock:
            expired_keys = [
                run_key
                for run_key, started_at in framework_stage_runs.items()
                if now - float(started_at or 0) > FRAMEWORK_STAGE_RUN_TTL_SECONDS
            ]
            for run_key in expired_keys:
                framework_stage_runs.pop(run_key, None)
            started_at = framework_stage_runs.get(key)
            if started_at is not None:
                logger.warning(
                    "framework-to-script stage already running: user_id=%s asset_id=%s stage=%s age_seconds=%s",
                    user_id,
                    key[1],
                    key[2],
                    int(now - float(started_at or now)),
                )
                return False
            framework_stage_runs[key] = now
        return True

    def _end_framework_stage(user_id: int, asset_id: str, stage: str) -> None:
        key = (int(user_id), str(asset_id or "").strip(), str(stage or "").strip())
        with framework_stage_runs_lock:
            framework_stage_runs.pop(key, None)
        try:
            project_id = int(str(asset_id or "").strip())
        except Exception:
            project_id = 0
        if project_id <= 0:
            return
        snapshot = task_manager.get_project_snapshot(project_id, user_id=user_id, public_view=False)
        if not snapshot:
            return
        artifacts = snapshot.get("artifacts")
        workspace_state = artifacts.get("framework_to_script_state") if isinstance(artifacts, dict) else None
        if not isinstance(workspace_state, dict):
            return
        persisted_stage = str(
            workspace_state.get("runningStage") or workspace_state.get("running_stage") or ""
        ).strip()
        if persisted_stage and persisted_stage != str(stage or "").strip():
            return
        workspace_state["runningStage"] = ""
        workspace_state["running_stage"] = ""
        workspace_state["updated_at"] = _now_iso()
        snapshot["updated_at"] = workspace_state["updated_at"]
        record = task_manager._projects.get(project_id)
        if record:
            with record.lock:
                record.snapshot = snapshot
            task_manager._persist_snapshot(record)
        else:
            task_manager._persist_project_snapshot_data(project_id, snapshot)

    def _active_framework_stage(user_id: int, asset_id: str) -> tuple[str, int]:
        asset_key = str(asset_id or "").strip()
        if not asset_key:
            return "", 0
        now = time.monotonic()
        with framework_stage_runs_lock:
            expired_keys = [
                run_key
                for run_key, started_at in framework_stage_runs.items()
                if now - float(started_at or 0) > FRAMEWORK_STAGE_RUN_TTL_SECONDS
            ]
            for run_key in expired_keys:
                framework_stage_runs.pop(run_key, None)
            active = [
                (run_key, started_at)
                for run_key, started_at in framework_stage_runs.items()
                if run_key[0] == int(user_id) and run_key[1] == asset_key
            ]
        if not active:
            return "", 0
        run_key, started_at = max(active, key=lambda item: float(item[1] or 0))
        return str(run_key[2] or ""), max(0, int(now - float(started_at or now)))

    def _save_framework_to_script_stage(
        *,
        user_id: int,
        asset_id: str,
        stage_key: str,
        output: dict,
        status: str = "completed",
    ) -> None:
        try:
            project_id = int(str(asset_id or "").strip())
        except Exception:
            return
        if project_id <= 0 or not isinstance(output, dict):
            return
        snapshot = task_manager.get_project_snapshot(project_id, user_id=user_id, public_view=False)
        if not snapshot or str(snapshot.get("asset_kind") or "").strip() != "framework_planner":
            return
        now = _now_iso()
        clean_output = _strip_raw_fastgpt_fields(copy.deepcopy(output))
        artifacts = snapshot.setdefault("artifacts", {})
        if not isinstance(artifacts, dict):
            artifacts = {}
            snapshot["artifacts"] = artifacts
        workspace_state = artifacts.get("framework_to_script_state")
        if not isinstance(workspace_state, dict):
            workspace_state = {"scriptStages": {}}
        script_stages = workspace_state.get("scriptStages")
        if not isinstance(script_stages, dict):
            script_stages = {}
        cascade = {
            "stage08": ("stage09", "stage10", "stage11", "stage12"),
            "stage09": ("stage10", "stage11", "stage12"),
            "stage10": ("stage11", "stage12"),
            "stage11": ("stage12",),
        }
        for downstream_stage in cascade.get(str(stage_key), ()):
            script_stages.pop(downstream_stage, None)
        stage_output_aliases = {
            "stage08": {
                "framework_scene_dictionary": clean_output,
                "sceneDictionary": clean_output.get("sceneDictionary"),
                "scriptWorldRulesDigest": clean_output.get("scriptWorldRulesDigest"),
            },
            "stage09": {
                "framework_appearanceMapping": clean_output,
                "appearanceMapping": clean_output.get("appearanceMapping"),
            },
            "stage10": {
                "framework_enriched_episode_plan": clean_output,
                "allEnrichedEpisodePlan": clean_output.get("allEnrichedEpisodePlan") or clean_output.get("enrichedEpisodePlan"),
                "allEnrichedEpisodePlanText": clean_output.get("allEnrichedEpisodePlanText") or clean_output.get("enrichedEpisodePlanText"),
                "batchEnrichedEpisodePlan": clean_output.get("batchEnrichedEpisodePlan")
                or clean_output.get("allEnrichedEpisodePlan")
                or clean_output.get("enrichedEpisodePlan"),
            },
            "stage11": {
                "framework_causal_conflict_plan": clean_output,
                "batchCausalConflictPlan": clean_output.get("batchCausalConflictPlan"),
                "conflictMemory": clean_output.get("conflictMemory"),
            },
            "stage12": {
                "framework_script_text": clean_output,
                "batchScriptText": clean_output.get("batchScriptText"),
                "scriptMemory": clean_output.get("scriptMemory"),
            },
        }
        clean_output["updated_at"] = now
        script_stages[str(stage_key)] = clean_output
        stage_outputs = workspace_state.get("stageOutputs")
        if not isinstance(stage_outputs, dict):
            stage_outputs = {}
        for key, value in stage_output_aliases.get(str(stage_key), {}).items():
            if _framework_value_present(value):
                stage_outputs[key] = _strip_raw_fastgpt_fields(copy.deepcopy(value))
        for downstream_stage in cascade.get(str(stage_key), ()):
            output_key = downstream_stage.replace("stage", "")
            for key in tuple(stage_outputs.keys()):
                if output_key in str(key):
                    stage_outputs.pop(key, None)
        stages_state = workspace_state.get("stages")
        if not isinstance(stages_state, dict):
            stages_state = {}
        stage_number = str(stage_key).replace("stage", "")
        normalized_status = str(status or "completed").strip() or "completed"
        if normalized_status == "completed" and str(stage_key) in {"stage11", "stage12"}:
            stage10 = script_stages.get("stage10") if isinstance(script_stages.get("stage10"), dict) else {}
            plan = (
                stage10.get("allEnrichedEpisodePlan")
                or stage10.get("enrichedEpisodePlan")
                or stage10.get("batchEnrichedEpisodePlan")
                or []
            )
            nested_plan = stage10.get("framework_enriched_episode_plan")
            if not isinstance(plan, list) and isinstance(nested_plan, dict):
                plan = (
                    nested_plan.get("allEnrichedEpisodePlan")
                    or nested_plan.get("enrichedEpisodePlan")
                    or nested_plan.get("batchEnrichedEpisodePlan")
                    or []
                )
            episode_limit = _positive_int(
                snapshot.get("episodes_per_season")
                or snapshot.get("total_episodes")
                or workspace_state.get("episodes_per_season")
                or workspace_state.get("total_episodes"),
                0,
            )
            expected_batch_starts = _framework_expected_batch_starts(plan, episode_limit)
            batches = clean_output.get("batches") if isinstance(clean_output.get("batches"), dict) else {}
            valid_batch_starts = set()
            for batch_start, batch in batches.items():
                if not isinstance(batch, dict):
                    continue
                if str(stage_key) == "stage11":
                    batch_value = batch.get("batchCausalConflictPlan") or batch.get("batch_causal_conflict_plan")
                else:
                    batch_value = batch.get("batchScriptText") or batch.get("batch_script_text")
                if _framework_value_present(batch_value):
                    valid_batch_starts.add(str(batch_start))
            if expected_batch_starts and any(
                batch_start not in valid_batch_starts for batch_start in expected_batch_starts
            ):
                normalized_status = "partial"
        stages_state[stage_number] = {
            "status": normalized_status,
            "stage_key": str(stage_key),
            "updated_at": now,
        }
        if normalized_status == "completed":
            for downstream_stage in cascade.get(str(stage_key), ()):
                stages_state[str(downstream_stage).replace("stage", "")] = {
                    "status": "pending",
                    "stage_key": str(downstream_stage),
                    "updated_at": now,
                }
        completed_stages = workspace_state.get("completedStages")
        if not isinstance(completed_stages, list):
            completed_stages = []
        completed_set = {str(item) for item in completed_stages}
        if normalized_status == "completed":
            completed_set.add(stage_number)
            for downstream_stage in cascade.get(str(stage_key), ()):
                completed_set.discard(str(downstream_stage).replace("stage", ""))
        else:
            completed_set.discard(stage_number)
        stage_drafts = workspace_state.get("stageDrafts")
        if not isinstance(stage_drafts, dict):
            stage_drafts = {}
        if normalized_status == "completed":
            stage_drafts.pop(str(stage_key), None)
            for downstream_stage in cascade.get(str(stage_key), ()):
                stage_drafts.pop(str(downstream_stage), None)
            for failed_stage_key in ("lastFailedStage", "last_failed_stage"):
                if str(workspace_state.get(failed_stage_key) or "") == stage_number:
                    workspace_state[failed_stage_key] = ""
        workspace_state["completedStages"] = sorted(completed_set, key=lambda item: int(item) if item.isdigit() else 999)
        workspace_state["scriptStages"] = script_stages
        workspace_state["stageOutputs"] = stage_outputs
        workspace_state["stages"] = stages_state
        workspace_state["stageDrafts"] = stage_drafts
        workspace_state["framework_asset_id"] = str(asset_id)
        workspace_state["project_id"] = project_id
        workspace_state["updated_at"] = now
        artifacts["framework_to_script_state"] = workspace_state
        snapshot["updated_at"] = now
        record = task_manager._projects.get(project_id)
        if record:
            with record.lock:
                record.snapshot = snapshot
            task_manager._persist_snapshot(record)
        else:
            try:
                task_manager._persist_project_snapshot_data(project_id, snapshot)
            except Exception:
                logger.exception("framework-to-script stage snapshot persist failed project_id=%s", project_id)

    def _save_framework_to_script_stage_draft(
        *,
        user_id: int,
        asset_id: str,
        stage_key: str,
        draft: dict,
        status: str = "running",
    ) -> None:
        try:
            project_id = int(str(asset_id or "").strip())
        except Exception:
            return
        if project_id <= 0 or not isinstance(draft, dict):
            return
        snapshot = task_manager.get_project_snapshot(project_id, user_id=user_id, public_view=False)
        if not snapshot or str(snapshot.get("asset_kind") or "").strip() != "framework_planner":
            return
        now = _now_iso()
        clean_draft = _strip_raw_fastgpt_fields(copy.deepcopy(draft))
        clean_draft["updated_at"] = now
        clean_draft["status"] = str(status or clean_draft.get("status") or "running")
        artifacts = snapshot.setdefault("artifacts", {})
        if not isinstance(artifacts, dict):
            artifacts = {}
            snapshot["artifacts"] = artifacts
        workspace_state = artifacts.get("framework_to_script_state")
        if not isinstance(workspace_state, dict):
            workspace_state = {"scriptStages": {}}
        stage_drafts = workspace_state.get("stageDrafts")
        if not isinstance(stage_drafts, dict):
            stage_drafts = {}
        stage_drafts[str(stage_key)] = clean_draft
        stages_state = workspace_state.get("stages")
        if not isinstance(stages_state, dict):
            stages_state = {}
        stage_number = str(stage_key).replace("stage", "")
        stages_state[stage_number] = {
            "status": clean_draft["status"],
            "stage_key": str(stage_key),
            "updated_at": now,
            "draft_only": True,
        }
        completed_stages = workspace_state.get("completedStages")
        if not isinstance(completed_stages, list):
            completed_stages = []
        completed_set = {str(item) for item in completed_stages}
        completed_set.discard(stage_number)
        workspace_state["completedStages"] = sorted(completed_set, key=lambda item: int(item) if item.isdigit() else 999)
        workspace_state["stageDrafts"] = stage_drafts
        workspace_state["stages"] = stages_state
        workspace_state["framework_asset_id"] = str(asset_id)
        workspace_state["project_id"] = project_id
        workspace_state["updated_at"] = now
        artifacts["framework_to_script_state"] = workspace_state
        snapshot["updated_at"] = now
        record = task_manager._projects.get(project_id)
        if record:
            with record.lock:
                record.snapshot = snapshot
            task_manager._persist_snapshot(record)
        else:
            try:
                task_manager._persist_project_snapshot_data(project_id, snapshot)
            except Exception:
                logger.exception("framework-to-script draft snapshot persist failed project_id=%s", project_id)

    def _inject_framework_asset(data: dict, user_id: int) -> tuple[dict, dict | None]:
        asset_id = str(data.get("framework_asset_id") or data.get("asset_id") or "").strip()
        if not asset_id:
            return data, None
        asset = _load_framework_asset_for_user(asset_id, user_id)
        if not asset:
            raise ValueError("框架资产不存在或当前账号无权访问。")
        workspace_state = asset.get("framework_to_script_state") if isinstance(asset.get("framework_to_script_state"), dict) else {}
        has_script_state = bool(
            asset.get("has_framework_to_script_state")
            or (isinstance(workspace_state.get("scriptStages"), dict) and workspace_state.get("scriptStages"))
            or (isinstance(workspace_state.get("stageOutputs"), dict) and workspace_state.get("stageOutputs"))
        )
        if not asset.get("can_import") and not has_script_state:
            reason = str(asset.get("import_disabled_reason") or "尚未生成可恢复的 07 最终策划包。").strip()
            raise ValueError(reason)
        merged = dict(data)
        package = asset.get("framework_plan_package") if isinstance(asset.get("framework_plan_package"), dict) else {}
        stage_outputs = asset.get("stage_outputs") if isinstance(asset.get("stage_outputs"), dict) else {}
        merged["framework_asset_id"] = asset.get("asset_id")
        merged["source_framework_project_id"] = asset.get("project_id")
        merged["framework_plan_package"] = package
        package_stage_prompts = _stage_prompt_lookup_from_mapping(package.get("stage_prompts")) if isinstance(package, dict) else {}
        if not isinstance(merged.get("preference_snapshot"), dict) or not merged.get("preference_snapshot"):
            merged["preference_snapshot"] = (
                asset.get("preference_snapshot")
                if isinstance(asset.get("preference_snapshot"), dict)
                else {}
            )
        merged["stage_prompts"] = _merge_stage_prompts_non_empty(
            package_stage_prompts,
            asset.get("stage_prompts") if isinstance(asset.get("stage_prompts"), dict) else {},
            merged.get("stage_prompts") if isinstance(merged.get("stage_prompts"), dict) else {},
        )
        prompt_preferences = merged.get("prompt_preferences") if isinstance(merged.get("prompt_preferences"), dict) else {}
        prompt_preferences = dict(prompt_preferences)
        prompt_preferences["stage_prompts"] = _merge_stage_prompts_non_empty(
            merged.get("stage_prompts"),
            prompt_preferences.get("stage_prompts"),
        )
        merged["prompt_preferences"] = prompt_preferences
        logger.info(
            "framework-to-script asset preference handoff: asset_id=%s non_empty_keys=%s",
            asset.get("asset_id"),
            [key for key, value in merged["stage_prompts"].items() if _coerce_prompt_text(value)],
        )
        for key, value in stage_outputs.items():
            merged.setdefault(key, value)
        return merged, asset

    def _login_required(view_func):
        @wraps(view_func)
        def wrapper(*args, **kwargs):
            if not _current_user():
                return _json_error("请先登录。", 401)
            return view_func(*args, **kwargs)

        return wrapper

    @app.route("/api/framework-to-script/running-stage", methods=["GET", "POST"])
    @_login_required
    def mark_framework_to_script_running_stage():
        data = request.args.to_dict() if request.method == "GET" else (request.get_json(silent=True) or {})
        user_id = _require_user_id()
        asset_id = str(data.get("framework_asset_id") or data.get("asset_id") or "").strip()
        running_stage = str(data.get("running_stage") or data.get("runningStage") or "").strip()
        try:
            project_id = int(asset_id)
        except Exception:
            project_id = 0
        if project_id <= 0:
            return _json_error("缺少有效 framework_asset_id。", status=400)
        snapshot = task_manager.get_project_snapshot(project_id, user_id=user_id, public_view=False)
        if not snapshot or str(snapshot.get("asset_kind") or "").strip() != "framework_planner":
            return _json_error("框架资产不存在。", status=404)
        now = _now_iso()
        artifacts = snapshot.setdefault("artifacts", {})
        if not isinstance(artifacts, dict):
            artifacts = {}
            snapshot["artifacts"] = artifacts
        workspace_state = artifacts.get("framework_to_script_state")
        if not isinstance(workspace_state, dict):
            workspace_state = {"scriptStages": {}, "stageOutputs": {}}
        active_stage, active_age_seconds = _active_framework_stage(user_id, asset_id)
        persisted_stage = str(workspace_state.get("runningStage") or workspace_state.get("running_stage") or "").strip()
        if request.method == "GET":
            if persisted_stage and not active_stage:
                persisted_stage = ""
                workspace_state["runningStage"] = ""
                workspace_state["running_stage"] = ""
                workspace_state["updated_at"] = now
                artifacts["framework_to_script_state"] = workspace_state
                snapshot["updated_at"] = now
                record = task_manager._projects.get(project_id)
                if record:
                    with record.lock:
                        record.snapshot = snapshot
                    task_manager._persist_snapshot(record)
                else:
                    task_manager._persist_project_snapshot_data(project_id, snapshot)
            return _json_ok(
                running_stage=active_stage or persisted_stage,
                backend_running=bool(active_stage),
                active_age_seconds=active_age_seconds,
                project_id=project_id,
            )
        if active_stage:
            running_stage = active_stage
        workspace_state["runningStage"] = running_stage
        workspace_state["running_stage"] = running_stage
        if "last_failed_stage" in data or "lastFailedStage" in data:
            last_failed_stage = str(data.get("last_failed_stage") or data.get("lastFailedStage") or "").strip()
            workspace_state["lastFailedStage"] = last_failed_stage
            workspace_state["last_failed_stage"] = last_failed_stage
        workspace_state["framework_asset_id"] = asset_id
        workspace_state["project_id"] = project_id
        workspace_state["updated_at"] = now
        artifacts["framework_to_script_state"] = workspace_state
        snapshot["updated_at"] = now
        record = task_manager._projects.get(project_id)
        if record:
            with record.lock:
                record.snapshot = snapshot
            task_manager._persist_snapshot(record)
        else:
            task_manager._persist_project_snapshot_data(project_id, snapshot)
        return _json_ok(
            running_stage=running_stage,
            backend_running=bool(active_stage),
            active_age_seconds=active_age_seconds,
            project_id=project_id,
        )

    @app.post("/api/framework-to-script/save-progress")
    @_login_required
    def save_framework_to_script_progress():
        data = request.get_json(silent=True) or {}
        user_id = _require_user_id()
        asset_id = str(data.get("framework_asset_id") or data.get("asset_id") or "").strip()
        try:
            project_id = int(asset_id)
        except Exception:
            project_id = 0
        if project_id <= 0:
            return _json_error("缺少有效 framework_asset_id，无法保存剧本进度。", status=400)

        snapshot = task_manager.get_project_snapshot(project_id, user_id=user_id, public_view=False)
        if not snapshot or str(snapshot.get("asset_kind") or "").strip() != "framework_planner":
            return _json_error("框架资产不存在或当前账号无权访问。", status=404)

        now = _now_iso()
        artifacts = snapshot.setdefault("artifacts", {})
        if not isinstance(artifacts, dict):
            artifacts = {}
            snapshot["artifacts"] = artifacts
        workspace_state = artifacts.get("framework_to_script_state")
        if not isinstance(workspace_state, dict):
            workspace_state = {"scriptStages": {}, "stageOutputs": {}}
        else:
            workspace_state = copy.deepcopy(workspace_state)

        for key in ("scriptStages", "stageOutputs", "completedStages", "stages", "settings"):
            value = data.get(key)
            if isinstance(value, (dict, list)):
                workspace_state[key] = _strip_raw_fastgpt_fields(copy.deepcopy(value))

        incoming_drafts = data.get("stageDrafts") or data.get("stage_drafts")
        if isinstance(incoming_drafts, dict):
            workspace_state["stageDrafts"] = _strip_raw_fastgpt_fields(copy.deepcopy(incoming_drafts))
        elif not isinstance(workspace_state.get("stageDrafts"), dict):
            workspace_state["stageDrafts"] = {}

        running_stage = str(data.get("runningStage") or data.get("running_stage") or workspace_state.get("runningStage") or "").strip()
        last_failed_stage = str(data.get("lastFailedStage") or data.get("last_failed_stage") or workspace_state.get("lastFailedStage") or "").strip()
        if running_stage:
            script_stages = workspace_state.get("scriptStages")
            completed_stages = workspace_state.get("completedStages")
            completed_set = {str(item).zfill(2) for item in completed_stages} if isinstance(completed_stages, list) else set()
            existing_stage = script_stages.get(f"stage{running_stage.zfill(2)}") if isinstance(script_stages, dict) else {}
            existing_batches = existing_stage.get("batches") if isinstance(existing_stage, dict) and isinstance(existing_stage.get("batches"), dict) else {}
            if running_stage.zfill(2) in completed_set and (existing_batches or _framework_value_present(existing_stage)):
                running_stage = ""
        workspace_state["runningStage"] = running_stage
        workspace_state["running_stage"] = running_stage
        workspace_state["lastFailedStage"] = last_failed_stage
        workspace_state["last_failed_stage"] = last_failed_stage
        running_started_at = str(data.get("runningStartedAt") or data.get("running_started_at") or workspace_state.get("runningStartedAt") or "").strip()
        running_progress = data.get("runningProgress") if isinstance(data.get("runningProgress"), dict) else data.get("running_progress")
        if isinstance(running_progress, dict):
            running_progress = _strip_raw_fastgpt_fields(copy.deepcopy(running_progress))
        elif not isinstance(workspace_state.get("runningProgress"), dict):
            running_progress = None
        else:
            running_progress = workspace_state.get("runningProgress")
        running_retry_message = str(data.get("runningRetryMessage") or data.get("running_retry_message") or workspace_state.get("runningRetryMessage") or "").strip()
        running_retry_countdown = data.get("runningRetryCountdown", data.get("running_retry_countdown", workspace_state.get("runningRetryCountdown", 0)))
        workspace_state["runningStartedAt"] = running_started_at
        workspace_state["running_started_at"] = running_started_at
        workspace_state["runningProgress"] = running_progress
        workspace_state["running_progress"] = running_progress
        workspace_state["runningRetryMessage"] = running_retry_message
        workspace_state["running_retry_message"] = running_retry_message
        workspace_state["runningRetryCountdown"] = _positive_int(running_retry_countdown, 0)
        workspace_state["running_retry_countdown"] = _positive_int(running_retry_countdown, 0)
        workspace_state["framework_asset_id"] = asset_id
        workspace_state["project_id"] = project_id
        workspace_state["updated_at"] = now
        artifacts["framework_to_script_state"] = workspace_state
        snapshot["updated_at"] = now

        record = task_manager._projects.get(project_id)
        if record:
            with record.lock:
                record.snapshot = snapshot
            task_manager._persist_snapshot(record)
        else:
            task_manager._persist_project_snapshot_data(project_id, snapshot)

        return _json_ok(
            message="剧本进度已保存到当前框架资产。",
            project_id=project_id,
            framework_asset=_framework_asset_payload(snapshot, include_detail=True),
            framework_to_script_state=_strip_raw_fastgpt_fields(copy.deepcopy(workspace_state)),
        )

    def _framework_script_stage_cache(framework_asset: dict | None, stage_key: str) -> dict:
        if not isinstance(framework_asset, dict):
            return {}
        stages = framework_asset.get("scriptStages")
        if not isinstance(stages, dict):
            state = framework_asset.get("framework_to_script_state")
            stages = state.get("scriptStages") if isinstance(state, dict) else {}
        if not isinstance(stages, dict):
            return {}
        cached = stages.get(stage_key)
        if isinstance(cached, dict) and cached:
            return cached
        stage_outputs = framework_asset.get("stage_outputs") if isinstance(framework_asset.get("stage_outputs"), dict) else {}
        if stage_key == "stage10" and isinstance(stage_outputs, dict):
            framework_output = stage_outputs.get("framework_enriched_episode_plan")
            if isinstance(framework_output, dict) and framework_output:
                return framework_output
            plan = (
                stage_outputs.get("allEnrichedEpisodePlan")
                or stage_outputs.get("batchEnrichedEpisodePlan")
                or stage_outputs.get("all_enriched_episode_plan")
                or []
            )
            text = stage_outputs.get("allEnrichedEpisodePlanText") or stage_outputs.get("all_enriched_episode_plan_text") or ""
            if _framework_value_present(plan) or _framework_value_present(text):
                return {
                    "allEnrichedEpisodePlan": plan,
                    "enrichedEpisodePlan": plan,
                    "batchEnrichedEpisodePlan": plan,
                    "allEnrichedEpisodePlanText": text,
                    "enrichedEpisodePlanText": text,
                }
        if stage_key == "stage11" and isinstance(stage_outputs, dict):
            framework_output = stage_outputs.get("framework_causal_conflict_plan")
            framework_output = framework_output if isinstance(framework_output, dict) else {}
            batch_plan = (
                framework_output.get("batchCausalConflictPlan")
                or framework_output.get("batch_causal_conflict_plan")
                or stage_outputs.get("batchCausalConflictPlan")
                or stage_outputs.get("batch_causal_conflict_plan")
                or {}
            )
            batches = framework_output.get("batches") if isinstance(framework_output.get("batches"), dict) else {}
            if not batches and _framework_value_present(batch_plan):
                start_episode = (
                    framework_output.get("batchStartEpisode")
                    or framework_output.get("batch_start_episode")
                    or stage_outputs.get("batchStartEpisode")
                    or stage_outputs.get("batch_start_episode")
                    or 1
                )
                batches = {
                    str(start_episode): {
                        "batchStartEpisode": start_episode,
                        "batchEndEpisode": framework_output.get("batchEndEpisode")
                        or framework_output.get("batch_end_episode")
                        or stage_outputs.get("batchEndEpisode")
                        or stage_outputs.get("batch_end_episode")
                        or "",
                        "batchCausalConflictPlan": batch_plan,
                        "batch_causal_conflict_plan": batch_plan,
                    }
                }
            if _framework_value_present(batch_plan) or batches:
                return {
                    **framework_output,
                    "batchCausalConflictPlan": batch_plan,
                    "batch_causal_conflict_plan": framework_output.get("batch_causal_conflict_plan") or batch_plan,
                    "conflictMemory": framework_output.get("conflictMemory")
                    or framework_output.get("conflict_memory")
                    or stage_outputs.get("conflictMemory")
                    or "",
                    "batches": batches,
                }
        if stage_key == "stage12" and isinstance(stage_outputs, dict):
            framework_output = stage_outputs.get("framework_script_text")
            framework_output = framework_output if isinstance(framework_output, dict) else {}
            batch_text = (
                framework_output.get("batchScriptText")
                or framework_output.get("batch_script_text")
                or stage_outputs.get("batchScriptText")
                or stage_outputs.get("batch_script_text")
                or ""
            )
            batches = framework_output.get("batches") if isinstance(framework_output.get("batches"), dict) else {}
            if not batches and str(batch_text or "").strip():
                stage11_output = stage_outputs.get("framework_causal_conflict_plan")
                stage11_batches = stage11_output.get("batches") if isinstance(stage11_output, dict) and isinstance(stage11_output.get("batches"), dict) else {}
                stage11_keys = sorted(
                    [str(key) for key in stage11_batches.keys() if str(key).isdigit()],
                    key=lambda item: int(item),
                )
                start_episode = (
                    framework_output.get("batchStartEpisode")
                    or framework_output.get("batch_start_episode")
                    or stage_outputs.get("batchStartEpisode")
                    or stage_outputs.get("batch_start_episode")
                    or (stage11_keys[0] if stage11_keys else 1)
                )
                try:
                    end_episode = int(start_episode) + 4
                except Exception:
                    end_episode = ""
                batches = {
                    str(start_episode): {
                        "batchStartEpisode": start_episode,
                        "batchEndEpisode": framework_output.get("batchEndEpisode")
                        or framework_output.get("batch_end_episode")
                        or stage_outputs.get("batchEndEpisode")
                        or stage_outputs.get("batch_end_episode")
                        or end_episode,
                        "batchScriptText": batch_text,
                        "batch_script_text": batch_text,
                    }
                }
            if str(batch_text or "").strip() or batches:
                return {
                    **framework_output,
                    "batchScriptText": batch_text,
                    "batch_script_text": framework_output.get("batch_script_text") or batch_text,
                    "scriptMemory": framework_output.get("scriptMemory")
                    or framework_output.get("script_memory")
                    or stage_outputs.get("scriptMemory")
                    or "",
                    "batches": batches,
                }
        return {}

    def _positive_int(value, default: int) -> int:
        try:
            number = int(value)
            return number if number > 0 else default
        except Exception:
            return default

    def _first_present(mapping, *keys, default=None):
        if not isinstance(mapping, dict):
            return default
        for key in keys:
            if key in mapping and mapping.get(key) is not None:
                return mapping.get(key)
        return default

    def _get_bool_alias(mapping, *keys, default=None):
        value = _first_present(mapping, *keys, default=default)
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            normalized = value.strip().lower()
            if normalized in {"true", "1", "yes", "y", "passed", "pass", "通过", "是"}:
                return True
            if normalized in {"false", "0", "no", "n", "failed", "fail", "不通过", "否"}:
                return False
        return default if value is None else bool(value)

    def _get_list_alias(mapping, *keys):
        value = _first_present(mapping, *keys, default=[])
        return value if isinstance(value, list) else []

    def _get_dict_alias(mapping, *keys):
        value = _first_present(mapping, *keys, default={})
        return value if isinstance(value, dict) else {}

    def _merge_aliases(base_vars, aliases):
        if not isinstance(base_vars, dict) or not isinstance(aliases, dict):
            return base_vars
        for source_key, alias_keys in aliases.items():
            if source_key not in base_vars:
                continue
            for alias_key in alias_keys:
                base_vars.setdefault(alias_key, base_vars[source_key])
        return base_vars

    def _unwrap_dict_alias(mapping, *keys):
        value = _get_dict_alias(mapping, *keys)
        for key in keys:
            if isinstance(value, dict) and isinstance(value.get(key), dict):
                return value.get(key), True
        return value, False

    def _normalize_dict_output_alias(mapping, *keys):
        value = _first_present(mapping, *keys, default=None)
        unwrapped = False
        if isinstance(value, str):
            text = value.strip()
            if text.startswith("```"):
                text = text.strip("`").strip()
                if text.lower().startswith("json"):
                    text = text[4:].strip()
            try:
                value = json.loads(text)
            except Exception as exc:
                return {}, unwrapped, f"string JSON parse failed: {exc}"
        for key in keys:
            if isinstance(value, dict) and isinstance(value.get(key), dict):
                value = value.get(key)
                unwrapped = True
                break
            if isinstance(value, dict) and isinstance(value.get(key), str):
                try:
                    value = json.loads(value.get(key).strip())
                    unwrapped = True
                    break
                except Exception as exc:
                    return {}, unwrapped, f"wrapped string JSON parse failed: {exc}"
        if isinstance(value, dict) and value:
            return value, unwrapped, ""
        if value is None:
            return {}, unwrapped, "missing output alias"
        return {}, unwrapped, f"output alias is {type(value).__name__}, expected object"

    def _validate_stage11_causal_conflict_plan(plan, *, start_episode: int, end_episode: int) -> list[str]:
        if not isinstance(plan, dict) or not plan:
            return ["batchCausalConflictPlan must be a non-empty JSON object"]
        issues = []
        required_root = ("batch_meta", "global_conflict_engine", "episodes")
        missing_root = [key for key in required_root if key not in plan]
        if missing_root:
            issues.append(f"batchCausalConflictPlan missing root fields: {', '.join(missing_root)}")
        episodes = plan.get("episodes")
        if not isinstance(episodes, list) or not episodes:
            issues.append("batchCausalConflictPlan.episodes must be a non-empty array")
            return issues
        expected_episodes = set(range(int(start_episode), int(end_episode) + 1))
        actual_episodes = set()
        required_episode_fields = (
            "episode",
            "episode_title",
            "active_characters",
            "scene_refs",
            "carry_in",
            "why_now",
            "character_motivation",
            "emotional_precondition",
            "scene_cause_chain",
            "non_conflict_moment",
            "natural_transition",
            "opening_image",
            "opening_action",
            "current_goal",
            "core_obstacle",
            "episode_state_change",
            "ending_hook",
            "dialogue_strategy",
        )
        if len(episodes) != len(expected_episodes) and len(episodes) > 5:
            issues.append("batchCausalConflictPlan.episodes count must match the current batch or be no more than 5")
        for index, episode_payload in enumerate(episodes, start=1):
            if not isinstance(episode_payload, dict):
                issues.append(f"batchCausalConflictPlan.episodes[{index}] must be an object")
                continue
            episode_no = _positive_int(episode_payload.get("episode"), 0)
            if episode_no > 0:
                actual_episodes.add(episode_no)
            missing_fields = [key for key in required_episode_fields if key not in episode_payload]
            if missing_fields:
                issues.append(
                    f"batchCausalConflictPlan.episodes[{index}] missing fields: {', '.join(missing_fields)}"
                )
        missing_episodes = sorted(expected_episodes - actual_episodes)
        if missing_episodes:
            issues.append(f"batchCausalConflictPlan missing episodes: {missing_episodes}")
        return issues

    def _framework_batch_from_plan(plan: list, requested_start=None, completed_starts=None) -> tuple[int, int, list]:
        completed = {int(item) for item in (completed_starts or []) if str(item).strip().isdigit()}
        episode_numbers = []
        for index, item in enumerate(plan, start=1):
            if not isinstance(item, dict):
                continue
            episode_numbers.append(
                _positive_int(
                    item.get("episode")
                    or item.get("episodeNumber")
                    or item.get("episode_number")
                    or item.get("episode_index")
                    or item.get("ep")
                    or index,
                    index,
                )
            )
        if requested_start:
            start_episode = _positive_int(requested_start, 1)
        else:
            starts = sorted({((episode - 1) // 5) * 5 + 1 for episode in episode_numbers if episode > 0})
            start_episode = next((start for start in starts if start not in completed), starts[0] if starts else 1)
        end_episode = start_episode + 4
        batch = []
        batch_episode_numbers = []
        for index, item in enumerate(plan, start=1):
            if not isinstance(item, dict):
                continue
            episode = _positive_int(
                item.get("episode")
                or item.get("episodeNumber")
                or item.get("episode_number")
                or item.get("episode_index")
                or item.get("ep")
                or index,
                index,
            )
            if start_episode <= episode <= end_episode:
                batch.append(item)
                batch_episode_numbers.append(episode)
        if batch:
            end_episode = max(batch_episode_numbers)
        return start_episode, end_episode, batch

    def _framework_expected_batch_starts(plan: object, max_episodes: int = 0) -> list[str]:
        starts = set()
        episode_limit = _positive_int(max_episodes, 0)
        for index, item in enumerate(plan if isinstance(plan, list) else [], start=1):
            if not isinstance(item, dict):
                continue
            episode = _positive_int(
                item.get("episode")
                or item.get("episodeNumber")
                or item.get("episode_number")
                or item.get("episode_index")
                or item.get("ep")
                or index,
                index,
            )
            if episode > 0 and (not episode_limit or episode <= episode_limit):
                starts.add(((episode - 1) // 5) * 5 + 1)
        return [str(start) for start in sorted(starts)]

    def _stage12_chinese_episode_number(value: str) -> int:
        """Parse common Chinese episode numerals so model formatting cannot hide valid output."""
        digits = {
            "零": 0,
            "〇": 0,
            "一": 1,
            "二": 2,
            "两": 2,
            "三": 3,
            "四": 4,
            "五": 5,
            "六": 6,
            "七": 7,
            "八": 8,
            "九": 9,
        }
        text = str(value or "").strip()
        if not text:
            return 0
        if text.isdigit():
            return int(text)
        if "百" in text:
            left, right = text.split("百", 1)
            hundreds = digits.get(left, 1)
            return hundreds * 100 + _stage12_chinese_episode_number(right)
        if "十" in text:
            left, right = text.split("十", 1)
            tens = digits.get(left, 1) if left else 1
            return tens * 10 + digits.get(right, 0)
        if all(char in digits for char in text):
            return int("".join(str(digits[char]) for char in text))
        return 0

    def _normalize_stage12_episode_headings(script_text: object) -> str:
        """Normalize Chinese or Arabic episode headings to the frontend's `第N集：` contract."""
        text = str(script_text or "").replace("\r\n", "\n").replace("\r", "\n")
        pattern = re.compile(
            r"(?m)^(?P<indent>[ \t]*)(?:[#＃]{1,6}[ \t]*)?"
            r"(?:[*_]{1,3})?"
            r"第[ \t]*(?P<number>\d+|[零〇一二两三四五六七八九十百]+)[ \t]*集"
            r"(?P<suffix>\s*(?:[:：]\s*)?[^\n]*)$"
        )

        def replace(match: re.Match) -> str:
            episode_no = _stage12_chinese_episode_number(match.group("number"))
            if episode_no <= 0:
                return match.group(0)
            suffix = str(match.group("suffix") or "").strip()
            suffix = re.sub(r"[*_]{1,3}\s*$", "", suffix).strip()
            suffix = re.sub(r"^[:：]\s*", "", suffix).strip()
            title = f"：{suffix}" if suffix else "："
            return f"{match.group('indent')}第{episode_no}集{title}"

        return pattern.sub(replace, text)

    def _stage12_script_episode_numbers(script_text: object) -> set[int]:
        """Return episode headings that are actually present in a generated batch."""
        text = _normalize_stage12_episode_headings(script_text)
        return {
            int(match.group(1))
            for match in re.finditer(r"(?m)^\s*第\s*(\d+)\s*集(?:\s*[:：]|\s|$)", text)
        }

    def _stage12_missing_script_episodes(
        script_text: object,
        start_episode: int,
        end_episode: int,
    ) -> list[int]:
        expected = set(range(int(start_episode), int(end_episode) + 1))
        return sorted(expected - _stage12_script_episode_numbers(script_text))

    def _stage12_opening_quality_issues(script_text: object) -> list[str]:
        """Catch weak cold opens that a remote reviewer may incorrectly approve."""
        text = _normalize_stage12_episode_headings(script_text)
        episode_matches = [
            match
            for match in re.finditer(
                r"(?m)^[ \t]*第[ \t]*(\d+)[ \t]*集(?:[ \t]*[：:][ \t]*(?P<title>[^\n]*))?[ \t]*$",
                text,
            )
            if str(match.group("title") or "").strip() not in {"完", "结束", "本集完"}
        ]
        if not episode_matches:
            return ["正文无法识别分集标题，无法验证每集开头钩子。"]

        metadata_prefixes = (
            "场次：", "场次:", "场次",
            "角色：", "角色:", "人物：", "人物:",
            "场景：", "场景:", "道具：", "道具:",
        )
        generic_speakers = {
            "同事", "众人", "广播", "系统", "手机", "电视", "新闻", "画外音", "旁白", "人群",
        }
        procedural_terms = (
            "报表", "审批单", "文件", "资料", "账本", "记录", "邮件", "合同", "会议",
            "汇报", "系统数据", "工作台", "电脑屏幕", "环境", "街道", "办公室里",
        )
        immediate_stakes = (
            "枪", "刀", "血", "尸体", "爆炸", "着火", "坠落", "绑架", "追杀", "窒息",
            "停职", "开除", "逮捕", "抓捕", "离婚", "退婚", "死亡", "死了", "失踪",
            "直播", "当众", "曝光", "倒计时", "最后一分钟", "只剩", "来不及",
        )
        collision_terms = (
            "对视", "认出", "掐住", "拦住", "推开", "夺过", "扇", "跪", "吻", "抱住",
            "指认", "质问", "威胁", "背叛", "前任", "前夫", "前妻", "未婚夫", "未婚妻",
        )
        issues: list[str] = []

        for index, match in enumerate(episode_matches):
            episode_no = int(match.group(1))
            end = episode_matches[index + 1].start() if index + 1 < len(episode_matches) else len(text)
            lines = []
            for raw_line in text[match.end():end].split("\n"):
                line = raw_line.strip()
                if not line or line.startswith(metadata_prefixes):
                    continue
                if re.match(r"^\d+\s*[-－—]\s*\d+\b", line):
                    continue
                lines.append(line)
                if len(lines) >= 4:
                    break
            if not lines:
                issues.append(f"第{episode_no}集场次资料后没有可验证的有效戏。")
                continue

            opening = "\n".join(lines)
            has_immediate_stakes = any(term in opening for term in immediate_stakes)
            has_collision_action = any(term in opening for term in collision_terms)
            has_named_dialogue = False
            for line in lines:
                dialogue = re.match(r"^([^▲△\s：:]{1,12})(?:\([^)]*\))?[：:]", line)
                if not dialogue:
                    continue
                speaker = re.sub(r"\([^)]*\)$", "", dialogue.group(1)).strip()
                if speaker not in generic_speakers:
                    has_named_dialogue = True
                    break

            first_line_is_procedural = any(term in lines[0] for term in procedural_terms)
            if first_line_is_procedural and not (has_named_dialogue or has_immediate_stakes or has_collision_action):
                issues.append(
                    f"第{episode_no}集以报表/文件/环境等程序性信息起笔，前4个有效拍点内没有形成"
                    "具名人物碰撞、公开且不可逆的代价或迫近危险；这属于铺垫，不是强钩子。"
                )

        return issues

    def _stage12_episode_plan_slice(plan: object, episode_no: int) -> list[dict]:
        if not isinstance(plan, list):
            return []
        selected: list[dict] = []
        for index, item in enumerate(plan, start=1):
            if not isinstance(item, dict):
                continue
            item_episode = _positive_int(
                item.get("episode")
                or item.get("episodeNumber")
                or item.get("episode_number")
                or item.get("episode_index")
                or item.get("ep"),
                index,
            )
            if item_episode == episode_no:
                selected.append(item)
        return selected

    def _stage12_conflict_plan_slice(plan: object, episode_no: int) -> dict:
        if not isinstance(plan, dict):
            return {}
        result = copy.deepcopy(plan)
        episodes = result.get("episodes")
        if isinstance(episodes, list):
            result["episodes"] = [
                item
                for index, item in enumerate(episodes, start=1)
                if isinstance(item, dict)
                and _positive_int(item.get("episode"), index) == episode_no
            ]
        return result

    def _sorted_numeric_batch_keys(batches: dict) -> list[str]:
        if not isinstance(batches, dict):
            return []
        return sorted(
            [str(key) for key in batches.keys() if str(key).strip().isdigit()],
            key=lambda item: int(item),
        )

    def _txt_label(key) -> str:
        return readable_label(key)

    def _txt_scalar(value) -> str:
        return readable_scalar(value)

    def _txt_readable(value, indent: int = 0) -> str:
        return readable_text(value, indent)

    def _framework_script_text_from_batches(stage12: dict) -> str:
        if not isinstance(stage12, dict):
            return ""
        batches = stage12.get("batches") if isinstance(stage12.get("batches"), dict) else {}
        parts = []
        for key in _sorted_numeric_batch_keys(batches):
            batch = batches.get(key)
            if not isinstance(batch, dict):
                continue
            text = _first_present(batch, "batchScriptText", "batch_script_text", default="")
            if str(text or "").strip():
                parts.append(str(text).replace("\r\n", "\n").replace("\r", "\n").strip())
        if parts:
            return "\n\n".join(parts)
        return str(_first_present(stage12, "batchScriptText", "batch_script_text", default="") or "").strip()


    # FRAMEWORK_TO_SCRIPT_EXPORT_PRUNE_NOISE_V1
    def _clean_framework_to_script_export_text(text: str) -> str:
        """Remove internal workflow notes from exported TXT/DOCX content."""
        lines = str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
        result = []
        skip_indent = None

        for line in lines:
            stripped = line.strip()

            # Remove internal adaptation-guide meta text:
            # "??????????????????..."
            if (
                stripped.startswith("\u6587\u672c\uff1a\u672c\u6b21\u6574\u4f53\u6539\u7f16\u6307\u5f15")
                or stripped.startswith("\u6587\u672c:\u672c\u6b21\u6574\u4f53\u6539\u7f16\u6307\u5f15")
            ):
                continue

            # Remove source-trace blocks under scenes:
            # "?????" and all child lines until the next same/lower-indent field.
            if (
                stripped.startswith("\u6765\u6e90\u4f9d\u636e\uff1a")
                or stripped.startswith("\u6765\u6e90\u4f9d\u636e:")
            ):
                skip_indent = len(line) - len(line.lstrip(" "))
                continue

            if skip_indent is not None:
                if not stripped:
                    continue
                indent = len(line) - len(line.lstrip(" "))
                if indent > skip_indent:
                    continue
                skip_indent = None

            result.append(line.rstrip())

        cleaned = "\n".join(result).strip()
        while "\n\n\n" in cleaned:
            cleaned = cleaned.replace("\n\n\n", "\n\n")
        return cleaned + "\n"

    def _framework_to_script_txt(asset: dict) -> str:
        package = asset.get("framework_plan_package") if isinstance(asset.get("framework_plan_package"), dict) else {}
        stage_outputs = asset.get("stage_outputs") if isinstance(asset.get("stage_outputs"), dict) else {}
        workspace_state = asset.get("framework_to_script_state") if isinstance(asset.get("framework_to_script_state"), dict) else {}
        script_stages = asset.get("scriptStages") if isinstance(asset.get("scriptStages"), dict) else {}
        if not script_stages and isinstance(workspace_state.get("scriptStages"), dict):
            script_stages = workspace_state.get("scriptStages")
        stage03 = _get_dict_alias(stage_outputs, "stage03")
        stage06 = _get_dict_alias(stage_outputs, "stage06")
        stage07 = _get_dict_alias(stage_outputs, "stage07")
        stage08 = _get_dict_alias(script_stages, "stage08")
        stage12 = _get_dict_alias(script_stages, "stage12")

        stage07_package = _get_dict_alias(stage07, "frameworkPlanPackage", "framework_plan_package")
        basic_config = _get_dict_alias(package, "basic_config")
        source_brief = _first_present(package, "source_brief", "sourceBrief", default=None)
        source_brief_dict = source_brief if isinstance(source_brief, dict) else {}
        adaptation_guide = _first_present(package, "adaptation_guide", "adaptationGuide", default=None)
        story = (
            _first_present(package, "story_synopsis", "synopsis", "summary", default=None)
            or _first_present(stage07_package, "story_synopsis", "synopsis", "summary", default=None)
            or _first_present(source_brief_dict, "story_synopsis", "synopsis", "summary", "source_summary", "source_text", default=None)
            or _first_present(stage_outputs, "source_brief", "sourceBrief", default=None)
            or _first_present(stage06, "overallAdaptationGuide", "overall_adaptation_guide", default=None)
            or adaptation_guide
            or _first_present(stage_outputs, "overallAdaptationGuide", "overall_adaptation_guide", "adaptation_guide", default=None)
            or _first_present(basic_config, "source_text", "adaptation_direction", "user_requirements", default=None)
            or "暂无"
        )
        characters = (
            _first_present(package, "characterPlan", "character_plan", default=None)
            or _first_present(stage03, "characterPlan", "character_plan", default=None)
            or _first_present(stage_outputs, "characterPlan", "character_plan", default=None)
            or _first_present(package, "character_storylines", "characterStorylines", default=None)
            or _first_present(stage_outputs, "character_storylines", "characterStorylines", default=None)
            or "暂无"
        )
        scenes = (
            _first_present(stage08, "sceneDictionary", "scene_dictionary", default=None)
            or _first_present(package, "sceneDictionary", "scene_dictionary", "coreScenes", "core_scenes", default=None)
            or _first_present(stage_outputs, "sceneDictionary", "scene_dictionary", "coreScenes", "core_scenes", default=None)
            or _first_present(source_brief_dict, "core_scenes", "coreScenes", "key_scenes", default=None)
            or "暂无"
        )
        script_text = _txt_readable(_framework_script_text_from_batches(stage12) or "暂无")
        title = _txt_scalar(asset.get("title") or package.get("title") or package.get("project_title") or "未命名框架剧本")
        export_text = "\n\n".join(
            [
                f"《{title}》",
                "一、故事梗概\n" + _txt_readable(story),
                "二、人物小传\n" + _txt_readable(characters),
                "三、核心场景\n" + _txt_readable(scenes),
                "四、剧本正文\n" + script_text,
            ]
        ) + "\n"
        return _clean_framework_to_script_export_text(export_text)

    def _framework_to_script_export_filename(asset: dict, extension: str) -> str:
        package = asset.get("framework_plan_package") if isinstance(asset.get("framework_plan_package"), dict) else {}
        title = _txt_scalar(
            asset.get("title")
            or asset.get("source_title")
            or package.get("title")
            or package.get("project_title")
            or "完整剧本"
        )
        total_episodes = _positive_int(
            asset.get("episodes_per_season")
            or asset.get("total_episodes")
            or package.get("episodes_per_season")
            or package.get("total_episodes"),
            0,
        )
        safe_title = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in str(title or "完整剧本"))
        episode_text = f"{total_episodes}集" if total_episodes else ""
        return f"{(safe_title[:48] or '完整剧本')}{episode_text}完整剧本.{extension}"

    def _framework_review_needs_rewrite(review: dict) -> bool:
        if not isinstance(review, dict):
            return False
        if bool(review.get("rewriteRequired") or review.get("rewrite_required")):
            return True
        if (
            review.get("reviewPassed") is False
            or review.get("passed") is False
            or review.get("approved") is False
        ):
            return True
        return False

    def _request_auth_token() -> str:
        auth_header = str(request.headers.get("Authorization") or "").strip()
        if auth_header.lower().startswith("bearer "):
            return auth_header[7:].strip()
        return str(
            request.args.get("auth_token")
            or request.form.get("auth_token")
            or session.get("auth_token")
            or ""
        ).strip()

    def _current_user():
        token = _request_auth_token()
        if token:
            user = auth_store.get_user_by_token(token)
            if user:
                return user
        user_id = session.get("user_id")
        if user_id is None:
            return None
        try:
            return auth_store.get_user(int(user_id))
        except (TypeError, ValueError):
            return None

    def _current_auth_token() -> str:
        return _request_auth_token()

    def _login_user(user) -> str:
        session.clear()
        auth_token = auth_store.create_session_token(user.id)
        session["user_id"] = int(user.id)
        session["auth_token"] = auth_token
        return auth_token

    def _logout_user() -> None:
        session.clear()

    def _login_required(view):
        @wraps(view)
        def wrapper(*args, **kwargs):
            if not _current_user():
                if request.path.startswith("/api/"):
                    return _json_error("请先登录", status=401)
                return redirect(url_for("login_page"))
            return view(*args, **kwargs)

        return wrapper

    def _require_user_id() -> int:
        user = _current_user()
        if not user:
            raise ValueError("请先登录")
        return int(user.id)

    def _request_ip_address() -> str:
        forwarded = str(request.headers.get("X-Forwarded-For") or "").strip()
        if forwarded:
            return forwarded.split(",", 1)[0].strip()[:80]
        return str(request.remote_addr or "").strip()[:80]

    def _admin_required(view):
        @wraps(view)
        def wrapper(*args, **kwargs):
            user = _current_user()
            if not user:
                if request.path.startswith("/api/"):
                    return _json_error("请先登录", status=401)
                return redirect(url_for("login_page", next="admin"))
            if not admin_store.is_admin(user.id, user.username):
                if request.path.startswith("/api/"):
                    return _json_error("没有管理员权限", status=403)
                return render_template("admin_forbidden.html", current_user=user), 403
            return view(*args, **kwargs)

        return wrapper

    @app.context_processor
    def _inject_admin_template_context():
        user = _current_user()
        return {
            "current_user_is_admin": bool(user and admin_store.is_admin(user.id, user.username)),
            "registration_enabled": bool(app.config["REGISTRATION_ENABLED"]),
            "admin_only_login": bool(app.config["ADMIN_ONLY_LOGIN"]),
        }

    def _workflow_request_identity(path: str) -> tuple[str, str] | None:
        value = str(path or "")
        if value.startswith("/api/tools/") and value.endswith("/run"):
            key = value.removeprefix("/api/tools/").removesuffix("/run").strip("/")
            labels = {
                "hot_review": "爆款文审核",
                "reskin": "换皮",
                "punchup": "增加爽感",
                "character_reskin": "只换人设",
                "sitcom": "情景剧生成",
            }
            return f"tool:{key}", labels.get(key, key)
        if value == "/api/workflows/start":
            return "script_generation", "剧本生成"
        if value.startswith("/api/framework-planner/stage/"):
            stage = value.rsplit("/", 1)[-1]
            return f"framework_stage_{stage}", f"框架流程 {stage}"
        if value.startswith("/api/framework-to-script/stage/"):
            stage = value.rsplit("/", 1)[-1]
            return f"script_stage_{stage}", f"剧本流程 {stage}"
        if value == "/api/framework-planner/generate-script":
            return "framework_generate_script", "框架生成剧本"
        if value == "/api/tools/new-framework":
            return "tool:new_framework", "新框架生成"
        if value == "/api/workbuddy/doctor/run":
            return "workbuddy_doctor", "AI剧本医生"
        return None

    @app.before_request
    def _begin_admin_audit_capture():
        if request.method not in {"POST", "PATCH", "PUT", "DELETE"}:
            return None
        if request.path in {"/login", "/register", "/logout"}:
            return None
        user = _current_user()
        g.admin_audit_started = time.perf_counter()
        g.admin_audit_user_id = int(user.id) if user else None
        g.admin_audit_username = str(user.username) if user else ""
        workflow = _workflow_request_identity(request.path)
        if workflow:
            workflow_key, workflow_label = workflow
            try:
                g.admin_workflow_run_id = admin_store.start_workflow_run(
                    user_id=g.admin_audit_user_id,
                    username=g.admin_audit_username,
                    workflow_key=workflow_key,
                    workflow_label=workflow_label,
                    http_method=request.method,
                    path=request.path,
                    ip_address=_request_ip_address(),
                    request_bytes=int(request.content_length or 0),
                )
            except Exception:
                logger.exception("admin workflow run start capture failed path=%s", request.path)
        return None

    @app.after_request
    def _finish_admin_audit_capture(response):
        started = getattr(g, "admin_audit_started", None)
        if started is None:
            return response
        duration_ms = max(0, int((time.perf_counter() - started) * 1000))
        status = "success" if response.status_code < 400 else "failed"
        try:
            admin_store.record_event(
                user_id=getattr(g, "admin_audit_user_id", None),
                username=getattr(g, "admin_audit_username", ""),
                category="request",
                action=f"{request.method} {request.path}",
                status=status,
                target_type="endpoint",
                target_id=request.path,
                http_method=request.method,
                path=request.path,
                ip_address=_request_ip_address(),
                user_agent=str(request.user_agent or ""),
                duration_ms=duration_ms,
                metadata={"http_status": response.status_code},
            )
        except Exception:
            logger.exception("admin audit capture failed path=%s", request.path)
        run_id = getattr(g, "admin_workflow_run_id", "")
        if run_id:
            run_status = "accepted" if request.path == "/api/workflows/start" and response.status_code < 400 else status
            try:
                admin_store.finish_workflow_run(
                    run_id,
                    status=run_status,
                    duration_ms=duration_ms,
                    http_status=response.status_code,
                    response_bytes=int(response.calculate_content_length() or 0),
                    error_code="" if response.status_code < 400 else f"HTTP_{response.status_code}",
                )
            except Exception:
                logger.exception("admin workflow run finish capture failed run_id=%s", run_id)
        return response

    def _resolve_spec_path(data: dict) -> str:
        custom = str(data.get("workflow_spec_path") or "").strip()
        return custom or str(app.config["WORKFLOW_SPEC_PATH"])

    def _coerce_string_list(value) -> list[str]:
        if value is None:
            return []
        if isinstance(value, str):
            return [value] if value.strip() else []
        if not isinstance(value, list):
            return []
        result: list[str] = []
        for item in value:
            text = str(item.get("id") if isinstance(item, dict) else item or "").strip()
            if text and text not in result:
                result.append(text)
        return result

    def _coerce_tag_list(value) -> list[dict]:
        if not isinstance(value, list):
            return []
        result: list[dict] = []
        for item in value:
            if isinstance(item, dict):
                tag = {
                    "id": str(item.get("id") or "").strip(),
                    "name": str(item.get("name") or "").strip(),
                    "category": str(item.get("category") or "").strip(),
                    "builtin": bool(item.get("builtin")),
                    "group": str(item.get("group") or "").strip(),
                    "group_label": str(item.get("group_label") or "").strip(),
                    "source": str(item.get("source") or "").strip(),
                    "type": str(item.get("type") or "").strip(),
                    "is_default": bool(item.get("is_default")) if "is_default" in item else bool(item.get("builtin")),
                    "is_user_editable": item.get("is_user_editable") is not False,
                    "description": str(item.get("description") or "").strip(),
                    "prompt_text": str(item.get("prompt_text") or "").strip(),
                    "stage_prompts": _coerce_stage_prompts(item.get("stage_prompts")),
                }
                if tag["id"] or tag["name"]:
                    result.append(tag)
            else:
                text = str(item or "").strip()
                if text:
                    result.append({"id": text, "name": text, "category": "", "builtin": False, "description": "", "prompt_text": "", "stage_prompts": _coerce_stage_prompts({})})
        return result

    _USER_KNOWLEDGE_STAGE_LABELS = {
        "basic": "01 原文提取偏好",
        "worldview": "02 世界观偏好",
        "character": "03 人设偏好",
        "beat": "04 节拍规划偏好",
        "storylines": "05 人物故事线偏好",
        "guide": "06 改编指引偏好",
        "package": "07 框架校验偏好",
        "scene": "08 场景字典偏好",
        "appearance": "09 确定角色外观偏好",
        "episode": "10 分集细化偏好",
        "conflict": "11 开头冲突钩子偏好",
        "script_text": "12 正文写作偏好",
    }

    def _coerce_stage_prompts(value) -> dict:
        source = value if isinstance(value, dict) else {}
        return {
            key: _coerce_prompt_text(source.get(key))
            for key in _USER_KNOWLEDGE_STAGE_LABELS
        }

    def _merge_stage_prompts_non_empty(*sources) -> dict:
        result = _coerce_stage_prompts({})
        for source in sources:
            normalized = _stage_prompt_lookup_from_mapping(source) if isinstance(source, dict) else _coerce_stage_prompts({})
            for key, value in normalized.items():
                text = _coerce_prompt_text(value)
                if text:
                    result[key] = text
        return result

    def _coerce_prompt_text(value) -> str:
        if value is None:
            return ""
        if isinstance(value, str):
            return value.strip()
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False)
        return str(value).strip()

    def _stage_prompt_lookup_from_mapping(value) -> dict:
        source = value if isinstance(value, dict) else {}
        result = {}
        stage_numbers = {
            "basic": "01",
            "worldview": "02",
            "character": "03",
            "beat": "04",
            "storylines": "05",
            "guide": "06",
            "package": "07",
            "scene": "08",
            "appearance": "09",
            "episode": "10",
            "conflict": "11",
            "script_text": "12",
        }
        for stage_key, label in _USER_KNOWLEDGE_STAGE_LABELS.items():
            stage_no = stage_numbers.get(stage_key, "")
            result[stage_key] = _coerce_prompt_text(
                source.get(stage_key)
                or source.get(stage_no)
                or source.get(label)
            )
        return result

    def _stage_prompts_from_snapshot(snapshot) -> dict:
        if not isinstance(snapshot, dict):
            return {}
        result = {}
        for key in ("stage_prompts", "stagePrompts", "stage_preferences", "stagePreferences"):
            prompts = snapshot.get(key)
            if isinstance(prompts, dict):
                normalized = _stage_prompt_lookup_from_mapping(prompts)
                for stage_key, text in normalized.items():
                    if text and not result.get(stage_key):
                        result[stage_key] = text
        return result

    def _stage_prompts_from_payload(data: dict) -> dict:
        if not isinstance(data, dict):
            return {}
        sources = [
            data.get("stage_prompts"),
            data.get("stagePrompts"),
            data.get("stagePreferences"),
            data.get("stage_preferences"),
            data.get("user_knowledge_stage_prompts"),
        ]
        prompt_preferences = data.get("prompt_preferences") if isinstance(data.get("prompt_preferences"), dict) else {}
        sources.append(prompt_preferences.get("stage_prompts"))
        framework_plan_package = data.get("framework_plan_package") or data.get("frameworkPlanPackage")
        if isinstance(framework_plan_package, dict):
            sources.append(framework_plan_package.get("stage_prompts"))
            sources.append(framework_plan_package.get("user_knowledge_stage_prompts"))
            package_prompt_preferences = framework_plan_package.get("prompt_preferences")
            if isinstance(package_prompt_preferences, dict):
                sources.append(package_prompt_preferences.get("stage_prompts"))
        return _merge_stage_prompts_non_empty(*sources)

    def _framework_asset_stage_prompts(framework_asset: dict | None) -> dict:
        if not isinstance(framework_asset, dict):
            return {}
        direct = _stage_prompt_lookup_from_mapping(framework_asset.get("stage_prompts"))
        if any(direct.values()):
            return direct
        return _stage_prompts_from_snapshot(framework_asset.get("preference_snapshot"))

    def _framework_context_vars(data: dict, framework_plan_package: dict | None) -> dict:
        package = framework_plan_package if isinstance(framework_plan_package, dict) else {}
        worldview_plan = (
            data.get("worldview_plan")
            or data.get("worldviewPlan")
            or package.get("worldview_plan")
            or package.get("worldviewPlan")
            or {}
        )
        character_plan = (
            data.get("character_plan")
            or data.get("characterPlan")
            or package.get("character_plan")
            or package.get("characterPlan")
            or {}
        )
        beat_checkpoint_timeline = (
            data.get("beat_checkpoint_timeline")
            or data.get("beatCheckpointTimeline")
            or package.get("beat_checkpoint_timeline")
            or package.get("beatCheckpointTimeline")
            or []
        )
        character_storylines = (
            data.get("character_storylines")
            or data.get("characterStorylines")
            or package.get("character_storylines")
            or package.get("characterStorylines")
            or []
        )
        alias_rules = (
            data.get("character_alias_naming_rules")
            or data.get("characterAliasNamingRules")
            or package.get("character_alias_naming_rules")
            or package.get("characterAliasNamingRules")
            or ""
        )
        core_scene_input = (
            data.get("core_scene_input")
            or data.get("coreSceneInput")
            or package.get("core_scene_input")
            or package.get("coreSceneInput")
            or ""
        )
        return {
            "frameworkPlanPackage": package,
            "framework_plan_package": package,
            "worldviewPlan": worldview_plan,
            "worldview_plan": worldview_plan,
            "characterPlan": character_plan,
            "character_plan": character_plan,
            "beatCheckpointTimeline": beat_checkpoint_timeline,
            "beat_checkpoint_timeline": beat_checkpoint_timeline,
            "characterStorylines": character_storylines,
            "character_storylines": character_storylines,
            "characterAliasNamingRules": alias_rules,
            "character_alias_naming_rules": alias_rules,
            "coreSceneInput": core_scene_input,
            "core_scene_input": core_scene_input,
        }

    def _framework_planner_stage_key(stage) -> str:
        return {
            "01": "basic",
            "02": "worldview",
            "03": "character",
            "04": "beat",
            "05": "storylines",
            "06": "guide",
            "07": "package",
        }.get(str(stage or "").zfill(2), "")

    def _user_knowledge_stage_key(stage) -> str:
        return {
            "01": "basic",
            "02": "worldview",
            "03": "character",
            "04": "beat",
            "05": "storylines",
            "06": "guide",
            "07": "package",
            "08": "scene",
            "09": "appearance",
            "10": "episode",
            "11": "conflict",
            "12": "script_text",
        }.get(str(stage or "").zfill(2), "")

    def _stage_preference_from_tags(stage: str, selected_tags: list[dict]) -> tuple[str, int]:
        stage_key = _user_knowledge_stage_key(stage)
        if not stage_key:
            return "", 0
        sections = []
        for tag in selected_tags or []:
            prompts = tag.get("stage_prompts") if isinstance(tag.get("stage_prompts"), dict) else {}
            text = _coerce_prompt_text(prompts.get(stage_key))
            if not text:
                continue
            name = str(tag.get("name") or tag.get("id") or "").strip()
            label = _USER_KNOWLEDGE_STAGE_LABELS.get(stage_key, stage_key)
            sections.append(f"【智慧库标签偏好：{name} / {label}】\n{text}")
        return "\n\n".join(sections), len(sections)

    def _resolve_stage_preference(data: dict, stage: str, framework_asset: dict | None = None) -> tuple[str, str]:
        stage_no = str(stage or "").zfill(2)
        stage_key = _user_knowledge_stage_key(stage_no)
        request_stage_prompts = _stage_prompts_from_payload(data)
        preference = _coerce_prompt_text(request_stage_prompts.get(stage_key))
        if preference:
            return preference, "request_stage_prompts"

        asset_stage_prompts = _framework_asset_stage_prompts(framework_asset)
        preference = _coerce_prompt_text(asset_stage_prompts.get(stage_key))
        if preference:
            return preference, "framework_asset_stage_prompts"

        snapshot = data.get("preference_snapshot") if isinstance(data.get("preference_snapshot"), dict) else {}
        if not snapshot and isinstance(data.get("preferenceSnapshot"), dict):
            snapshot = data.get("preferenceSnapshot")
        if not snapshot and isinstance(data.get("metadata"), dict):
            metadata_snapshot = data["metadata"].get("preference_snapshot") or data["metadata"].get("preferenceSnapshot")
            snapshot = metadata_snapshot if isinstance(metadata_snapshot, dict) else {}
        snapshot_stage_prompts = _stage_prompts_from_snapshot(snapshot)
        preference = _coerce_prompt_text(snapshot_stage_prompts.get(stage_key))
        if preference:
            return preference, "source_framework_stage_prompts"

        for key in (
            "stagePreference",
            "stage_preference",
            "stage_preference_prompt",
            "user_stage_preference_prompt",
            "userPreference",
            "user_preferences",
            "userPreferences",
            "userRequirements",
            "user_constraints",
        ):
            if key in data:
                preference = _coerce_prompt_text(data.get(key))
                if preference:
                    return preference, "request_legacy_preference"

        preference = _coerce_prompt_text(data.get("prompt_text"))
        if preference:
            return preference, "prompt_text_fallback"
        return "", "none"

    def _inject_snapshot_stage_preference(
        variables: dict,
        data: dict,
        stage: str,
        *,
        framework_asset: dict | None,
        workflow_stage: str | None = None,
    ) -> dict:
        preference, source = _resolve_stage_preference(data, stage, framework_asset)
        stage_no = str(stage or "").zfill(2)
        stage_key = _user_knowledge_stage_key(stage_no)
        stage_prompt_key = stage_prompt_key_for(workflow_stage or stage_no) or stage_key
        workflow_preference_keys = tuple(
            FRAMEWORK_TO_SCRIPT_STAGE_PREFS.get(workflow_stage or stage_no, {}).get("workflow_keys")
            or preference_keys_for(workflow_stage or stage_no)
        )
        asset_id = str(
            (framework_asset or {}).get("asset_id")
            or data.get("framework_asset_id")
            or data.get("asset_id")
            or ""
        ).strip()
        inject_stage_preference(variables, preference, workflow_preference_keys)
        logger.info(
            "framework_to_script stage preference injected",
            extra={
                "stage": workflow_stage or stage_no,
                "stage_prompt_key": stage_prompt_key,
                "has_stage_preference": bool(preference),
                "preference_length": len(preference or ""),
                "workflow_preference_keys": list(workflow_preference_keys),
                "preference_source": source,
                "asset_id": asset_id,
            },
        )
        return variables

    def _resolve_saved_preference_tags(user_id: int | str, selected_ids: list[str]) -> list[dict]:
        if not selected_ids:
            try:
                selected_ids = _coerce_string_list(
                    user_knowledge_store.get_preferences(user_id).get("selected_preference_tag_ids")
                )
            except Exception:
                selected_ids = []
        if not selected_ids:
            return []
        tags_by_id = {str(tag.get("id") or ""): tag for tag in user_knowledge_store.list_tags(enabled_only=True, user_id=user_id)}
        result = []
        for tag_id in selected_ids:
            tag = tags_by_id.get(str(tag_id))
            if tag:
                result.append(tag)
        return result

    def _inject_stage_preference_variables(variables: dict, data: dict, stage: str, *, user_id: int | str) -> dict:
        selected_tags = _coerce_tag_list(data.get("selected_preference_tags"))
        selected_ids = _coerce_string_list(data.get("selected_preference_tag_ids"))
        if selected_tags and not selected_ids:
            selected_ids = [tag["id"] for tag in selected_tags if tag.get("id")]
        if not selected_tags:
            selected_tags = _resolve_saved_preference_tags(user_id, selected_ids)
            if not selected_ids:
                selected_ids = [tag["id"] for tag in selected_tags if tag.get("id")]
        stage_key = _user_knowledge_stage_key(stage)
        stage_label = _USER_KNOWLEDGE_STAGE_LABELS.get(stage_key, "")
        stage_preference_prompt, tags_with_preference_count = _stage_preference_from_tags(stage, selected_tags)
        if stage_preference_prompt:
            source_names = "、".join(str(tag.get("name") or tag.get("id") or "").strip() for tag in selected_tags if tag.get("id"))
            header = f"【当前阶段智慧库偏好：{stage_label}】\n来源标签：{source_names}\n"
            stage_preference_prompt = header + stage_preference_prompt
            for key in (
                "stagePreference",
                "stage_preference",
                "stage_preference_prompt",
                "user_stage_preference_prompt",
                "user_preferences",
                "userPreferences",
                "userRequirements",
                "user_constraints",
            ):
                if key not in variables or not _coerce_prompt_text(variables.get(key)):
                    variables[key] = stage_preference_prompt
                elif key in {"user_preferences", "userPreferences", "userRequirements", "user_constraints"}:
                    variables[key] = f"{_coerce_prompt_text(variables.get(key))}\n\n{stage_preference_prompt}"
        logger.info(
            "stage preference injection: stage_key=%s stage_name=%s preference_source=智慧库标签 preference_stage_key=%s preference_stage_label=%s selected_tag_count=%s tags_with_stage_preference_count=%s has_stage_preference=%s preference_length=%s",
            str(stage or "").zfill(2),
            stage_label,
            stage_key,
            stage_label,
            len(selected_tags or selected_ids),
            tags_with_preference_count,
            bool(stage_preference_prompt),
            len(stage_preference_prompt),
        )
        return variables

    def _attach_user_knowledge_payload(payload: dict, data: dict, stage: str | None = None) -> None:
        selected_tags = _coerce_tag_list(data.get("selected_preference_tags"))
        selected_ids = _coerce_string_list(data.get("selected_preference_tag_ids"))
        if not selected_ids and selected_tags:
            selected_ids = [tag["id"] for tag in selected_tags if tag.get("id")]
        if not selected_tags and selected_ids:
            current_user = _current_user()
            selected_tags = _resolve_saved_preference_tags(current_user.id if current_user else "", selected_ids)
        payload["selected_preference_tags"] = selected_tags
        payload["selected_preference_tag_ids"] = selected_ids
        payload["user_preference_prompt"] = _coerce_prompt_text(data.get("user_preference_prompt"))
        payload["user_knowledge_tag_prompt"] = _coerce_prompt_text(data.get("user_knowledge_tag_prompt"))
        prompt_preferences = data.get("prompt_preferences") if isinstance(data.get("prompt_preferences"), dict) else {}
        prompt_preferences = dict(prompt_preferences)
        tag_stage_prompts = {}
        if selected_tags:
            for key in _USER_KNOWLEDGE_STAGE_LABELS:
                sections = []
                for tag in selected_tags:
                    prompts = tag.get("stage_prompts") if isinstance(tag.get("stage_prompts"), dict) else {}
                    text = _coerce_prompt_text(prompts.get(key))
                    if not text:
                        continue
                    name = str(tag.get("name") or tag.get("id") or "").strip()
                    label = _USER_KNOWLEDGE_STAGE_LABELS.get(key, key)
                    sections.append(f"【智慧库标签偏好：{name} / {label}】\n{text}")
                if sections:
                    tag_stage_prompts[key] = "\n\n".join(sections)
        merged_stage_prompts = _merge_stage_prompts_non_empty(
            prompt_preferences.get("stage_prompts"),
            data.get("stage_prompts"),
            data.get("stagePrompts"),
            data.get("stage_preferences"),
            data.get("stagePreferences"),
            data.get("user_knowledge_stage_prompts"),
            tag_stage_prompts,
        )
        prompt_preferences["stage_prompts"] = merged_stage_prompts
        payload["user_knowledge_stage_prompts"] = merged_stage_prompts
        payload["prompt_preferences"] = prompt_preferences
        stage_key = _framework_planner_stage_key(stage)
        current_stage_prompt = payload["user_knowledge_stage_prompts"].get(stage_key, "") if stage_key else ""
        if current_stage_prompt:
            payload["stage_preference_prompt"] = current_stage_prompt
            payload["user_stage_preference_prompt"] = current_stage_prompt
            payload["user_preference_prompt"] = current_stage_prompt
        if any(key in data for key in ("selected_preference_tags", "selected_preference_tag_ids", "user_preference_prompt", "user_knowledge_tag_prompt", "user_knowledge_stage_prompts", "prompt_preferences")):
            selected_prompt, tags_with_preference_count = _stage_preference_from_tags(stage or "", selected_tags)
            logger.info(
                "workflow user knowledge fields: stage_key=%s stage_name=%s preference_source=智慧库标签 preference_stage_key=%s preference_stage_label=%s selected_tag_count=%s tags_with_stage_preference_count=%s has_stage_preference=%s preference_length=%s",
                str(stage or "").zfill(2) if stage else "",
                _USER_KNOWLEDGE_STAGE_LABELS.get(_user_knowledge_stage_key(stage), ""),
                stage_key,
                _USER_KNOWLEDGE_STAGE_LABELS.get(stage_key, ""),
                len(selected_tags or selected_ids),
                tags_with_preference_count,
                bool(current_stage_prompt or selected_prompt),
                len(current_stage_prompt or selected_prompt),
            )

    @app.get("/")
    def index():
        if _current_user():
            return redirect(url_for("workspace_page", auth_token=_current_auth_token()))
        return render_template(
            "home.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
            community_assets=task_manager.list_public_assets(),
        )

    @app.get("/workspace")
    def workspace_page():
        return render_template(
            "index.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
        )

    @app.get("/workbuddy-studio")
    @_login_required
    def workbuddy_studio_page():
        return render_template(
            "workbuddy_studio.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
        )

    @app.get("/agent-studio")
    @_login_required
    def agent_studio_page():
        return render_template(
            "agent_studio.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
        )

    @app.get("/api/agent/status")
    @_login_required
    def agent_status_api():
        return _json_ok(agent=platform_conversation_agent.status())

    def _public_agent_message(message: dict[str, Any] | None) -> dict[str, Any]:
        public_message = copy.deepcopy(message) if isinstance(message, dict) else {}
        metadata = public_message.get("metadata") if isinstance(public_message.get("metadata"), dict) else {}
        if "model" in metadata:
            metadata["model"] = "剧本 Agent"
        public_message["metadata"] = metadata
        return public_message

    @app.route("/api/agent/conversations", methods=["GET", "POST"])
    @_login_required
    def agent_conversations_api():
        user = _current_user()
        if request.method == "GET":
            return _json_ok(conversations=agent_conversation_store.list(user.id))
        data = request.get_json(silent=True) or {}
        conversation = agent_conversation_store.create(
            user.id,
            title=str(data.get("title") or "新的创作对话"),
        )
        welcome = agent_conversation_store.add_message(
            user.id,
            conversation["id"],
            role="assistant",
            content="告诉我你想创作什么剧本。我会从对话里提取题材、集数、角色和风格，只追问缺少的必要信息，确认后直接调度现有剧本平台。",
            metadata={"kind": "welcome", "model": "剧本 Agent"},
        )
        return _json_ok(conversation=conversation, messages=[welcome])

    @app.route("/api/agent/conversations/<conversation_id>", methods=["GET", "DELETE"])
    @_login_required
    def agent_conversation_api(conversation_id: str):
        user = _current_user()
        conversation = agent_conversation_store.get(user.id, conversation_id)
        if not conversation:
            return _json_error("对话不存在或无权访问。", status=404)
        if request.method == "DELETE":
            agent_conversation_store.delete(user.id, conversation_id)
            return _json_ok(conversation_id=conversation_id)
        return _json_ok(
            conversation=conversation,
            messages=[
                _public_agent_message(item)
                for item in agent_conversation_store.messages(user.id, conversation_id)
            ],
            context=platform_conversation_agent._conversation_context(conversation, user.id),
        )

    @app.post("/api/agent/conversations/<conversation_id>/messages")
    @_login_required
    def agent_conversation_message_api(conversation_id: str):
        user = _current_user()
        data = request.get_json(silent=True) or {}
        content = str(data.get("content") or "").strip()
        selected_skill = str(data.get("selected_skill") or "").strip()
        distilled_skill_id = str(data.get("distilled_skill_id") or "").strip()
        distilled_skill_version_id = str(data.get("distilled_skill_version_id") or "").strip()
        selected_knowledge_tag_ids = data.get("selected_knowledge_tag_ids")
        if not isinstance(selected_knowledge_tag_ids, list):
            selected_knowledge_tag_ids = []
        attachment_id = str(data.get("attachment_id") or "").strip()
        request_id = str(data.get("request_id") or uuid.uuid4().hex).strip()
        if not content:
            return _json_error("请输入要交给智能体的内容。", status=400)
        if len(content) > 12000:
            return _json_error("单条消息不能超过12000字。", status=400)
        if not platform_conversation_agent.status().get("configured"):
            return _json_error("剧本 Agent 尚未配置完成。", status=501)
        try:
            result = platform_conversation_agent.handle_message(
                user_id=user.id,
                username=user.username,
                conversation_id=conversation_id,
                content=content,
                request_id=request_id,
                selected_skill=selected_skill,
                distilled_skill_id=distilled_skill_id,
                distilled_skill_version_id=distilled_skill_version_id,
                selected_knowledge_tag_ids=selected_knowledge_tag_ids,
                attachment_id=attachment_id,
                internal_api_base_url=f"http://127.0.0.1:{int(request.environ.get('SERVER_PORT') or 5002)}",
                internal_auth_token=_current_auth_token() or auth_store.create_session_token(user.id),
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        if not result.get("success"):
            return jsonify(result), 502
        public_result = copy.deepcopy(result)
        public_result["model"] = "剧本 Agent"
        if isinstance(public_result.get("user_message"), dict):
            public_result["user_message"] = _public_agent_message(public_result["user_message"])
        if isinstance(public_result.get("assistant_message"), dict):
            public_result["assistant_message"] = _public_agent_message(public_result["assistant_message"])
        return jsonify(public_result)

    def _request_codebuddy_api_key(data: dict[str, Any] | None = None) -> str | None:
        value = (
            request.headers.get("X-CodeBuddy-Api-Key")
            or request.headers.get("X-Codebuddy-Api-Key")
            or (data or {}).get("codebuddy_api_key")
            or ""
        )
        value = str(value).strip()
        return value or None

    def _request_codebuddy_internet_env(data: dict[str, Any] | None = None) -> str | None:
        value = (
            request.headers.get("X-CodeBuddy-Internet-Environment")
            or request.headers.get("X-Codebuddy-Internet-Environment")
            or (data or {}).get("codebuddy_internet_environment")
            or ""
        )
        value = str(value).strip()
        return value or None

    def _request_codebuddy_model(data: dict[str, Any] | None = None) -> str | None:
        value = request.headers.get("X-CodeBuddy-Model") or (data or {}).get("codebuddy_model") or ""
        value = str(value).strip()
        return value or None

    @app.get("/api/workbuddy/status")
    @_login_required
    def workbuddy_status():
        include_metadata = str(request.args.get("metadata") or "1").lower() not in {"0", "false", "no"}
        return jsonify(
            workbuddy_config_status(
                include_metadata=include_metadata,
                api_key=_request_codebuddy_api_key(),
                internet_env=_request_codebuddy_internet_env(),
                model=_request_codebuddy_model(),
            )
        )

    def _public_workbuddy_history_entry(entry: dict[str, Any] | None) -> dict[str, Any]:
        public_entry = copy.deepcopy(entry) if isinstance(entry, dict) else {}
        public_entry["provider"] = "script_agent"
        public_entry["model"] = "剧本 Agent"
        result = public_entry.get("result") if isinstance(public_entry.get("result"), dict) else None
        if result is not None:
            result["provider"] = "script_agent"
            result["model"] = "剧本 Agent"
        return public_entry

    @app.get("/api/workbuddy/doctor/history")
    @_login_required
    def workbuddy_doctor_history_list():
        user = _current_user()
        try:
            limit = min(max(int(request.args.get("limit") or 50), 1), 100)
        except (TypeError, ValueError):
            limit = 50
        items = [_public_workbuddy_history_entry(item) for item in list_workbuddy_history(user.id, limit=limit)]
        return jsonify({"ok": True, "items": items})

    @app.delete("/api/workbuddy/doctor/history")
    @_login_required
    def workbuddy_doctor_history_clear():
        user = _current_user()
        deleted = clear_workbuddy_history(user.id)
        return jsonify({"ok": True, "deleted": deleted, "items": []})

    @app.get("/api/workbuddy/doctor/history/<entry_id>")
    @_login_required
    def workbuddy_doctor_history_get(entry_id: str):
        user = _current_user()
        entry = load_workbuddy_history(user.id, entry_id)
        if not entry:
            return jsonify({"ok": False, "error": "历史记录不存在或已删除。"}), 404
        return jsonify({"ok": True, "entry": _public_workbuddy_history_entry(entry)})

    @app.delete("/api/workbuddy/doctor/history/<entry_id>")
    @_login_required
    def workbuddy_doctor_history_delete(entry_id: str):
        user = _current_user()
        deleted = delete_workbuddy_history(user.id, entry_id)
        items = [_public_workbuddy_history_entry(item) for item in list_workbuddy_history(user.id)]
        return jsonify({"ok": True, "deleted": deleted, "items": items})

    @app.post("/api/workbuddy/framework/generate")
    @_login_required
    def workbuddy_generate_framework():
        return jsonify(
            {
                "ok": False,
                "error": "该入口已调整为 AI剧本医生实验室。请使用 /api/workbuddy/doctor/run 做剧本二次质检。",
            }
        ), 410

    @app.post("/api/workbuddy/review/run")
    @_login_required
    def workbuddy_run_review():
        return jsonify(
            {
                "ok": False,
                "error": "该入口已调整为 AI剧本医生实验室。请使用 /api/workbuddy/doctor/run 做剧本二次质检。",
            }
        ), 410

    @app.post("/api/workbuddy/doctor/run")
    @_login_required
    def workbuddy_doctor_run():
        data = request.get_json(silent=True) or {}
        payload, error = validate_doctor_payload(data)
        if error:
            return jsonify({"ok": False, "error": error}), 400

        user = _current_user()
        if payload.get("source_document_id"):
            source = load_workbuddy_source(user.id, payload["source_document_id"])
            if not source:
                return jsonify({"ok": False, "error": "原始 Word 文件已失效，请重新上传 DOCX 后再运行体检。"}), 400
            payload["source_filename"] = str(source.get("original_filename") or "原始剧本.docx")

        codebuddy_api_key = _request_codebuddy_api_key(data)
        codebuddy_internet_env = _request_codebuddy_internet_env(data)
        codebuddy_model = payload.get("model") or _request_codebuddy_model(data)
        status = workbuddy_config_status(
            api_key=codebuddy_api_key,
            internet_env=codebuddy_internet_env,
            model=codebuddy_model,
        )
        prompt = build_doctor_prompt(
            payload["skill"],
            payload["title"],
            payload["script_text"],
            payload["user_goal"],
        )
        if not status["configured"]:
            return jsonify(
                {
                    "ok": False,
                    "error": "剧本 Agent 尚未配置完成，暂不能调用 AI 剧本医生。",
                    "missing": status["missing"],
                    "accepted_payload": {
                        "title": payload["title"],
                        "skill": payload["skill"],
                        "model": "剧本 Agent",
                        "script_length": len(payload["script_text"]),
                        "user_goal": payload["user_goal"],
                    },
                    "prompt_preview": prompt[:1600],
                }
            ), 501

        timeout_seconds = doctor_timeout_seconds(len(payload["script_text"]), payload["skill"])
        try:
            result = run_codebuddy_doctor(
                prompt,
                project_dir=Path.cwd(),
                timeout_seconds=timeout_seconds,
                model=codebuddy_model or None,
                api_key=codebuddy_api_key,
                internet_env=codebuddy_internet_env,
            )
        except Exception as exc:
            formatted_error = format_doctor_exception(exc, timeout_seconds=timeout_seconds)
            return jsonify(
                {
                    "ok": False,
                    "error": formatted_error["message"],
                    "error_type": formatted_error["type"],
                    "accepted_payload": {
                        "title": payload["title"],
                        "skill": payload["skill"],
                        "model": "剧本 Agent",
                        "script_length": len(payload["script_text"]),
                        "timeout_seconds": timeout_seconds,
                    },
                    "prompt_preview": prompt[:1200],
                }
            ), 502

        response_payload = {
            "ok": True,
            "provider": "script_agent",
            "title": payload["title"],
            "skill": payload["skill"],
            "script_length": len(payload["script_text"]),
            "timeout_seconds": timeout_seconds,
            "model": result.get("model"),
            "session_id": result.get("session_id"),
            "usage": result.get("usage"),
            "report": result.get("content") or "",
            "structured_output": result.get("structured_output"),
        }
        try:
            response_payload["history_entry"] = save_workbuddy_history(
                user.id,
                username=user.username,
                payload={**payload, "script_length": len(payload["script_text"])},
                result=response_payload,
            )
        except Exception as exc:
            logger.warning("WorkBuddy 历史保存失败: %s", exc)
            response_payload["history_warning"] = "报告已生成，但历史记录保存失败。"

        response_payload["model"] = "剧本 Agent"
        response_payload["provider"] = "script_agent"
        if isinstance(response_payload.get("history_entry"), dict):
            response_payload["history_entry"] = _public_workbuddy_history_entry(response_payload["history_entry"])

        return jsonify(response_payload)

    @app.post("/api/workbuddy/doctor/history/<entry_id>/optimize")
    @_login_required
    def workbuddy_doctor_optimize(entry_id: str):
        user = _current_user()
        entry = load_workbuddy_history(user.id, entry_id)
        if not entry or not entry.get("ok"):
            return jsonify({"ok": False, "error": "审查记录不存在或尚未成功完成。"}), 404

        source_document_id = str(entry.get("source_document_id") or "")
        source = load_workbuddy_source(user.id, source_document_id)
        if not source:
            return jsonify({"ok": False, "error": "这条记录没有可回写的原始 Word，请重新上传 DOCX 并运行审查。"}), 400

        report = extract_doctor_report(entry.get("result") or {})
        if not report:
            return jsonify({"ok": False, "error": "审查报告无法解析，请重新运行审查后再优化。"}), 400

        data = request.get_json(silent=True) or {}
        codebuddy_api_key = _request_codebuddy_api_key(data)
        codebuddy_internet_env = _request_codebuddy_internet_env(data)
        codebuddy_model = str(data.get("model") or _request_codebuddy_model(data) or "").strip()
        status = workbuddy_config_status(
            api_key=codebuddy_api_key,
            internet_env=codebuddy_internet_env,
            model=codebuddy_model,
        )
        if not status["configured"]:
            return jsonify(
                {
                    "ok": False,
                    "error": "剧本 Agent 尚未配置完成，暂不能执行一键优化。",
                    "missing": status["missing"],
                }
            ), 501

        conversation_id = str(data.get("conversation_id") or "").strip()
        background = bool(data.get("background")) and bool(conversation_id)
        if background:
            conversation = agent_conversation_store.get(user.id, conversation_id)
            if not conversation:
                return jsonify({"ok": False, "error": "对话不存在或无权访问。"}), 404
            state = dict(conversation.get("state") or {})
            existing_operation = state.get("active_operation") if isinstance(state.get("active_operation"), dict) else {}
            if existing_operation.get("type") == "word_optimization" and existing_operation.get("status") == "running":
                return jsonify({
                    "ok": True,
                    "accepted": True,
                    "already_running": True,
                    "operation": existing_operation,
                    "context": platform_conversation_agent._conversation_context(conversation, user.id),
                }), 202

            operation_id = uuid.uuid4().hex
            started_at = datetime.now(timezone.utc).astimezone().isoformat()
            operation = {
                "request_id": operation_id,
                "type": "word_optimization",
                "status": "running",
                "progress_percent": 20,
                "stage": "preparing",
                "message": "正在读取原始 Word 和剧本医生报告",
                "filename": str(source.get("original_filename") or entry.get("source_filename") or "原始剧本.docx"),
                "char_count": int(entry.get("script_length") or 0),
                "skill_key": str(entry.get("skill") or ""),
                "skill_name": str(entry.get("skill_name") or entry.get("skill") or "剧本医生"),
                "history_entry_id": entry_id,
                "started_at": started_at,
            }
            state["active_operation"] = operation
            platform_conversation_agent.remember_operation(state, operation)
            updated_conversation = agent_conversation_store.update(user.id, conversation_id, state=state) or conversation
            agent_conversation_store.add_message(
                user.id,
                conversation_id,
                role="assistant",
                content="Word 一键优化任务已经开始，现已转入后台。刷新或离开页面不会中止，完成后会在本对话提供下载按钮。",
                metadata={"background_operation_started": True, "operation_type": "word_optimization", "request_id": operation_id},
            )
            local_port = int(request.environ.get("SERVER_PORT") or 5002)
            local_url = f"http://127.0.0.1:{local_port}/api/workbuddy/doctor/history/{quote(entry_id, safe='')}/optimize"
            local_auth_token = _current_auth_token() or auth_store.create_session_token(user.id)

            def _run_background_word_optimization() -> None:
                current = agent_conversation_store.get(user.id, conversation_id) or updated_conversation
                current_state = dict(current.get("state") or {})
                active = dict(current_state.get("active_operation") or {})
                if str(active.get("request_id") or "") == operation_id:
                    active.update({
                        "progress_percent": 45,
                        "stage": "ai_optimize",
                        "message": "剧本 Agent 正在生成并校验 Word 修改方案",
                    })
                    current_state["active_operation"] = active
                    agent_conversation_store.update(user.id, conversation_id, state=current_state)
                try:
                    response = requests.post(
                        local_url,
                        headers={"Authorization": f"Bearer {local_auth_token}"},
                        json={"model": codebuddy_model} if codebuddy_model else {},
                        timeout=720,
                    )
                    try:
                        result_payload = response.json()
                    except Exception:
                        result_payload = {}
                    succeeded = response.ok and bool(result_payload.get("ok"))
                    if not succeeded:
                        result_payload = {
                            "ok": False,
                            "error": str(result_payload.get("error") or result_payload.get("message") or f"优化请求失败：HTTP {response.status_code}"),
                        }
                except Exception as exc:
                    succeeded = False
                    result_payload = {"ok": False, "error": f"Word 优化任务失败：{exc}"}

                current = agent_conversation_store.get(user.id, conversation_id) or current
                current_state = dict(current.get("state") or {})
                current_active = dict(current_state.get("active_operation") or {})
                if str(current_active.get("request_id") or "") == operation_id:
                    current_state.pop("active_operation", None)
                message = (
                    f"Word 优化完成，已应用 {int(result_payload.get('applied_count') or 0)} 处修改。"
                    if succeeded
                    else str(result_payload.get("error") or "Word 优化任务失败。")
                )
                current_state["last_operation"] = {
                    **operation,
                    "status": "completed" if succeeded else "failed",
                    "progress_percent": 100 if succeeded else 45,
                    "stage": "completed" if succeeded else "failed",
                    "message": message,
                    "finished_at": datetime.now(timezone.utc).astimezone().isoformat(),
                    "download_url": str(result_payload.get("download_url") or ""),
                    "download_name": str(result_payload.get("download_name") or ""),
                }
                platform_conversation_agent.remember_operation(current_state, current_state["last_operation"])
                agent_conversation_store.update(user.id, conversation_id, state=current_state)
                if succeeded:
                    event_result = {
                        "ok": True,
                        "filename": str(result_payload.get("download_name") or "AI优化版剧本.docx"),
                        "download_url": str(result_payload.get("download_url") or ""),
                        "applied_count": int(result_payload.get("applied_count") or 0),
                        "summary": str(result_payload.get("summary") or ""),
                        "ui": {"kind": "download"},
                    }
                    assistant_content = message + " 点击下方按钮即可下载修改后的 Word。"
                else:
                    event_result = {"ok": False, "error": message}
                    assistant_content = f"一键优化没有完成：{message}"
                agent_conversation_store.add_message(
                    user.id,
                    conversation_id,
                    role="assistant",
                    content=assistant_content,
                    metadata={
                        "events": [{"tool": "optimize_word", "result": event_result}],
                        "background_operation_completed": succeeded,
                        "background_operation_failed": not succeeded,
                        "request_id": operation_id,
                    },
                )

            thread = threading.Thread(
                target=_run_background_word_optimization,
                daemon=False,
                name=f"agent-word-optimize-{operation_id[:8]}",
            )
            thread.start()
            return jsonify({
                "ok": True,
                "accepted": True,
                "operation": operation,
                "context": platform_conversation_agent._conversation_context(updated_conversation, user.id),
            }), 202

        try:
            indexed_paragraphs = indexed_docx_text(source["path"])
        except Exception:
            logger.exception("WorkBuddy 原始 Word 读取失败: source=%s", source_document_id)
            return jsonify({"ok": False, "error": "原始 Word 文件损坏或无法读取，请重新上传。"}), 400

        prompt = build_doctor_optimization_prompt(
            skill_name=str(entry.get("skill_name") or entry.get("skill") or "剧本医生"),
            title=str(entry.get("title") or "未命名剧本"),
            report=report,
            indexed_paragraphs=indexed_paragraphs,
            user_goal=str(entry.get("user_goal") or ""),
        )
        timeout_seconds = min(600, max(420, doctor_timeout_seconds(len(indexed_paragraphs), str(entry.get("skill") or ""))))
        try:
            result = run_codebuddy_doctor(
                prompt,
                project_dir=Path.cwd(),
                timeout_seconds=timeout_seconds,
                model=codebuddy_model or None,
                api_key=codebuddy_api_key,
                internet_env=codebuddy_internet_env,
            )
            optimization = parse_doctor_optimization_result(result)
            optimized_raw, applied, skipped = apply_docx_replacements(
                source["path"],
                optimization["operations"],
            )
        except ValueError as exc:
            return jsonify({"ok": False, "error": str(exc)}), 422
        except Exception as exc:
            logger.exception("WorkBuddy 一键优化失败: history=%s", entry_id)
            formatted_error = format_doctor_exception(exc, timeout_seconds=timeout_seconds)
            return jsonify({"ok": False, "error": formatted_error["message"]}), 502

        if not applied:
            return jsonify(
                {
                    "ok": False,
                    "error": "AI 返回了修改建议，但没有一项通过原文匹配校验，Word 未被改动。请重新审查后再试。",
                    "skipped": skipped[:10],
                }
            ), 422

        output = save_workbuddy_optimized_document(
            user.id,
            history_entry_id=entry_id,
            original_filename=str(source.get("original_filename") or "原始剧本.docx"),
            raw=optimized_raw,
            applied_count=len(applied),
            summary=optimization["summary"],
        )
        return jsonify(
            {
                "ok": True,
                "summary": optimization["summary"],
                "applied_count": len(applied),
                "skipped_count": len(skipped),
                "download_name": output["download_name"],
                "download_url": url_for("workbuddy_doctor_download_optimized", output_id=output["output_id"]),
            }
        )

    @app.get("/api/workbuddy/doctor/optimized/<output_id>/download")
    @_login_required
    def workbuddy_doctor_download_optimized(output_id: str):
        user = _current_user()
        output = load_workbuddy_optimized_document(user.id, output_id)
        if not output:
            return jsonify({"ok": False, "error": "优化后的 Word 不存在或已失效。"}), 404
        return send_file(
            output["path"],
            as_attachment=True,
            download_name=str(output.get("download_name") or "AI优化版剧本.docx"),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    @app.get("/assets/framework")
    @_login_required
    def framework_assets_page():
        return render_template(
            "asset_library.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
            asset_page_mode="framework",
            asset_page_title="框架资产",
            asset_page_subtitle="集中管理已完成的 01-07 框架策划资产，并继续进入框架到剧本。",
        )

    @app.get("/assets/hot-review")
    @_login_required
    def hot_review_assets_page():
        return render_template(
            "asset_library.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
            asset_page_mode="hot_review",
            asset_page_title="爆款文审核资产",
            asset_page_subtitle="查看已保存的爆款文审核结果、评分心电图和修改建议。",
        )

    @app.get("/framework-planner")
    @_login_required
    def framework_planner_page():
        return render_template(
            "framework_planner.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
            framework_backend_ready=framework_planner_backend_ready(),
            framework_planner_storage_key=FRAMEWORK_PLANNER_STORAGE_KEY,
            workflow_line=PRODUCTION_WORKFLOW_LINE,
            framework_to_script_url="/framework-to-script",
            page_title="剧本框架策划工作台",
        )

    @app.get("/framework-planner-test")
    @_login_required
    def framework_planner_test_page():
        return redirect(url_for("new_workflow_test_page", **request.args.to_dict(flat=True)))

    @app.get("/new-workflow-test")
    @_login_required
    def new_workflow_test_page():
        return render_template(
            "new_workflow_test.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
        )

    @app.get("/distillation-lab")
    @_login_required
    def distillation_lab_page():
        return render_template(
            "distillation_lab.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
        )

    @app.get("/compliance-review")
    @_login_required
    def compliance_review_page():
        return render_template(
            "compliance_review.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
        )

    @app.get("/community/<int:project_id>")
    def community_detail_page(project_id: int):
        asset = task_manager.get_public_asset(project_id)
        if not asset:
            return render_template(
                "community_detail.html",
                current_user=_current_user(),
                current_auth_token=_current_auth_token(),
                asset=None,
            ), 404
        return render_template(
            "community_detail.html",
            current_user=_current_user(),
            current_auth_token=_current_auth_token(),
            asset=asset,
        )

    @app.get("/login")
    def login_page():
        if _current_user():
            return redirect(url_for("workspace_page", auth_token=_current_auth_token()))
        return render_template("login.html")

    @app.post("/login")
    def login_submit():
        username = str(request.form.get("username") or "").strip()
        password = str(request.form.get("password") or "")
        user = auth_store.authenticate(username, password)
        allowed_usernames = app.config["LOGIN_ALLOWED_USERNAMES"]
        allowed_user_ids = app.config["LOGIN_ALLOWED_USER_IDS"]
        if user and (allowed_usernames or allowed_user_ids):
            if user.username not in allowed_usernames and user.id not in allowed_user_ids:
                user = None
        if user and app.config["ADMIN_ONLY_LOGIN"] and not admin_store.is_admin(user.id, user.username):
            user = None
        if not user:
            admin_store.record_event(
                user_id=None,
                username=username,
                category="auth",
                action="login",
                status="failed",
                http_method=request.method,
                path=request.path,
                ip_address=_request_ip_address(),
                user_agent=str(request.user_agent or ""),
            )
            return render_template("login.html", error="用户名或密码错误", username=username), 400
        auth_token = _login_user(user)
        admin_store.record_event(
            user_id=user.id,
            username=user.username,
            category="auth",
            action="login",
            status="success",
            http_method=request.method,
            path=request.path,
            ip_address=_request_ip_address(),
            user_agent=str(request.user_agent or ""),
        )
        destination = "admin_dashboard_page" if admin_store.is_admin(user.id, user.username) and request.args.get("next") == "admin" else "workspace_page"
        return redirect(url_for(destination, auth_token=auth_token))

    @app.get("/register")
    def register_page():
        if not app.config["REGISTRATION_ENABLED"]:
            return render_template(
                "login.html",
                error="当前为个人专用平台，仅授权账号可以登录。",
            ), 403
        if _current_user():
            return redirect(url_for("workspace_page", auth_token=_current_auth_token()))
        return render_template("register.html")

    @app.post("/register")
    def register_submit():
        if not app.config["REGISTRATION_ENABLED"]:
            return render_template(
                "login.html",
                error="当前为个人专用平台，注册功能已关闭。",
            ), 403
        username = str(request.form.get("username") or "").strip()
        password = str(request.form.get("password") or "")
        confirm_password = str(request.form.get("confirm_password") or "")
        if password != confirm_password:
            return render_template(
                "register.html",
                error="两次输入的密码不一致",
                username=username,
            ), 400
        try:
            user = auth_store.register_user(username, password)
        except ValueError as exc:
            return render_template(
                "register.html",
                error=str(exc),
                username=username,
            ), 400
        auth_token = _login_user(user)
        admin_store.record_event(
            user_id=user.id,
            username=user.username,
            category="auth",
            action="register",
            status="success",
            http_method=request.method,
            path=request.path,
            ip_address=_request_ip_address(),
            user_agent=str(request.user_agent or ""),
        )
        return redirect(url_for("workspace_page", auth_token=auth_token))

    @app.get("/logout")
    def logout():
        user = _current_user()
        if user:
            admin_store.record_event(
                user_id=user.id,
                username=user.username,
                category="auth",
                action="logout",
                status="success",
                http_method=request.method,
                path=request.path,
                ip_address=_request_ip_address(),
                user_agent=str(request.user_agent or ""),
            )
        _logout_user()
        return redirect(url_for("login_page"))

    @app.get("/api/me")
    @_login_required
    def current_user_api():
        user = _current_user()
        return _json_ok(user={"id": user.id, "username": user.username})

    @app.get("/admin")
    @_admin_required
    def admin_dashboard_page():
        user = _current_user()
        return render_template(
            "admin_dashboard.html",
            current_user=user,
            current_auth_token=_current_auth_token(),
        )

    @app.get("/api/admin/overview")
    @_admin_required
    def admin_overview_api():
        return _json_ok(overview=admin_store.overview())

    @app.get("/api/admin/users")
    @_admin_required
    def admin_users_api():
        try:
            page = int(request.args.get("page") or 1)
            page_size = int(request.args.get("page_size") or 25)
        except (TypeError, ValueError):
            return _json_error("分页参数无效", status=400)
        return _json_ok(
            **admin_store.list_users(
                search=str(request.args.get("search") or "").strip(),
                page=page,
                page_size=page_size,
            )
        )

    @app.get("/api/admin/audit-events")
    @_admin_required
    def admin_audit_events_api():
        try:
            page = int(request.args.get("page") or 1)
            page_size = int(request.args.get("page_size") or 30)
        except (TypeError, ValueError):
            return _json_error("分页参数无效", status=400)
        return _json_ok(
            **admin_store.list_events(
                search=str(request.args.get("search") or "").strip(),
                category=str(request.args.get("category") or "").strip(),
                status=str(request.args.get("status") or "").strip(),
                page=page,
                page_size=page_size,
            )
        )

    @app.get("/api/admin/workflow-runs")
    @_admin_required
    def admin_workflow_runs_api():
        try:
            page = int(request.args.get("page") or 1)
            page_size = int(request.args.get("page_size") or 30)
        except (TypeError, ValueError):
            return _json_error("分页参数无效", status=400)
        return _json_ok(
            **admin_store.list_workflow_runs(
                search=str(request.args.get("search") or "").strip(),
                status=str(request.args.get("status") or "").strip(),
                page=page,
                page_size=page_size,
            )
        )

    @app.patch("/api/me/username")
    @_login_required
    def update_username_api():
        data = request.get_json(silent=True) or {}
        try:
            user = auth_store.update_username(
                _require_user_id(),
                str(data.get("username") or ""),
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(user={"id": user.id, "username": user.username})

    @app.patch("/api/me/password")
    @_login_required
    def update_password_api():
        data = request.get_json(silent=True) or {}
        new_password = str(data.get("new_password") or "")
        confirm_password = str(data.get("confirm_password") or "")
        if new_password != confirm_password:
            return _json_error("两次输入的新密码不一致", status=400)
        try:
            auth_store.update_password(
                _require_user_id(),
                str(data.get("current_password") or ""),
                new_password,
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(message="密码已修改")

    @app.get("/api/models")
    @_login_required
    def list_models():
        spec_path = str(request.args.get("workflow_spec_path") or app.config["WORKFLOW_SPEC_PATH"])
        try:
            models = task_manager.list_model_options(spec_path)
        except Exception as exc:
            return _json_error(str(exc), status=500, fallback="模型列表加载失败，请稍后重试。")
        return _json_ok(models=models, workflow_spec_path=spec_path)


    # HOT_REVIEW_ASSET_PERSIST_V5
    def _script_audit_asset_title_from_payload(payload: dict, request_payload: dict | None = None) -> str:
        if not isinstance(payload, dict):
            return ""
        request_payload = request_payload if isinstance(request_payload, dict) else {}
        audit = payload.get("audit") if isinstance(payload.get("audit"), dict) else {}
        view = payload.get("view") if isinstance(payload.get("view"), dict) else {}
        meta = audit.get("meta") if isinstance(audit.get("meta"), dict) else {}
        view_meta = view.get("meta") if isinstance(view.get("meta"), dict) else {}
        direct = (
            meta.get("script_title")
            or view_meta.get("script_title")
            or audit.get("script_title")
            or view.get("script_title")
            or payload.get("script_title")
            or ""
        )
        if str(direct or "").strip():
            return str(direct).strip()
        filename = str(
            payload.get("filename")
            or request_payload.get("filename")
            or request_payload.get("file_name")
            or request_payload.get("source_filename")
            or ""
        ).replace("\\", "/").rsplit("/", 1)[-1].strip()
        if filename:
            return Path(filename).stem.strip()[:120]
        source_text = str(
            request_payload.get("text")
            or request_payload.get("review_text")
            or payload.get("text")
            or payload.get("answer_text")
            or ""
        )
        for line in source_text.splitlines():
            cleaned = line.strip().strip("#").strip()
            if cleaned:
                return cleaned[:80]
        raw = str(payload.get("answer_text") or payload.get("text") or payload.get("raw_text") or "")
        match = re.search(r"""["']script_title["']\s*:\s*["']([^"']{1,120})["']""", raw)
        return match.group(1).strip() if match else ""

    def _script_audit_asset_summary_from_payload(payload: dict) -> str:
        if not isinstance(payload, dict):
            return "结构化爆款文审核结果"
        audit = payload.get("audit") if isinstance(payload.get("audit"), dict) else {}
        overall = audit.get("overall") if isinstance(audit.get("overall"), dict) else {}
        return str(
            overall.get("core_judgement")
            or overall.get("final_judgement")
            or overall.get("summary")
            or payload.get("summary")
            or "结构化爆款文审核结果"
        ).strip()

    def _persist_script_audit_ecg_asset_metadata(saved_asset, result, flattened, request_payload, user_id: int):
        if not isinstance(saved_asset, dict):
            return saved_asset
        if not isinstance(result, dict):
            result = {}
        if not isinstance(flattened, dict):
            flattened = {}

        merged = {}
        merged.update(result)
        merged.update(flattened)

        is_audit = str(merged.get("result_type") or merged.get("resultType") or "").strip() == "script_audit_ecg"
        audit = merged.get("audit") if isinstance(merged.get("audit"), dict) else None
        view = merged.get("view") if isinstance(merged.get("view"), dict) else None
        visualization = merged.get("visualization") if isinstance(merged.get("visualization"), dict) else None
        if not is_audit and not audit and not view:
            return saved_asset

        project_id = saved_asset.get("project_id") or saved_asset.get("id")
        try:
            project_id_int = int(project_id)
        except Exception:
            return saved_asset

        snapshot = task_manager.get_project_snapshot(project_id_int, user_id=user_id, public_view=False)
        if not isinstance(snapshot, dict):
            return saved_asset

        script_title = _script_audit_asset_title_from_payload(merged, request_payload)
        title = f"爆款文审核｜{script_title}" if script_title else str(saved_asset.get("title") or "爆款文审核").strip()
        summary = _script_audit_asset_summary_from_payload(merged)

        artifacts = snapshot.get("artifacts") if isinstance(snapshot.get("artifacts"), dict) else {}
        artifacts = copy.deepcopy(artifacts)

        if audit:
            artifacts["audit"] = copy.deepcopy(audit)
            snapshot["audit"] = copy.deepcopy(audit)
        if view:
            artifacts["view"] = copy.deepcopy(view)
            snapshot["view"] = copy.deepcopy(view)
        if visualization:
            artifacts["visualization"] = copy.deepcopy(visualization)
            snapshot["visualization"] = copy.deepcopy(visualization)

        tool_result = {
            "tool_key": "hot_review",
            "asset_type": "hot_review",
            "asset_kind": "tool_result",
            "category": "hot_review",
            "result_type": "script_audit_ecg",
            "resultType": "script_audit_ecg",
            "parsed": True,
            "title": title,
            "script_title": script_title,
            "summary": summary,
            "text": str(merged.get("text") or "").strip(),
            "answer_text": str(merged.get("answer_text") or merged.get("answerText") or merged.get("text") or "").strip(),
            "filename": str(merged.get("filename") or saved_asset.get("tool_filename") or "").strip(),
            "parse_warnings": merged.get("parse_warnings") or merged.get("parseWarnings") or [],
            "warnings": merged.get("warnings") or merged.get("parse_warnings") or merged.get("parseWarnings") or [],
        }
        if audit:
            tool_result["audit"] = copy.deepcopy(audit)
        if view:
            tool_result["view"] = copy.deepcopy(view)
        if visualization:
            tool_result["visualization"] = copy.deepcopy(visualization)

        artifacts["tool_result"] = tool_result
        artifacts["result"] = tool_result
        raw_candidates = [
            merged.get("raw_json"),
            merged.get("raw_model_json"),
            merged.get("raw_text"),
            merged.get("answer_text"),
            merged.get("answerText"),
            merged.get("text"),
        ]
        raw_model_json = next((str(item).strip() for item in raw_candidates if str(item or "").strip()), "")
        if raw_model_json:
            tool_result["raw_json"] = raw_model_json
            tool_result["raw_model_json"] = raw_model_json
            artifacts["raw_json"] = raw_model_json
            artifacts["raw_model_json"] = raw_model_json
        if tool_result["answer_text"]:
            artifacts["final_output_text"] = tool_result["answer_text"]
        elif tool_result["text"]:
            artifacts["final_output_text"] = tool_result["text"]
        snapshot["artifacts"] = artifacts

        input_payload = snapshot.get("input_payload") if isinstance(snapshot.get("input_payload"), dict) else {}
        input_payload = copy.deepcopy(input_payload)
        if isinstance(request_payload, dict):
            input_payload.update(copy.deepcopy(request_payload))
        snapshot["input_payload"] = input_payload

        snapshot["tool_key"] = "hot_review"
        snapshot["tool_label"] = f"爆款文审核｜{script_title}" if script_title else "爆款文审核"
        snapshot["tool_output_type"] = "script_audit_ecg"
        snapshot["asset_kind"] = "tool_result"
        snapshot["asset_type"] = "hot_review"
        snapshot["category"] = "hot_review"
        snapshot["workflow_type"] = "hot_review"
        snapshot["result_type"] = "script_audit_ecg"
        snapshot["resultType"] = "script_audit_ecg"
        snapshot["parsed"] = True
        snapshot["title"] = title
        snapshot["summary"] = summary
        if script_title:
            snapshot["script_title"] = script_title

        snapshot["updated_at"] = _now_iso()

        record = task_manager._projects.get(project_id_int) or task_manager._projects.get(str(project_id_int))
        if record:
            with record.lock:
                record.snapshot = snapshot
            task_manager._persist_snapshot(record)
        else:
            task_manager._persist_project_snapshot_data(project_id_int, snapshot)

        return snapshot

    def _save_hot_review_asset_manual(
        *,
        user_id: int,
        request_payload: dict,
        result_payload: dict,
        asset_id: str = "",
    ) -> tuple[dict, dict, dict]:
        if not isinstance(request_payload, dict):
            request_payload = {}
        if not isinstance(result_payload, dict):
            raise ValueError("缺少可保存的爆款文审核结果。")

        result, flattened = _maybe_enrich_script_audit_ecg_tool_result(
            "hot_review",
            copy.deepcopy(result_payload),
            copy.deepcopy(result_payload),
        )
        raw_model_json = str(
            result_payload.get("raw_json")
            or result_payload.get("raw_model_json")
            or result_payload.get("raw_text")
            or result_payload.get("answer_text")
            or result_payload.get("answerText")
            or result_payload.get("text")
            or ""
        ).strip()
        if not raw_model_json:
            try:
                raw_model_json = json.dumps(
                    {
                        key: copy.deepcopy(result_payload.get(key))
                        for key in ("audit", "view", "visualization", "warnings", "parse_warnings")
                        if result_payload.get(key) is not None
                    },
                    ensure_ascii=False,
                    default=str,
                ).strip()
            except Exception:
                raw_model_json = ""
        for target in (result, flattened):
            if not isinstance(target, dict):
                continue
            target["tool_key"] = "hot_review"
            target["asset_kind"] = "tool_result"
            target["asset_type"] = "hot_review"
            target["category"] = "hot_review"
            target["workflow_type"] = "hot_review"
            target["result_type"] = "script_audit_ecg"
            target["resultType"] = "script_audit_ecg"
            if raw_model_json:
                target["raw_json"] = raw_model_json
                target["raw_model_json"] = raw_model_json
            if not str(target.get("text") or "").strip():
                target["text"] = raw_model_json or str(target.get("answer_text") or target.get("answerText") or "").strip()
            if not str(target.get("answer_text") or "").strip():
                target["answer_text"] = raw_model_json or str(target.get("text") or "").strip()

        existing_asset = None
        project_id = 0
        try:
            project_id = int(str(asset_id or "").strip())
        except Exception:
            project_id = 0
        if project_id > 0:
            existing_asset = task_manager.get_project_snapshot(project_id, user_id=user_id, public_view=False)
            if not existing_asset:
                raise PermissionError("爆款文审核资产不存在或当前账号无权访问。")
            existing_kind = str(existing_asset.get("asset_type") or existing_asset.get("category") or existing_asset.get("workflow_type") or "").strip()
            existing_tool = str(existing_asset.get("tool_key") or "").strip()
            if existing_kind != "hot_review" and existing_tool != "hot_review":
                raise ValueError("只能更新爆款文审核资产，不能覆盖其他资产。")
            saved_asset = {
                "project_id": project_id,
                "id": project_id,
                "title": existing_asset.get("title") or "爆款文审核",
                "tool_filename": existing_asset.get("tool_filename") or "",
            }
        else:
            saved_asset = task_manager.save_auxiliary_asset(
                user_id=user_id,
                tool_key="hot_review",
                request_payload=request_payload,
                result=result,
            )

        saved_asset = _persist_script_audit_ecg_asset_metadata(
            saved_asset,
            result,
            flattened,
            request_payload,
            user_id,
        )
        return saved_asset, result, flattened



    # HOT_REVIEW_FILE_UPLOAD_V1
    HOT_REVIEW_UPLOAD_MAX_BYTES = 20 * 1024 * 1024
    HOT_REVIEW_UPLOAD_MAX_CHARS = 500_000
    HOT_REVIEW_UPLOAD_EXTENSIONS = {
        ".txt",
        ".md",
        ".json",
        ".docx",
        ".pdf",
    }

    def _decode_hot_review_text_file(raw: bytes) -> str:
        for encoding in (
            "utf-8-sig",
            "utf-8",
            "gb18030",
            "gbk",
        ):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue

        raise ValueError(
            "无法识别文本文件编码，请将文件转换为 UTF-8 后重试。"
        )

    def _extract_hot_review_docx_text(raw: bytes) -> str:
        from docx import Document

        document = Document(BytesIO(raw))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = str(paragraph.text or "").strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [
                    str(cell.text or "").strip()
                    for cell in row.cells
                    if str(cell.text or "").strip()
                ]
                if cells:
                    parts.append("\t".join(cells))

        return "\n".join(parts).strip()

    def _extract_hot_review_pdf_text(raw: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw))

        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ValueError("暂不支持加密 PDF。") from exc

        pages: list[str] = []

        for page_number, page in enumerate(reader.pages, start=1):
            text = str(page.extract_text() or "").strip()

            if not text:
                continue

            pages.append(
                f"【第 {page_number} 页】\n{text}"
            )

        return "\n\n".join(pages).strip()

    def _extract_hot_review_uploaded_text(
        raw: bytes,
        extension: str,
    ) -> str:
        if extension in {".txt", ".md", ".json"}:
            return _decode_hot_review_text_file(raw)

        if extension == ".docx":
            return _extract_hot_review_docx_text(raw)

        if extension == ".pdf":
            return _extract_hot_review_pdf_text(raw)

        raise ValueError("暂不支持该文件格式。")

    def _uploaded_text_response(uploaded, *, log_label: str):
        _require_user_id()

        if uploaded is None:
            return _json_error(
                "没有收到上传文件。",
                status=400,
            )

        original_filename = str(
            uploaded.filename or ""
        ).replace("\\", "/").rsplit("/", 1)[-1].strip()

        if not original_filename:
            return _json_error(
                "文件名无效。",
                status=400,
            )

        extension = Path(original_filename).suffix.lower()

        if extension not in HOT_REVIEW_UPLOAD_EXTENSIONS:
            return _json_error(
                "暂不支持该文件格式，请上传 TXT、MD、JSON、DOCX 或 PDF。",
                status=400,
            )

        raw = uploaded.read(
            HOT_REVIEW_UPLOAD_MAX_BYTES + 1
        )

        if not raw:
            return _json_error(
                "上传文件为空。",
                status=400,
            )

        if len(raw) > HOT_REVIEW_UPLOAD_MAX_BYTES:
            return _json_error(
                "文件超过 20 MB，无法上传。",
                status=413,
            )

        try:
            text = _extract_hot_review_uploaded_text(
                raw,
                extension,
            )
        except ValueError as exc:
            return _json_error(
                str(exc),
                status=400,
            )
        except Exception:
            logger.exception(
                "%s 文件解析失败：filename=%s extension=%s",
                log_label,
                original_filename,
                extension,
            )
            return _json_error(
                "文件解析失败，请检查文件是否损坏。",
                status=400,
            )

        text = str(text or "").strip()

        if not text:
            if extension == ".pdf":
                message = (
                    "没有从 PDF 中提取到文字。"
                    "该文件可能是扫描版 PDF，目前暂不执行 OCR。"
                )
            else:
                message = "没有从文件中提取到可用文字。"

            return _json_error(
                message,
                status=400,
            )

        if len(text) > HOT_REVIEW_UPLOAD_MAX_CHARS:
            return _json_error(
                (
                    f"文件提取后共有 {len(text)} 个字符，"
                    f"超过 {HOT_REVIEW_UPLOAD_MAX_CHARS} 字符限制。"
                ),
                status=400,
            )

        return _json_ok(
            filename=original_filename,
            extension=extension,
            mime_type=str(uploaded.mimetype or ""),
            text=text,
            char_count=len(text),
            byte_count=len(raw),
        )

    @app.post("/api/files/extract-text")
    def extract_uploaded_text_file():
        return _uploaded_text_response(
            request.files.get("file"),
            log_label="通用上传文本提取",
        )

    @app.post("/api/workbuddy/doctor/source")
    @_login_required
    def upload_workbuddy_doctor_source():
        user = _current_user()
        uploaded = request.files.get("file")
        if uploaded is None:
            return _json_error("没有收到上传文件。", status=400)
        original_filename = str(uploaded.filename or "").replace("\\", "/").rsplit("/", 1)[-1].strip()
        if Path(original_filename).suffix.lower() != ".docx":
            return _json_error("一键优化只支持 Word DOCX 原文件。", status=400)
        raw = uploaded.read(HOT_REVIEW_UPLOAD_MAX_BYTES + 1)
        if not raw:
            return _json_error("上传文件为空。", status=400)
        if len(raw) > HOT_REVIEW_UPLOAD_MAX_BYTES:
            return _json_error("文件超过 20 MB，无法上传。", status=413)
        try:
            text = _extract_hot_review_docx_text(raw)
            if not text:
                raise ValueError("没有从 Word 中提取到可用文字。")
            if len(text) > HOT_REVIEW_UPLOAD_MAX_CHARS:
                raise ValueError(f"Word 共有 {len(text)} 个字符，超过 {HOT_REVIEW_UPLOAD_MAX_CHARS} 字符限制。")
            source = save_workbuddy_source(user.id, filename=original_filename, raw=raw)
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception:
            logger.exception("AI 剧本医生 Word 上传失败: filename=%s", original_filename)
            return _json_error("Word 文件解析或暂存失败，请检查文件是否损坏。", status=400)
        return _json_ok(
            filename=source["original_filename"],
            extension=".docx",
            mime_type=str(uploaded.mimetype or ""),
            text=text,
            char_count=len(text),
            byte_count=len(raw),
            source_document_id=source["source_document_id"],
            can_optimize=True,
        )

    @app.post("/api/agent/conversations/<conversation_id>/attachments")
    @_login_required
    def upload_agent_conversation_attachment(conversation_id: str):
        user = _current_user()
        if not agent_conversation_store.get(user.id, conversation_id):
            return _json_error("对话不存在或无权访问。", status=404)
        uploaded = request.files.get("file")
        if uploaded is None:
            return _json_error("没有收到上传文件。", status=400)
        original_filename = str(uploaded.filename or "").replace("\\", "/").rsplit("/", 1)[-1].strip()
        extension = Path(original_filename).suffix.lower()
        if extension not in HOT_REVIEW_UPLOAD_EXTENSIONS:
            return _json_error("暂不支持该文件格式，请上传 TXT、MD、JSON、DOCX 或 PDF。", status=400)
        raw = uploaded.read(HOT_REVIEW_UPLOAD_MAX_BYTES + 1)
        if not raw:
            return _json_error("上传文件为空。", status=400)
        if len(raw) > HOT_REVIEW_UPLOAD_MAX_BYTES:
            return _json_error("文件超过 20 MB，无法上传。", status=413)
        try:
            text = str(_extract_hot_review_uploaded_text(raw, extension) or "").strip()
            if not text:
                if extension == ".pdf":
                    raise ValueError("没有从 PDF 中提取到文字；扫描版 PDF 暂不支持 OCR。")
                raise ValueError("没有从文件中提取到可用文字。")
            if len(text) > HOT_REVIEW_UPLOAD_MAX_CHARS:
                raise ValueError(f"文件提取后共有 {len(text)} 个字符，超过 {HOT_REVIEW_UPLOAD_MAX_CHARS} 字符限制。")

            converted_to_docx = extension != ".docx"
            if converted_to_docx:
                from docx import Document

                document = Document()
                for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
                    document.add_paragraph(line)
                output = BytesIO()
                document.save(output)
                source_raw = output.getvalue()
                source_filename = f"{Path(original_filename).stem or '剧本'}_工作副本.docx"
            else:
                source_raw = raw
                source_filename = original_filename
            source = save_workbuddy_source(user.id, filename=source_filename, raw=source_raw)
            attachment = agent_conversation_store.save_attachment(
                user.id,
                conversation_id,
                filename=original_filename,
                extension=extension,
                script_text=text,
                source_document_id=str(source.get("source_document_id") or ""),
                metadata={
                    "char_count": len(text),
                    "byte_count": len(raw),
                    "converted_to_docx": converted_to_docx,
                    "source_filename": str(source.get("original_filename") or source_filename),
                },
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception:
            logger.exception("Agent 剧本附件解析失败: filename=%s", original_filename)
            return _json_error("文件解析失败，请检查文件是否损坏。", status=400)
        return _json_ok(
            attachment={
                "id": attachment.get("id"),
                "filename": attachment.get("filename"),
                "extension": attachment.get("extension"),
                "char_count": len(text),
                "converted_to_docx": converted_to_docx,
                "can_optimize": True,
            }
        )

    @app.post("/api/tools/hot_review/extract-file")
    def extract_hot_review_file():
        return _uploaded_text_response(
            request.files.get("file"),
            log_label="爆款文审核",
        )

    @app.get("/api/tools")
    @_login_required
    def list_tools():
        try:
            tools = list_simple_tools()
            for tool in tools:
                if str(tool.get("tool_id") or "") != "hot_review":
                    continue
                fields = tool.get("fields") or []
                for field in fields:
                    if str(field.get("name") or "") == "review_text":
                        field["name"] = "text"
        except Exception as exc:
            return _json_error(str(exc), status=500, fallback="辅助工具列表加载失败，请稍后重试。")
        return _json_ok(ok=True, tools=tools)

    def _run_tool_request(tool_key: str, data: dict):
        try:
            result = run_simple_tool(tool_key, data)
        except ToolExecutionError as exc:
            return jsonify(
                {
                    "success": False,
                    "ok": False,
                    "tool_id": exc.tool_id or tool_key,
                    "message": str(exc),
                    "debug": exc.debug,
                }
            ), int(exc.status_code or 400)
        except FastGPTTransientError as exc:
            return jsonify(
                {
                    "success": False,
                    "ok": False,
                    "tool_id": tool_key,
                    "message": "工具上游暂时不可用，请稍后重试。",
                    "debug": {
                        "stage_name": exc.stage_name,
                        "status_code": exc.status_code,
                        "url": exc.url,
                    },
                }
            ), 503
        except Exception as exc:
            return _json_error(str(exc), status=500, fallback="工具执行失败，请稍后重试。")
        result, flattened = _maybe_enrich_script_audit_ecg_tool_result(tool_key, result, dict(result))
        # HOT_REVIEW_ASSET_SAVE_V4
        if str(flattened.get("result_type") or result.get("result_type") or "").strip() == "script_audit_ecg":
            import re as _script_audit_re

            def _hot_review_script_title_from_payload(value):
                if not isinstance(value, dict):
                    return ""
                audit = value.get("audit") if isinstance(value.get("audit"), dict) else {}
                view = value.get("view") if isinstance(value.get("view"), dict) else {}
                meta = audit.get("meta") if isinstance(audit.get("meta"), dict) else {}
                view_meta = view.get("meta") if isinstance(view.get("meta"), dict) else {}
                direct = (
                    meta.get("script_title")
                    or view_meta.get("script_title")
                    or audit.get("script_title")
                    or view.get("script_title")
                    or value.get("script_title")
                    or ""
                )
                if str(direct or "").strip():
                    return str(direct).strip()
                filename = str(
                    value.get("filename")
                    or data.get("filename")
                    or data.get("file_name")
                    or data.get("source_filename")
                    or ""
                ).replace("\\", "/").rsplit("/", 1)[-1].strip()
                if filename:
                    return Path(filename).stem.strip()[:120]
                source_text = str(data.get("text") or data.get("review_text") or value.get("text") or "")
                for line in source_text.splitlines():
                    cleaned = line.strip().strip("#").strip()
                    if cleaned:
                        return cleaned[:80]
                raw = str(value.get("answer_text") or value.get("text") or "")
                match = _script_audit_re.search(r"""["']script_title["']\s*:\s*["']([^"']{1,120})["']""", raw)
                return match.group(1).strip() if match else ""

            _audit_payload = flattened.get("audit") if isinstance(flattened.get("audit"), dict) else {}
            _overall = _audit_payload.get("overall") if isinstance(_audit_payload.get("overall"), dict) else {}
            _script_title = _hot_review_script_title_from_payload(flattened) or _hot_review_script_title_from_payload(result)
            _asset_title = f"爆款文审核｜{_script_title}" if _script_title else "爆款文审核"
            _asset_summary = str(
                _overall.get("core_judgement")
                or _overall.get("final_judgement")
                or _overall.get("summary")
                or "结构化爆款文审核结果"
            ).strip()

            for _target in (result, flattened):
                if not isinstance(_target, dict):
                    continue
                _target["tool_key"] = "hot_review"
                _target["asset_kind"] = "tool_result"
                _target["asset_type"] = "hot_review"
                _target["category"] = "hot_review"
                _target["workflow_type"] = "hot_review"
                _target["result_type"] = "script_audit_ecg"
                _target["resultType"] = "script_audit_ecg"
                _target["parsed"] = True
                if _script_title:
                    _target["script_title"] = _script_title
                _target["title"] = _asset_title
                _target["summary"] = _asset_summary

        asset_saved = False
        saved_asset = None
        asset_save_error = ""
        if str(flattened.get("text") or flattened.get("answer_text") or "").strip() or str(flattened.get("result_type") or "") == "script_audit_ecg":
            try:
                saved_asset = task_manager.save_auxiliary_asset(
                    user_id=_require_user_id(),
                    tool_key=tool_key,
                    request_payload=data if isinstance(data, dict) else {},
                    result=result,
                )
                asset_saved = True
                saved_asset = _persist_script_audit_ecg_asset_metadata(
                    saved_asset,
                    result,
                    flattened,
                    data if isinstance(data, dict) else {},
                    _require_user_id(),
                )
            except Exception as exc:
                asset_save_error = _sanitize_error_message(
                    str(exc),
                    status=500,
                    fallback="结果已生成，但写入用户资产失败，请稍后重试。",
                )
        flattened["asset_saved"] = asset_saved
        if saved_asset is not None:
            flattened["saved_asset"] = saved_asset
        if asset_save_error:
            flattened["asset_save_error"] = asset_save_error
        result["asset_saved"] = asset_saved
        if saved_asset is not None:
            result["saved_asset"] = saved_asset
        if asset_save_error:
            result["asset_save_error"] = asset_save_error
        ok = bool(flattened.pop("ok", True))
        return _json_ok(ok=ok, result=result, **flattened)

    @app.post("/api/tools/<tool_key>/run")
    @_login_required
    def run_tool(tool_key: str):
        data = request.get_json(silent=True) or {}
        return _run_tool_request(tool_key, data)

    @app.post("/api/tools/hot_review/save")
    @_login_required
    def save_hot_review_tool_result():
        data = request.get_json(silent=True) or {}
        user_id = _require_user_id()
        result_payload = data.get("result") if isinstance(data.get("result"), dict) else data
        request_payload = data.get("request_payload") if isinstance(data.get("request_payload"), dict) else {}
        saved_asset_payload = data.get("saved_asset") if isinstance(data.get("saved_asset"), dict) else {}
        asset_id = str(
            data.get("asset_id")
            or data.get("project_id")
            or data.get("saved_asset_id")
            or saved_asset_payload.get("project_id")
            or saved_asset_payload.get("id")
            or ""
        ).strip()
        try:
            saved_asset, result, flattened = _save_hot_review_asset_manual(
                user_id=user_id,
                request_payload=request_payload,
                result_payload=result_payload,
                asset_id=asset_id,
            )
        except PermissionError as exc:
            return _json_error(str(exc), status=403)
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception:
            logger.exception("manual hot_review asset save failed user_id=%s asset_id=%s", user_id, asset_id)
            return _json_error("保存爆款文审核资产失败，请稍后重试。", status=500)

        flattened["asset_saved"] = True
        flattened["saved_asset"] = saved_asset
        result["asset_saved"] = True
        result["saved_asset"] = saved_asset
        response_payload = dict(flattened)
        response_payload.update(
            {
                "ok": True,
                "asset_saved": True,
                "saved_asset": saved_asset,
                "result": result,
            }
        )
        return _json_ok(**response_payload)

    @app.post("/api/tools/hot_review/export/docx")
    @_login_required
    def export_hot_review_docx():
        data = request.get_json(silent=True) or {}
        title = str(data.get("title") or "爆款文审核报告").strip()[:120] or "爆款文审核报告"
        text = str(data.get("text") or "").strip()
        if not text:
            return _json_error("缺少可导出的报告正文。", status=400)
        try:
            from docx import Document
        except ModuleNotFoundError as exc:
            if exc.name == "docx":
                return _json_error("当前环境缺少 python-docx，暂时无法导出 DOCX。", status=500)
            raise
        document = Document()
        document.add_heading(title, level=1)
        for block in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
            line = block.strip()
            if not line:
                document.add_paragraph("")
            elif line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("- "):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                document.add_paragraph(line)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        safe_title = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in title)[:48] or "hot_review"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{safe_title}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    @app.post("/api/tools/new-framework")
    @_login_required
    def run_new_framework_tool():
        data = request.get_json(silent=True) or {}
        return _run_tool_request("new_framework", data)

    @app.get("/api/user-knowledge/tags")
    @_login_required
    def list_user_knowledge_tags_api():
        try:
            return _json_ok(tags=user_knowledge_store.list_tags(enabled_only=True, user_id=_require_user_id()))
        except Exception as exc:
            logger.exception("user knowledge tags list failed")
            return _json_error(str(exc), status=500, fallback="智慧库标签加载失败，请稍后重试。")

    @app.post("/api/user-knowledge/tags")
    @_login_required
    def create_user_knowledge_tag_api():
        data = request.get_json(silent=True) or {}
        try:
            tag = user_knowledge_store.create_tag(data if isinstance(data, dict) else {}, user_id=_require_user_id())
            return _json_ok(tag=tag)
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception as exc:
            logger.exception("user knowledge tag create failed")
            return _json_error(str(exc), status=500, fallback="智慧库标签创建失败，请稍后重试。")

    @app.patch("/api/user-knowledge/tags/<tag_id>")
    @_login_required
    def update_user_knowledge_tag_api(tag_id: str):
        data = request.get_json(silent=True) or {}
        try:
            tag = user_knowledge_store.update_tag(tag_id, data if isinstance(data, dict) else {}, user_id=_require_user_id())
            return _json_ok(tag=tag)
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception as exc:
            logger.exception("user knowledge tag update failed")
            return _json_error(str(exc), status=500, fallback="智慧库标签更新失败，请稍后重试。")

    @app.delete("/api/user-knowledge/tags/<tag_id>")
    @_login_required
    def delete_user_knowledge_tag_api(tag_id: str):
        try:
            tag = user_knowledge_store.delete_tag(tag_id, user_id=_require_user_id())
            return _json_ok(tag=tag)
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception as exc:
            logger.exception("user knowledge tag delete failed")
            return _json_error(str(exc), status=500, fallback="智慧库标签删除失败，请稍后重试。")

    @app.get("/api/user-knowledge/preferences")
    @_login_required
    def get_user_knowledge_preferences_api():
        try:
            return _json_ok(preferences=user_knowledge_store.get_preferences(_require_user_id()))
        except Exception as exc:
            logger.exception("user knowledge preferences get failed")
            return _json_error(str(exc), status=500, fallback="用户偏好加载失败，请稍后重试。")

    @app.put("/api/user-knowledge/preferences")
    @_login_required
    def save_user_knowledge_preferences_api():
        data = request.get_json(silent=True) or {}
        try:
            preferences = user_knowledge_store.save_preferences(_require_user_id(), data if isinstance(data, dict) else {})
            return _json_ok(preferences=preferences)
        except Exception as exc:
            logger.exception("user knowledge preferences save failed")
            return _json_error(str(exc), status=500, fallback="用户偏好保存失败，请稍后重试。")

    @app.post("/api/user-knowledge/apply-tags")
    @_login_required
    def apply_user_knowledge_tags_api():
        data = request.get_json(silent=True) or {}
        if not isinstance(data, dict):
            data = {}
        try:
            result = user_knowledge_store.apply_tags(
                data.get("selected_tag_ids") if "selected_tag_ids" in data else data.get("selected_preference_tag_ids"),
                existing_user_preference=data.get("existing_user_preference") or data.get("user_preference_prompt") or "",
                user_id=_require_user_id(),
            )
            return _json_ok(**result)
        except Exception as exc:
            logger.exception("user knowledge apply tags failed")
            return _json_error(str(exc), status=500, fallback="智慧库标签应用失败，请稍后重试。")

    def _framework_planner_error(
        stage: str,
        message: str,
        *,
        status: int = 400,
        detail: dict | None = None,
    ):
        return jsonify(
            {
                "ok": False,
                "stage": stage,
                "error": message,
                "detail": detail or {},
            }
        ), status

    @app.get("/api/framework-planner/diagnostics/fastgpt")
    @_login_required
    def framework_planner_fastgpt_diagnostics_api():
        stage = str(request.args.get("stage") or "05").zfill(2)
        try:
            payload = framework_planner_fastgpt_diagnostics(stage)
        except FrameworkPlannerStageError as exc:
            return _framework_planner_error(
                exc.stage,
                str(exc),
                status=exc.status_code,
                detail=exc.detail,
            )
        return jsonify(payload)

    @app.post("/api/framework-planner/stage/04/score")
    @_login_required
    def run_framework_planner_stage_score():
        data = request.get_json(silent=True) or {}
        project_id = data.get("project_id") if isinstance(data, dict) else None
        try:
            payload = run_framework_planner_score(data)
            payload["history"] = save_framework_stage_history(
                project_id=project_id,
                stage="stage04_score",
                payload=data,
                output=payload.get("data") or {},
                status="success",
            )
        except FrameworkPlannerStageError as exc:
            write_framework_stage_exception_log(
                project_id=project_id,
                stage=exc.stage or "stage04_score",
                payload=data,
                exc_type=type(exc).__name__,
                message=str(exc),
                status_code=exc.status_code,
            )
            save_framework_stage_history(
                project_id=project_id,
                stage=exc.stage or "stage04_score",
                payload=data,
                output={},
                status="failed",
                error={"type": type(exc).__name__, "message": str(exc)},
            )
            return _framework_planner_error(
                exc.stage,
                str(exc),
                status=exc.status_code,
                detail=exc.detail,
            )
        except Exception as exc:
            write_framework_stage_exception_log(
                project_id=project_id,
                stage="stage04_score",
                payload=data,
                exc_type=type(exc).__name__,
                message=str(exc),
                status_code=500,
            )
            save_framework_stage_history(
                project_id=project_id,
                stage="stage04_score",
                payload=data,
                output={},
                status="failed",
                error={"type": type(exc).__name__, "message": str(exc)},
            )
            return _framework_planner_error(
                "04",
                "评分接口执行失败，请稍后重试。",
                status=500,
                detail={"message": str(exc)},
            )
        return jsonify(payload)

    @app.post("/api/framework-planner/stage/<stage>")
    @_login_required
    def run_framework_planner_stage_api(stage: str):
        data = request.get_json(silent=True) or {}
        user_id = _require_user_id()
        if isinstance(data, dict):
            data = _hydrate_framework_planner_stage_payload(data, user_id)
        if isinstance(data, dict):
            _attach_user_knowledge_payload(data, data, stage)
        project_id = data.get("project_id") if isinstance(data, dict) else None
        source_text = ""
        if isinstance(data, dict):
            source_text = str(data.get("source_text") or "")
            if not source_text and isinstance(data.get("basic_config"), dict):
                source_text = str(data["basic_config"].get("source_text") or "")
        logger.info(
            "framework planner api request: stage=%s payload_keys=%s source_text_length=%s",
            str(stage).zfill(2),
            sorted(data.keys()) if isinstance(data, dict) else [],
            len(source_text),
        )
        try:
            payload = run_framework_planner_stage(stage, data)
            payload["history"] = save_framework_stage_history(
                project_id=project_id,
                stage=str(stage).zfill(2),
                payload=data,
                output=payload.get("data") or {},
                status="success",
            )
            try:
                autosaved_asset = _autosave_framework_planner_stage(
                    user_id=user_id,
                    stage=stage,
                    request_payload=data if isinstance(data, dict) else {},
                    stage_payload=payload,
                )
                if autosaved_asset:
                    payload["autosaved"] = True
                    payload["autosaved_asset"] = _strip_raw_fastgpt_fields(autosaved_asset)
                    payload["project_id"] = autosaved_asset.get("project_id")
                    if is_test_workflow_line():
                        _save_new_workflow_test_stage_output(
                            user_id=user_id,
                            project_id=payload.get("project_id"),
                            stage=stage,
                            output=payload.get("data"),
                        )
            except Exception as autosave_exc:
                logger.exception(
                    "framework planner stage autosave failed: stage=%s project_id=%s",
                    str(stage).zfill(2),
                    project_id,
                )
                payload["autosaved"] = False
                payload["autosave_error"] = str(autosave_exc)
        except FrameworkPlannerStageError as exc:
            write_framework_stage_exception_log(
                project_id=project_id,
                stage=exc.stage or str(stage).zfill(2),
                payload=data,
                exc_type=type(exc).__name__,
                message=str(exc),
                status_code=exc.status_code,
            )
            save_framework_stage_history(
                project_id=project_id,
                stage=exc.stage or str(stage).zfill(2),
                payload=data,
                output={},
                status="failed",
                error={"type": type(exc).__name__, "message": str(exc)},
            )
            return _framework_planner_error(
                exc.stage,
                str(exc),
                status=exc.status_code,
                detail=exc.detail,
            )
        except Exception as exc:
            logger.exception(
                "framework planner api unexpected error: stage=%s payload_keys=%s",
                str(stage).zfill(2),
                sorted(data.keys()) if isinstance(data, dict) else [],
            )
            write_framework_stage_exception_log(
                project_id=project_id,
                stage=str(stage).zfill(2),
                payload=data,
                exc_type=type(exc).__name__,
                message=str(exc),
                status_code=500,
            )
            save_framework_stage_history(
                project_id=project_id,
                stage=str(stage).zfill(2),
                payload=data,
                output={},
                status="failed",
                error={"type": type(exc).__name__, "message": str(exc)},
            )
            return _framework_planner_error(
                str(stage).zfill(2),
                "模型暂时不可用，请稍后重试。",
                status=500,
                detail={"message": str(exc)},
            )
        return jsonify(payload)

    @app.get("/api/framework-planner/history")
    @_login_required
    def list_framework_planner_history_api():
        # print("[framework-planner-history] route hit", flush=True)
        project_id = request.args.get("project_id") or "unsaved"
        project_name = _framework_project_history_name(project_id, request.args.get("project_name") or "")
        stage = request.args.get("stage") or ""
        return jsonify(list_framework_stage_history(project_id, stage or None, project_name=project_name))

    @app.get("/api/framework-planner/history/<project_id>/<filename>")
    @_login_required
    def load_framework_planner_history_api(project_id: str, filename: str):
        try:
            project_name = _framework_project_history_name(project_id, request.args.get("project_name") or "")
            return jsonify(load_framework_stage_history(project_id, filename, project_name=project_name))
        except FrameworkPlannerStageError as exc:
            return _framework_planner_error(
                exc.stage,
                str(exc),
                status=exc.status_code,
                detail=exc.detail,
            )

    @app.post("/api/framework-planner/debug/frontend")
    @_login_required
    def framework_planner_frontend_debug_api():
        data = request.get_json(silent=True) or {}
        project_id = data.get("project_id") if isinstance(data, dict) else None
        event = str(data.get("event") or "frontend_event") if isinstance(data, dict) else "frontend_event"
        payload = data.get("payload") if isinstance(data, dict) and isinstance(data.get("payload"), dict) else {}
        detail = data.get("detail") if isinstance(data, dict) and isinstance(data.get("detail"), dict) else {}
        try:
            result = write_framework_frontend_debug_event(
                project_id=project_id,
                event=event,
                payload=payload,
                detail=detail,
            )
            return jsonify(result)
        except Exception as exc:
            logger.exception("framework planner frontend debug write failed")
            return _json_error(str(exc), status=500, fallback="前端调试日志写入失败")

    @app.get("/api/projects/latest")
    @_login_required
    def latest_project():
        snapshot = task_manager.latest_project_snapshot(user_id=_require_user_id())
        return _json_ok(project=snapshot)

    @app.get("/api/projects")
    @_login_required
    def list_projects():
        projects = task_manager.list_user_projects(user_id=_require_user_id())
        return _json_ok(projects=projects)

    @app.get("/api/assets")
    @_login_required
    def list_assets():
        assets = task_manager.list_user_assets(user_id=_require_user_id())
        return _json_ok(assets=assets)

    @app.get("/api/framework-assets")
    @_login_required
    def list_framework_assets_api():
        user_id = _require_user_id()
        projects = [
            task_manager._public_snapshot(snapshot)
            for snapshot in task_manager._all_project_snapshots()
            if task_manager._snapshot_belongs_to_user(snapshot, user_id)
            and str(snapshot.get("asset_kind") or "").strip() == "framework_planner"
        ]
        assets = [
            _framework_asset_payload(project, include_detail=False)
            for project in projects
        ]
        assets.sort(key=lambda item: str(item.get("updated_at") or item.get("created_at") or ""), reverse=True)
        return _json_ok(assets=_strip_raw_fastgpt_fields(assets))

    @app.get("/api/framework-assets/<asset_id>")
    @_login_required
    def get_framework_asset_api(asset_id: str):
        asset = _load_framework_asset_for_user(asset_id, _require_user_id())
        if not asset:
            return _json_error("框架资产不存在，或尚未完成 07 最终策划包。", status=404)
        return _json_ok(asset=_strip_raw_fastgpt_fields(asset))

    @app.post("/api/framework-planner/assets")
    @_login_required
    def create_framework_planner_asset_api():
        data = request.get_json(silent=True) or {}
        try:
            asset = task_manager.create_framework_planner_asset(
                user_id=_require_user_id(),
                title=str(data.get("title") or data.get("project_title") or ""),
                season_count=int(data.get("season_count") or 1),
                episodes_per_season=int(data.get("episodes_per_season") or data.get("total_episodes") or 60),
                target_format=str(data.get("target_format") or data.get("genre") or "短剧"),
                style=str(data.get("style") or ""),
                description=str(data.get("description") or data.get("story_outline") or ""),
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception as exc:
            logger.exception("framework planner asset create failed")
            return _json_error(str(exc), status=500, fallback="新建剧本失败，请稍后重试。")
        return _json_ok(asset=asset)



    @app.post("/api/framework-to-script/stage/08")
    @_login_required
    def run_framework_to_script_stage08_api():
        """单独运行 08 提炼核心场景。只跑 08，不继续 09/10。"""
        import json as _json

        data = request.get_json(silent=True) or {}
        if not isinstance(data, dict):
            return _json_error("请求体必须是 JSON object。", status=400)
        try:
            user_id = _require_user_id()
            data, framework_asset = _inject_framework_asset(data, user_id)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

        framework_plan_package = (
            data.get("framework_plan_package")
            or data.get("frameworkPlanPackage")
            or {}
        )
        if not isinstance(framework_plan_package, dict) or not framework_plan_package:
            return _json_error(
                "缺少 framework_plan_package，请先从 07 最终策划包进入框架转剧本工作台。",
                status=400,
            )

        worldview_plan = (
            data.get("worldview_plan")
            or data.get("worldviewPlan")
            or framework_plan_package.get("worldview_plan")
            or framework_plan_package.get("worldviewPlan")
            or {}
        )
        beat_checkpoint_timeline = (
            data.get("beat_checkpoint_timeline")
            or data.get("beatCheckpointTimeline")
            or framework_plan_package.get("beat_checkpoint_timeline")
            or framework_plan_package.get("beatCheckpointTimeline")
            or []
        )
        character_storylines = (
            data.get("character_storylines")
            or data.get("characterStorylines")
            or framework_plan_package.get("character_storylines")
            or framework_plan_package.get("characterStorylines")
            or []
        )

        variables = {
            "frameworkPlanPackage": framework_plan_package,
            "framework_plan_package": framework_plan_package,
            "worldviewPlan": worldview_plan,
            "worldview_plan": worldview_plan,
            "beatCheckpointTimeline": beat_checkpoint_timeline,
            "beat_checkpoint_timeline": beat_checkpoint_timeline,
            "characterStorylines": character_storylines,
            "character_storylines": character_storylines,
            "sourceFrameworkProjectId": data.get("source_framework_project_id") or data.get("sourceFrameworkProjectId") or "",
        }
        variables = _inject_snapshot_stage_preference(
            variables,
            data,
            "08",
            framework_asset=framework_asset,
            workflow_stage="08",
        )
        def _try_parse_json_text(value):
            if not isinstance(value, str):
                return None
            text = value.strip()
            if not text:
                return None
            if text.startswith("```"):
                text = text.strip("`").strip()
                if text.lower().startswith("json"):
                    text = text[4:].strip()
            try:
                return _json.loads(text)
            except Exception:
                return None

        def _extract_update_vars(payload):
            found = {}
            if not isinstance(payload, dict):
                return found
            response_data = payload.get("responseData")
            if not isinstance(response_data, dict):
                return found
            update_results = response_data.get("updateVarResult")
            if not isinstance(update_results, list):
                return found
            for item in update_results:
                if not isinstance(item, dict):
                    continue
                variable = item.get("variable")
                value = item.get("value")
                key = ""
                if isinstance(variable, list) and variable:
                    key = str(variable[-1] or "")
                elif isinstance(variable, str):
                    key = variable
                if key:
                    found[key] = value
            return found

        def _extract_scene_payload(payload):
            candidates = []

            def visit(obj):
                if isinstance(obj, dict):
                    candidates.append(obj)

                    parsed_answer = _try_parse_json_text(obj.get("answerText"))
                    if parsed_answer is not None:
                        visit(parsed_answer)

                    parsed_text = _try_parse_json_text(obj.get("text"))
                    if parsed_text is not None:
                        visit(parsed_text)

                    update_vars = _extract_update_vars(obj)
                    if update_vars:
                        candidates.append(update_vars)
                        for val in update_vars.values():
                            parsed = _try_parse_json_text(val)
                            if parsed is not None:
                                visit(parsed)

                    for key in ("data", "result", "output", "response", "responseData"):
                        val = obj.get(key)
                        if isinstance(val, (dict, list)):
                            visit(val)
                        else:
                            parsed = _try_parse_json_text(val)
                            if parsed is not None:
                                visit(parsed)

                elif isinstance(obj, list):
                    for item in obj:
                        visit(item)
                else:
                    parsed = _try_parse_json_text(obj)
                    if parsed is not None:
                        visit(parsed)

            visit(payload)

            for item in candidates:
                if not isinstance(item, dict):
                    continue
                scene_dictionary = (
                    item.get("sceneDictionary")
                    or item.get("scene_dictionary")
                    or item.get("scene_dictionary_result")
                )
                rules_digest = (
                    item.get("scriptWorldRulesDigest")
                    or item.get("script_world_rules_digest")
                    or item.get("worldRulesDigest")
                )
                if scene_dictionary and rules_digest:
                    return scene_dictionary, rules_digest

            return None, None

        asset_id = str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip()
        if not _try_begin_framework_stage(user_id, asset_id, "08"):
            return _json_error("08 正在运行中，请稍后刷新页面，已完成输出会自动恢复。", status=409)
        try:
            try:
                from .services.fastgpt_client import fastgpt_client
                from .services.fastgpt_contracts import STAGE_FRAMEWORK_SCENE_DICTIONARY

                raw_output = fastgpt_client.run_stage(
                    STAGE_FRAMEWORK_SCENE_DICTIONARY,
                    variables,
                )
            except Exception as exc:
                if is_test_workflow_line():
                    return _json_error(
                        str(exc),
                        status=502,
                        fallback="08 测试工作流远端请求失败，未获得新结果，请稍后重试。",
                    )
                return _json_error(
                    str(exc),
                    status=500,
                    fallback="08 提炼核心场景调用失败，请检查 FASTGPT_FRAMEWORK_SCENE_DICTIONARY_API_KEY 和工作流变量。",
                )

            scene_dictionary, rules_digest = _extract_scene_payload(raw_output)
            if not scene_dictionary or not rules_digest:
                return _json_error(
                    "08 场景字典阶段输出缺少 sceneDictionary 或 scriptWorldRulesDigest。",
                    status=500,
                    fallback="请检查 08_提炼核心场景.json 是否把 sceneDictionary 和 scriptWorldRulesDigest 写入变量或 answerText JSON。",
                )
            if asset_id:
                _save_framework_to_script_stage(
                    user_id=user_id,
                    asset_id=asset_id,
                    stage_key="stage08",
                    output={
                        "sceneDictionary": scene_dictionary,
                        "scriptWorldRulesDigest": rules_digest,
                    },
                )
        finally:
            _end_framework_stage(user_id, asset_id, "08")

        return _json_ok(
            stage="08",
            framework_asset_id=asset_id,
            sceneDictionary=scene_dictionary,
            scriptWorldRulesDigest=rules_digest,
        )



    @app.post("/api/framework-to-script/stage/09")
    @_login_required
    def run_framework_to_script_stage09_api():
        """单独运行 09 人设服装 alias 映射。只跑 09，不继续 10。"""
        import json as _json

        data = request.get_json(silent=True) or {}
        if not isinstance(data, dict):
            return _json_error("请求体必须是 JSON object。", status=400)
        try:
            user_id = _require_user_id()
            data, framework_asset = _inject_framework_asset(data, user_id)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

        framework_plan_package = (
            data.get("framework_plan_package")
            or data.get("frameworkPlanPackage")
            or {}
        )
        if not isinstance(framework_plan_package, dict) or not framework_plan_package:
            return _json_error(
                "缺少 framework_plan_package，请先从 07 最终策划包进入框架转剧本工作台。",
                status=400,
            )

        character_plan = (
            data.get("character_plan")
            or data.get("characterPlan")
            or framework_plan_package.get("character_plan")
            or framework_plan_package.get("characterPlan")
            or {}
        )
        if not isinstance(character_plan, dict) or not character_plan:
            return _json_error(
                "缺少 character_plan，无法运行 09 人设服装 alias 映射。",
                status=400,
            )

        stage08 = _framework_script_stage_cache(framework_asset, "stage08")
        scene_dictionary = (
            data.get("sceneDictionary")
            or data.get("scene_dictionary")
            or stage08.get("sceneDictionary")
            or {}
        )
        if not isinstance(scene_dictionary, dict) or not scene_dictionary:
            return _json_error(
                "缺少 sceneDictionary，请先运行并确认 08 提炼核心场景。",
                status=400,
            )

        beat_checkpoint_timeline = (
            data.get("beat_checkpoint_timeline")
            or data.get("beatCheckpointTimeline")
            or framework_plan_package.get("beat_checkpoint_timeline")
            or framework_plan_package.get("beatCheckpointTimeline")
            or []
        )

        variables = {
            "frameworkPlanPackage": framework_plan_package,
            "framework_plan_package": framework_plan_package,
            "characterPlan": character_plan,
            "character_plan": character_plan,
            "sceneDictionary": scene_dictionary,
            "scene_dictionary": scene_dictionary,
            "beatCheckpointTimeline": beat_checkpoint_timeline,
            "beat_checkpoint_timeline": beat_checkpoint_timeline,
            "sourceFrameworkProjectId": data.get("source_framework_project_id") or data.get("sourceFrameworkProjectId") or "",
        }
        variables = _inject_snapshot_stage_preference(
            variables,
            data,
            "09",
            framework_asset=framework_asset,
            workflow_stage="09",
        )
        def _try_parse_json_text(value):
            if not isinstance(value, str):
                return None
            text = value.strip()
            if not text:
                return None
            if text.startswith("```"):
                text = text.strip("`").strip()
                if text.lower().startswith("json"):
                    text = text[4:].strip()
            try:
                return _json.loads(text)
            except Exception:
                return None

        def _extract_json_object_candidates(value):
            if not isinstance(value, str):
                return []
            text = value.strip()
            if not text:
                return []
            parsed = _try_parse_json_text(text)
            candidates = []
            seen = set()
            if isinstance(parsed, dict):
                fingerprint = _json.dumps(parsed, ensure_ascii=False, sort_keys=True, default=str)
                seen.add(fingerprint)
                candidates.append(parsed)

            in_string = False
            escaping = False
            depth = 0
            start = None
            for index, char in enumerate(text):
                if depth > 0:
                    if escaping:
                        escaping = False
                        continue
                    if char == "\\":
                        escaping = True
                        continue
                    if char == '"':
                        in_string = not in_string
                        continue
                    if in_string:
                        continue
                    if char == "{":
                        depth += 1
                        continue
                    if char == "}":
                        depth -= 1
                        if depth == 0 and start is not None:
                            parsed = _try_parse_json_text(text[start : index + 1])
                            if isinstance(parsed, dict):
                                fingerprint = _json.dumps(parsed, ensure_ascii=False, sort_keys=True, default=str)
                                if fingerprint not in seen:
                                    seen.add(fingerprint)
                                    candidates.append(parsed)
                            start = None
                        continue
                elif char == "{":
                    depth = 1
                    start = index
                    in_string = False
                    escaping = False
            return candidates

        def _extract_update_vars(payload):
            found = {}
            if isinstance(payload, list):
                for item in payload:
                    found.update(_extract_update_vars(item))
                return found
            if not isinstance(payload, dict):
                return found
            update_results = payload.get("updateVarResult")
            if not isinstance(update_results, list):
                response_data = payload.get("responseData")
                if isinstance(response_data, (dict, list)):
                    found.update(_extract_update_vars(response_data))
                return found
            for item in update_results:
                if not isinstance(item, dict):
                    continue
                variable = item.get("variable")
                value = item.get("value")
                key = ""
                if isinstance(variable, list) and variable:
                    key = str(variable[-1] or "")
                elif isinstance(variable, str):
                    key = variable
                if key:
                    found[key] = value
            return found

        def _extract_appearance_payload(payload):
            candidates = []

            def _list_count(value):
                return len(value) if isinstance(value, list) else 0

            def _mapping_from_candidate(item):
                if not isinstance(item, dict):
                    return None
                mapping = (
                    item.get("appearanceMapping")
                    or item.get("appearance_mapping")
                    or item.get("appearanceMappingResult")
                    or item.get("appearance_mapping_result")
                )
                if isinstance(mapping, str):
                    parsed = _try_parse_json_text(mapping)
                    mapping = parsed if isinstance(parsed, dict) else mapping
                if isinstance(mapping, dict) and isinstance(mapping.get("appearanceMapping"), dict):
                    mapping = mapping.get("appearanceMapping")
                if isinstance(mapping, dict):
                    return mapping
                if isinstance(item.get("characters"), list):
                    return item
                return None

            def _rich_counts(mapping):
                if not isinstance(mapping, dict):
                    return {
                        "characters": 0,
                        "outfit_versions": 0,
                        "outfit_variants": 0,
                        "alias_rules": 0,
                        "scene_trigger_rules": 0,
                        "episode_usage_plan": 0,
                        "episode_level_usage_plan": 0,
                        "global_alias_rules": 0,
                    }

                counts = {
                    "characters": 0,
                    "outfit_versions": 0,
                    "outfit_variants": 0,
                    "alias_rules": 0,
                    "scene_trigger_rules": _list_count(mapping.get("scene_level_usage_plan")),
                    "episode_usage_plan": _list_count(mapping.get("episode_level_usage_plan")),
                    "episode_level_usage_plan": _list_count(mapping.get("episode_level_usage_plan")),
                    "global_alias_rules": _list_count(mapping.get("global_alias_rules")),
                }

                characters = mapping.get("characters")
                if isinstance(characters, list):
                    counts["characters"] = len(characters)
                    for character in characters:
                        if not isinstance(character, dict):
                            continue
                        counts["outfit_versions"] += _list_count(character.get("outfit_versions"))
                        counts["outfit_variants"] += _list_count(character.get("outfit_variants"))
                        counts["alias_rules"] += _list_count(character.get("alias_rules"))
                        counts["scene_trigger_rules"] += _list_count(character.get("scene_trigger_rules"))
                        counts["episode_usage_plan"] += _list_count(character.get("episode_usage_plan"))
                        counts["episode_level_usage_plan"] += _list_count(character.get("episode_level_usage_plan"))

                return counts

            def _score_mapping(mapping, source):
                counts = _rich_counts(mapping)
                return (
                    counts["outfit_versions"] + counts["outfit_variants"],
                    counts["outfit_versions"],
                    counts["outfit_variants"],
                    counts["alias_rules"],
                    counts["scene_trigger_rules"],
                    counts["episode_usage_plan"] + counts["episode_level_usage_plan"],
                    counts["global_alias_rules"],
                    counts["characters"],
                    1 if any(key in str(source).lower() for key in ("textoutput", "answertext", "choices")) else 0,
                )

            def _normalize_mapping(mapping):
                if not isinstance(mapping, dict):
                    return mapping

                normalized = copy.deepcopy(mapping)
                characters = normalized.get("characters")
                if not isinstance(characters, list):
                    return normalized

                normalized_characters = []
                for character in characters:
                    if not isinstance(character, dict):
                        normalized_characters.append(character)
                        continue

                    item = copy.deepcopy(character)
                    outfit_versions = item.get("outfit_versions")
                    outfit_variants = item.get("outfit_variants")

                    # 09 ???????? outfit_versions?
                    # ???????? outfit_variants?????????????????????
                    if isinstance(outfit_versions, list) and outfit_versions and not outfit_variants:
                        converted = []
                        default_name = (
                            item.get("default_name")
                            or item.get("name")
                            or item.get("canonical_name")
                            or item.get("character_id")
                            or ""
                        )
                        for version in outfit_versions:
                            if isinstance(version, dict):
                                version_item = copy.deepcopy(version)
                                version_id = (
                                    version_item.get("variant_id")
                                    or version_item.get("version_id")
                                    or version_item.get("linked_outfit_version")
                                    or ""
                                )
                                version_item.setdefault("variant_id", version_id)
                                version_item.setdefault("linked_outfit_version", version_id)
                                version_item.setdefault("alias_name", version_item.get("alias_name") or default_name)
                                version_item.setdefault(
                                    "outfit_description",
                                    version_item.get("outfit_description")
                                    or version_item.get("clothing")
                                    or version_item.get("description")
                                    or "",
                                )
                                version_item.setdefault(
                                    "usage_rule",
                                    version_item.get("usage_rule")
                                    or version_item.get("trigger_condition")
                                    or "",
                                )
                                if "scene_trigger_rules" not in version_item:
                                    scene_refs = version_item.get("scene_refs")
                                    version_item["scene_trigger_rules"] = scene_refs if isinstance(scene_refs, list) else []
                                converted.append(version_item)
                            elif isinstance(version, str) and version.strip():
                                converted.append({
                                    "variant_id": version.strip(),
                                    "linked_outfit_version": version.strip(),
                                    "alias_name": default_name,
                                    "outfit_description": "",
                                    "usage_rule": "",
                                    "scene_trigger_rules": [],
                                })
                        item["outfit_variants"] = converted

                    if "episode_level_usage_plan" not in item and isinstance(item.get("episode_usage_plan"), list):
                        item["episode_level_usage_plan"] = item.get("episode_usage_plan")

                    normalized_characters.append(item)

                normalized["characters"] = normalized_characters
                return normalized

            def _add_candidate(source, obj):
                mapping = _mapping_from_candidate(obj)
                if isinstance(mapping, dict) and mapping:
                    candidates.append((source, mapping))

            def visit(obj, source="root"):
                if isinstance(obj, dict):
                    _add_candidate(source, obj)

                    for text_key in ("textOutput", "answerText", "text", "content"):
                        text_value = obj.get(text_key)
                        for candidate_index, parsed in enumerate(_extract_json_object_candidates(text_value)):
                            visit(parsed, f"{source}.{text_key}" if candidate_index == 0 else f"{source}.{text_key}[json_object:{candidate_index}]")

                    choices = obj.get("choices")
                    if isinstance(choices, list):
                        for index, choice in enumerate(choices):
                            if not isinstance(choice, dict):
                                continue
                            message = choice.get("message")
                            if isinstance(message, dict):
                                for candidate_index, parsed in enumerate(_extract_json_object_candidates(message.get("content"))):
                                    visit(parsed, f"{source}.choices[{index}].message.content" if candidate_index == 0 else f"{source}.choices[{index}].message.content[json_object:{candidate_index}]")

                    update_vars = _extract_update_vars(obj)
                    if update_vars:
                        _add_candidate(f"{source}.updateVars", update_vars)
                        for key, val in update_vars.items():
                            for candidate_index, parsed in enumerate(_extract_json_object_candidates(val)):
                                visit(parsed, f"{source}.updateVars.{key}" if candidate_index == 0 else f"{source}.updateVars.{key}[json_object:{candidate_index}]")

                    for key in (
                        "data",
                        "result",
                        "output",
                        "outputs",
                        "response",
                        "responseData",
                        "newVariables",
                        "updateVarResult",
                    ):
                        val = obj.get(key)
                        if isinstance(val, (dict, list)):
                            visit(val, f"{source}.{key}")
                        else:
                            for candidate_index, parsed in enumerate(_extract_json_object_candidates(val)):
                                visit(parsed, f"{source}.{key}" if candidate_index == 0 else f"{source}.{key}[json_object:{candidate_index}]")

                elif isinstance(obj, list):
                    for index, item in enumerate(obj):
                        visit(item, f"{source}[{index}]")
                else:
                    for parsed in _extract_json_object_candidates(obj):
                        visit(parsed, source)

            visit(payload)

            if not candidates:
                return None

            scored = []
            for source, mapping in candidates:
                normalized = _normalize_mapping(mapping)
                counts = _rich_counts(normalized)
                score = _score_mapping(normalized, source)
                scored.append((score, source, counts, normalized))

            scored.sort(key=lambda item: item[0], reverse=True)
            best_score, best_source, best_counts, best_mapping = scored[0]

            logger.info(
                "stage09 appearanceMapping selected source=%s counts=%s",
                best_source,
                best_counts,
            )

            return best_mapping



        asset_id = str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip()
        if not _try_begin_framework_stage(user_id, asset_id, "09"):
            return _json_error(
                "09 正在运行中，请稍后刷新页面，已完成输出会自动恢复。",
                status=409,
                detail={
                    "stage": "09",
                    "asset_id": asset_id,
                    "user_id": user_id,
                    "failure_reason": "stage_already_running",
                },
            )
        try:
            try:
                from .services.fastgpt_client import fastgpt_client
                from .services.fastgpt_contracts import STAGE_FRAMEWORK_APPEARANCE_MAPPING

                raw_output = fastgpt_client.run_stage(
                    STAGE_FRAMEWORK_APPEARANCE_MAPPING,
                    variables,
                )
            except Exception as exc:
                if is_test_workflow_line():
                    return _json_error(
                        str(exc),
                        status=502,
                        fallback="09 测试工作流远端请求失败，未获得新结果，请稍后重试。",
                    )
                return _json_error(
                    str(exc),
                    status=500,
                    fallback="09 人设服装 alias 映射调用失败，请检查 FASTGPT_FRAMEWORK_APPEARANCE_MAPPING_API_KEY 和工作流变量。",
                )

            if is_test_workflow_line():
                if asset_id:
                    _save_framework_to_script_stage(
                        user_id=user_id,
                        asset_id=asset_id,
                        stage_key="stage09",
                        output=raw_output,
                    )
                return _json_ok(
                    stage="09",
                    framework_asset_id=asset_id,
                    data=raw_output,
                    output=raw_output,
                )

            appearanceMapping = _extract_appearance_payload(raw_output)
            if not appearanceMapping:
                return _json_error(
                    "09 人设服装 alias 映射输出缺少 appearanceMapping。",
                    status=500,
                    fallback="请检查 09_人设服装alias映射.json 是否把 appearanceMapping 写入变量或 answerText JSON。",
                )

            characters = appearanceMapping.get("characters")
            if not isinstance(characters, list) or not characters:
                return _json_error(
                    "09 人设服装 alias 映射输出缺少 appearanceMapping.characters。",
                    status=500,
                    fallback="请检查 09 工作流输出 schema。",
                )
            if asset_id:
                _save_framework_to_script_stage(
                    user_id=user_id,
                    asset_id=asset_id,
                    stage_key="stage09",
                    output={"appearanceMapping": appearanceMapping},
                )
        finally:
            _end_framework_stage(user_id, asset_id, "09")

        return _json_ok(
            stage="09",
            framework_asset_id=asset_id,
            appearanceMapping=appearanceMapping,
        )


    @app.post("/api/framework-to-script/stage/10")
    @_login_required
    def run_framework_to_script_stage10_api():
        """单独运行 10 优化分集计划。只跑 10，不继续后续因果冲突。"""
        import json as _json

        data = request.get_json(silent=True) or {}
        if not isinstance(data, dict):
            return _json_error("请求体必须是 JSON object。", status=400)
        try:
            user_id = _require_user_id()
            data, framework_asset = _inject_framework_asset(data, user_id)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

        framework_plan_package = (
            data.get("framework_plan_package")
            or data.get("frameworkPlanPackage")
            or {}
        )
        if not isinstance(framework_plan_package, dict) or not framework_plan_package:
            return _json_error("缺少 framework_plan_package，请先导入框架资产。", status=400)

        stage08 = _framework_script_stage_cache(framework_asset, "stage08")
        stage09 = _framework_script_stage_cache(framework_asset, "stage09")
        scene_dictionary = data.get("sceneDictionary") or data.get("scene_dictionary") or stage08.get("sceneDictionary") or {}
        if not isinstance(scene_dictionary, dict) or not scene_dictionary:
            return _json_error("缺少 sceneDictionary，请先完成 08 提炼核心场景。", status=400)

        rules_digest = (
            data.get("scriptWorldRulesDigest")
            or data.get("script_world_rules_digest")
            or stage08.get("scriptWorldRulesDigest")
            or {}
        )
        if not isinstance(rules_digest, dict) or not rules_digest:
            return _json_error("缺少 scriptWorldRulesDigest，请先完成 08 提炼核心场景。", status=400)

        appearance_mapping = (
            data.get("appearanceMapping")
            or data.get("appearance_mapping")
            or stage09.get("appearanceMapping")
            or {}
        )
        if not isinstance(appearance_mapping, dict) or not appearance_mapping:
            return _json_error("缺少 appearanceMapping，请先完成 09 确定角色外观。", status=400)

        beat_checkpoint_timeline = (
            data.get("beat_checkpoint_timeline")
            or data.get("beatCheckpointTimeline")
            or framework_plan_package.get("beat_checkpoint_timeline")
            or framework_plan_package.get("beatCheckpointTimeline")
            or []
        )
        character_storylines = (
            data.get("character_storylines")
            or data.get("characterStorylines")
            or framework_plan_package.get("character_storylines")
            or framework_plan_package.get("characterStorylines")
            or []
        )

        variables = {
            "frameworkPlanPackage": framework_plan_package,
            "framework_plan_package": framework_plan_package,
            "sceneDictionary": scene_dictionary,
            "scene_dictionary": scene_dictionary,
            "scriptWorldRulesDigest": rules_digest,
            "script_world_rules_digest": rules_digest,
            "appearanceMapping": appearance_mapping,
            "appearance_mapping": appearance_mapping,
            "beatCheckpointTimeline": beat_checkpoint_timeline,
            "beat_checkpoint_timeline": beat_checkpoint_timeline,
            "characterStorylines": character_storylines,
            "character_storylines": character_storylines,
            "sourceFrameworkProjectId": data.get("source_framework_project_id") or data.get("sourceFrameworkProjectId") or "",
        }
        variables = _inject_snapshot_stage_preference(
            variables,
            data,
            "10",
            framework_asset=framework_asset,
            workflow_stage="10",
        )
        def _try_parse_json_text(value):
            if not isinstance(value, str):
                return None
            text = value.strip()
            if not text:
                return None
            if text.startswith("```"):
                text = text.strip("`").strip()
                if text.lower().startswith("json"):
                    text = text[4:].strip()
            try:
                return _json.loads(text)
            except Exception:
                return None

        def _extract_update_vars(payload):
            found = {}
            if not isinstance(payload, dict):
                return found
            response_data = payload.get("responseData")
            if not isinstance(response_data, dict):
                return found
            update_results = response_data.get("updateVarResult")
            if not isinstance(update_results, list):
                return found
            for item in update_results:
                if not isinstance(item, dict):
                    continue
                variable = item.get("variable")
                value = item.get("value")
                key = ""
                if isinstance(variable, list) and variable:
                    key = str(variable[-1] or "")
                elif isinstance(variable, str):
                    key = variable
                if key:
                    found[key] = value
            return found

        def _extract_enriched_payload(payload):
            candidates = []

            def visit(obj):
                if isinstance(obj, dict):
                    candidates.append(obj)
                    for text_key in ("answerText", "text", "content"):
                        parsed = _try_parse_json_text(obj.get(text_key))
                        if parsed is not None:
                            visit(parsed)
                    update_vars = _extract_update_vars(obj)
                    if update_vars:
                        candidates.append(update_vars)
                        for value in update_vars.values():
                            parsed = _try_parse_json_text(value)
                            if parsed is not None:
                                visit(parsed)
                    for key in (
                        "data",
                        "result",
                        "output",
                        "payload",
                        "message",
                        "response",
                        "responseData",
                        "enrichedEpisodePlanResult",
                        "enriched_episode_plan_result",
                    ):
                        value = obj.get(key)
                        if isinstance(value, (dict, list)):
                            visit(value)
                        else:
                            parsed = _try_parse_json_text(value)
                            if parsed is not None:
                                visit(parsed)
                elif isinstance(obj, list):
                    candidates.append({"allEnrichedEpisodePlan": obj})
                    for item in obj:
                        visit(item)
                else:
                    parsed = _try_parse_json_text(obj)
                    if parsed is not None:
                        visit(parsed)

            visit(payload)

            for item in candidates:
                if not isinstance(item, dict):
                    continue
                nested = item.get("enrichedEpisodePlanResult") or item.get("enriched_episode_plan_result")
                if isinstance(nested, str):
                    nested = _try_parse_json_text(nested) or {}
                if isinstance(nested, dict):
                    candidates.append(nested)

            for item in candidates:
                if not isinstance(item, dict):
                    continue
                plan = (
                    item.get("allEnrichedEpisodePlan")
                    or item.get("enrichedEpisodePlan")
                    or item.get("all_enriched_episode_plan")
                    or item.get("enriched_episode_plan")
                    or item.get("episode_plans")
                    or item.get("episodePlans")
                    or item.get("episodes")
                )
                text = (
                    item.get("allEnrichedEpisodePlanText")
                    or item.get("enrichedEpisodePlanText")
                    or item.get("all_enriched_episode_plan_text")
                    or item.get("enriched_episode_plan_text")
                    or item.get("episode_plans_text")
                    or item.get("episodePlansText")
                    or ""
                )
                if isinstance(plan, str):
                    parsed_plan = _try_parse_json_text(plan)
                    if isinstance(parsed_plan, list):
                        plan = parsed_plan
                    elif isinstance(parsed_plan, dict):
                        plan = (
                            parsed_plan.get("episode_plans")
                            or parsed_plan.get("episodePlans")
                            or parsed_plan.get("episodes")
                            or parsed_plan
                        )
                if isinstance(plan, dict):
                    nested_plan = (
                        plan.get("episode_plans")
                        or plan.get("episodePlans")
                        or plan.get("episodes")
                    )
                    if isinstance(nested_plan, list):
                        plan = nested_plan
                    elif plan and all(str(key).strip().isdigit() for key in plan):
                        plan = list(plan.values())
                if isinstance(plan, list) and plan:
                    plan_text = str(text or "").strip()
                    if not plan_text:
                        plan_text = _json.dumps(plan, ensure_ascii=False, indent=2)
                    return plan, plan_text
            return None, ""

        asset_id = str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip()
        if not _try_begin_framework_stage(user_id, asset_id, "10"):
            return _json_error("10 正在运行中，请稍后刷新页面，已完成输出会自动恢复。", status=409)
        generation_source = "fastgpt"

        def _run_compact_direct_stage10():
            from .services.framework_stage10_fallback import generate_stage10_plan

            direct_plan, direct_text = generate_stage10_plan(
                framework_plan_package=framework_plan_package,
                scene_dictionary=scene_dictionary,
                appearance_mapping=appearance_mapping,
                preference_prompt=str(variables.get("shmRs8OT") or ""),
            )
            return {
                "allEnrichedEpisodePlan": direct_plan,
                "allEnrichedEpisodePlanText": direct_text,
            }

        try:
            basic_config = framework_plan_package.get("basic_config")
            basic_config = basic_config if isinstance(basic_config, dict) else {}
            episode_count = _positive_int(
                basic_config.get("episodes_per_season")
                or basic_config.get("total_episodes")
                or (framework_asset or {}).get("episodes_per_season")
                or (framework_asset or {}).get("total_episodes"),
                0,
            )
            import os as _os

            direct_threshold = _positive_int(
                _os.getenv("FRAMEWORK_STAGE10_DIRECT_THRESHOLD_EPISODES"),
                20,
            )
            if episode_count >= direct_threshold:
                logger.info(
                    "framework-to-script stage10 using compact direct generation: "
                    "asset_id=%s episodes=%s threshold=%s",
                    asset_id,
                    episode_count,
                    direct_threshold,
                )
                try:
                    raw_output = _run_compact_direct_stage10()
                except Exception as fallback_exc:
                    logger.exception(
                        "framework-to-script stage10 compact direct generation failed: asset_id=%s",
                        asset_id,
                    )
                    return _json_error(
                        str(fallback_exc),
                        status=500,
                        fallback="10 优化分集计划精简生成失败，请查看后端诊断日志。",
                    )
                generation_source = "deepseek_direct_compact"
            else:
                try:
                    from .services.fastgpt_client import fastgpt_client
                    from .services.fastgpt_contracts import (
                        STAGE_FRAMEWORK_ENRICHED_EPISODE_PLAN,
                    )

                    raw_output = fastgpt_client.run_stage(
                        STAGE_FRAMEWORK_ENRICHED_EPISODE_PLAN,
                        variables,
                    )
                except Exception as exc:
                    try:
                        debug_info = (
                            fastgpt_client.get_last_stage_debug_info(
                                STAGE_FRAMEWORK_ENRICHED_EPISODE_PLAN
                            )
                            if "fastgpt_client" in locals()
                            and "STAGE_FRAMEWORK_ENRICHED_EPISODE_PLAN" in locals()
                            else {}
                        )
                    except Exception:
                        debug_info = {}
                    diagnostics = str(
                        (debug_info or {}).get("failure_diagnostics") or ""
                    )
                    is_zero_token_timeout = (
                        "Request timed out" in diagnostics
                        and '"inputTokens": 0' in diagnostics
                        and '"outputTokens": 0' in diagnostics
                    )
                    if not is_zero_token_timeout:
                        return _json_error(
                            str(exc),
                            status=500,
                            fallback="10 优化分集计划调用失败，请检查 FASTGPT_FRAMEWORK_ENRICHED_EPISODE_PLAN_API_KEY 和工作流变量。",
                        )
                    logger.warning(
                        "framework-to-script stage10 FastGPT zero-token timeout; "
                        "using compact direct fallback: asset_id=%s diagnostics=%s",
                        asset_id,
                        diagnostics,
                    )
                    try:
                        raw_output = _run_compact_direct_stage10()
                    except Exception as fallback_exc:
                        logger.exception(
                            "framework-to-script stage10 direct fallback failed: asset_id=%s",
                            asset_id,
                        )
                        return _json_error(
                            str(fallback_exc),
                            status=500,
                            fallback="10 优化分集计划的 FastGPT 与精简兜底均执行失败，请查看后端诊断日志。",
                        )
                    generation_source = "deepseek_direct_compact"

            plan, plan_text = _extract_enriched_payload(raw_output)
            if not plan:
                return _json_error(
                    "10 优化分集计划输出缺少 allEnrichedEpisodePlan。",
                    status=500,
                    fallback="请检查 10 工作流是否把 allEnrichedEpisodePlan 写入变量或 answerText JSON。",
                )
            episode_limit = _positive_int(
                (framework_asset or {}).get("episodes_per_season")
                or (framework_asset or {}).get("total_episodes")
                or data.get("episodes_per_season")
                or data.get("total_episodes"),
                0,
            )
            if episode_limit:
                original_plan = plan
                plan = [
                    item
                    for index, item in enumerate(original_plan, start=1)
                    if isinstance(item, dict)
                    and _positive_int(
                        item.get("episode")
                        or item.get("episodeNumber")
                        or item.get("episode_number")
                        or item.get("episode_index")
                        or item.get("ep")
                        or index,
                        index,
                    ) <= episode_limit
                ]
                if len(plan) != len(original_plan):
                    logger.warning(
                        "framework-to-script stage10 trimmed out-of-range episodes: asset_id=%s limit=%s before=%s after=%s",
                        asset_id,
                        episode_limit,
                        len(original_plan),
                        len(plan),
                    )
                    plan_text = _json.dumps(plan, ensure_ascii=False, indent=2)
            if asset_id:
                _save_framework_to_script_stage(
                    user_id=user_id,
                    asset_id=asset_id,
                    stage_key="stage10",
                    output={
                        "framework_enriched_episode_plan": {
                            "allEnrichedEpisodePlan": plan,
                            "allEnrichedEpisodePlanText": plan_text,
                            "batchEnrichedEpisodePlan": plan,
                            "generationSource": generation_source,
                        },
                        "allEnrichedEpisodePlan": plan,
                        "allEnrichedEpisodePlanText": plan_text,
                        "batchEnrichedEpisodePlan": plan,
                        "enrichedEpisodePlan": plan,
                        "enrichedEpisodePlanText": plan_text,
                        "generationSource": generation_source,
                    },
                )
        finally:
            _end_framework_stage(user_id, asset_id, "10")

        return _json_ok(
            stage="10",
            framework_asset_id=asset_id,
            framework_enriched_episode_plan={
                "allEnrichedEpisodePlan": plan,
                "allEnrichedEpisodePlanText": plan_text,
                "batchEnrichedEpisodePlan": plan,
            },
            allEnrichedEpisodePlan=plan,
            allEnrichedEpisodePlanText=plan_text,
            batchEnrichedEpisodePlan=plan,
            enrichedEpisodePlan=plan,
            enrichedEpisodePlanText=plan_text,
            generationSource=generation_source,
            stageOutputs={
                "framework_enriched_episode_plan": {
                    "allEnrichedEpisodePlan": plan,
                    "allEnrichedEpisodePlanText": plan_text,
                    "batchEnrichedEpisodePlan": plan,
                },
                "allEnrichedEpisodePlan": plan,
                "allEnrichedEpisodePlanText": plan_text,
                "batchEnrichedEpisodePlan": plan,
            },
            stages={"10": {"status": "completed"}},
            completedStages=["10"],
        )

    @app.post("/api/framework-to-script/stage/11")
    @_login_required
    def run_framework_to_script_stage11_api():
        """单独运行 11 当前批次因果冲突：write -> review -> rewrite(必要时) -> memory。"""
        data = request.get_json(silent=True) or {}
        if not isinstance(data, dict):
            return _json_error("请求体必须是 JSON object。", status=400)
        try:
            user_id = _require_user_id()
            data, framework_asset = _inject_framework_asset(data, user_id)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

        framework_plan_package = (
            data.get("framework_plan_package")
            or data.get("frameworkPlanPackage")
            or {}
        )
        asset_id = str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip()

        def _stage11_400(message: str, reason: str, status: int = 400, **extra):
            detail = {
                "stage": "11",
                "failure_reason": reason,
                "asset_id": asset_id,
                "user_id": user_id,
                "requested_batch_start": data.get("batchStartEpisode") or data.get("batch_start_episode") or "",
            }
            try:
                detail["plan_count"] = len(plan) if isinstance(plan, list) else 0
            except Exception:
                detail["plan_count"] = 0
            try:
                detail["existing_batch_keys"] = _sorted_numeric_batch_keys(existing_batches) if isinstance(existing_batches, dict) else []
            except Exception:
                detail["existing_batch_keys"] = []
            try:
                detail["has_stage08"] = bool(stage08)
                detail["has_stage09"] = bool(stage09)
                detail["has_stage10"] = bool(stage10)
            except Exception:
                pass
            detail.update({key: value for key, value in extra.items() if value is not None})
            logger.warning("framework-to-script stage11 rejected: %s", detail)
            return _json_error(message, status=status, detail=detail)

        stage10 = _framework_script_stage_cache(framework_asset, "stage10")
        plan = (
            data.get("allEnrichedEpisodePlan")
            or data.get("enrichedEpisodePlan")
            or stage10.get("allEnrichedEpisodePlan")
            or stage10.get("enrichedEpisodePlan")
            or []
        )
        if not isinstance(plan, list) or not plan:
            return _stage11_400("缺少 allEnrichedEpisodePlan，请先完成 10 优化分集计划。", "missing_all_enriched_episode_plan")
        episode_limit = _positive_int(
            (framework_asset or {}).get("episodes_per_season")
            or (framework_asset or {}).get("total_episodes")
            or data.get("episodes_per_season")
            or data.get("total_episodes"),
            0,
        )
        expected_batch_starts = _framework_expected_batch_starts(plan, episode_limit)

        stage08 = _framework_script_stage_cache(framework_asset, "stage08")
        stage09 = _framework_script_stage_cache(framework_asset, "stage09")
        scene_dictionary = data.get("sceneDictionary") or stage08.get("sceneDictionary") or {}
        rules_digest = data.get("scriptWorldRulesDigest") or stage08.get("scriptWorldRulesDigest") or {}
        appearance_mapping = (
            data.get("appearanceMapping")
            or data.get("appearance_mapping")
            or stage09.get("appearanceMapping")
            or stage09.get("appearance_mapping")
            or {}
        )
        if is_test_workflow_line() and not appearance_mapping and isinstance(stage09, dict):
            appearance_mapping = stage09
        if not isinstance(scene_dictionary, dict) or not scene_dictionary:
            return _stage11_400("缺少 sceneDictionary，请先完成 08 提炼核心场景。", "missing_scene_dictionary")
        if not isinstance(appearance_mapping, dict) or not appearance_mapping:
            return _stage11_400("缺少 appearanceMapping，请先完成 09 确定角色外观。", "missing_appearance_mapping")
        if not isinstance(rules_digest, dict) or not rules_digest:
            return _stage11_400("缺少 scriptWorldRulesDigest，请先完成 08 提炼核心场景。", "missing_script_world_rules_digest")

        existing_stage11 = _framework_script_stage_cache(framework_asset, "stage11")
        existing_batches = existing_stage11.get("batches") if isinstance(existing_stage11.get("batches"), dict) else {}
        reset_stage11 = bool(data.get("reset_stage11") or data.get("resetStage11"))
        if reset_stage11:
            existing_batches = {}
            existing_stage11 = {}
            reset_asset_id = str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip()
            if reset_asset_id:
                _save_framework_to_script_stage(
                    user_id=user_id,
                    asset_id=reset_asset_id,
                    stage_key="stage12",
                    output={"batches": {}},
                    status="pending",
                )
        else:
            valid_existing_batches = {}
            for existing_key, existing_batch in existing_batches.items():
                if not isinstance(existing_batch, dict):
                    continue
                existing_start = _positive_int(
                    existing_batch.get("batchStartEpisode") or existing_key,
                    _positive_int(existing_key, 1),
                )
                if str(existing_start) not in expected_batch_starts:
                    logger.warning(
                        "framework-to-script stage11 dropping stale cached batch outside stage10 plan: "
                        "asset_id=%s batchStartEpisode=%s expected_batch_starts=%s",
                        asset_id,
                        existing_start,
                        expected_batch_starts,
                    )
                    continue
                existing_end = _positive_int(existing_batch.get("batchEndEpisode"), existing_start + 4)
                existing_plan, _ = _unwrap_dict_alias(
                    existing_batch,
                    "batchCausalConflictPlan",
                    "batch_causal_conflict_plan",
                )
                existing_issues = _validate_stage11_causal_conflict_plan(
                    existing_plan,
                    start_episode=existing_start,
                    end_episode=existing_end,
                )
                if existing_issues:
                    logger.warning(
                        "framework-to-script stage11 dropping invalid cached batch before retry: "
                        "asset_id=%s batchStartEpisode=%s batchEndEpisode=%s reason=%s",
                        str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip(),
                        existing_start,
                        existing_end,
                        "; ".join(existing_issues[:8]),
                    )
                    continue
                valid_existing_batches[str(existing_start)] = existing_batch
            existing_batches = valid_existing_batches
        requested_start_value = data.get("batchStartEpisode") or data.get("batch_start_episode")
        requested_batch_key = str(_positive_int(requested_start_value, 0)) if requested_start_value else ""
        if requested_batch_key and requested_batch_key in existing_batches and not reset_stage11:
            existing_output = existing_batches.get(requested_batch_key) if isinstance(existing_batches.get(requested_batch_key), dict) else {}
            logger.info(
                "framework-to-script stage11 explicit batch already completed: asset_id=%s batchStartEpisode=%s existing_batch_keys=%s",
                asset_id,
                requested_batch_key,
                _sorted_numeric_batch_keys(existing_batches),
            )
            return _json_ok(
                stage="11",
                framework_asset_id=asset_id,
                skipped=True,
                skip_reason="batch_already_completed",
                batches=existing_batches,
                **existing_output,
            )
        start_episode, end_episode, batch_plan = _framework_batch_from_plan(
            plan,
            requested_start_value,
            completed_starts=existing_batches.keys(),
        )
        if not batch_plan:
            return _stage11_400(
                "当前批次缺少 batchEnrichedEpisodePlan，请检查 10 输出集数。",
                "empty_batch_enriched_episode_plan",
                expected_batch_starts=expected_batch_starts,
            )

        def _autosave_stage11_draft(
            sub_stage: str,
            *,
            conflict_plan_value=None,
            conflict_review_value=None,
            conflict_memory_value=None,
            status: str = "running",
        ) -> None:
            if not asset_id:
                return
            draft = {
                "batchStartEpisode": start_episode,
                "batchEndEpisode": end_episode,
                "batchEnrichedEpisodePlan": batch_plan,
                "subStage": sub_stage,
                "sub_stage": sub_stage,
                "status": status,
                "updated_at": _now_iso(),
            }
            if isinstance(conflict_plan_value, dict) and conflict_plan_value:
                draft["batchCausalConflictPlan"] = conflict_plan_value
                draft["batch_causal_conflict_plan"] = conflict_plan_value
            if isinstance(conflict_review_value, dict) and conflict_review_value:
                draft["batchCausalConflictReview"] = conflict_review_value
                draft["batch_causal_conflict_review"] = conflict_review_value
            if conflict_memory_value is not None:
                draft["conflictMemory"] = str(conflict_memory_value)
                draft["conflict_memory"] = str(conflict_memory_value)
            _save_framework_to_script_stage_draft(
                user_id=user_id,
                asset_id=asset_id,
                stage_key="stage11",
                draft=draft,
                status=status,
            )

        if not _try_begin_framework_stage(user_id, asset_id, "11"):
            return _json_error("11 正在运行中，请稍后刷新页面，已完成输出会自动恢复。", status=409)
        try:
            failed_sub_stage = "stage11_prepare"
            base_vars = {}
            total_episodes = 0
            try:
                from .services.fastgpt_client import FastGPTStageFormatError, fastgpt_client
                from .services.fastgpt_contracts import (
                    STAGE_FRAMEWORK_CAUSAL_CONFLICT_MEMORY,
                    STAGE_FRAMEWORK_CAUSAL_CONFLICT_REVIEW,
                    STAGE_FRAMEWORK_CAUSAL_CONFLICT_REWRITE,
                    STAGE_FRAMEWORK_CAUSAL_CONFLICT_WRITE,
                )

                total_episodes = _positive_int(
                    data.get("total_episodes") or (framework_asset or {}).get("episodes_per_season"),
                    len(plan),
                )
                conflict_memory = str(
                    _first_present(data, "conflictMemory", "conflict_memory", default=None)
                    or _first_present(existing_stage11, "conflictMemory", "conflict_memory", default="")
                    or ""
                )
                if reset_stage11 or start_episode <= 1:
                    conflict_memory = ""
                base_vars = {
                    **_framework_context_vars(data, framework_plan_package),
                    "totalEpisodes": total_episodes,
                    "conflictStartEpisode": start_episode,
                    "batchEnrichedEpisodePlan": batch_plan,
                    "sceneDictionary": scene_dictionary,
                    "appearanceMapping": appearance_mapping,
                    "scriptWorldRulesDigest": rules_digest,
                    "conflictMemory": conflict_memory,
                }
                _merge_aliases(
                    base_vars,
                    {
                        "totalEpisodes": ("total_episodes",),
                        "conflictMemory": ("conflict_memory",),
                        "batchEnrichedEpisodePlan": ("batch_enriched_episode_plan",),
                        "sceneDictionary": ("scene_dictionary",),
                        "appearanceMapping": ("appearance_mapping",),
                        "scriptWorldRulesDigest": ("script_world_rules_digest",),
                    },
                )
                base_vars = _inject_snapshot_stage_preference(
                    base_vars,
                    data,
                    "11",
                    framework_asset=framework_asset,
                    workflow_stage="11_write",
                )
                first_batch_item = batch_plan[0] if batch_plan and isinstance(batch_plan[0], dict) else {}
                appearance_characters = appearance_mapping.get("characters") if isinstance(appearance_mapping, dict) else []
                appearance_character_count = (
                    len(appearance_characters)
                    if isinstance(appearance_characters, list)
                    else len(appearance_mapping)
                    if isinstance(appearance_mapping, dict)
                    else 0
                )
                logger.info(
                    "framework-to-script stage11 start: asset_id=%s start_episode=%s end_episode=%s "
                    "batch_plan_count=%s total_episodes=%s has_sceneDictionary=%s "
                    "has_scriptWorldRulesDigest=%s has_appearanceMapping=%s appearanceMapping.characters_count=%s "
                    "first_batch_episode=%s first_batch_title=%s first_batch_characters=%s stage_names=%s",
                    asset_id,
                    start_episode,
                    end_episode,
                    len(batch_plan),
                    total_episodes,
                    bool(scene_dictionary),
                    bool(rules_digest),
                    bool(appearance_mapping),
                    appearance_character_count,
                    first_batch_item.get("episode"),
                    first_batch_item.get("title"),
                    first_batch_item.get("characters"),
                    {
                        "write": STAGE_FRAMEWORK_CAUSAL_CONFLICT_WRITE,
                        "review": STAGE_FRAMEWORK_CAUSAL_CONFLICT_REVIEW,
                        "rewrite": STAGE_FRAMEWORK_CAUSAL_CONFLICT_REWRITE,
                        "memory": STAGE_FRAMEWORK_CAUSAL_CONFLICT_MEMORY,
                    },
                )
                failed_sub_stage = "causal_conflict_write"
                write_output_data = {}
                conflict_plan = {}
                conflict_plan_unwrapped = False
                write_failure_reason = ""
                write_output_keys = []
                write_retry_count = 0
                direct_test_write = False
                max_write_retries = 0 if is_test_workflow_line() else 3
                for retry_count in range(max_write_retries + 1):
                    write_retry_count = retry_count
                    if retry_count:
                        logger.warning(
                            "framework-to-script stage11 causal_conflict_write retry: asset_id=%s retry_count=%s "
                            "start_episode=%s end_episode=%s reason=%s",
                            asset_id,
                            retry_count,
                            start_episode,
                            end_episode,
                            write_failure_reason,
                        )
                    try:
                        write_output = fastgpt_client.run_stage(STAGE_FRAMEWORK_CAUSAL_CONFLICT_WRITE, base_vars)
                        write_output_data = write_output if isinstance(write_output, dict) else {}
                        write_output_keys = sorted(write_output_data.keys())
                        conflict_plan, conflict_plan_unwrapped, write_failure_reason = _normalize_dict_output_alias(
                            write_output_data,
                            "batchCausalConflictPlan",
                            "batch_causal_conflict_plan",
                        )
                        direct_episodes = None
                        if is_test_workflow_line():
                            for direct_candidate in (
                                write_output_data.get("output"),
                                write_output_data.get("batchCausalConflictPlan"),
                                write_output_data.get("batch_causal_conflict_plan"),
                                conflict_plan,
                            ):
                                if isinstance(direct_candidate, list) and direct_candidate:
                                    direct_episodes = direct_candidate
                                    break
                                if isinstance(direct_candidate, dict):
                                    nested_episodes = (
                                        direct_candidate.get("output")
                                        or direct_candidate.get("episode_plans")
                                        or direct_candidate.get("episodes")
                                    )
                                    if isinstance(nested_episodes, list) and nested_episodes:
                                        direct_episodes = nested_episodes
                                        break
                        if isinstance(direct_episodes, list) and direct_episodes:
                            direct_test_write = True
                            conflict_plan = {
                                "batch_meta": {
                                    "start_episode": start_episode,
                                    "end_episode": end_episode,
                                    "source": "test_workflow_direct_output",
                                },
                                "episodes": direct_episodes,
                            }
                            conflict_plan_unwrapped = True
                            write_failure_reason = ""
                    except FastGPTStageFormatError as exc:
                        write_output_data = {}
                        write_output_keys = []
                        write_failure_reason = (
                            f"{exc.failure_reason}; missing_fields={list(exc.missing_fields)}; "
                            f"probable_truncated_json={exc.probable_truncated_json}; "
                            f"raw_output_source={exc.raw_output_source}"
                        )
                        if retry_count >= max_write_retries:
                            break
                        continue
                    except Exception as exc:
                        write_failure_reason = f"{type(exc).__name__}: {exc}"
                        if retry_count >= max_write_retries:
                            break
                        continue
                    if isinstance(conflict_plan, dict) and conflict_plan:
                        if direct_test_write:
                            break
                        contract_issues = _validate_stage11_causal_conflict_plan(
                            conflict_plan,
                            start_episode=start_episode,
                            end_episode=end_episode,
                        )
                        if not contract_issues:
                            break
                        write_failure_reason = "contract validation failed: " + "; ".join(contract_issues[:8])
                        conflict_plan = {}
                        if retry_count >= max_write_retries:
                            break
                        continue
                    if not write_failure_reason:
                        write_failure_reason = "normalized batchCausalConflictPlan is empty"
                    if retry_count >= max_write_retries:
                        break
                conflict_episodes = conflict_plan.get("episodes") if isinstance(conflict_plan, dict) else []
                logger.info(
                    "framework-to-script stage11 write done: asset_id=%s write_output_keys=%s "
                    "conflict_plan_type=%s conflict_plan_empty=%s normalized_keys=%s unwrapped=%s "
                    "batchCausalConflictPlan_episodes_count=%s input_keys=%s write_failure_reason=%s",
                    asset_id,
                    write_output_keys,
                    type(conflict_plan).__name__,
                    not bool(conflict_plan),
                    sorted(conflict_plan.keys()) if isinstance(conflict_plan, dict) else [],
                    conflict_plan_unwrapped,
                    len(conflict_episodes) if isinstance(conflict_episodes, list) else 0,
                    sorted(base_vars.keys()),
                    write_failure_reason,
                )
                if not isinstance(conflict_plan, dict) or not conflict_plan:
                    logger.error(
                        "framework-to-script stage11 causal_conflict_write failed after retries: asset_id=%s "
                        "start_episode=%s end_episode=%s retry_count=%s max_write_retries=%s "
                        "reason=%s write_output_keys=%s input_keys=%s",
                        asset_id,
                        start_episode,
                        end_episode,
                        write_retry_count,
                        max_write_retries,
                        write_failure_reason,
                        write_output_keys,
                        sorted(base_vars.keys()),
                    )
                    return jsonify(
                        {
                            "success": False,
                            "message": (
                                "11 write 阶段未返回可用的因果计划。"
                                if is_test_workflow_line()
                                else "11 write 阶段未返回可用的 batchCausalConflictPlan，已自动重试 3 次；请查看后端调试终端日志。"
                            ),
                            "detail": {
                                "failed_sub_stage": "causal_conflict_write",
                                "retry_count": write_retry_count,
                                "max_write_retries": max_write_retries,
                                "start_episode": start_episode,
                                "end_episode": end_episode,
                                "write_failure_reason": write_failure_reason,
                                "write_output_keys": write_output_keys,
                            },
                        }
                    ), 500
                _autosave_stage11_draft(
                    "causal_conflict_write",
                    conflict_plan_value=conflict_plan,
                    conflict_memory_value=conflict_memory,
                )
                max_review_rounds = 0 if direct_test_write else 3
                conflict_review = (
                    {
                        "reviewPassed": True,
                        "passed": True,
                        "rewriteRequired": False,
                        "rewrite_required": False,
                        "blockingIssues": [],
                        "blocking_issues": [],
                        "nonBlockingIssues": [],
                        "non_blocking_issues": [],
                        "rewriteBrief": "",
                        "rewrite_brief": "",
                        "qualityGateStatus": "test_direct_output",
                        "quality_gate_status": "test_direct_output",
                    }
                    if direct_test_write
                    else {}
                )
                rewrite_round = 0
                for review_round in range(1, max_review_rounds + 1):
                    failed_sub_stage = "causal_conflict_review"
                    try:
                        review_output = fastgpt_client.run_stage(
                            STAGE_FRAMEWORK_CAUSAL_CONFLICT_REVIEW,
                            {
                                **base_vars,
                                "batchCausalConflictPlan": conflict_plan,
                            },
                        )
                    except FastGPTStageFormatError as exc:
                        logger.warning(
                            "framework-to-script stage11 review missing fields, using defaults: "
                            "asset_id=%s missing_fields=%s error=%s",
                            asset_id,
                            list(exc.missing_fields),
                            str(exc),
                        )
                        review_output = {}
                    review_output_data = review_output if isinstance(review_output, dict) else {}
                    review_passed = _get_bool_alias(review_output_data, "reviewPassed", "passed", default=None)
                    rewrite_required = _get_bool_alias(review_output_data, "rewriteRequired", "rewrite_required", default=None)
                    blocking_issues = _get_list_alias(review_output_data, "blockingIssues", "blocking_issues")
                    non_blocking_issues = _get_list_alias(review_output_data, "nonBlockingIssues", "non_blocking_issues")
                    rewrite_brief = _first_present(review_output_data, "rewriteBrief", "rewrite_brief", default="")
                    if review_passed is None and rewrite_required is None:
                        logger.warning(
                            "framework-to-script stage11 review output missing reviewPassed/passed and rewriteRequired/rewrite_required; "
                            "treating as rewrite needed: "
                            "asset_id=%s review_output_keys=%s",
                            asset_id,
                            sorted(review_output_data.keys()),
                        )
                        review_passed = False
                        rewrite_required = True
                        blocking_issues = []
                    conflict_review = {
                        "reviewPassed": review_passed,
                        "passed": review_passed,
                        "rewriteRequired": rewrite_required,
                        "rewrite_required": rewrite_required,
                        "blockingIssues": blocking_issues,
                        "blocking_issues": blocking_issues,
                        "nonBlockingIssues": non_blocking_issues,
                        "non_blocking_issues": non_blocking_issues,
                        "rewriteBrief": rewrite_brief,
                        "rewrite_brief": rewrite_brief,
                    }
                    _autosave_stage11_draft(
                        "causal_conflict_review",
                        conflict_plan_value=conflict_plan,
                        conflict_review_value=conflict_review,
                        conflict_memory_value=conflict_memory,
                    )
                    rewrite_triggered = _framework_review_needs_rewrite(conflict_review)
                    logger.info(
                        "framework-to-script stage11 review loop: stage=%s batchStartEpisode=%s review_round=%s "
                        "rewrite_round=%s reviewPassed=%s rewriteRequired=%s blockingIssues_count=%s "
                        "asset_id=%s review_output_keys=%s nonBlockingIssues_count=%s rewriteBrief_length=%s "
                        "rewrite_triggered=%s input_keys=%s",
                        "stage11",
                        start_episode,
                        review_round,
                        rewrite_round,
                        review_passed,
                        rewrite_required,
                        len(blocking_issues),
                        asset_id,
                        sorted(review_output_data.keys()),
                        len(non_blocking_issues),
                        len(str(rewrite_brief or "")),
                        rewrite_triggered,
                        sorted(base_vars.keys()),
                    )
                    if review_passed is True and rewrite_required is False:
                        break
                    if review_round >= max_review_rounds:
                        conflict_review = {
                            **(conflict_review if isinstance(conflict_review, dict) else {}),
                            "acceptedAfterReviewLimit": True,
                            "accepted_after_review_limit": True,
                            "qualityGateStatus": "accepted_with_warnings",
                            "quality_gate_status": "accepted_with_warnings",
                        }
                        _autosave_stage11_draft(
                            "causal_conflict_review_limit_reached",
                            conflict_plan_value=conflict_plan,
                            conflict_review_value=conflict_review,
                            conflict_memory_value=conflict_memory,
                            status="accepted_with_warnings",
                        )
                        logger.warning(
                            "framework-to-script stage11 review limit reached; saving latest revision with warnings: "
                            "asset_id=%s batch=%s-%s rounds=%s blocking=%s",
                            asset_id,
                            start_episode,
                            end_episode,
                            max_review_rounds,
                            len(blocking_issues),
                        )
                        break
                    failed_sub_stage = "causal_conflict_rewrite"
                    rewrite_round += 1
                    logger.info(
                        "framework-to-script stage11 rewrite loop: stage=%s batchStartEpisode=%s review_round=%s "
                        "rewrite_round=%s reviewPassed=%s rewriteRequired=%s blockingIssues_count=%s asset_id=%s",
                        "stage11",
                        start_episode,
                        review_round,
                        rewrite_round,
                        review_passed,
                        rewrite_required,
                        len(blocking_issues),
                        asset_id,
                    )
                    rewrite_vars = _inject_snapshot_stage_preference(
                        {
                            **base_vars,
                            "batchCausalConflictPlan": conflict_plan,
                            "batchCausalConflictReview": conflict_review,
                        },
                        data,
                        "11",
                        framework_asset=framework_asset,
                        workflow_stage="11_rewrite",
                    )
                    rewrite_output = fastgpt_client.run_stage(
                        STAGE_FRAMEWORK_CAUSAL_CONFLICT_REWRITE,
                        rewrite_vars,
                    )
                    logger.info(
                        "framework-to-script stage11 rewrite done: asset_id=%s rewrite_output_keys=%s",
                        asset_id,
                        sorted(rewrite_output.keys()) if isinstance(rewrite_output, dict) else [],
                    )
                    rewrite_output_data = rewrite_output if isinstance(rewrite_output, dict) else {}
                    rewrite_conflict_plan, rewrite_unwrapped = _unwrap_dict_alias(
                        rewrite_output_data,
                        "batchCausalConflictPlan",
                        "batch_causal_conflict_plan",
                    )
                    conflict_plan = rewrite_conflict_plan or conflict_plan
                    _autosave_stage11_draft(
                        "causal_conflict_rewrite",
                        conflict_plan_value=conflict_plan,
                        conflict_review_value=conflict_review,
                        conflict_memory_value=conflict_memory,
                    )
                    rewrite_episodes = conflict_plan.get("episodes") if isinstance(conflict_plan, dict) else []
                    logger.info(
                        "framework-to-script stage11 rewrite normalized: asset_id=%s normalized_keys=%s "
                        "unwrapped=%s batchCausalConflictPlan_episodes_count=%s",
                        asset_id,
                        sorted(conflict_plan.keys()) if isinstance(conflict_plan, dict) else [],
                        rewrite_unwrapped,
                        len(rewrite_episodes) if isinstance(rewrite_episodes, list) else 0,
                    )
                failed_sub_stage = "causal_conflict_memory"
                try:
                    if direct_test_write:
                        memory_output = {}
                        last_episode = (
                            conflict_plan.get("episodes", [])[-1]
                            if isinstance(conflict_plan.get("episodes"), list)
                            and conflict_plan.get("episodes")
                            else {}
                        )
                        conflict_memory = json.dumps(
                            {
                                "last_episode": last_episode,
                                "next_batch_start_episode": end_episode + 1,
                            },
                            ensure_ascii=False,
                        )
                    else:
                        memory_output = fastgpt_client.run_stage(
                            STAGE_FRAMEWORK_CAUSAL_CONFLICT_MEMORY,
                            {
                                **base_vars,
                                "batchCausalConflictPlan": conflict_plan,
                                "conflictMemory": conflict_memory,
                                "conflictStartEpisode": start_episode,
                            },
                        )
                except FastGPTStageFormatError as exc:
                    logger.warning(
                        "framework-to-script stage11 memory missing conflictMemory, keeping previous memory: "
                        "asset_id=%s missing_fields=%s error=%s",
                        asset_id,
                        list(exc.missing_fields),
                        str(exc),
                    )
                    memory_output = {}
                memory_output_data = memory_output if isinstance(memory_output, dict) else {}
                memory_value = _first_present(memory_output_data, "conflictMemory", "conflict_memory", default=None)
                if memory_value is None:
                    logger.warning(
                        "framework-to-script stage11 memory output missing conflictMemory/conflict_memory, keeping previous memory: "
                        "asset_id=%s memory_output_keys=%s",
                        asset_id,
                        sorted(memory_output_data.keys()),
                    )
                else:
                    conflict_memory = str(memory_value)
                _autosave_stage11_draft(
                    "causal_conflict_memory",
                    conflict_plan_value=conflict_plan,
                    conflict_review_value=conflict_review,
                    conflict_memory_value=conflict_memory,
                )
                logger.info(
                    "framework-to-script stage11 memory done: asset_id=%s memory_output_keys=%s normalized_keys=%s conflictMemory_length=%s",
                    asset_id,
                    sorted(memory_output_data.keys()),
                    ["conflictMemory", "conflict_memory"],
                    len(conflict_memory),
                )
            except Exception as exc:
                logger.exception(
                    "framework-to-script stage11 failed: failed_sub_stage=%s asset_id=%s start_episode=%s "
                    "end_episode=%s batch_plan_count=%s base_vars_keys=%s error_type=%s error=%s",
                    failed_sub_stage,
                    asset_id,
                    start_episode,
                    end_episode,
                    len(batch_plan) if isinstance(batch_plan, list) else 0,
                    sorted(base_vars.keys()) if isinstance(base_vars, dict) else [],
                    type(exc).__name__,
                    str(exc),
                )
                return jsonify(
                    {
                        "success": False,
                        "message": "11 因果冲突批次调用失败",
                        "detail": {
                            "failed_sub_stage": failed_sub_stage,
                            "error_type": type(exc).__name__,
                            "error_message": str(exc),
                            "asset_id": asset_id,
                            "start_episode": start_episode,
                            "end_episode": end_episode,
                            "batch_plan_count": len(batch_plan) if isinstance(batch_plan, list) else 0,
                            "base_var_keys": sorted(base_vars.keys()) if isinstance(base_vars, dict) else [],
                        },
                    }
                ), 500

            batch_key = str(start_episode)
            batches = dict(existing_batches)
            batch_output = {
                "batchStartEpisode": start_episode,
                "batchEndEpisode": end_episode,
                "batchEnrichedEpisodePlan": batch_plan,
                "batchCausalConflictPlan": conflict_plan,
                "batch_causal_conflict_plan": conflict_plan,
                "batchCausalConflictReview": conflict_review,
                "batch_causal_conflict_review": conflict_review,
                "conflictMemory": conflict_memory,
                "conflict_memory": conflict_memory,
            }
            batches[batch_key] = batch_output
            output = {**batch_output, "batches": batches}
            if asset_id:
                _save_framework_to_script_stage(
                    user_id=user_id,
                    asset_id=asset_id,
                    stage_key="stage11",
                    output=output,
                )
        finally:
            _end_framework_stage(user_id, asset_id, "11")

        return _json_ok(stage="11", framework_asset_id=asset_id, **output)

    @app.post("/api/framework-to-script/stage/12")
    @_login_required
    def run_framework_to_script_stage12_api():
        """单独运行 12 当前批次正文：write -> review -> rewrite(必要时) -> memory。"""
        data = request.get_json(silent=True) or {}
        if not isinstance(data, dict):
            return _json_error("请求体必须是 JSON object。", status=400)
        try:
            user_id = _require_user_id()
            data, framework_asset = _inject_framework_asset(data, user_id)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

        framework_plan_package = (
            data.get("framework_plan_package")
            or data.get("frameworkPlanPackage")
            or {}
        )
        stage10 = _framework_script_stage_cache(framework_asset, "stage10")
        all_episode_plan = (
            data.get("allEnrichedEpisodePlan")
            or data.get("enrichedEpisodePlan")
            or stage10.get("allEnrichedEpisodePlan")
            or stage10.get("enrichedEpisodePlan")
            or []
        )
        episode_limit = _positive_int(
            (framework_asset or {}).get("episodes_per_season")
            or (framework_asset or {}).get("total_episodes")
            or data.get("episodes_per_season")
            or data.get("total_episodes"),
            0,
        )
        expected_batch_starts = _framework_expected_batch_starts(all_episode_plan, episode_limit)
        if not expected_batch_starts:
            return _json_error("缺少第10阶段有效分集计划，请先重新运行10。", status=400)
        stage08 = data.get("stage08") if isinstance(data.get("stage08"), dict) else _framework_script_stage_cache(framework_asset, "stage08")
        stage09 = data.get("stage09") if isinstance(data.get("stage09"), dict) else _framework_script_stage_cache(framework_asset, "stage09")
        stage11 = data.get("stage11") if isinstance(data.get("stage11"), dict) else _framework_script_stage_cache(framework_asset, "stage11")
        stage11_batches = stage11.get("batches") if isinstance(stage11.get("batches"), dict) else {}
        existing_stage12 = data.get("stage12") if isinstance(data.get("stage12"), dict) else _framework_script_stage_cache(framework_asset, "stage12")
        existing_batches = existing_stage12.get("batches") if isinstance(existing_stage12.get("batches"), dict) else {}
        stage11_batches = {
            str(key): value
            for key, value in stage11_batches.items()
            if str(key) in expected_batch_starts and isinstance(value, dict)
        }
        existing_batches = {
            str(key): value
            for key, value in existing_batches.items()
            if str(key) in expected_batch_starts and isinstance(value, dict)
        }
        reset_stage12 = bool(data.get("reset_stage12") or data.get("resetStage12"))
        if reset_stage12:
            existing_stage12 = {}
            existing_batches = {}
        requested_start = data.get("batchStartEpisode") or data.get("batch_start_episode")
        selected_batch_source = ""
        if requested_start:
            batch_key = str(_positive_int(requested_start, 1))
            selected_batch_source = "explicit_request"
            if batch_key not in expected_batch_starts:
                return _json_error(
                    f"请求的第 {batch_key} 集批次不在第10阶段分集计划中。",
                    status=400,
                    detail={"expected_batch_starts": expected_batch_starts, "requested_batch_start": batch_key},
                )
            top_level_start = str(_positive_int(stage11.get("batchStartEpisode"), 0))
            if (
                not stage11_batches
                and top_level_start == batch_key
                and _first_present(stage11, "batchCausalConflictPlan", "batch_causal_conflict_plan", default=None)
            ):
                stage11_batches = {batch_key: stage11}
                logger.warning(
                    "framework-to-script stage12 stage11.batches empty, falling back to top-level stage11 batch: "
                    "asset_id=%s batchStartEpisode=%s",
                    str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip(),
                    batch_key,
                )
        else:
            stage11_batch_keys = [key for key in expected_batch_starts if key in stage11_batches]
            existing_stage12_keys = set(_sorted_numeric_batch_keys(existing_batches))
            batch_key = next((key for key in stage11_batch_keys if key not in existing_stage12_keys), "")
            selected_batch_source = "first_missing_from_stage11_batches"
            if not batch_key and not stage11_batch_keys and _first_present(stage11, "batchCausalConflictPlan", "batch_causal_conflict_plan", default=None):
                top_level_key = str(_positive_int(stage11.get("batchStartEpisode"), 0))
                if top_level_key in expected_batch_starts and top_level_key not in existing_stage12_keys:
                    batch_key = top_level_key
                    stage11_batches = {batch_key: stage11}
                    selected_batch_source = "fallback_top_level"
                    logger.warning(
                        "framework-to-script stage12 stage11.batches empty, falling back to top-level stage11 batch: "
                        "asset_id=%s batchStartEpisode=%s",
                        str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip(),
                        batch_key,
                    )
        if batch_key not in expected_batch_starts:
            return _json_error(
                "第11阶段没有与第10阶段分集计划匹配的可运行批次。",
                status=400,
                detail={"expected_batch_starts": expected_batch_starts, "selected_batch_start": batch_key},
            )
        stage11_batch = stage11_batches.get(batch_key) if isinstance(stage11_batches, dict) else None
        logger.info(
            "framework-to-script stage12 selected batchStartEpisode=%s source=%s stage11_batch_keys=%s stage12_batch_keys=%s",
            batch_key,
            selected_batch_source,
            _sorted_numeric_batch_keys(stage11_batches),
            _sorted_numeric_batch_keys(existing_batches),
        )
        if requested_start and batch_key in existing_batches and not reset_stage12:
            existing_output = existing_batches.get(batch_key) if isinstance(existing_batches.get(batch_key), dict) else {}
            logger.info(
                "framework-to-script stage12 explicit batch already completed: asset_id=%s batchStartEpisode=%s existing_batch_keys=%s",
                str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip(),
                batch_key,
                _sorted_numeric_batch_keys(existing_batches),
            )
            return _json_ok(
                stage="12",
                framework_asset_id=str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip(),
                skipped=True,
                skip_reason="batch_already_completed",
                batches=existing_batches,
                **existing_output,
            )
        if not isinstance(stage11_batch, dict):
            return _json_error("请先完成 11 当前批次因果冲突。", status=400)

        batch_plan = _first_present(stage11_batch, "batchEnrichedEpisodePlan", "batch_enriched_episode_plan", default=[]) or []
        conflict_plan, conflict_plan_unwrapped = _unwrap_dict_alias(
            stage11_batch,
            "batchCausalConflictPlan",
            "batch_causal_conflict_plan",
        )
        scene_dictionary = stage08.get("sceneDictionary") or {}
        appearance_mapping = (
            stage09.get("appearanceMapping")
            or stage09.get("appearance_mapping")
            or (stage09 if is_test_workflow_line() else {})
        )
        if not isinstance(batch_plan, list) or not batch_plan:
            return _json_error("缺少第11阶段批次分集计划，请先重新运行11。", status=400)
        if not isinstance(conflict_plan, dict) or not conflict_plan:
            return _json_error("缺少第11阶段因果冲突计划，请先重新运行11。", status=400)
        if not scene_dictionary:
            return _json_error("缺少第08阶段场景字典，请先重新运行08。", status=400)
        if not appearance_mapping:
            return _json_error("缺少第09阶段人设服装映射，请先重新运行09。", status=400)

        asset_id = str((framework_asset or {}).get("asset_id") or data.get("framework_asset_id") or "").strip()
        if not _try_begin_framework_stage(user_id, asset_id, "12"):
            return _json_error("12 正在运行中，请稍后刷新页面，已完成输出会自动恢复。", status=409)
        try:
            failed_sub_stage = "stage12_prepare"
            start_episode = None
            end_episode = None
            base_vars = {}
            debug_record = {
                "project_id": asset_id or data.get("project_id") or data.get("source_framework_project_id"),
                "project_title": _stage12_debug_project_title(data, framework_asset),
                "stage": 12,
                "status": "started",
                "selected_batch_source": selected_batch_source,
                "batch_key": batch_key,
                "attempt": 1,
                "rewrite_attempt": 0,
                "review_attempt": 0,
                "events": [],
                "created_at": _now_iso(),
                "updated_at": _now_iso(),
            }
            debug_path = ""
            try:
                from .services.fastgpt_client import fastgpt_client
                from .services.fastgpt_contracts import (
                    STAGE_CONTRACTS,
                    STAGE_FRAMEWORK_SCRIPT_MEMORY,
                    STAGE_FRAMEWORK_SCRIPT_REVIEW,
                    STAGE_FRAMEWORK_SCRIPT_REWRITE,
                    STAGE_FRAMEWORK_SCRIPT_WRITE,
                )

                start_episode = _positive_int(stage11_batch.get("batchStartEpisode"), _positive_int(batch_key, 1))
                end_episode = _positive_int(stage11_batch.get("batchEndEpisode"), start_episode + 4)
                debug_record.update(
                    {
                        "batch_start_episode": start_episode,
                        "batch_end_episode": end_episode,
                        "episode_count": max(0, end_episode - start_episode + 1),
                        "updated_at": _now_iso(),
                    }
                )
                debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                episode_word_count = _positive_int(
                    data.get("episode_word_count") or (framework_asset or {}).get("episode_word_count"),
                    600,
                )
                total_episodes = _positive_int(
                    data.get("total_episodes") or (framework_asset or {}).get("episodes_per_season"),
                    end_episode,
                )
                script_memory = str(
                    _first_present(data, "scriptMemory", "script_memory", default=None)
                    or _first_present(existing_stage12, "scriptMemory", "script_memory", default="")
                    or ""
                )
                if reset_stage12 or start_episode <= 1:
                    script_memory = ""
                base_vars = {
                    **_framework_context_vars(data, framework_plan_package),
                    "totalEpisodes": total_episodes,
                    "scriptStartEpisode": start_episode,
                    "episodeWordCount": episode_word_count,
                    "batchEnrichedEpisodePlan": batch_plan,
                    "batchCausalConflictPlan": conflict_plan,
                    "sceneDictionary": scene_dictionary,
                    "appearanceMapping": appearance_mapping,
                    "scriptWorldRulesDigest": stage08.get("scriptWorldRulesDigest") or {},
                    "scriptMemory": script_memory,
                }
                _merge_aliases(
                    base_vars,
                    {
                        "totalEpisodes": ("total_episodes",),
                        "episodeWordCount": ("episode_word_count",),
                        "scriptMemory": ("script_memory",),
                        "scriptStartEpisode": ("script_start_episode",),
                        "batchEnrichedEpisodePlan": ("batch_enriched_episode_plan",),
                        "batchCausalConflictPlan": ("batch_causal_conflict_plan",),
                        "sceneDictionary": ("scene_dictionary",),
                        "appearanceMapping": ("appearance_mapping",),
                        "scriptWorldRulesDigest": ("script_world_rules_digest",),
                    },
                )
                base_vars = _inject_snapshot_stage_preference(
                    base_vars,
                    data,
                    "12",
                    framework_asset=framework_asset,
                    workflow_stage="12_write",
                )
                preference_snapshot = data.get("preference_snapshot") if isinstance(data.get("preference_snapshot"), dict) else {}
                prompt_preferences = data.get("prompt_preferences") if isinstance(data.get("prompt_preferences"), dict) else {}
                user_preference_payload = {
                    "preference_snapshot": preference_snapshot,
                    "prompt_preferences": prompt_preferences,
                }

                def _autosave_stage12_draft(
                    sub_stage: str,
                    *,
                    batch_script_value: str = "",
                    script_review_value=None,
                    script_memory_value=None,
                    status: str = "running",
                    debug_path_value: str = "",
                ) -> None:
                    if not asset_id:
                        return
                    draft = {
                        "batchStartEpisode": start_episode,
                        "batchEndEpisode": end_episode,
                        "batchEnrichedEpisodePlan": batch_plan,
                        "batchCausalConflictPlan": conflict_plan,
                        "batch_causal_conflict_plan": conflict_plan,
                        "subStage": sub_stage,
                        "sub_stage": sub_stage,
                        "status": status,
                        "updated_at": _now_iso(),
                    }
                    if str(batch_script_value or "").strip():
                        draft["batchScriptText"] = str(batch_script_value)
                        draft["batch_script_text"] = str(batch_script_value)
                    if isinstance(script_review_value, dict) and script_review_value:
                        draft["batchScriptReview"] = script_review_value
                        draft["batch_script_review"] = script_review_value
                    if script_memory_value is not None:
                        draft["scriptMemory"] = str(script_memory_value)
                        draft["script_memory"] = str(script_memory_value)
                    if debug_path_value:
                        draft["stage12DebugPath"] = debug_path_value
                    _save_framework_to_script_stage_draft(
                        user_id=user_id,
                        asset_id=asset_id,
                        stage_key="stage12",
                        draft=draft,
                        status=status,
                    )

                debug_record.update(
                    {
                        "status": "prepared",
                        "request_variable_keys_before_fastgpt": sorted(base_vars.keys()),
                        "variable_length_summary": {
                            key: _stage12_debug_summary(base_vars.get(key))
                            for key in sorted(base_vars.keys())
                            if key
                            in {
                                "totalEpisodes",
                                "scriptStartEpisode",
                                "episodeWordCount",
                                "batchEnrichedEpisodePlan",
                                "batchCausalConflictPlan",
                                "sceneDictionary",
                                "appearanceMapping",
                                "scriptWorldRulesDigest",
                                "scriptMemory",
                                "userPreference",
                                "userPreferences",
                                "preferenceSnapshot",
                                "promptPreferences",
                            }
                        },
                        "scriptMemory_length": len(script_memory),
                        "scriptMemory_preview": _stage12_debug_preview(script_memory),
                        "batchEnrichedEpisodePlan_length": len(json.dumps(batch_plan, ensure_ascii=False, default=str)),
                        "batchEnrichedEpisodePlan_preview": _stage12_debug_preview(batch_plan),
                        "batchCausalConflictPlan_length": len(json.dumps(conflict_plan, ensure_ascii=False, default=str)),
                        "batchCausalConflictPlan_preview": _stage12_debug_preview(conflict_plan),
                        "appearanceMapping_exists": bool(appearance_mapping),
                        "scriptWorldRulesDigest_exists": bool(stage08.get("scriptWorldRulesDigest") or {}),
                        "user_preference_exists": bool(preference_snapshot or prompt_preferences),
                        "user_preference_length": len(json.dumps(user_preference_payload, ensure_ascii=False, default=str)),
                        "user_preference_preview": _stage12_debug_preview(user_preference_payload),
                        "updated_at": _now_iso(),
                    }
                )
                debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                first_plan = batch_plan[0] if isinstance(batch_plan[0], dict) else {}
                first_plan_summary = {
                    "episode": first_plan.get("episode"),
                    "title": first_plan.get("title"),
                    "characters": first_plan.get("characters"),
                }
                appearance_characters = (
                    appearance_mapping.get("characters")
                    if isinstance(appearance_mapping, dict)
                    else []
                )
                conflict_episodes = conflict_plan.get("episodes") if isinstance(conflict_plan, dict) else []
                stage_names = {
                    "script_write": getattr(STAGE_CONTRACTS.get(STAGE_FRAMEWORK_SCRIPT_WRITE), "stage_name", STAGE_FRAMEWORK_SCRIPT_WRITE),
                    "script_review": getattr(STAGE_CONTRACTS.get(STAGE_FRAMEWORK_SCRIPT_REVIEW), "stage_name", STAGE_FRAMEWORK_SCRIPT_REVIEW),
                    "script_rewrite": getattr(STAGE_CONTRACTS.get(STAGE_FRAMEWORK_SCRIPT_REWRITE), "stage_name", STAGE_FRAMEWORK_SCRIPT_REWRITE),
                    "script_memory": getattr(STAGE_CONTRACTS.get(STAGE_FRAMEWORK_SCRIPT_MEMORY), "stage_name", STAGE_FRAMEWORK_SCRIPT_MEMORY),
                }
                logger.info(
                    "framework-to-script stage12 entering FastGPT asset_id=%s start_episode=%s end_episode=%s "
                    "batch_plan_count=%s has_batchCausalConflictPlan=%s has_sceneDictionary=%s "
                    "has_appearanceMapping=%s appearanceMapping.characters_count=%s base_vars_keys=%s "
                    "first_batchEnrichedEpisodePlan=%s batchCausalConflictPlan_episodes_count=%s "
                    "batchCausalConflictPlan_unwrapped=%s normalized_conflict_keys=%s stage_names=%s",
                    asset_id,
                    start_episode,
                    end_episode,
                    len(batch_plan),
                    isinstance(conflict_plan, dict) and bool(conflict_plan),
                    bool(scene_dictionary),
                    bool(appearance_mapping),
                    len(appearance_characters) if isinstance(appearance_characters, list) else 0,
                    sorted(base_vars.keys()),
                    first_plan_summary,
                    len(conflict_episodes) if isinstance(conflict_episodes, list) else 0,
                    conflict_plan_unwrapped,
                    sorted(conflict_plan.keys()) if isinstance(conflict_plan, dict) else [],
                    stage_names,
                )
                failed_sub_stage = "script_write"
                write_event = {
                    "sub_stage": "script_write",
                    "attempt": 1,
                    "review_attempt": 0,
                    "rewrite_attempt": 0,
                    "fastgpt_request_started_at": _now_iso(),
                    "request_variable_keys": sorted(base_vars.keys()),
                }
                debug_record["events"].append(write_event)
                debug_record.update({"status": "requesting_fastgpt", "failed_sub_stage": failed_sub_stage, "updated_at": _now_iso()})
                debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                write_started = time.monotonic()
                if is_test_workflow_line() and start_episode < end_episode:
                    episode_scripts: list[str] = []
                    for episode_no in range(start_episode, end_episode + 1):
                        episode_plan = _stage12_episode_plan_slice(batch_plan, episode_no)
                        episode_conflict_plan = _stage12_conflict_plan_slice(
                            conflict_plan,
                            episode_no,
                        )
                        if not episode_plan or not episode_conflict_plan.get("episodes"):
                            raise RuntimeError(
                                f"12 测试流程无法生成第 {episode_no} 集：缺少该集计划。"
                            )
                        episode_output = fastgpt_client.run_stage(
                            STAGE_FRAMEWORK_SCRIPT_WRITE,
                            {
                                **base_vars,
                                "totalEpisodes": total_episodes,
                                "total_episodes": total_episodes,
                                "scriptStartEpisode": episode_no,
                                "script_start_episode": episode_no,
                                "batchEnrichedEpisodePlan": episode_plan,
                                "batch_enriched_episode_plan": episode_plan,
                                "batchCausalConflictPlan": episode_conflict_plan,
                                "batch_causal_conflict_plan": episode_conflict_plan,
                            },
                        )
                        episode_script = _normalize_stage12_episode_headings(
                            _first_present(
                                episode_output,
                                "batchScriptText",
                                "batch_script_text",
                                default="",
                            )
                        ).strip()
                        returned_episodes = _stage12_script_episode_numbers(
                            episode_script
                        )
                        if episode_no not in returned_episodes:
                            if returned_episodes:
                                raise RuntimeError(
                                    f"12 测试流程第 {episode_no} 集返回不匹配："
                                    f"实际识别集号 {sorted(returned_episodes)}。"
                                )
                            if not episode_script:
                                raise RuntimeError(
                                    f"12 测试流程第 {episode_no} 集返回空正文。"
                                )
                            episode_source = (
                                episode_plan[0]
                                if episode_plan and isinstance(episode_plan[0], dict)
                                else {}
                            )
                            episode_title = str(
                                episode_source.get("episode_title")
                                or episode_source.get("title")
                                or ""
                            ).strip()
                            heading = f"第{episode_no}集"
                            if episode_title:
                                heading += f"：{episode_title}"
                            episode_script = f"{heading}\n\n{episode_script}"
                        episode_scripts.append(episode_script)
                        _autosave_stage12_draft(
                            f"script_write_episode_{episode_no}",
                            batch_script_value="\n\n".join(episode_scripts),
                            script_memory_value=script_memory,
                            debug_path_value=debug_path,
                        )
                    write_output = {
                        "batchScriptText": "\n\n".join(episode_scripts)
                    }
                else:
                    write_output = fastgpt_client.run_stage(
                        STAGE_FRAMEWORK_SCRIPT_WRITE,
                        base_vars,
                    )
                write_event["fastgpt_request_ended_at"] = _now_iso()
                write_event["duration_ms"] = int((time.monotonic() - write_started) * 1000)
                write_event["fastgpt_debug"] = _stage12_fastgpt_debug_summary(fastgpt_client, STAGE_FRAMEWORK_SCRIPT_WRITE)
                write_keys = sorted(write_output.keys()) if isinstance(write_output, dict) else []
                batch_script_value = _first_present(write_output, "batchScriptText", "batch_script_text", default=None)
                batch_script = _normalize_stage12_episode_headings(batch_script_value)
                write_event["raw_response_summary"] = _stage12_debug_summary(write_output, preview_limit=600)
                write_event["parsed_fields"] = write_keys
                write_event["batchScriptText_length"] = len(batch_script)
                write_event["parse_failure_reason"] = "" if batch_script.strip() else "missing_or_empty_batchScriptText"
                debug_record.update(
                    {
                        "status": "script_write_done" if batch_script.strip() else "parse_failed",
                        "parsed_fields": write_keys,
                        "raw_response_summary": _stage12_debug_summary(write_output, preview_limit=600),
                        "parse_failure_reason": write_event["parse_failure_reason"],
                        "updated_at": _now_iso(),
                    }
                )
                debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                logger.info(
                    "framework-to-script stage12 script_write output write_output_keys=%s batchScriptText_type=%s "
                    "batchScriptText_length=%s batchScriptText_empty=%s normalized_keys=%s input_keys=%s",
                    write_keys,
                    type(batch_script_value).__name__,
                    len(batch_script),
                    not bool(batch_script.strip()),
                    ["batchScriptText", "batch_script_text"],
                    sorted(base_vars.keys()),
                )
                if not batch_script.strip():
                    debug_record.update(
                        {
                            "status": "failed",
                            "failure_phase": "解析",
                            "failed_sub_stage": "script_write",
                            "exception_type": "",
                            "exception_message": "12 write/rewrite 阶段未返回 batchScriptText",
                            "updated_at": _now_iso(),
                        }
                    )
                    debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                    return jsonify(
                        {
                            "success": False,
                            "message": f"12 阶段第 {start_episode}-{end_episode} 集生成失败/解析失败：write 未返回正文。",
                            "detail": {
                                "failed_sub_stage": "script_write",
                                "error_message": "12 write/rewrite 阶段未返回 batchScriptText",
                                "write_output_keys": write_keys,
                                "debug_path": debug_path,
                            },
                        }
                    ), 500

                missing_script_episodes = _stage12_missing_script_episodes(
                    batch_script,
                    start_episode,
                    end_episode,
                )
                write_event["generated_episode_numbers"] = sorted(
                    _stage12_script_episode_numbers(batch_script)
                )
                write_event["missing_episode_numbers"] = missing_script_episodes
                if (
                    missing_script_episodes
                    and start_episode < end_episode
                    and not is_test_workflow_line()
                ):
                    # The batch workflow occasionally returns only the first episode despite
                    # receiving a complete five-episode plan. Regenerate the batch as explicit
                    # one-episode calls before asking the reviewer to judge it.
                    logger.warning(
                        "framework-to-script stage12 write omitted episodes; switching to one-episode recovery "
                        "asset_id=%s start_episode=%s end_episode=%s missing=%s",
                        asset_id,
                        start_episode,
                        end_episode,
                        missing_script_episodes,
                    )
                    recovery_event = {
                        "sub_stage": "script_write_recovery",
                        "reason": "initial_batch_missing_episode_headings",
                        "missing_episode_numbers_before": missing_script_episodes,
                        "episodes": [],
                        "fastgpt_request_started_at": _now_iso(),
                    }
                    debug_record["events"].append(recovery_event)
                    debug_record.update(
                        {
                            "status": "requesting_fastgpt",
                            "failed_sub_stage": "script_write_recovery",
                            "missing_episode_numbers": missing_script_episodes,
                            "updated_at": _now_iso(),
                        }
                    )
                    debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                    recovered_scripts: list[str] = []
                    for episode_no in range(start_episode, end_episode + 1):
                        episode_plan = _stage12_episode_plan_slice(batch_plan, episode_no)
                        episode_conflict_plan = _stage12_conflict_plan_slice(conflict_plan, episode_no)
                        if not episode_plan or not episode_conflict_plan.get("episodes"):
                            raise RuntimeError(
                                f"12 阶段无法补齐第 {episode_no} 集：缺少该集的分集计划或因果冲突计划。"
                            )
                        episode_vars = {
                            **base_vars,
                            "totalEpisodes": episode_no,
                            "total_episodes": episode_no,
                            "scriptStartEpisode": episode_no,
                            "script_start_episode": episode_no,
                            "batchEnrichedEpisodePlan": episode_plan,
                            "batch_enriched_episode_plan": episode_plan,
                            "batchCausalConflictPlan": episode_conflict_plan,
                            "batch_causal_conflict_plan": episode_conflict_plan,
                        }
                        episode_started = time.monotonic()
                        episode_output = fastgpt_client.run_stage(
                            STAGE_FRAMEWORK_SCRIPT_WRITE,
                            episode_vars,
                        )
                        episode_script = _normalize_stage12_episode_headings(
                            _first_present(
                                episode_output,
                                "batchScriptText",
                                "batch_script_text",
                                default="",
                            )
                        ).strip()
                        episode_event = {
                            "episode": episode_no,
                            "duration_ms": int((time.monotonic() - episode_started) * 1000),
                            "batchScriptText_length": len(episode_script),
                            "generated_episode_numbers": sorted(
                                _stage12_script_episode_numbers(episode_script)
                            ),
                        }
                        recovery_event["episodes"].append(episode_event)
                        if episode_no not in _stage12_script_episode_numbers(episode_script):
                            raise RuntimeError(
                                f"12 阶段第 {episode_no} 集补齐失败：工作流未返回该集正文标题。"
                            )
                        recovered_scripts.append(episode_script)
                    batch_script = "\n\n".join(recovered_scripts)
                    missing_script_episodes = _stage12_missing_script_episodes(
                        batch_script,
                        start_episode,
                        end_episode,
                    )
                    recovery_event.update(
                        {
                            "fastgpt_request_ended_at": _now_iso(),
                            "batchScriptText_length_after": len(batch_script),
                            "missing_episode_numbers_after": missing_script_episodes,
                        }
                    )
                    debug_record.update(
                        {
                            "status": "script_write_recovery_done",
                            "failed_sub_stage": "",
                            "missing_episode_numbers": missing_script_episodes,
                            "updated_at": _now_iso(),
                        }
                    )
                    debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                _autosave_stage12_draft(
                    "script_write",
                    batch_script_value=batch_script,
                    script_memory_value=script_memory,
                    debug_path_value=debug_path,
                )
                max_review_rounds = 0 if is_test_workflow_line() else 3
                script_review = (
                    {
                        "reviewPassed": True,
                        "passed": True,
                        "rewriteRequired": False,
                        "rewrite_required": False,
                        "blockingIssues": [],
                        "blocking_issues": [],
                        "nonBlockingIssues": [],
                        "non_blocking_issues": [],
                        "rewriteBrief": "",
                        "rewrite_brief": "",
                        "qualityGateStatus": "test_direct_output",
                        "quality_gate_status": "test_direct_output",
                    }
                    if is_test_workflow_line()
                    else {}
                )
                rewrite_round = 0
                for review_round in range(1, max_review_rounds + 1):
                    failed_sub_stage = "script_review"
                    debug_record.update(
                        {
                            "status": "requesting_fastgpt",
                            "failed_sub_stage": failed_sub_stage,
                            "review_attempt": review_round,
                            "rewrite_attempt": rewrite_round,
                            "updated_at": _now_iso(),
                        }
                    )
                    review_vars = {
                        **base_vars,
                        "batchScriptText": batch_script,
                    }
                    review_event = {
                        "sub_stage": "script_review",
                        "attempt": 1,
                        "review_attempt": review_round,
                        "rewrite_attempt": rewrite_round,
                        "fastgpt_request_started_at": _now_iso(),
                        "request_variable_keys": sorted(review_vars.keys()),
                        "batchScriptText_length": len(batch_script),
                        "batchScriptText_preview": _stage12_debug_preview(batch_script),
                    }
                    debug_record["events"].append(review_event)
                    debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                    review_started = time.monotonic()
                    review_output = fastgpt_client.run_stage(
                        STAGE_FRAMEWORK_SCRIPT_REVIEW,
                        review_vars,
                    )
                    review_event["fastgpt_request_ended_at"] = _now_iso()
                    review_event["duration_ms"] = int((time.monotonic() - review_started) * 1000)
                    review_event["fastgpt_debug"] = _stage12_fastgpt_debug_summary(fastgpt_client, STAGE_FRAMEWORK_SCRIPT_REVIEW)
                    review_event["http_status"] = review_event["fastgpt_debug"].get("http_status")
                    review_keys = sorted(review_output.keys()) if isinstance(review_output, dict) else []
                    if not isinstance(review_output, dict):
                        review_output = {}
                    review_passed = _get_bool_alias(review_output, "reviewPassed", "passed", default=None)
                    rewrite_required = _get_bool_alias(review_output, "rewriteRequired", "rewrite_required", default=None)
                    blocking_issues = _get_list_alias(review_output, "blockingIssues", "blocking_issues")
                    non_blocking_issues = _get_list_alias(review_output, "nonBlockingIssues", "non_blocking_issues")
                    rewrite_brief = _first_present(review_output, "rewriteBrief", "rewrite_brief", default="")
                    if review_passed is None and rewrite_required is None:
                        logger.warning(
                            "framework-to-script stage12 script_review missing reviewPassed/passed and rewriteRequired/rewrite_required; "
                            "treating as rewrite needed review_output_keys=%s",
                            review_keys,
                        )
                        review_passed = False
                        rewrite_required = True
                    local_opening_issues = _stage12_opening_quality_issues(batch_script)
                    if local_opening_issues:
                        blocking_issues = list(dict.fromkeys([*blocking_issues, *local_opening_issues]))
                        review_passed = False
                        rewrite_required = True
                        local_brief = "；".join(local_opening_issues)
                        rewrite_brief = "；".join(
                            part for part in (str(rewrite_brief or "").strip(), local_brief) if part
                        )
                        logger.warning(
                            "framework-to-script stage12 local opening gate rejected remote approval: issues=%s",
                            local_opening_issues,
                        )
                    script_review = {
                        "reviewPassed": review_passed,
                        "passed": review_passed,
                        "rewriteRequired": rewrite_required,
                        "rewrite_required": rewrite_required,
                        "blockingIssues": blocking_issues,
                        "blocking_issues": blocking_issues,
                        "nonBlockingIssues": non_blocking_issues,
                        "non_blocking_issues": non_blocking_issues,
                        "rewriteBrief": rewrite_brief,
                        "rewrite_brief": rewrite_brief,
                    }
                    rewrite_triggered = _framework_review_needs_rewrite(script_review)
                    review_result_summary = {
                        "reviewPassed": review_passed,
                        "rewriteRequired": rewrite_required,
                        "blockingIssues_count": len(blocking_issues),
                        "blockingIssues_preview": _stage12_debug_preview(blocking_issues),
                        "nonBlockingIssues_count": len(non_blocking_issues),
                        "nonBlockingIssues_preview": _stage12_debug_preview(non_blocking_issues),
                        "rewriteBrief_length": len(str(rewrite_brief or "")),
                        "rewriteBrief_preview": _stage12_debug_preview(rewrite_brief),
                        "rewrite_triggered": rewrite_triggered,
                    }
                    review_event.update(
                        {
                            "raw_response_summary": _stage12_debug_summary(review_output, preview_limit=600),
                            "parsed_fields": review_keys,
                            "review_result": review_result_summary,
                        }
                    )
                    debug_record.update(
                        {
                            "status": "script_review_done",
                            "parsed_fields": review_keys,
                            "review_result": review_result_summary,
                            "updated_at": _now_iso(),
                        }
                    )
                    debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                    _autosave_stage12_draft(
                        "script_review",
                        batch_script_value=batch_script,
                        script_review_value=script_review,
                        script_memory_value=script_memory,
                        debug_path_value=debug_path,
                    )
                    logger.info(
                        "framework-to-script stage12 review loop: stage=%s batchStartEpisode=%s review_round=%s "
                        "rewrite_round=%s reviewPassed=%s rewriteRequired=%s blockingIssues_count=%s "
                        "review_output_keys=%s batchScriptReview_type=%s nonBlockingIssues_count=%s "
                        "rewriteBrief_length=%s rewrite_triggered=%s",
                        "stage12",
                        start_episode,
                        review_round,
                        rewrite_round,
                        review_passed,
                        rewrite_required,
                        len(blocking_issues),
                        review_keys,
                        type(script_review).__name__,
                        len(non_blocking_issues),
                        len(str(rewrite_brief or "")),
                        rewrite_triggered,
                    )
                    if review_passed is True and rewrite_required is False:
                        break
                    if review_round >= max_review_rounds:
                        script_review = {
                            **(script_review if isinstance(script_review, dict) else {}),
                            "acceptedAfterReviewLimit": True,
                            "accepted_after_review_limit": True,
                            "qualityGateStatus": "accepted_with_warnings",
                            "quality_gate_status": "accepted_with_warnings",
                        }
                        debug_record.update(
                            {
                                "status": "review_limit_reached",
                                "failure_phase": "",
                                "failed_sub_stage": "",
                                "review_attempt": review_round,
                                "rewrite_attempt": rewrite_round,
                                "rewrite_reason": _stage12_debug_preview(blocking_issues or rewrite_brief),
                                "updated_at": _now_iso(),
                            }
                        )
                        debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                        _autosave_stage12_draft(
                            "script_review_limit_reached",
                            batch_script_value=batch_script,
                            script_review_value=script_review,
                            script_memory_value=script_memory,
                            status="accepted_with_warnings",
                            debug_path_value=debug_path,
                        )
                        logger.warning(
                            "framework-to-script stage12 review limit reached; saving latest revision with warnings: "
                            "asset_id=%s batch=%s-%s rounds=%s blocking=%s",
                            asset_id,
                            start_episode,
                            end_episode,
                            max_review_rounds,
                            len(blocking_issues),
                        )
                        break
                    failed_sub_stage = "script_rewrite"
                    rewrite_round += 1
                    debug_record.update(
                        {
                            "status": "requesting_fastgpt",
                            "failed_sub_stage": failed_sub_stage,
                            "review_attempt": review_round,
                            "rewrite_attempt": rewrite_round,
                            "rewrite_reason": _stage12_debug_preview(blocking_issues or rewrite_brief),
                            "updated_at": _now_iso(),
                        }
                    )
                    logger.info(
                        "framework-to-script stage12 rewrite loop: stage=%s batchStartEpisode=%s review_round=%s "
                        "rewrite_round=%s reviewPassed=%s rewriteRequired=%s blockingIssues_count=%s",
                        "stage12",
                        start_episode,
                        review_round,
                        rewrite_round,
                        review_passed,
                        rewrite_required,
                        len(blocking_issues),
                    )
                    rewrite_vars = _inject_snapshot_stage_preference(
                        {
                            **base_vars,
                            "batchScriptText": batch_script,
                            "batchScriptReview": script_review,
                        },
                        data,
                        "12",
                        framework_asset=framework_asset,
                        workflow_stage="12_rewrite",
                    )
                    rewrite_event = {
                        "sub_stage": "script_rewrite",
                        "attempt": 1,
                        "review_attempt": review_round,
                        "rewrite_attempt": rewrite_round,
                        "rewrite_reason": _stage12_debug_preview(blocking_issues or rewrite_brief),
                        "fastgpt_request_started_at": _now_iso(),
                        "request_variable_keys": sorted(rewrite_vars.keys()),
                        "batchScriptText_length": len(batch_script),
                        "batchScriptText_preview": _stage12_debug_preview(batch_script),
                        "batchScriptReview_summary": _stage12_debug_summary(script_review),
                    }
                    debug_record["events"].append(rewrite_event)
                    debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                    rewrite_started = time.monotonic()
                    rewrite_output = fastgpt_client.run_stage(
                        STAGE_FRAMEWORK_SCRIPT_REWRITE,
                        rewrite_vars,
                    )
                    rewrite_event["fastgpt_request_ended_at"] = _now_iso()
                    rewrite_event["duration_ms"] = int((time.monotonic() - rewrite_started) * 1000)
                    rewrite_event["fastgpt_debug"] = _stage12_fastgpt_debug_summary(fastgpt_client, STAGE_FRAMEWORK_SCRIPT_REWRITE)
                    rewrite_event["http_status"] = rewrite_event["fastgpt_debug"].get("http_status")
                    rewrite_keys = sorted(rewrite_output.keys()) if isinstance(rewrite_output, dict) else []
                    rewrite_script_value = _first_present(rewrite_output, "batchScriptText", "batch_script_text", default=None)
                    batch_script = _normalize_stage12_episode_headings(rewrite_script_value)
                    rewrite_event["raw_response_summary"] = _stage12_debug_summary(rewrite_output, preview_limit=600)
                    rewrite_event["parsed_fields"] = rewrite_keys
                    rewrite_event["batchScriptText_length_after"] = len(batch_script)
                    rewrite_event["parse_failure_reason"] = "" if batch_script.strip() else "missing_or_empty_batchScriptText"
                    debug_record.update(
                        {
                            "status": "script_rewrite_done" if batch_script.strip() else "parse_failed",
                            "parsed_fields": rewrite_keys,
                            "parse_failure_reason": rewrite_event["parse_failure_reason"],
                            "raw_response_summary": _stage12_debug_summary(rewrite_output, preview_limit=600),
                            "updated_at": _now_iso(),
                        }
                    )
                    debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                    logger.info(
                        "framework-to-script stage12 script_rewrite output rewrite_output_keys=%s batchScriptText_length=%s "
                        "normalized_keys=%s",
                        rewrite_keys,
                        len(batch_script),
                        ["batchScriptText", "batch_script_text"],
                    )
                    if not batch_script.strip():
                        debug_record.update(
                            {
                                "status": "failed",
                                "failure_phase": "解析",
                                "failed_sub_stage": "script_rewrite",
                                "exception_type": "",
                                "exception_message": "12 write/rewrite 阶段未返回 batchScriptText",
                                "updated_at": _now_iso(),
                            }
                        )
                        debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                        return jsonify(
                            {
                                "success": False,
                                "message": f"12 阶段第 {start_episode}-{end_episode} 集生成失败/解析失败：rewrite 未返回正文。",
                                "detail": {
                                    "failed_sub_stage": "script_rewrite",
                                    "error_message": "12 write/rewrite 阶段未返回 batchScriptText",
                                    "rewrite_output_keys": rewrite_keys,
                                    "debug_path": debug_path,
                                },
                            }
                        ), 500
                    _autosave_stage12_draft(
                        "script_rewrite",
                        batch_script_value=batch_script,
                        script_review_value=script_review,
                        script_memory_value=script_memory,
                        debug_path_value=debug_path,
                    )
                failed_sub_stage = "script_memory"
                memory_vars = {
                    **base_vars,
                    "batchScriptText": batch_script,
                    "scriptMemory": script_memory,
                    "scriptStartEpisode": start_episode,
                }
                debug_record.update(
                    {
                        "status": "requesting_fastgpt",
                        "failed_sub_stage": failed_sub_stage,
                        "review_attempt": review_round if "review_round" in locals() else 0,
                        "rewrite_attempt": rewrite_round,
                        "updated_at": _now_iso(),
                    }
                )
                memory_event = {
                    "sub_stage": "script_memory",
                    "attempt": 1,
                    "review_attempt": review_round if "review_round" in locals() else 0,
                    "rewrite_attempt": rewrite_round,
                    "fastgpt_request_started_at": _now_iso(),
                    "request_variable_keys": sorted(memory_vars.keys()),
                    "batchScriptText_length": len(batch_script),
                    "batchScriptText_preview": _stage12_debug_preview(batch_script),
                    "scriptMemory_length_before": len(script_memory),
                    "scriptMemory_preview_before": _stage12_debug_preview(script_memory),
                }
                debug_record["events"].append(memory_event)
                debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                memory_started = time.monotonic()
                if is_test_workflow_line():
                    memory_output = {}
                    script_memory = json.dumps(
                        {
                            "batch_start_episode": start_episode,
                            "batch_end_episode": end_episode,
                            "script_tail": batch_script[-2000:],
                        },
                        ensure_ascii=False,
                    )
                else:
                    memory_output = fastgpt_client.run_stage(
                        STAGE_FRAMEWORK_SCRIPT_MEMORY,
                        memory_vars,
                    )
                memory_event["fastgpt_request_ended_at"] = _now_iso()
                memory_event["duration_ms"] = int((time.monotonic() - memory_started) * 1000)
                memory_event["fastgpt_debug"] = _stage12_fastgpt_debug_summary(fastgpt_client, STAGE_FRAMEWORK_SCRIPT_MEMORY)
                memory_event["http_status"] = memory_event["fastgpt_debug"].get("http_status")
                memory_keys = sorted(memory_output.keys()) if isinstance(memory_output, dict) else []
                memory_value = _first_present(memory_output, "scriptMemory", "script_memory", default=None)
                if memory_value is None:
                    logger.warning(
                        "framework-to-script stage12 script_memory missing scriptMemory/script_memory, keeping previous memory memory_output_keys=%s",
                        memory_keys,
                    )
                else:
                    script_memory = str(memory_value)
                _autosave_stage12_draft(
                    "script_memory",
                    batch_script_value=batch_script,
                    script_review_value=script_review,
                    script_memory_value=script_memory,
                    debug_path_value=debug_path,
                )
                memory_event.update(
                    {
                        "raw_response_summary": _stage12_debug_summary(memory_output, preview_limit=600),
                        "parsed_fields": memory_keys,
                        "scriptMemory_length_after": len(script_memory),
                        "scriptMemory_preview_after": _stage12_debug_preview(script_memory),
                    }
                )
                debug_record.update(
                    {
                        "status": "script_memory_done",
                        "parsed_fields": memory_keys,
                        "scriptMemory_length": len(script_memory),
                        "updated_at": _now_iso(),
                    }
                )
                debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                logger.info(
                    "framework-to-script stage12 script_memory output memory_output_keys=%s normalized_keys=%s scriptMemory_length=%s",
                    memory_keys,
                    ["scriptMemory", "script_memory"],
                    len(script_memory),
                )
            except Exception as exc:
                debug_record.update(
                    {
                        "status": "failed",
                        "failure_phase": failed_sub_stage,
                        "failed_sub_stage": failed_sub_stage,
                        "exception_type": type(exc).__name__,
                        "exception_message": str(exc),
                        "traceback": traceback.format_exc(),
                        "updated_at": _now_iso(),
                    }
                )
                debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                logger.exception(
                    "framework-to-script stage12 failed failed_sub_stage=%s asset_id=%s start_episode=%s "
                    "end_episode=%s batch_plan_count=%s base_vars_keys=%s error_type=%s error_message=%s",
                    failed_sub_stage,
                    asset_id,
                    start_episode,
                    end_episode,
                    len(batch_plan),
                    sorted(base_vars.keys()),
                    type(exc).__name__,
                    str(exc),
                )
                return jsonify(
                    {
                        "success": False,
                        "message": f"12 阶段第 {start_episode or batch_key}-{end_episode or '?'} 集生成失败：{failed_sub_stage}",
                        "detail": {
                            "failed_sub_stage": failed_sub_stage,
                            "error_type": type(exc).__name__,
                            "error_message": str(exc),
                            "asset_id": asset_id,
                            "start_episode": start_episode,
                            "end_episode": end_episode,
                            "batch_plan_count": len(batch_plan),
                            "base_var_keys": sorted(base_vars.keys()),
                            "debug_path": debug_path,
                        },
                    }
                ), 500

            batches = dict(existing_batches)
            batch_output = {
                "batchStartEpisode": start_episode,
                "batchEndEpisode": end_episode,
                "batchEnrichedEpisodePlan": batch_plan,
                "batchCausalConflictPlan": conflict_plan,
                "batch_causal_conflict_plan": conflict_plan,
                "batchScriptText": batch_script,
                "batch_script_text": batch_script,
                "batchScriptReview": script_review,
                "batch_script_review": script_review,
                "scriptMemory": script_memory,
                "script_memory": script_memory,
                "stage12DebugPath": debug_path,
            }
            batches[str(start_episode)] = batch_output
            output = {**batch_output, "batches": batches}
            if asset_id:
                debug_record.update(
                    {
                        "status": "saving",
                        "final_save_fields": sorted(batch_output.keys()),
                        "final_batchScriptText_length": len(batch_script),
                        "final_scriptMemory_length": len(script_memory),
                        "updated_at": _now_iso(),
                    }
                )
                debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
                _save_framework_to_script_stage(
                    user_id=user_id,
                    asset_id=asset_id,
                    stage_key="stage12",
                    output=output,
                )
            debug_record.update(
                {
                    "status": "success",
                    "failure_phase": "",
                    "failed_sub_stage": "",
                    "final_save_fields": sorted(batch_output.keys()),
                    "final_batchScriptText_length": len(batch_script),
                    "final_scriptMemory_length": len(script_memory),
                    "updated_at": _now_iso(),
                }
            )
            debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset)
            success_debug_path = _write_stage12_debug_file(debug_record, data=data, framework_asset=framework_asset, success=True)
            if debug_path:
                output["stage12DebugPath"] = debug_path
            if success_debug_path:
                output["stage12SuccessDebugPath"] = success_debug_path
        finally:
            _end_framework_stage(user_id, asset_id, "12")

        return _json_ok(stage="12", framework_asset_id=asset_id, **output)


    @app.route("/api/framework-to-script/export/txt", methods=["GET", "POST"])
    @_login_required
    def export_framework_to_script_txt_api():
        data = request.get_json(silent=True) if request.method == "POST" else {}
        data = data if isinstance(data, dict) else {}
        asset_id = str(
            request.args.get("framework_asset_id")
            or data.get("framework_asset_id")
            or request.args.get("asset_id")
            or data.get("asset_id")
            or ""
        ).strip()
        if not asset_id:
            return _json_error("缺少 framework_asset_id。", status=400)
        asset = _load_framework_asset_for_user(asset_id, _require_user_id())
        if not asset:
            return _json_error("框架资产不存在、无权访问，或尚未生成 07 最终策划包。", status=404)
        text = _framework_to_script_txt(asset)
        filename = _framework_to_script_export_filename(asset, "txt")
        return Response(
            text,
            content_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
        )

    @app.route("/api/framework-to-script/export/docx", methods=["GET", "POST"])
    @_login_required
    def export_framework_to_script_docx_api():
        data = request.get_json(silent=True) if request.method == "POST" else {}
        data = data if isinstance(data, dict) else {}
        asset_id = str(
            request.args.get("framework_asset_id")
            or data.get("framework_asset_id")
            or request.args.get("asset_id")
            or data.get("asset_id")
            or ""
        ).strip()
        if not asset_id:
            return _json_error("缺少 framework_asset_id。", status=400)
        asset = _load_framework_asset_for_user(asset_id, _require_user_id())
        if not asset:
            return _json_error("框架资产不存在、无权访问，或尚未生成 07 最终策划包。", status=404)
        text = _framework_to_script_txt(asset)
        filename = _framework_to_script_export_filename(asset, "docx")
        try:
            from .utils.txt_to_docx import convert as convert_txt_to_docx
            with tempfile.TemporaryDirectory() as tmp_dir:
                tmp_path = Path(tmp_dir)
                txt_path = tmp_path / "framework_script.txt"
                docx_path = tmp_path / "framework_script.docx"
                txt_path.write_text(text, encoding="utf-8")
                convert_txt_to_docx(str(txt_path), str(docx_path))
                docx_bytes = docx_path.read_bytes()
                return send_file(
                    BytesIO(docx_bytes),
                    as_attachment=True,
                    download_name=filename,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        except ModuleNotFoundError as exc:
            if exc.name == "docx":
                return _json_error("当前环境缺少 python-docx，暂时无法导出 DOCX。", status=500)
            raise
        except Exception as exc:
            logger.exception("framework-to-script docx export failed")
            return _json_error(str(exc), status=500, fallback="DOCX 导出失败，请稍后重试。")


    @app.get("/framework-to-script")
    @_login_required
    def framework_to_script_workspace():
        return render_template(
            "framework_to_script.html",
            workflow_line=PRODUCTION_WORKFLOW_LINE,
            framework_planner_url="/framework-planner",
            page_title="框架转剧本工作台",
        )

    @app.get("/framework-to-script-test")
    @_login_required
    def framework_to_script_test_workspace():
        return render_template(
            "framework_to_script.html",
            workflow_line=TEST_WORKFLOW_LINE,
            framework_planner_url="/framework-planner-test",
            page_title="新工作流测试线路 08-12",
        )

    @app.get("/api/test-workflow/status")
    @_login_required
    def test_workflow_status_api():
        return _json_ok(**test_workflow_status())

    @app.get("/api/compliance-review/catalog")
    @_login_required
    def compliance_review_catalog_api():
        return _json_ok(**compliance_catalog())

    @app.post("/api/compliance-review/check")
    @_login_required
    def compliance_review_check_api():
        payload = request.get_json(silent=True) or {}
        try:
            report = run_compliance_review(payload)
            filename = str(payload.get("filename") or "").strip()
            title = str(payload.get("title") or "").strip()
            if not title and filename:
                title = Path(filename).stem
            if not title:
                first_line = next(
                    (
                        line.strip(" #《》\t")
                        for line in str(payload.get("text") or "").splitlines()
                        if line.strip()
                    ),
                    "未命名检测",
                )
                title = first_line[:100] or "未命名检测"
            history_record = None
            history_warning = ""
            try:
                history_record = compliance_review_store.save(
                    _require_user_id(),
                    report=report,
                    title=title,
                    filename=filename,
                    text=str(payload.get("text") or ""),
                )
            except Exception:
                logger.exception("script compliance history save failed")
                history_warning = "检测已完成，但记录保存失败。"
            return _json_ok(
                report=report,
                history_record=history_record,
                history_warning=history_warning,
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception:
            logger.exception("script compliance review failed")
            return _json_error("合规检测暂时失败，请稍后重试。", status=500)

    @app.get("/api/compliance-review/history")
    @_login_required
    def compliance_review_history_api():
        try:
            limit = int(request.args.get("limit") or 50)
        except (TypeError, ValueError):
            limit = 50
        return _json_ok(
            records=compliance_review_store.list(_require_user_id(), limit=limit)
        )

    @app.get("/api/compliance-review/history/<entry_id>")
    @_login_required
    def compliance_review_history_get_api(entry_id: str):
        record = compliance_review_store.get(_require_user_id(), entry_id)
        if not record:
            return _json_error("检测记录不存在。", status=404)
        return _json_ok(record=record)

    @app.delete("/api/compliance-review/history/<entry_id>")
    @_login_required
    def compliance_review_history_delete_api(entry_id: str):
        deleted = compliance_review_store.delete(_require_user_id(), entry_id)
        if not deleted:
            return _json_error("检测记录不存在。", status=404)
        return _json_ok(deleted=True)

    @app.delete("/api/compliance-review/history")
    @_login_required
    def compliance_review_history_clear_api():
        return _json_ok(deleted=compliance_review_store.clear(_require_user_id()))

    @app.post("/api/script-rating/check")
    @_login_required
    def script_rating_check_api():
        payload = request.get_json(silent=True) or {}
        try:
            report = run_script_rating(payload)
            text = str(payload.get("text") or "")
            filename = str(payload.get("filename") or "").strip()
            title = str(payload.get("title") or "").strip()
            if not title and filename:
                title = Path(filename).stem
            if not title:
                title = next((line.strip(" #《》\t") for line in text.splitlines() if line.strip()), "未命名评级")[:100]
            record = script_rating_store.save(
                _require_user_id(), report=report, title=title,
                filename=filename, text=text,
            )
            return _json_ok(report=report, history_record=record)
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception:
            logger.exception("script rating failed")
            return _json_error("剧本评级暂时失败，请稍后重试。", status=500)

    @app.get("/api/script-rating/history")
    @_login_required
    def script_rating_history_api():
        try:
            limit = int(request.args.get("limit") or 50)
        except (TypeError, ValueError):
            limit = 50
        return _json_ok(records=script_rating_store.list(_require_user_id(), limit))

    @app.get("/api/script-rating/history/<entry_id>")
    @_login_required
    def script_rating_history_get_api(entry_id: str):
        record = script_rating_store.get(_require_user_id(), entry_id)
        if not record:
            return _json_error("评级记录不存在。", status=404)
        return _json_ok(record=record)

    @app.post("/api/script-rating/repair")
    @_login_required
    def script_rating_repair_api():
        payload = request.get_json(silent=True) or {}
        try:
            return _json_ok(repair=run_script_rating_repair(payload))
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception:
            logger.exception("script rating repair failed")
            return _json_error("AI精准修复暂时失败，请稍后重试。", status=500)

    @app.delete("/api/script-rating/history/<entry_id>")
    @_login_required
    def script_rating_history_delete_api(entry_id: str):
        if not script_rating_store.delete(_require_user_id(), entry_id):
            return _json_error("评级记录不存在。", status=404)
        return _json_ok(deleted=True)

    @app.delete("/api/script-rating/history")
    @_login_required
    def script_rating_history_clear_api():
        return _json_ok(deleted=script_rating_store.clear(_require_user_id()))

    @app.get("/api/distillation-lab/overview")
    @_login_required
    def distillation_lab_overview_api():
        return _json_ok(**distillation_lab_store.overview(_require_user_id()))

    @app.post("/api/distillation-lab/projects")
    @_login_required
    def distillation_lab_create_project_api():
        payload = request.get_json(silent=True) or {}
        try:
            project = distillation_lab_store.create_project(_require_user_id(), payload)
            return _json_ok(project=project)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

    @app.get("/api/distillation-lab/projects/<project_id>")
    @_login_required
    def distillation_lab_project_api(project_id: str):
        try:
            return _json_ok(
                project=distillation_lab_store.get_project(_require_user_id(), project_id)
            )
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)

    @app.put("/api/distillation-lab/projects/<project_id>")
    @_login_required
    def distillation_lab_update_project_api(project_id: str):
        payload = request.get_json(silent=True) or {}
        try:
            project = distillation_lab_store.update_project(
                _require_user_id(), project_id, payload
            )
            return _json_ok(project=project)
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

    @app.delete("/api/distillation-lab/projects/<project_id>")
    @_login_required
    def distillation_lab_delete_project_api(project_id: str):
        try:
            distillation_lab_store.delete_project(_require_user_id(), project_id)
            return _json_ok(deleted=True)
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)

    @app.post("/api/distillation-lab/projects/<project_id>/sources")
    @_login_required
    def distillation_lab_upload_sources_api(project_id: str):
        files = request.files.getlist("files")
        if not files:
            return _json_error("请选择要上传的素材文件。", status=400)
        polarity = str(request.form.get("polarity") or "positive")
        try:
            weight = float(request.form.get("weight") or 1.0)
        except ValueError:
            weight = 1.0
        created: list[dict[str, Any]] = []
        errors: list[dict[str, str]] = []
        for uploaded in files:
            try:
                created.append(
                    distillation_lab_store.add_source(
                        _require_user_id(),
                        project_id,
                        uploaded.filename or "source.txt",
                        uploaded.read(),
                        polarity=polarity,
                        weight=weight,
                    )
                )
            except (KeyError, ValueError) as exc:
                errors.append({"name": uploaded.filename or "未命名文件", "error": str(exc).strip("'")})
        if not created:
            return _json_error(errors[0]["error"] if errors else "素材上传失败。", status=400)
        return _json_ok(sources=created, errors=errors)

    @app.delete("/api/distillation-lab/projects/<project_id>/sources/<source_id>")
    @_login_required
    def distillation_lab_delete_source_api(project_id: str, source_id: str):
        try:
            distillation_lab_store.delete_source(_require_user_id(), project_id, source_id)
            return _json_ok(deleted=True)
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)

    @app.post("/api/distillation-lab/projects/<project_id>/runs")
    @_login_required
    def distillation_lab_start_run_api(project_id: str):
        try:
            run = distillation_lab_store.start_run(_require_user_id(), project_id)
            return _json_ok(run=run)
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

    @app.get("/api/distillation-lab/runs/<run_id>")
    @_login_required
    def distillation_lab_run_api(run_id: str):
        try:
            return _json_ok(run=distillation_lab_store.get_run(_require_user_id(), run_id))
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)

    @app.put("/api/distillation-lab/projects/<project_id>/versions/<version_id>")
    @_login_required
    def distillation_lab_update_version_api(project_id: str, version_id: str):
        payload = request.get_json(silent=True) or {}
        try:
            version = distillation_lab_store.update_version(
                _require_user_id(), project_id, version_id, payload
            )
            return _json_ok(version=version)
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

    @app.post("/api/distillation-lab/projects/<project_id>/versions/<version_id>/publish")
    @_login_required
    def distillation_lab_publish_version_api(project_id: str, version_id: str):
        try:
            version = distillation_lab_store.publish_version(
                _require_user_id(), project_id, version_id
            )
            return _json_ok(version=version)
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

    @app.post("/api/distillation-lab/projects/<project_id>/versions/<version_id>/optimize")
    @_login_required
    def distillation_lab_optimize_version_api(project_id: str, version_id: str):
        try:
            version = distillation_lab_store.optimize_version_with_ai(
                _require_user_id(), project_id, version_id
            )
            return _json_ok(version=version)
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)
        except (ValueError, DeepSeekAgentError) as exc:
            return _json_error(str(exc), status=400)

    @app.post("/api/distillation-lab/projects/<project_id>/versions/<version_id>/unpublish")
    @_login_required
    def distillation_lab_unpublish_version_api(project_id: str, version_id: str):
        try:
            version = distillation_lab_store.unpublish_version(
                _require_user_id(), project_id, version_id
            )
            return _json_ok(version=version)
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)
        except ValueError as exc:
            return _json_error(str(exc), status=400)

    @app.delete("/api/distillation-lab/projects/<project_id>/versions/<version_id>")
    @_login_required
    def distillation_lab_delete_version_api(project_id: str, version_id: str):
        try:
            distillation_lab_store.delete_version(_require_user_id(), project_id, version_id)
            return _json_ok(deleted=True)
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)

    @app.get("/api/distillation-lab/projects/<project_id>/versions/<version_id>/export")
    @_login_required
    def distillation_lab_export_version_api(project_id: str, version_id: str):
        try:
            buffer, filename = distillation_lab_store.export_version(
                _require_user_id(), project_id, version_id
            )
            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype="application/zip",
            )
        except KeyError as exc:
            return _json_error(str(exc).strip("'"), status=404)

    @app.get("/api/new-workflow-test/npc/config")
    @_login_required
    def codebuddy_npc_config_api():
        return _json_ok(config=codebuddy_npc_config.public_status())

    @app.get("/api/new-workflow-test/skills")
    @_login_required
    def new_workflow_skill_catalog_api():
        return _json_ok(skills=distillation_lab_store.list_published_skills(_require_user_id()))

    @app.post("/api/new-workflow-test/npc/jobs")
    @_login_required
    def codebuddy_npc_create_job_api():
        payload = request.get_json(silent=True) or {}
        if not isinstance(payload, dict):
            return _json_error("NPC 剧本任务格式不正确。", status=400)
        try:
            skill_id = str(payload.get("distilled_skill_id") or "").strip()
            if skill_id:
                payload["distilled_skill_snapshot"] = distillation_lab_store.resolve_runtime_skill(
                    _require_user_id(),
                    skill_id,
                    str(payload.get("distilled_skill_version_id") or "").strip(),
                )
            job = codebuddy_npc_jobs.create(
                user_id=_require_user_id(),
                request_payload=payload,
            )
            continue_after = job.get("execution_mode") == "auto"
            job = codebuddy_npc_jobs.save(job)
            try:
                job = _submit_remote_npc_stage(
                    job,
                    stage="showrunner",
                    continue_after=continue_after,
                    retry_count=0,
                )
            except CodeBuddyNpcError as remote_exc:
                job = codebuddy_npc_stage_runner.start(
                    job_id=str(job["job_id"]),
                    user_id=_require_user_id(),
                    stage="showrunner",
                    continue_after=continue_after,
                )
                job["fallback_reason"] = str(remote_exc)
                job["status_text"] = "CNB 暂不可用，已切换本地兜底运行总编剧"
                job = codebuddy_npc_jobs.save(job)
        except (ValueError, KeyError) as exc:
            return _json_error(str(exc).strip("'"), status=400)
        except CodeBuddyNpcError as exc:
            logger.warning("codebuddy npc create failed: %s", exc)
            return _json_error(
                str(exc),
                status=exc.status_code,
                detail=exc.detail if isinstance(exc.detail, dict) else None,
            )
        except Exception as exc:
            logger.exception("codebuddy npc create failed")
            return _json_error(str(exc), status=500, fallback="NPC 剧本任务创建失败。")
        return _json_ok(job=public_job(job))

    @app.get("/api/new-workflow-test/npc/jobs/latest")
    @_login_required
    def codebuddy_npc_latest_job_api():
        job = codebuddy_npc_jobs.latest(user_id=_require_user_id())
        return _json_ok(job=public_job(job) if job else None)

    @app.get("/api/new-workflow-test/npc/jobs")
    @_login_required
    def codebuddy_npc_list_jobs_api():
        jobs = codebuddy_npc_jobs.list(user_id=_require_user_id())
        history = []
        for item in jobs:
            item = _refresh_remote_npc_job(item)
            request_data = item.get("request") if isinstance(item.get("request"), dict) else {}
            history.append(
                {
                    "job_id": item.get("job_id"),
                    "project_title": request_data.get("project_title") or "未命名剧本",
                    "mode": request_data.get("mode") or "原创",
                    "episodes": request_data.get("episodes"),
                    "episode_start": request_data.get("episode_start") or 1,
                    "episode_end": request_data.get("episode_end") or request_data.get("episodes"),
                    "production_type": request_data.get("production_type"),
                    "status": item.get("status"),
                    "status_text": item.get("status_text"),
                    "progress": item.get("progress"),
                    "active_stage": item.get("active_stage"),
                    "execution_mode": item.get("execution_mode"),
                    "execution_target": item.get("execution_target"),
                    "build_log_url": (item.get("build") or {}).get("build_log_url"),
                    "has_final_script": bool(str(item.get("final_script") or "").strip()),
                    "created_at": item.get("created_at"),
                    "updated_at": item.get("updated_at"),
                }
            )
        return _json_ok(jobs=history)

    @app.get("/api/new-workflow-test/npc/jobs/<job_id>")
    @_login_required
    def codebuddy_npc_job_api(job_id: str):
        job = codebuddy_npc_jobs.load(job_id, user_id=_require_user_id())
        if not job:
            return _json_error("NPC 剧本任务不存在。", status=404)
        job = _refresh_remote_npc_job(job)
        return _json_ok(job=public_job(job))

    @app.delete("/api/new-workflow-test/npc/jobs/<job_id>")
    @_login_required
    def codebuddy_npc_delete_job_api(job_id: str):
        if codebuddy_npc_stage_runner.is_running(job_id):
            return _json_error("节点正在运行，请先停止后再删除。", status=409)
        deleted = codebuddy_npc_jobs.delete(job_id, user_id=_require_user_id())
        if not deleted:
            return _json_error("NPC 剧本任务不存在。", status=404)
        return _json_ok(deleted=True)

    @app.get("/api/new-workflow-test/npc/jobs/<job_id>/export/docx")
    @_login_required
    def codebuddy_npc_export_docx_api(job_id: str):
        job = codebuddy_npc_jobs.load(job_id, user_id=_require_user_id())
        if not job:
            return _json_error("NPC 剧本任务不存在。", status=404)
        final_script = str(job.get("final_script") or "").strip()
        if not final_script:
            return _json_error("当前任务还没有可导出的最终剧本。", status=409)
        request_data = job.get("request") if isinstance(job.get("request"), dict) else {}
        title = str(request_data.get("project_title") or "完整剧本").strip()
        safe_title = re.sub(r'[\\/:*?"<>|\r\n]+', "_", title).strip(" ._") or "完整剧本"
        try:
            from .services.script_delivery import build_delivery_script
            from .utils.txt_to_docx import convert_script_team

            with tempfile.TemporaryDirectory(prefix="npc-script-docx-") as tmp_dir:
                tmp_path = Path(tmp_dir)
                txt_path = tmp_path / "final_script.txt"
                docx_path = tmp_path / "final_script.docx"
                txt_path.write_text(build_delivery_script(job), encoding="utf-8")
                convert_script_team(
                    str(txt_path),
                    str(docx_path),
                    title=title,
                )
                docx_bytes = docx_path.read_bytes()
        except ModuleNotFoundError as exc:
            if exc.name == "docx":
                return _json_error("当前环境缺少 python-docx，暂时无法导出 Word。", status=500)
            raise
        except Exception as exc:
            logger.exception("codebuddy npc docx export failed: job_id=%s", job_id)
            return _json_error(str(exc), status=500, fallback="Word 导出失败，请稍后重试。")
        return send_file(
            BytesIO(docx_bytes),
            as_attachment=True,
            download_name=f"{safe_title}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    @app.post("/api/new-workflow-test/npc/jobs/<job_id>/stages/<stage>/run")
    @_login_required
    def codebuddy_npc_run_stage_api(job_id: str, stage: str):
        payload = request.get_json(silent=True) or {}
        try:
            current_job = (
                codebuddy_npc_jobs.load(job_id, user_id=_require_user_id())
                if stage == "state_recorder"
                else None
            )
            if stage == "state_recorder" and current_job and _state_recorder_uses_code(current_job):
                job = _complete_code_state_recorder(
                    current_job,
                    continue_after=bool(payload.get("continue_after")),
                )
            elif bool(payload.get("local_fallback")):
                job = codebuddy_npc_stage_runner.start(
                    job_id=job_id,
                    user_id=_require_user_id(),
                    stage=stage,
                    feedback=str(payload.get("feedback") or ""),
                    continue_after=bool(payload.get("continue_after")),
                )
            else:
                job = codebuddy_npc_stage_runner.prepare_remote(
                    job_id=job_id,
                    user_id=_require_user_id(),
                    stage=stage,
                )
                try:
                    job = _submit_remote_npc_stage(
                        job,
                        stage=stage,
                        feedback=str(payload.get("feedback") or ""),
                        continue_after=bool(payload.get("continue_after")),
                    )
                except CodeBuddyNpcError as remote_exc:
                    job = codebuddy_npc_stage_runner.start(
                        job_id=job_id,
                        user_id=_require_user_id(),
                        stage=stage,
                        feedback=str(payload.get("feedback") or ""),
                        continue_after=bool(payload.get("continue_after")),
                    )
                    job["fallback_reason"] = str(remote_exc)
                    job["status_text"] = (
                        f"CNB 暂不可用，已切换本地兜底运行{STAGE_NAMES[stage]}"
                    )
                    job = codebuddy_npc_jobs.save(job)
        except CodeBuddyNpcError as exc:
            return _json_error(str(exc), status=exc.status_code)
        return _json_ok(job=public_job(job))

    @app.post("/api/new-workflow-test/npc/jobs/<job_id>/cancel")
    @_login_required
    def codebuddy_npc_cancel_stage_api(job_id: str):
        try:
            current = codebuddy_npc_jobs.load(job_id, user_id=_require_user_id())
            build_sn = str(((current or {}).get("build") or {}).get("sn") or "")
            job = codebuddy_npc_stage_runner.request_cancel(
                job_id=job_id,
                user_id=_require_user_id(),
            )
            if current and str(current.get("execution_target") or "") == "remote_cnb":
                if build_sn:
                    try:
                        codebuddy_npc_client.stop_build(build_sn)
                    except CodeBuddyNpcError:
                        logger.warning("failed to stop remote CNB build: %s", build_sn)
        except CodeBuddyNpcError as exc:
            return _json_error(str(exc), status=exc.status_code)
        return _json_ok(job=public_job(job))

    @app.put("/api/new-workflow-test/npc/jobs/<job_id>/artifacts/<artifact_key>")
    @_login_required
    def codebuddy_npc_edit_artifact_api(job_id: str, artifact_key: str):
        payload = request.get_json(silent=True) or {}
        try:
            job = codebuddy_npc_stage_runner.edit_artifact(
                job_id=job_id,
                user_id=_require_user_id(),
                artifact_key=artifact_key,
                content=str(payload.get("content") or ""),
            )
        except CodeBuddyNpcError as exc:
            return _json_error(str(exc), status=exc.status_code)
        return _json_ok(job=public_job(job))

    @app.post("/api/new-workflow-test/npc/jobs/<job_id>/recover")
    @_login_required
    def codebuddy_npc_recover_job_api(job_id: str):
        job = codebuddy_npc_jobs.load(job_id, user_id=_require_user_id())
        if not job:
            return _json_error("NPC 剧本任务不存在。", status=404)
        try:
            job = codebuddy_npc_client.refresh(job)
            job = codebuddy_npc_jobs.save(job)
        except CodeBuddyNpcError as exc:
            return _json_error(
                str(exc),
                status=exc.status_code,
                detail=exc.detail if isinstance(exc.detail, dict) else None,
            )
        except Exception as exc:
            logger.exception("codebuddy npc recovery failed: job_id=%s", job_id)
            return _json_error(str(exc), status=500, fallback="NPC 任务产物恢复失败。")
        if not job.get("final_script"):
            return _json_error(
                "已检查成功阶段日志，但没有找到可恢复的完整最终稿。可使用原任务参数重新生成。",
                status=409,
                detail={
                    "recovered_files": sorted(
                        (job.get("recovered_files") or {}).keys()
                    )
                },
            )
        return _json_ok(job=public_job(job))

    @app.get("/api/new-workflow-test/projects/<int:project_id>/outputs")
    @_login_required
    def new_workflow_test_outputs_api(project_id: int):
        snapshot = task_manager.get_project_snapshot(
            project_id,
            user_id=_require_user_id(),
            public_view=False,
        )
        if not snapshot:
            return _json_error("测试项目不存在。", status=404)

        outputs: dict[str, Any] = {}
        artifacts = snapshot.get("artifacts") if isinstance(snapshot.get("artifacts"), dict) else {}
        test_state = artifacts.get("new_workflow_test_state") if isinstance(artifacts.get("new_workflow_test_state"), dict) else {}
        saved_outputs = test_state.get("outputs") if isinstance(test_state.get("outputs"), dict) else {}
        for stage, value in saved_outputs.items():
            if _framework_value_present(value):
                outputs[str(stage).zfill(2)] = _strip_raw_fastgpt_fields(copy.deepcopy(value))

        framework_state = _framework_state_from_project(snapshot)
        planner_fields = {
            "01": framework_state.get("source_brief"),
            "02": framework_state.get("worldview_plan"),
            "03": framework_state.get("character_plan"),
            "04": framework_state.get("beat_checkpoint_timeline"),
            "05": framework_state.get("character_storylines"),
            "06": framework_state.get("adaptation_guide"),
            "07": framework_state.get("framework_plan_package"),
        }
        for stage, value in planner_fields.items():
            if stage not in outputs and _framework_value_present(value):
                outputs[stage] = _strip_raw_fastgpt_fields(copy.deepcopy(value))

        def _has_test_business_content(value: Any) -> bool:
            if not isinstance(value, dict):
                return _framework_value_present(value)
            ignored = {
                "prompt_preferences",
                "stage_prompts",
                "user_knowledge_stage_prompts",
                "parse_warning",
                "warnings",
            }
            return any(
                key not in ignored and _framework_value_present(item)
                for key, item in value.items()
            )

        project_title = str(snapshot.get("title") or framework_state.get("project_title") or "").strip()
        cache_root = Path(__file__).resolve().parents[2] / "cache"
        safe_project_name = Path(project_title).name
        project_cache = (cache_root / safe_project_name).resolve()
        if safe_project_name and project_cache.parent == cache_root.resolve():
            for stage in ("01", "02", "03", "04", "05", "06", "07"):
                if _has_test_business_content(outputs.get(stage)):
                    continue
                cache_path = project_cache / f"latest_stage{stage}.json"
                try:
                    cached = json.loads(cache_path.read_text(encoding="utf-8"))
                except (OSError, ValueError):
                    continue
                cached_output = cached.get("output") if isinstance(cached, dict) else None
                if str(cached.get("status") or "") == "success" and _has_test_business_content(cached_output):
                    outputs[stage] = _strip_raw_fastgpt_fields(copy.deepcopy(cached_output))

        workspace = artifacts.get("framework_to_script_state") if isinstance(artifacts.get("framework_to_script_state"), dict) else {}
        script_stages = workspace.get("scriptStages") if isinstance(workspace.get("scriptStages"), dict) else {}
        for stage in ("08", "09", "10", "11", "12"):
            value = script_stages.get(f"stage{stage}")
            if _framework_value_present(value):
                outputs[stage] = _strip_raw_fastgpt_fields(copy.deepcopy(value))

        test_state["outputs"] = copy.deepcopy(outputs)
        test_state["updated_at"] = _now_iso()
        artifacts["new_workflow_test_state"] = test_state
        snapshot["artifacts"] = artifacts
        snapshot["updated_at"] = test_state["updated_at"]
        record = task_manager._projects.get(project_id)
        if record:
            with record.lock:
                record.snapshot = snapshot
            task_manager._persist_snapshot(record)
        else:
            task_manager._persist_project_snapshot_data(project_id, snapshot)

        return _json_ok(
            project_id=project_id,
            outputs=outputs,
            updated_at=test_state.get("updated_at") or workspace.get("updated_at") or snapshot.get("updated_at"),
        )

    @app.post("/api/framework-planner/assets/save")
    @_login_required
    def save_framework_planner_asset_api():
        data = request.get_json(silent=True) or {}
        if not isinstance(data, dict):
            return _json_error("保存内容格式不正确", status=400)
        try:
            asset = task_manager.save_framework_planner_asset(
                user_id=_require_user_id(),
                payload=data,
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        except Exception as exc:
            logger.exception("framework planner asset save failed")
            return _json_error(str(exc), status=500, fallback="当前框架保存失败，请稍后重试。")
        return _json_ok(asset=asset, project=asset, project_id=asset.get("project_id"))

    @app.get("/api/community")
    def community_assets():
        assets = task_manager.list_public_assets()
        return _json_ok(assets=assets)

    @app.get("/api/community/<int:project_id>")
    def community_asset_detail(project_id: int):
        asset = task_manager.get_public_asset(project_id)
        if not asset:
            return _json_error("公开作品不存在", status=404)
        return _json_ok(asset=asset)

    @app.get("/api/projects/<int:project_id>")
    @_login_required
    def get_project(project_id: int):
        snapshot = task_manager.get_project_snapshot(project_id, user_id=_require_user_id())
        if not snapshot:
            return _json_error("项目不存在", status=404)
        return _json_ok(project=snapshot)

    @app.patch("/api/projects/<int:project_id>")
    @_login_required
    def update_project(project_id: int):
        data = request.get_json(silent=True) or {}
        try:
            snapshot = task_manager.update_project_asset(
                project_id,
                user_id=_require_user_id(),
                changes=data,
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(project=snapshot)

    @app.post("/api/projects/<int:project_id>/confirm-completion")
    @_login_required
    def confirm_project_completion(project_id: int):
        try:
            snapshot = task_manager.confirm_project_completion(
                project_id,
                user_id=_require_user_id(),
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(project=snapshot)

    @app.post("/api/projects/<int:project_id>/rollback")
    @_login_required
    def rollback_project(project_id: int):
        data = request.get_json(silent=True) or {}
        try:
            snapshot = task_manager.rollback_project_to_stage(
                project_id,
                user_id=_require_user_id(),
                stage_key=str(data.get("stage_key") or ""),
                start_episode=data.get("start_episode"),
            )
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(task=snapshot)


    @app.post("/api/framework-planner/generate-script")
    @_login_required
    def generate_script_from_framework_planner():
        data = request.get_json(silent=True) or {}

        framework_plan_package = data.get("framework_plan_package")
        if not isinstance(framework_plan_package, dict) or not framework_plan_package:
            return _json_error(
                "缺少 framework_plan_package，请先完成并确认 07 最终策划包输出。",
                status=400,
            )

        basic_config = data.get("basic_config") if isinstance(data.get("basic_config"), dict) else {}
        title = (
            str(data.get("title") or "")
            or str(data.get("project_title") or "")
            or str(basic_config.get("project_title") or "")
            or str(basic_config.get("source_title") or "")
            or "未命名框架剧本"
        ).strip()

        def _safe_int(value, default):
            try:
                number = int(value)
                return number if number > 0 else default
            except Exception:
                return default

        total_episodes = _safe_int(
            data.get("total_episodes")
            or data.get("episodes_per_season")
            or basic_config.get("episodes_per_season"),
            60,
        )
        episode_word_count = _safe_int(
            data.get("episode_word_count") or basic_config.get("episode_word_count"),
            600,
        )

        expectation_parts = [
            str(data.get("user_expectation") or "").strip(),
            str(data.get("adaptation_direction") or basic_config.get("adaptation_direction") or "").strip(),
            str(data.get("user_requirements") or basic_config.get("user_requirements") or "").strip(),
        ]
        expectation = "\n".join([item for item in expectation_parts if item]).strip()
        if not expectation:
            expectation = f"基于《{title}》的三幕十五节拍框架策划包生成短剧正文。"

        payload = {
            "title": title,
            "project_title": title,
            "source_title": str(data.get("source_title") or basic_config.get("source_title") or title),
            "target_format": str(data.get("target_format") or basic_config.get("target_format") or "短剧"),
            "season_count": _safe_int(data.get("season_count") or basic_config.get("season_count"), 1),
            "episodes_per_season": total_episodes,
            "total_episodes": total_episodes,
            "episode_word_count": episode_word_count,
            "user_expectation": expectation,
            "user_requirements": str(data.get("user_requirements") or basic_config.get("user_requirements") or ""),
            "adaptation_direction": str(data.get("adaptation_direction") or basic_config.get("adaptation_direction") or ""),
            "basic_config": basic_config,
            "framework_plan_package": framework_plan_package,
            "source_framework_project_id": data.get("source_framework_project_id") or data.get("project_id"),
            "source_brief": data.get("source_brief") or framework_plan_package.get("source_brief") or {},
            "worldview_plan": data.get("worldview_plan") or framework_plan_package.get("worldview_plan") or {},
            "character_plan": data.get("character_plan") or framework_plan_package.get("character_plan") or {},
            "beat_checkpoint_timeline": data.get("beat_checkpoint_timeline") or framework_plan_package.get("beat_checkpoint_timeline") or [],
            "checkpoint_explanation": data.get("checkpoint_explanation") or framework_plan_package.get("checkpoint_explanation") or {},
            "character_storylines": data.get("character_storylines") or framework_plan_package.get("character_storylines") or [],
            "storyline_decisions": data.get("storyline_decisions") or framework_plan_package.get("storyline_decisions") or [],
            "adaptation_guide": data.get("adaptation_guide") or framework_plan_package.get("adaptation_guide") or {},
            "workflow_mode": "framework_to_script",
            "generation_chain": "framework_to_script",
            "framework_to_script": True,
        "script_format_mode": "framework_to_script",
            "framework_planner_source": True,
        }
        if isinstance(data.get("user_knowledge_step_prompts"), dict):
            payload["user_knowledge_step_prompts"] = data.get("user_knowledge_step_prompts")

        _attach_user_knowledge_payload(payload, data)

        try:
            snapshot = task_manager.start_task(
                user_id=_require_user_id(),
                input_payload=payload,
                workflow_spec_path=_resolve_spec_path(data),
                model_selection_id=data.get("model_selection_id"),
            )
        except Exception as exc:
            return _json_error(
                str(exc),
                status=400,
                fallback="框架转剧本任务创建失败，请检查新链路 FastGPT API Key 和工作流配置。",
            )

        return _json_ok(task=snapshot)


    @app.post("/api/workflows/start")
    @_login_required
    def start_workflow():
        data = request.get_json(silent=True) or {}
        spec_path = _resolve_spec_path(data)
        expectation = str(data.get("user_expectation", ""))
        payload = {
            "title": data.get("title", "") or derive_script_title_content(expectation),
            "episode_word_count": data.get("episode_word_count", 600) or 600,
            "total_episodes": data.get("total_episodes", 0),
            "user_expectation": expectation,
            "character_count": data.get("character_count", 0),
        }
        for optional_key in (
            "character_appearance_requirements",
            "character_alias_naming_rules",
            "outfit_switch_rules",
            "story_outline",
            "core_scene_input",
            "character_bios",
            "episode_plan",
            "script_format_mode",
            "framework_plan_package",
            "worldview_plan",
            "beat_checkpoint_timeline",
            "character_storylines",
            "character_plan",
        ):
            value = data.get(optional_key)
            if value not in (None, "", [], {}):
                payload[optional_key] = value
        _attach_user_knowledge_payload(payload, data)
        try:
            snapshot = task_manager.start_task(
                user_id=_require_user_id(),
                input_payload=payload,
                workflow_spec_path=spec_path,
                model_selection_id=data.get("model_selection_id"),
            )
        except Exception as exc:
            return _json_error(str(exc), status=400, fallback="任务创建失败，请稍后重试。")
        return _json_ok(task=snapshot)

    @app.post("/api/projects/<int:project_id>/restart")
    @_login_required
    def restart_project(project_id: int):
        data = request.get_json(silent=True) or {}
        spec_path = _resolve_spec_path(data)
        expectation = str(data.get("user_expectation", ""))
        payload = {
            "title": data.get("title", "") or derive_script_title_content(expectation),
            "episode_word_count": data.get("episode_word_count", 600) or 600,
            "total_episodes": data.get("total_episodes", 0),
            "user_expectation": expectation,
            "character_count": data.get("character_count", 0),
        }
        for optional_key in (
            "character_appearance_requirements",
            "character_alias_naming_rules",
            "outfit_switch_rules",
            "story_outline",
            "core_scene_input",
            "character_bios",
            "episode_plan",
            "script_format_mode",
            "framework_plan_package",
            "worldview_plan",
            "beat_checkpoint_timeline",
            "character_storylines",
            "character_plan",
        ):
            value = data.get(optional_key)
            if value not in (None, "", [], {}):
                payload[optional_key] = value
        _attach_user_knowledge_payload(payload, data)
        try:
            snapshot = task_manager.restart_project(
                project_id,
                user_id=_require_user_id(),
                input_payload=payload,
                workflow_spec_path=spec_path,
                model_selection_id=data.get("model_selection_id"),
            )
        except Exception as exc:
            return _json_error(str(exc), status=400, fallback="重新开始失败，请稍后重试。")
        return _json_ok(task=snapshot)

    @app.get("/api/tasks/<task_id>")
    @_login_required
    def get_task(task_id: str):
        snapshot = task_manager.get_task_snapshot(task_id, user_id=_require_user_id())
        if not snapshot:
            return _json_error("任务不存在", status=404)
        return _json_ok(task=snapshot)

    @app.get("/api/tasks/<task_id>/debug")
    @_login_required
    def get_task_debug(task_id: str):
        snapshot = task_manager.get_task_snapshot(
            task_id,
            user_id=_require_user_id(),
            public_view=False,
        )
        if not snapshot:
            return _json_error("任务不存在", status=404)

        return _json_ok(
            debug={
                "task_id": snapshot.get("task_id"),
                "project_id": snapshot.get("project_id"),
                "status": snapshot.get("status"),
                "message": snapshot.get("message"),
                "error": snapshot.get("error"),
                "current_stage": snapshot.get("current_stage"),
                "current_stage_label": snapshot.get("current_stage_label"),
                "current_node_id": snapshot.get("current_node_id"),
                "current_node_name": snapshot.get("current_node_name"),
                "current_batch": snapshot.get("current_batch"),
                "progress_percent": snapshot.get("progress_percent"),
                "generated_episodes": snapshot.get("generated_episodes"),
                "cache_retained": snapshot.get("cache_retained"),
                "awaiting_user_confirmation": snapshot.get("awaiting_user_confirmation"),
                "runtime_cache_notice": snapshot.get("runtime_cache_notice"),
                "logs": snapshot.get("logs", []),
                "last_log": snapshot.get("last_log"),
                "resume_checkpoint_exists": isinstance(snapshot.get("_resume_checkpoint"), dict),
                "resume_checkpoint": snapshot.get("_resume_checkpoint"),
                "debug_state": snapshot.get("debug_state", {}),
                "prompt_fixes": snapshot.get("prompt_fixes", []),
            }
        )

    @app.post("/api/tasks/<task_id>/pause")
    @_login_required
    def pause_task(task_id: str):
        try:
            snapshot = task_manager.pause_task(task_id, user_id=_require_user_id())
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(task=snapshot)

    @app.post("/api/tasks/<task_id>/resume")
    @_login_required
    def resume_task(task_id: str):
        try:
            snapshot = task_manager.resume_task(task_id, user_id=_require_user_id())
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(task=snapshot)

    @app.post("/api/tasks/<task_id>/retry")
    @_login_required
    def retry_task(task_id: str):
        try:
            snapshot = task_manager.retry_task(task_id, user_id=_require_user_id())
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(task=snapshot)

    @app.post("/api/tasks/<task_id>/terminate")
    @_login_required
    def terminate_task(task_id: str):
        try:
            snapshot = task_manager.terminate_task(task_id, user_id=_require_user_id())
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(task=snapshot)

    @app.delete("/api/tasks/<task_id>")
    @_login_required
    def delete_task(task_id: str):
        try:
            task_manager.delete_task(task_id, user_id=_require_user_id())
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(task_id=task_id)

    @app.delete("/api/projects/<int:project_id>")
    @_login_required
    def clear_project(project_id: int):
        try:
            task_manager.clear_project(project_id, user_id=_require_user_id())
        except ValueError as exc:
            return _json_error(str(exc), status=400)
        return _json_ok(project_id=project_id)

    @app.post("/api/projects/<int:project_id>/save")
    @_login_required
    def save_project(project_id: int):
        try:
            path = task_manager.save_final_script(project_id, user_id=_require_user_id())
        except ValueError as exc:
            return _json_error(str(exc), status=400, fallback="当前剧本暂时无法导出，请确认已生成完成后再试。")
        return _json_ok(project_id=project_id, saved_file=str(path))

    @app.get("/api/projects/<int:project_id>/download")
    @_login_required
    def download_project(project_id: int):
        try:
            path = task_manager.save_final_script(project_id, user_id=_require_user_id())
        except ValueError as exc:
            return _json_error(str(exc), status=400, fallback="当前剧本暂时无法下载，请确认已生成完成后再试。")
        return send_file(
            path,
            as_attachment=True,
            download_name=path.name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    print(
        "[route-probe] framework history route registered =",
        any(str(rule) == "/api/framework-planner/history" for rule in app.url_map.iter_rules()),
        flush=True,
    )

    return app
    extract_doctor_report,
    load_workbuddy_optimized_document,
    load_workbuddy_source,
    parse_doctor_optimization_result,
    save_workbuddy_optimized_document,
    save_workbuddy_source,
