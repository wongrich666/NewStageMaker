"""
txt_to_docx.py
将包含 JSON 块的 txt 剧本文件解析并生成结构化 Word 文档。

文件结构支持：
  - 纯文本（标题、故事梗概等）
  - ```json ... ``` 代码块（人物小传、场景设定等 JSON 数据）
  - 剧本正文（带场景标记的对白）

用法：
  python txt_to_docx.py <input.txt> <output.docx>
  或直接调用 convert(input_path, output_path)
"""

import re
import json
import logging
import sys
from pathlib import Path
from typing import Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .user_visible_text import clean_user_visible_text, is_meaningful_text, normalize_user_visible_text

logger = logging.getLogger(__name__)

PLAIN_SECTION_HEADINGS = (
    "故事梗概",
    "世界观设定",
    "人物小传",
    "人物服饰说明",
    "核心场景",
    "分集计划",
    "剧本正文",
)


# ─────────────────────────── 样式辅助 ───────────────────────────

def set_run_color(run, hex_color: str):
    """设置文字颜色，hex_color 如 '2E75B6'"""
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)


def add_heading(doc: Document, text: str, level: int = 1):
    """添加标题段落"""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc: Document, text: str, bold=False, italic=False,
             size=11, color=None, indent=0):
    """添加普通段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        set_run_color(run, color)
    return p


def add_divider(doc: Document):
    """添加分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_kv(doc: Document, key: str, value: str):
    """添加键值对段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run_k = p.add_run(f"{key}：")
    run_k.bold = True
    run_k.font.size = Pt(10)
    # 去除 Markdown 粗体标记
    clean_val = re.sub(r'\*\*(.+?)\*\*', r'\1', str(value))
    run_v = p.add_run(clean_val)
    run_v.font.size = Pt(10)


def safe_get(obj: Any, key: str, default=None):
    """仅在 dict 上安全读取字段。"""
    if isinstance(obj, dict):
        return obj.get(key, default)
    return default


def _json_fallback(value: Any) -> str:
    try:
        return json.dumps(value, ensure_ascii=False, indent=2)
    except Exception:
        return str(value)


def normalize_list(value: Any, *, _depth: int = 0) -> list[str]:
    """把任意值尽量转换成可读字符串列表。"""
    if _depth > 5:
        text = clean_user_visible_text(value).strip()
        return [text] if text else []
    text = clean_user_visible_text(value).strip()
    if not text:
        return []
    return [line.strip() for line in text.splitlines() if line.strip()] or [text]


def normalize_text(value: Any, *, _depth: int = 0) -> str:
    """把任意值尽量转换成单段可读文本。"""
    if _depth > 5:
        return clean_user_visible_text(value).strip()
    return normalize_user_visible_text(value).strip()


def split_plain_sections(text: str) -> tuple[str, dict[str, str]]:
    """按导出用的一级中文标题切分纯文本章节，兼容新导出格式。"""
    source = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = source.split("\n")
    title = ""
    index = 0
    while index < len(lines):
        line = str(lines[index] or "").strip()
        if line:
            title = line
            index += 1
            break
        index += 1

    sections: dict[str, list[str]] = {}
    current_heading = ""
    buffer: list[str] = []
    known = set(PLAIN_SECTION_HEADINGS)
    for raw_line in lines[index:]:
        line = str(raw_line or "")
        stripped = line.strip()
        if stripped in known:
            if current_heading:
                sections[current_heading] = buffer[:]
            current_heading = stripped
            buffer = []
            continue
        if current_heading:
            buffer.append(line)
    if current_heading:
        sections[current_heading] = buffer[:]

    return title, {
        key: "\n".join(value).strip()
        for key, value in sections.items()
        if "\n".join(value).strip()
    }


def _section_is_code_fence(text: str) -> bool:
    content = str(text or "").strip()
    return content.startswith("```json") or content.startswith("```")


def render_plain_section(doc: Document, title: str, text: Any):
    content = normalize_text(text)
    if not content:
        return
    add_heading(doc, title, level=1)
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", content) if part.strip()]
    if not paragraphs:
        paragraphs = [line.strip() for line in str(content).splitlines() if line.strip()]
    for paragraph in paragraphs:
        add_para(doc, paragraph)
    add_divider(doc)


def render_field(doc: Document, label: str, value: Any):
    """按字段类型安全渲染键值、列表或复杂对象。"""
    items = normalize_list(value)
    text = normalize_text(value)
    if not items and not text:
        return
    if len(items) <= 1:
        add_kv(doc, label, items[0] if items else text)
        return
    add_para(doc, label, bold=True, size=10, indent=0.5)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"• {item}")
        run.font.size = Pt(10)


def _extract_script_text(value: Any, *, _depth: int = 0) -> str:
    """从任意正文载体中递归提取可读文本。"""
    if _depth > 6:
        return normalize_text(value)
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        for key in (
            "final_script",
            "script",
            "content",
            "text",
            "body",
            "dialogue",
            "dialogues",
        ):
            text = _extract_script_text(safe_get(value, key), _depth=_depth + 1)
            if text:
                return text
        parts = [
            _extract_script_text(item, _depth=_depth + 1)
            for item in value.values()
        ]
        return "\n".join(part for part in parts if part).strip()
    if isinstance(value, (list, tuple, set)):
        parts = [_extract_script_text(item, _depth=_depth + 1) for item in value]
        return "\n".join(part for part in parts if part).strip()
    return normalize_text(value)


# ─────────────────────────── 解析 txt ───────────────────────────

def fix_json_block(raw: str) -> str:
    """
    修复文件中常见的非标准 JSON 写法：

    1. 值以 ** 开头（Markdown 粗体替代了开头引号）
       "key": **文字...  →  "key": "文字...
    2. 值以中文左引号 " (U+201C) 开头
       "key": "文字  →  "key": "文字
    3. 行内剩余 ** 粗体标记（在字符串内部）→ 直接去除
    4. 字符串内部的中文弯引号 " " → 转义为 \"
    5. 行尾多余反斜杠：...\",  或  ...\"  →  ...",  或  ..."
    """
    lines = raw.split('\n')
    result = []
    for line in lines:
        # 修复1：值以 ** 开头 —— 添加缺失的开引号，移除 **
        line = re.sub(r'(:\s*)\*\*', r'\1"', line)

        # 修复2：值以中文左引号开头
        line = re.sub(r'(:\s*)\u201c', r'\1"', line)

        # 修复3：去除字符串内部剩余的 ** 标记
        line = line.replace('**', '')

        # 修复4：转义字符串内部的中文弯引号
        line = line.replace('\u201c', '\\"').replace('\u201d', '\\"')

        # 修复5：行尾多余反斜杠  ...\"  或  ...\",  →  ..."  或  ...",
        line = re.sub(r'\\"(\s*,?\s*)$', r'"\1', line)

        result.append(line)
    return '\n'.join(result)


def parse_txt(filepath: str) -> dict:
    """
    解析 txt 文件，返回结构化内容：
    {
      "title": str,
      "synopsis": str,
      "json_blocks": [{"label": str, "data": dict}, ...],
      "script": str
    }
    """
    text = Path(filepath).read_text(encoding='utf-8')
    title, plain_sections = split_plain_sections(text)

    # 提取所有 ```json ... ``` 块（含位置信息）
    json_pattern = re.compile(r'```json\s*(.*?)```', re.DOTALL)
    json_matches = list(json_pattern.finditer(text))

    # 移除 JSON 块后的纯文本部分
    plain_text = json_pattern.sub('<<<JSON_BLOCK>>>', text)
    plain_parts = plain_text.split('<<<JSON_BLOCK>>>')

    # 兼容旧格式：第一个纯文本块包含标题和故事梗概
    header_block = plain_parts[0].strip() if plain_parts else ""
    lines = header_block.splitlines()
    legacy_title = lines[0].strip() if lines else ""
    title = title or legacy_title
    synopsis_lines = []
    in_synopsis = False
    for line in lines[1:]:
        stripped = line.strip()
        if stripped == "故事梗概":
            in_synopsis = True
            continue
        if in_synopsis and stripped and any(
            stripped.startswith(h) for h in ["人物小传", "核心场景", "剧本正文"]
        ):
            break
        if in_synopsis:
            synopsis_lines.append(line)
    synopsis = plain_sections.get("故事梗概") or "\n".join(synopsis_lines).strip()

    # 找剧本正文（最后一个 JSON 块之后的文本）
    script_text = plain_sections.get("剧本正文", "")
    if not script_text and len(plain_parts) > len(json_matches):
        last_part = plain_parts[-1]
        script_match = re.search(r'剧本正文(.*)', last_part, re.DOTALL)
        if script_match:
            script_text = script_match.group(0).strip()

    # 解析每个 JSON 块并打标签
    json_blocks = []
    # 从 JSON 块前面的文本推断 label
    for i, m in enumerate(json_matches):
        # 取该 JSON 块前面那段纯文本的最后非空行作为 label
        label = ""
        if i < len(plain_parts):
            prev_text = plain_parts[i].strip()
            prev_lines = [l.strip() for l in prev_text.splitlines() if l.strip()]
            if prev_lines:
                label = prev_lines[-1]

        raw_json = m.group(1).strip()
        # 应用全面修复（中文引号、** 标记、行尾反斜杠等）
        raw_json = fix_json_block(raw_json)
        # 尝试解析
        try:
            data = json.loads(raw_json)
        except json.JSONDecodeError as e:
            data = {"_parse_error": str(e), "_raw": raw_json[:300]}

        json_blocks.append({"label": label, "data": data})

    return {
        "title": title,
        "synopsis": synopsis,
        "sections": plain_sections,
        "json_blocks": json_blocks,
        "script": script_text,
    }


# ─────────────────────── 渲染人物小传 ───────────────────────────

def render_character(doc: Document, char: Any):
    """渲染单个人物信息"""
    char_dict = char if isinstance(char, dict) else {}
    name = normalize_text(
        safe_get(char_dict, "character_name")
        or safe_get(char_dict, "name")
    )
    if not name:
        name = "角色信息" if not isinstance(char, dict) else "未知角色"
    add_heading(doc, f"👤 {name}", level=2)

    if not isinstance(char, dict):
        fallback_text = normalize_text(char)
        if fallback_text:
            render_field(doc, "人物信息", fallback_text)
        return

    simple_fields = {
        "故事角色": "story_role",
        "核心动机": "core_motivation",
        "外在目标": "external_goal",
        "内在需求": "inner_need",
        "深层恐惧": "deep_fear",
        "自我欺骗": "self_deception",
        "戏剧价值": "dramatic_value",
        "背景": "background",
        "动机": "motivation",
    }
    for label, key in simple_fields.items():
        render_field(doc, label, safe_get(char_dict, key))

    # 性格
    personality = safe_get(char_dict, "personality")
    if personality not in (None, "", [], {}):
        if isinstance(personality, dict):
            add_para(doc, "性格特征", bold=True, size=10, indent=0.5)
            render_field(doc, "特质", safe_get(personality, "traits"))
            render_field(doc, "外在印象", safe_get(personality, "surface_impression"))
            render_field(doc, "内在矛盾", safe_get(personality, "inner_contradiction"))
            for key, value in personality.items():
                if key in {"traits", "surface_impression", "inner_contradiction"}:
                    continue
                render_field(doc, key, value)
        else:
            render_field(doc, "性格特征", personality)

    # 家庭背景
    family = safe_get(char_dict, "family")
    if family not in (None, "", [], {}):
        if isinstance(family, dict):
            add_para(doc, "家庭背景", bold=True, size=10, indent=0.5)
            render_field(doc, "家世", safe_get(family, "family_background"))
            render_field(doc, "成长经历", safe_get(family, "upbringing"))
            render_field(doc, "关键影响", safe_get(family, "key_family_influence"))
            for key, value in family.items():
                if key in {"family_background", "upbringing", "key_family_influence"}:
                    continue
                render_field(doc, key, value)
        else:
            render_field(doc, "家庭背景", family)

    # 外貌
    appearance = safe_get(char_dict, "appearance")
    if appearance not in (None, "", [], {}):
        if isinstance(appearance, dict):
            add_para(doc, "外貌描述", bold=True, size=10, indent=0.5)
            render_field(doc, "整体形象", safe_get(appearance, "overall_look"))
            render_field(doc, "外貌效果", safe_get(appearance, "external_impression_effect"))
            render_field(doc, "标志性特征", safe_get(appearance, "recognizable_features"))
            for key, value in appearance.items():
                if key in {"overall_look", "external_impression_effect", "recognizable_features"}:
                    continue
                render_field(doc, key, value)
        else:
            render_field(doc, "外貌描述", appearance)

    # 行为模式
    behavior = safe_get(char_dict, "behavior")
    if behavior not in (None, "", [], {}):
        if isinstance(behavior, dict):
            add_para(doc, "行为模式", bold=True, size=10, indent=0.5)
            render_field(doc, "情感反应", safe_get(behavior, "emotional_response_pattern"))
            render_field(doc, "社交风格", safe_get(behavior, "social_interaction_style"))
            for key, value in behavior.items():
                if key in {"emotional_response_pattern", "social_interaction_style"}:
                    continue
                render_field(doc, key, value)
        else:
            render_field(doc, "行为模式", behavior)

    for label, key in (
        ("人物关系", "relationships"),
        ("说话风格", "speech_profile"),
        ("决策逻辑", "decision_logic"),
        ("关系模式", "relation_modes"),
        ("可演证据", "actable_evidence"),
        ("成长弧线", "growth_arc"),
        ("剧情功能", "plot_function"),
    ):
        render_field(doc, label, safe_get(char_dict, key))


# ─────────────────────── 渲染核心场景 ───────────────────────────

def render_scene(doc: Document, scene: Any):
    """渲染单个场景信息"""
    scene_dict = scene if isinstance(scene, dict) else {}
    name = normalize_text(
        safe_get(scene_dict, "scene_name")
        or safe_get(scene_dict, "name")
    ) or ("场景说明" if not isinstance(scene, dict) else "未知场景")
    s_type = normalize_text(safe_get(scene_dict, "scene_type"))
    add_heading(doc, f"🎬 {name}", level=2)
    if s_type:
        add_para(doc, s_type, italic=True, size=10, color="888888")

    if not isinstance(scene, dict):
        render_field(doc, "场景说明", scene)
        return

    field_map = [
        ("story_function", "场景功能"),
        ("environment_description", "环境描述"),
        ("atmosphere_description", "氛围描述"),
        ("character_interaction_effect", "人物互动效果"),
        ("worldview_support", "世界观支撑"),
    ]
    for key, label in field_map:
        render_field(doc, label, safe_get(scene_dict, key))

    render_field(doc, "视觉元素", safe_get(scene_dict, "visual_elements"))
    render_field(doc, "冲突潜力", safe_get(scene_dict, "conflict_potential"))


# ─────────────────────── 渲染剧本正文 ───────────────────────────

def render_script(doc: Document, script_text: Any):
    """渲染剧本正文"""
    script_body = _extract_script_text(script_text)
    if not script_body.strip():
        return

    heading_added = False
    stripped_body = script_body.lstrip()
    if stripped_body.startswith("剧本正文"):
        normalized_lines = stripped_body.splitlines()
        if normalized_lines and normalized_lines[0].strip() == "剧本正文":
            script_body = "\n".join(normalized_lines[1:]).lstrip()
    add_heading(doc, "剧本正文", level=1)
    heading_added = True

    lines = script_body.splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # 集标题（支持 Markdown 标题前缀）
        if re.match(r'^[#>\-\s]*第[0-9０-９一二三四五六七八九十百千万两零〇]+集[:：]', stripped):
            clean_heading = re.sub(r'^[#>\-\s]+', '', stripped).strip()
            add_heading(doc, clean_heading, level=2)

        # 场景标记（场景X-Y：...）
        elif re.match(r'^场景\d+', stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            set_run_color(run, "2E75B6")

        # 舞台指示（▲ 开头）
        elif stripped.startswith('▲'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(stripped)
            run.italic = True
            run.font.size = Pt(10)
            set_run_color(run, "666666")

        # 系统/角色对白（含冒号的台词行）
        elif re.match(r'^[\w\u4e00-\u9fff]+（', stripped) or \
             re.match(r'^[\w\u4e00-\u9fff]+：', stripped):
            # 分离角色名与台词
            m = re.match(r'^([\w\u4e00-\u9fff]+(?:（[^）]*）)?)[：:](.+)', stripped)
            if m:
                p = doc.add_paragraph()
                name_run = p.add_run(m.group(1) + "：")
                name_run.bold = True
                name_run.font.size = Pt(11)
                set_run_color(name_run, "1F4E79")
                dialogue_run = p.add_run(m.group(2).strip())
                dialogue_run.font.size = Pt(11)
            else:
                add_para(doc, stripped)

        # 剧本正文标题本身
        elif stripped == "剧本正文":
            if not heading_added:
                add_heading(doc, stripped, level=1)
                heading_added = True

        # 普通段落
        else:
            add_para(doc, stripped)


# ─────────────────────── 主转换函数 ───────────────────────────

# TXT_TO_DOCX_PLAIN_EXPORT_ALIGNED_V2
# TXT_TO_DOCX_FIXED_ONLY_HEADINGS_V3
def convert(input_path: str, output_path: str):
    """
    Plain TXT -> DOCX converter for framework-to-script export.

    Heading policy:
    - The first non-empty line is the document title.
    - Only the four fixed export sections become headings:
      1) story synopsis
      2) character biographies
      3) core scenes
      4) script body
    - Episode titles, scene numbers, dialogue, action lines and numbered lists stay as body text.
    """
    raw_text = Path(input_path).read_text(encoding="utf-8", errors="replace")
    text = str(raw_text or "").replace("\r\n", "\n").replace("\r", "\n")

    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    font_name = "Microsoft YaHei"

    def apply_font(run, size=11, bold=False):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bool(bold)
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        except Exception:
            pass

    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_name
        normal.font.size = Pt(11)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass

    fixed_headings = {
        "\u4e00\u3001\u6545\u4e8b\u6897\u6982",
        "\u4e8c\u3001\u4eba\u7269\u5c0f\u4f20",
        "\u4e09\u3001\u6838\u5fc3\u573a\u666f",
        "\u56db\u3001\u5267\u672c\u6b63\u6587",
    }

    def normalize_heading(value):
        return str(value or "").strip().lstrip("#> -").strip()

    def is_fixed_export_heading(value):
        return normalize_heading(value) in fixed_headings

    def add_title(value):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run(str(value or "").strip())
        apply_font(run, size=18, bold=True)

    def add_heading_plain(value):
        p = doc.add_heading(normalize_heading(value), level=1)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            apply_font(run, size=14, bold=True)

    def add_paragraph_plain(value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(str(value or ""))
        apply_font(run, size=11)

    lines = text.split("\n")
    first_content_index = None
    for idx, line in enumerate(lines):
        if line.strip():
            first_content_index = idx
            break

    if first_content_index is None:
        add_paragraph_plain("")
        doc.save(output_path)
        return

    for idx, line in enumerate(lines):
        stripped = line.strip()

        if idx == first_content_index:
            add_title(stripped)
            continue

        if not stripped:
            doc.add_paragraph()
            continue

        if is_fixed_export_heading(stripped):
            add_heading_plain(stripped)
        else:
            add_paragraph_plain(line.rstrip())

    doc.save(output_path)


SCRIPT_EPISODE_HEADER = re.compile(
    r"^\s*(?:#{1,6}\s*)?(?:《[^》\r\n]+》\s*[·\-—]?\s*)?"
    r"第\s*([0-9０-９一二三四五六七八九十百千万两零〇]+)\s*集"
    r"(?:\s*[：:｜|·\-—]?\s*(《[^》\r\n]+》|[^\r\n]+))?\s*$"
)
SCRIPT_SCENE_HEADER = re.compile(
    r"^\s*(?:场景\s*)?(\d{1,3})(?:\s*[-—]\s*(\d{1,3}))?"
    r"\s*[：:]\s*(.+?)\s*$"
)
SCRIPT_NUMBERED_SCENE_HEADER = re.compile(
    r"^\s*(\d{1,3})\s*[-—]\s*(\d{1,3})\s+(.+?)\s*$"
)
SCRIPT_PEOPLE_LINE = re.compile(r"^\s*人物\s*[：:]\s*(.+?)\s*$")
SCRIPT_DIALOGUE_LINE = re.compile(
    r"^\s*([\w\u3400-\u9fff·]{1,24}(?:OS)?)\s*[：:]\s*(.+?)\s*$",
    re.IGNORECASE,
)


def _set_script_run_font(run, *, size: float = 12, bold: bool = False, italic: bool = False):
    run.font.name = "SimSun"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")


def _clean_script_line(value: str) -> str:
    text = str(value or "").strip()
    text = re.sub(r"^\s*#{1,6}\s*", "", text)
    text = re.sub(r"^\s*[-*]\s+", "", text)
    text = text.replace("**", "")
    return text.strip()


def _episode_number(value: str) -> int | None:
    token = str(value or "").strip().translate(
        str.maketrans("０１２３４５６７８９", "0123456789")
    )
    if token.isdigit():
        number = int(token)
        return number if number > 0 else None

    digits = {
        "零": 0,
        "〇": 0,
        "一": 1,
        "二": 2,
        "两": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
    }
    units = {"十": 10, "百": 100, "千": 1000, "万": 10000}
    if not token or any(char not in digits and char not in units for char in token):
        return None
    total = 0
    section = 0
    current = 0
    for char in token:
        if char in digits:
            current = digits[char]
            continue
        unit = units[char]
        if unit == 10000:
            section = (section + current) * unit
            total += section
            section = 0
            current = 0
        else:
            section += (current or 1) * unit
            current = 0
    number = total + section + current
    return number if number > 0 else None


def convert_script_team(input_path: str, output_path: str, *, title: str = ""):
    """Convert the NPC team's final script into a navigable production-script DOCX."""
    raw_text = Path(input_path).read_text(encoding="utf-8", errors="replace")
    lines = str(raw_text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    document = Document()
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(12)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        style = document.styles[style_name]
        style.font.name = "SimSun"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
        style.paragraph_format.keep_with_next = True

    script_body_index = next(
        (
            index
            for index, line in enumerate(lines)
            if _clean_script_line(line) in {"剧本正文", "四、剧本正文"}
        ),
        None,
    )
    script_start_index = script_body_index + 1 if script_body_index is not None else 0
    first_episode_index = next(
        (
            index
            for index, line in enumerate(lines[script_start_index:], start=script_start_index)
            if SCRIPT_EPISODE_HEADER.match(_clean_script_line(line))
        ),
        None,
    )
    preserve_unrecognized_script = first_episode_index is None
    overview_end_index = (
        script_body_index
        if script_body_index is not None
        else first_episode_index
    )
    source_title = ""
    if overview_end_index is not None:
        for line in lines[:overview_end_index]:
            candidate = _clean_script_line(line)
            title_match = re.fullmatch(r"《[^》\r\n]+》", candidate)
            labeled_title = re.match(r"^作品名称\s*[：:]\s*(《[^》\r\n]+》)", candidate)
            if title_match:
                source_title = candidate
                break
            if labeled_title:
                source_title = labeled_title.group(1)
                break
    display_title = _clean_script_line(title) or source_title or "完整剧本"
    if not (display_title.startswith("《") and display_title.endswith("》")):
        display_title = f"《{display_title.strip('《》')}》"

    title_paragraph = document.add_paragraph()
    title_paragraph.style = document.styles["Title"]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(12)
    title_paragraph.paragraph_format.space_after = Pt(22)
    title_run = title_paragraph.add_run(display_title)
    _set_script_run_font(title_run, size=20, bold=True)

    has_overview = False
    if overview_end_index is not None:
        for raw_line in lines[:overview_end_index]:
            stripped = _clean_script_line(raw_line)
            if (
                not stripped
                or stripped in {"剧本正文", "四、剧本正文"}
                or re.fullmatch(r"[-*_]{3,}", stripped)
                or stripped == display_title
            ):
                continue
            heading_match = re.match(r"^\s*(#{1,6})\s*(.+?)\s*$", raw_line)
            if heading_match:
                heading_text = _clean_script_line(heading_match.group(2))
                if heading_text in {"剧本正文", "四、剧本正文"}:
                    continue
                heading_level = min(3, len(heading_match.group(1)))
                paragraph = document.add_paragraph(style=f"Heading {heading_level}")
                paragraph.paragraph_format.space_before = Pt(8)
                paragraph.paragraph_format.space_after = Pt(6)
                run = paragraph.add_run(heading_text)
                _set_script_run_font(
                    run,
                    size={1: 16, 2: 13, 3: 12}[heading_level],
                    bold=True,
                )
                has_overview = True
                continue
            is_list_item = bool(re.match(r"^\s*[-+*]\s+", raw_line))
            paragraph = document.add_paragraph(style="List Bullet" if is_list_item else None)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(stripped)
            _set_script_run_font(run, size=12)
            has_overview = True

    episode_number = 0
    scene_number = 0
    content_started = preserve_unrecognized_script
    previous_blank = False
    for raw_line in lines[script_start_index:]:
        stripped = _clean_script_line(raw_line)
        if not stripped:
            if content_started and not previous_blank:
                spacer = document.add_paragraph()
                spacer.paragraph_format.space_after = Pt(2)
            previous_blank = True
            continue
        if re.fullmatch(r"[-*_]{3,}", stripped):
            continue
        previous_blank = False

        episode_match = SCRIPT_EPISODE_HEADER.match(stripped)
        if episode_match:
            if content_started or has_overview:
                document.add_page_break()
            content_started = True
            parsed_episode_number = _episode_number(episode_match.group(1))
            if parsed_episode_number is None:
                parsed_episode_number = episode_number + 1 if episode_number else 1
            episode_number = parsed_episode_number
            scene_number = 0
            episode_heading = document.add_paragraph(style="Heading 1")
            episode_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            episode_heading.paragraph_format.space_before = Pt(0)
            episode_heading.paragraph_format.space_after = Pt(8)
            episode_run = episode_heading.add_run(f"第 {episode_number} 集")
            _set_script_run_font(episode_run, size=16, bold=True)
            episode_title = _clean_script_line(
                episode_match.group(2) or ""
            )
            if episode_title:
                subtitle = document.add_paragraph()
                subtitle.paragraph_format.space_after = Pt(14)
                subtitle_run = subtitle.add_run(episode_title)
                _set_script_run_font(subtitle_run, size=12, bold=True)
                set_run_color(subtitle_run, "4B5E66")
            continue

        if not content_started:
            continue

        if preserve_unrecognized_script and episode_number == 0 and re.search(
            r"(?:第\s*.+?\s*集|\bepisode\s+\w+\b)",
            stripped,
            flags=re.IGNORECASE,
        ):
            fallback_heading = document.add_paragraph(style="Heading 1")
            fallback_run = fallback_heading.add_run(stripped)
            _set_script_run_font(fallback_run, size=16, bold=True)
            continue

        numbered_scene = SCRIPT_NUMBERED_SCENE_HEADER.match(stripped)
        scene_match = SCRIPT_SCENE_HEADER.match(stripped)
        if numbered_scene:
            scene_number = int(numbered_scene.group(2))
            scene_text = numbered_scene.group(3).strip()
        elif scene_match:
            scene_number = int(scene_match.group(2) or scene_match.group(1))
            scene_text = scene_match.group(3).strip()
        else:
            scene_text = ""
        if scene_text:
            scene_heading = document.add_paragraph(style="Heading 2")
            scene_heading.paragraph_format.space_before = Pt(8)
            scene_heading.paragraph_format.space_after = Pt(8)
            scene_run = scene_heading.add_run(
                f"{episode_number}-{scene_number} {scene_text}"
            )
            _set_script_run_font(scene_run, size=13, bold=True)
            continue

        people_match = SCRIPT_PEOPLE_LINE.match(stripped)
        if people_match:
            people_paragraph = document.add_paragraph()
            people_paragraph.paragraph_format.keep_with_next = True
            people_paragraph.paragraph_format.space_after = Pt(10)
            people_run = people_paragraph.add_run(f"人物：{people_match.group(1).strip()}")
            _set_script_run_font(people_run, size=12, bold=True)
            continue

        dialogue_match = SCRIPT_DIALOGUE_LINE.match(stripped)
        if dialogue_match and not stripped.startswith(("场景", "人物")):
            dialogue_paragraph = document.add_paragraph()
            dialogue_paragraph.paragraph_format.left_indent = Cm(0.2)
            dialogue_paragraph.paragraph_format.space_before = Pt(3)
            dialogue_paragraph.paragraph_format.space_after = Pt(6)
            dialogue_run = dialogue_paragraph.add_run(
                f"{dialogue_match.group(1)}：{dialogue_match.group(2).strip()}"
            )
            _set_script_run_font(dialogue_run, size=12, bold=True)
            continue

        action_text = re.sub(r"^[▲△]\s*", "", stripped).strip()
        if action_text:
            action_paragraph = document.add_paragraph()
            action_paragraph.paragraph_format.left_indent = Cm(0.2)
            action_paragraph.paragraph_format.space_after = Pt(6)
            action_run = action_paragraph.add_run(f"△ {action_text}")
            _set_script_run_font(action_run, size=12)

    document.core_properties.title = display_title.strip("《》")
    document.core_properties.subject = "剧本正文"
    document.save(output_path)


if __name__ == "__main__":
    convert(r'C:\Users\Administrator\Desktop\txt_script\txt\西幻：葬爱武神.txt',
            r'C:\Users\Administrator\Desktop\西幻：葬爱武神.docx')
    # if len(sys.argv) < 3:
    #     print("用法: python txt_to_docx.py <input.txt> <output.docx>")
    #     sys.exit(1)
    # convert(sys.argv[1], sys.argv[2])
