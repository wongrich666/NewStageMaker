from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def build_character_reskin_docx(
    *,
    output_path: str | Path,
    title: str,
    profile_json: Any,
    character_profile: Any,
    script_batches: Any,
    final_output_text: str,
    core_scenes: Any = None,
) -> Path:
    doc = Document()
    _configure_doc(doc)

    _add_title(doc, str(title or "只换人设剧本").strip() or "只换人设剧本")

    for block in _ordered_output_blocks(
        profile_json=profile_json,
        character_profile=character_profile,
        script_batches=script_batches,
        final_output_text=final_output_text,
    ):
        _add_text_block(doc, block)

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "微软雅黑"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.styles["Normal"].font.size = Pt(11)


def _add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def _add_text_block(doc: Document, text: str) -> None:
    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        return
    if len(doc.paragraphs) > 1:
        doc.add_paragraph()
    for line in normalized.split("\n"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(line.strip())
        run.font.size = Pt(11)


def _ordered_output_blocks(
    *,
    profile_json: Any,
    character_profile: Any,
    script_batches: Any,
    final_output_text: str,
) -> list[str]:
    blocks: list[str] = []

    profile_text = _stringify_for_export(character_profile).strip()
    if profile_text:
        blocks.append(profile_text)
    else:
        profile_json_text = _stringify_for_export(profile_json).strip()
        if profile_json_text:
            blocks.append(profile_json_text)

    batch_texts = _script_batch_texts(script_batches)
    if batch_texts:
        blocks.extend(batch_texts)
    else:
        final_text = str(final_output_text or "").strip()
        if final_text:
            blocks.append(final_text)

    return blocks


def _script_batch_texts(script_batches: Any) -> list[str]:
    if isinstance(script_batches, list):
        return [
            text
            for text in (_extract_script_text(item).strip() for item in script_batches)
            if text
        ]
    if isinstance(script_batches, dict):
        return [
            text
            for text in (_extract_script_text(script_batches[key]).strip() for key in _sorted_batch_keys(script_batches))
            if text
        ]
    return []


def _sorted_batch_keys(value: dict[Any, Any]) -> list[Any]:
    def sort_key(key: Any) -> tuple[int, Any]:
        text = str(key)
        return (0, int(text)) if text.isdigit() else (1, text)

    return sorted(value, key=sort_key)


def _extract_script_text(value: Any) -> str:
    parsed = _coerce_json(value)
    if isinstance(parsed, str):
        return parsed
    if isinstance(parsed, dict):
        for key in (
            "text",
            "script_text",
            "scriptText",
            "batchScriptText",
            "batch_script_text",
            "final_output_text",
            "final_script",
            "script_body",
            "scriptBody",
            "juben_zhengwen",
        ):
            item = parsed.get(key)
            if item not in (None, "", [], {}):
                return _stringify_for_export(item)
        return ""
    return _stringify_for_export(parsed)


def _coerce_json(value: Any) -> Any:
    if not isinstance(value, str):
        return value
    text = value.strip()
    if not text:
        return ""
    try:
        return json.loads(text)
    except Exception:
        return value


def _stringify_for_export(value: Any) -> str:
    value = _coerce_json(value)
    if isinstance(value, str):
        return value
    try:
        return json.dumps(value, ensure_ascii=False, indent=2)
    except Exception:
        return str(value)
