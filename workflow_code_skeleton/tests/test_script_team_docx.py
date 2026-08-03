from __future__ import annotations

from pathlib import Path

from docx import Document

from workflow_code_skeleton.app.utils.txt_to_docx import convert_script_team


def test_script_team_docx_uses_episode_scene_and_script_styles(tmp_path: Path) -> None:
    source = tmp_path / "script.txt"
    output = tmp_path / "script.docx"
    source.write_text(
        "\n".join(
            [
                "《测试剧》",
                "",
                "第1集：《门外的人》",
                "场景1：出租屋｜夜｜内",
                "人物：艾拉、谢敛",
                "艾拉攥紧账单，门外传来三声敲门。",
                "艾拉：（压低声音，眉心微蹙）谁？",
                "艾拉OS：他不该知道这里。",
                "",
                "场景2：楼道｜夜｜内",
                "人物：谢敛",
                "▲谢敛把钥匙插进锁孔。",
                "",
                "第2集：《钥匙》",
                "场景1：楼道｜夜｜内",
                "人物：艾拉、谢敛",
                "谢敛：开门。",
            ]
        ),
        encoding="utf-8",
    )

    convert_script_team(str(source), str(output), title="测试剧")

    document = Document(output)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    by_text = {paragraph.text: paragraph for paragraph in paragraphs}

    assert "《测试剧》" in by_text
    assert by_text["第 1 集"].style.name == "Heading 1"
    assert by_text["1-1 出租屋｜夜｜内"].style.name == "Heading 2"
    assert by_text["1-2 楼道｜夜｜内"].style.name == "Heading 2"
    assert by_text["人物：艾拉、谢敛"].runs[0].bold is True
    assert by_text["△ 艾拉攥紧账单，门外传来三声敲门。"].runs[0].bold is not True
    assert by_text["艾拉：（压低声音，眉心微蹙）谁？"].runs[0].bold is True
    assert by_text["艾拉OS：他不该知道这里。"].runs[0].bold is True


def test_script_team_docx_accepts_markdown_combined_title_and_episode(tmp_path: Path) -> None:
    source = tmp_path / "markdown-script.txt"
    output = tmp_path / "markdown-script.docx"
    source.write_text(
        "\n".join(
            [
                "# 《暗夜吸血鬼不恋爱脑》第1集：《银刃》",
                "",
                "## 场景1：瓦伦丁集团顶层会议室｜夜｜内",
                "",
                "人物：莱奥诺尔、马尔科",
                "",
                "莱奥诺尔翻开董事会文件。",
                "",
                "**莱奥诺尔**：（未抬眼）背叛者不自裁。",
                "",
                "---",
                "",
                "## 场景2：集团地下车库｜夜｜内",
                "",
                "人物：莱奥诺尔",
                "",
                "▲莱奥诺尔停在银色轿车前。",
            ]
        ),
        encoding="utf-8",
    )

    convert_script_team(str(source), str(output), title="暗夜吸血鬼不恋爱脑")

    document = Document(output)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    by_text = {paragraph.text: paragraph for paragraph in paragraphs}

    assert len(paragraphs) > 5
    assert by_text["第 1 集"].style.name == "Heading 1"
    assert "《银刃》" in by_text
    assert by_text["1-1 瓦伦丁集团顶层会议室｜夜｜内"].style.name == "Heading 2"
    assert by_text["1-2 集团地下车库｜夜｜内"].style.name == "Heading 2"
    assert by_text["莱奥诺尔：（未抬眼）背叛者不自裁。"].runs[0].bold is True
    assert "---" not in by_text


def test_script_team_docx_accepts_episode_title_without_colon(tmp_path: Path) -> None:
    source = tmp_path / "episode-without-colon.txt"
    output = tmp_path / "episode-without-colon.docx"
    source.write_text(
        "\n".join(
            [
                "# 剧本正文",
                "第1集《从天而降的“锅”》",
                "",
                "场景1：会议室｜日｜内",
                "人物：打工鱼、波士鱼",
                "波士鱼：从今天起，你就是组长。",
                "",
                "第2集《半包纸巾》",
                "场景1：办公室｜日｜内",
                "人物：打工鱼",
                "打工鱼OS：这班还能上吗？",
            ]
        ),
        encoding="utf-8",
    )

    convert_script_team(str(source), str(output), title="今天也在努力打工呢")

    document = Document(output)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    by_text = {paragraph.text: paragraph for paragraph in paragraphs}

    assert by_text["第 1 集"].style.name == "Heading 1"
    assert "《从天而降的“锅”》" in by_text
    assert by_text["1-1 会议室｜日｜内"].style.name == "Heading 2"
    assert by_text["波士鱼：从今天起，你就是组长。"].runs[0].bold is True
    assert by_text["第 2 集"].style.name == "Heading 1"
    assert by_text["2-1 办公室｜日｜内"].style.name == "Heading 2"


def test_script_team_docx_accepts_pipe_separated_episode_title(tmp_path: Path) -> None:
    source = tmp_path / "pipe-separated-script.txt"
    output = tmp_path / "pipe-separated-script.docx"
    source.write_text(
        "\n".join(
            [
                "# 剧本正文",
                "# 第1集｜《从天而降的“锅”》",
                "",
                "**场景1：会议室｜日｜内**",
                "",
                "**人物：** 波士鱼、打工鱼",
                "",
                "△ 波士鱼的指挥棒敲在会议桌上。",
                "",
                "波士鱼：降本增效，谁有想法？",
            ]
        ),
        encoding="utf-8",
    )

    convert_script_team(str(source), str(output), title="今天也在努力打工呢")

    document = Document(output)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    by_text = {paragraph.text: paragraph for paragraph in paragraphs}

    assert by_text["第 1 集"].style.name == "Heading 1"
    assert "《从天而降的“锅”》" in by_text
    assert by_text["1-1 会议室｜日｜内"].style.name == "Heading 2"
    assert "△ 波士鱼的指挥棒敲在会议桌上。" in by_text
    assert by_text["波士鱼：降本增效，谁有想法？"].runs[0].bold is True


def test_script_team_docx_accepts_chinese_episode_number(tmp_path: Path) -> None:
    source = tmp_path / "chinese-episode-number.txt"
    output = tmp_path / "chinese-episode-number.docx"
    source.write_text(
        "\n".join(
            [
                "# 剧本正文",
                "第十二集·《旧账》",
                "12-1 古堡大厅｜夜｜内",
                "人物：伊莎贝拉、塞缪尔",
                "伊莎贝拉：你终于肯说真话了。",
            ]
        ),
        encoding="utf-8",
    )

    convert_script_team(str(source), str(output), title="旧账")

    document = Document(output)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    by_text = {paragraph.text: paragraph for paragraph in paragraphs}

    assert by_text["第 12 集"].style.name == "Heading 1"
    assert by_text["12-1 古堡大厅｜夜｜内"].style.name == "Heading 2"
    assert by_text["伊莎贝拉：你终于肯说真话了。"].runs[0].bold is True


def test_script_team_docx_preserves_body_when_episode_heading_is_unknown(tmp_path: Path) -> None:
    source = tmp_path / "unknown-heading-script.txt"
    output = tmp_path / "unknown-heading-script.docx"
    source.write_text(
        "\n".join(
            [
                "# 剧本正文",
                "EPISODE ONE / THE DOOR",
                "INT. HALLWAY - NIGHT",
                "MIRA: Do not open it.",
            ]
        ),
        encoding="utf-8",
    )

    convert_script_team(str(source), str(output), title="The Door")

    document = Document(output)
    body = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "EPISODE ONE / THE DOOR" in body
    assert "INT. HALLWAY - NIGHT" in body
    assert "MIRA：Do not open it." in body


def test_script_team_docx_renders_delivery_overview_before_script(tmp_path: Path) -> None:
    source = tmp_path / "delivery-script.txt"
    output = tmp_path / "delivery-script.docx"
    source.write_text(
        "\n".join(
            [
                "# 剧本大纲",
                "## 基本信息",
                "- **作品名称**：《我有99条命但没人信》",
                "- **总集数**：5集",
                "## 故事背景",
                "天阶城以命格划分人的价值。",
                "## 故事梗概",
                "林烬必须在三天内完成第99次死亡。",
                "第1集：主角完成第98次死亡，第2集继续追查。",
                "## 核心主线",
                "- **主角目标**：活过夺格大典。",
                "## 主要人物",
                "- **林烬——主角**：被全城当成废物的复活者。",
                "# 剧本正文",
                "第1集：《第98次死亡》",
                "场景1：暗巷｜夜｜外",
                "人物：林烬",
                "林烬睁开眼。",
            ]
        ),
        encoding="utf-8",
    )

    convert_script_team(str(source), str(output), title="我有99条命但没人信")

    document = Document(output)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    by_text = {paragraph.text: paragraph for paragraph in paragraphs}

    assert by_text["剧本大纲"].style.name == "Heading 1"
    assert by_text["故事背景"].style.name == "Heading 2"
    assert "天阶城以命格划分人的价值。" in by_text
    assert by_text["第 1 集"].style.name == "Heading 1"
    assert paragraphs.index(by_text["主要人物"]) < paragraphs.index(by_text["第 1 集"])
    assert [paragraph.text for paragraph in paragraphs].count("第 1 集") == 1
