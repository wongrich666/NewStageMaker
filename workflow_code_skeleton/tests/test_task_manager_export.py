from __future__ import annotations

import json
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from workflow_code_skeleton.app.services.task_manager import (
    APPEARANCE_NATURAL_LANGUAGE_ARTIFACT,
    EPISODE_PLAN_DISPLAY_ARTIFACT,
    TaskManager,
)
from workflow_code_skeleton.app.workflow_ids import (
    CHARACTER_BIOS_VAR,
    APPEARANCE_NATURAL_LANGUAGE_VAR,
    CHARACTER_NATURAL_LANGUAGE_VAR,
    UNSTRUCTURED_KIND_VAR,
    UNSTRUCTURED_SOURCE_VAR,
)
from workflow_code_skeleton.tests.test_support import WorkspaceTempDir


def _structured_characters_payload() -> str:
    return json.dumps(
        {
            "character_setting": {
                "characters": [
                    {
                        "character_name": "林夏",
                        "story_role": "主角",
                        "personality": "冷静克制",
                        "core_motivation": "查明真相并保住团队",
                        "growth_arc": "学会重新信任同伴",
                        "appearance_anchor": "深色风衣与冷色系妆发",
                    },
                    {
                        "character_name": "顾川",
                        "story_role": "关键对手",
                        "personality": "强势克制",
                        "core_motivation": "掩盖旧案真相",
                        "growth_arc": "在对抗中暴露真实恐惧",
                        "appearance_anchor": "利落西装与冷硬神情",
                    },
                    {
                        "character_name": "周临",
                        "story_role": "关键配角",
                        "personality": "圆滑机敏",
                        "core_motivation": "在两方博弈中求生",
                        "growth_arc": "最终选择站到真相一侧",
                        "appearance_anchor": "总带着未熄灭的烟味",
                    },
                ]
            }
        },
        ensure_ascii=False,
    )


def _appearanceMapping_payload() -> str:
    return json.dumps(
        {
            "appearanceMapping": {
                "characters": [
                    {
                        "character_name": "林夏",
                        "default_name": "林夏【日常】",
                        "same_person_anchor": "深色风衣与冷静眼神",
                        "forbidden_generic_names": ["女主", "配角"],
                        "outfit_variants": [
                            {
                                "alias_name": "林夏【会议室交锋态】",
                                "usage_rule": "高压谈判时使用",
                                "episode_range_hint": "第1集后段",
                            }
                        ],
                    }
                ]
            }
        },
        ensure_ascii=False,
    )


def _snapshot_with_export_artifacts() -> dict[str, object]:
    return {
        "user_id": 1,
        "project_id": 1,
        "task_id": "task-export",
        "status": "completed",
        "title": "长夜回潮",
        "message": "已完成",
        "created_at": "2026-04-29T12:00:00+08:00",
        "updated_at": "2026-04-29T12:00:00+08:00",
        "finished_at": "2026-04-29T12:00:00+08:00",
        "workflow_spec_path": "spec.json",
        "total_episodes": 1,
        "input_payload": {
            "title": "长夜回潮",
            "story_outline": "旧版结构化梗概",
            "total_episodes": 1,
        },
        "model_option": {},
        "completion_confirmed": True,
        "awaiting_user_confirmation": False,
        "cache_retained": False,
        "artifacts": {
            "script_title_content": "长夜回潮",
            "framework_natural_language": "剧本标题：长夜回潮\n故事梗概：故事从主角重返故乡展开，旧案迫使她再次入局。\n主要人物小传：林夏与顾川的关系贯穿始终。\n核心场景说明：旧码头与会议室构成主要舞台。",
            "worldview_natural_language": "故事发生在资源高度紧张的近未来港城，效率优先与身份等级并行。",
            "character_natural_language": "林夏：作为项目负责人，性格冷静克制，核心目标是保住团队，并在真相追索中完成自我修正。\n顾川：表面强势，实际被旧案束缚，是主角的重要对手与镜像。",
            APPEARANCE_NATURAL_LANGUAGE_ARTIFACT: "【角色】林夏\n默认称呼：林夏【日常】\n固定识别锚点：深色风衣与冷色系妆发\n服装版本与使用条件：会议室交锋态用于高压谈判场景\n禁止退回泛称：女主、配角",
            "scene_natural_language": "旧码头负责承载悬念与秘密交易，深夜会议室则集中呈现角色对峙与关系撕裂。",
            EPISODE_PLAN_DISPLAY_ARTIFACT: "第1集《风起》\n林夏回城后被迫接手旧案调查。",
            "final_script": "第1集：风起\n场景1：旧码头\n林夏：先查人，再查船。",
            "final_output_text": "第1集：风起\n场景1：旧码头\n林夏：先查人，再查船。",
        },
        "debug_state": {
            "variables": {
                APPEARANCE_NATURAL_LANGUAGE_VAR: "【角色】林夏\n默认称呼：林夏【日常】",
            },
            "node_outputs": {},
        },
    }


def _counterexample_snapshot() -> dict[str, object]:
    snapshot = _snapshot_with_export_artifacts()
    artifacts = snapshot["artifacts"]
    debug_variables = snapshot["debug_state"]["variables"]
    artifacts["framework_natural_language"] = (
        "剧本标题：长夜回潮\n"
        "故事梗概：林夏重返港城调查沉船旧案，旧同盟与旧敌人同时逼近。\n"
        "主要人物小传：林夏与顾川围绕真相与责任持续对抗。\n"
        "核心场景说明：旧港调度塔与封闭会议室构成主要冲突舞台。"
    )
    artifacts["worldview_natural_language"] = (
        "故事发生在资源高度集中的近未来港城，航运财团与城市治理深度绑定，"
        "每个人都在效率秩序与私人情感之间做选择。"
    )
    artifacts.pop("character_natural_language", None)
    artifacts.pop("character_summary", None)
    debug_variables.pop(CHARACTER_NATURAL_LANGUAGE_VAR, None)
    artifacts["characters"] = json.dumps(
        {
            "character_setting": {
                "characters": [
                    {
                        "character_name": "林夏",
                        "story_role": "主角",
                        "personality": "冷静克制，但在旧案面前会暴露迟疑",
                        "core_motivation": "查清沉船事故真相并保住自己的调查团队",
                        "relationship_to_protagonist": "她自己就是故事的行动轴心，必须带着队友一起推进调查",
                        "growth_arc": "从只信证据到重新学会信任同伴",
                        "appearance_anchor": "【待补全：补充人物定位】",
                    },
                    {
                        "character_name": "顾川",
                        "story_role": "关键对手",
                        "personality": "强势隐忍，总想把局势控制在自己手里",
                        "core_motivation": "掩盖旧案责任，守住自己在财团里的位置",
                        "relationship_to_protagonist": "与林夏互为镜像，也是她追查真相时最大的现实阻力",
                        "growth_arc": "在持续对抗中暴露真实恐惧与软肋",
                        "appearance_anchor": "利落西装与始终紧绷的站姿",
                    },
                ]
            }
        },
        ensure_ascii=False,
    )
    artifacts.pop("scene_natural_language", None)
    artifacts.pop("core_scene_summary", None)
    artifacts["scene_json"] = json.dumps(
        {
            "scene_setting": {
                "scenes": [
                    {
                        "scene_name": "旧港调度塔",
                        "scene_type": "悬疑调查场",
                        "story_function": "暴露沉船旧案的关键线索",
                        "environment_description": "狭窄高塔、玻璃窗外是被雨雾吞没的泊位灯火",
                        "atmosphere_description": "空气紧绷，像任何一句试探都可能触发翻盘",
                        "character_interaction_effect": "林夏与顾川在追问与反问之间不断拉高敌意",
                    }
                ]
            }
        },
        ensure_ascii=False,
    )
    artifacts["appearanceMapping"] = json.dumps(
        {
            "服装版本映射内容": {
                "characters": [
                    {
                        "character_name": "林夏",
                        "default_name": "林夏【夜巡常服】",
                        "outfit_variants": [
                            {
                                "alias_name": "林夏【会议室交锋态】",
                                "usage_rule": "高压对峙场景使用",
                            }
                        ],
                    }
                ]
            },
            "appearanceMapping": {
                "characters": [
                    {
                        "character_name": "林夏",
                        "default_name": "林夏【夜巡常服】",
                        "same_person_anchor": "深色风衣、冷调衬衫与始终不离身的记录本",
                        "forbidden_generic_names": ["女主", "调查员"],
                        "outfit_variants": [
                            {
                                "alias_name": "林夏【会议室交锋态】",
                                "usage_rule": "封闭会议室摊牌时使用",
                                "episode_range_hint": "第1集末段",
                            }
                        ],
                    }
                ]
            },
        },
        ensure_ascii=False,
    )
    artifacts[APPEARANCE_NATURAL_LANGUAGE_ARTIFACT] = (
        "服饰版本自然语言说明：林夏平时穿深色风衣与冷调衬衫，保持克制利落的职业感；"
        "进入封闭会议室摊牌时，会切换成更锋利的交锋态着装，以强化她在高压谈判里的压迫感。"
    )
    debug_variables[APPEARANCE_NATURAL_LANGUAGE_VAR] = artifacts[APPEARANCE_NATURAL_LANGUAGE_ARTIFACT]
    artifacts["final_script"] = (
        "第1集：风起\n"
        "1-1 旧港调度塔 夜\n"
        "林夏：先盯住泊位记录，再查是谁删了监控。\n"
        "顾川：你越往前走，越会发现这件事没有你想得那么干净。"
    )
    artifacts["final_output_text"] = artifacts["final_script"]
    return snapshot


class TaskManagerExportTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = WorkspaceTempDir(prefix="task-manager-export-")
        self.addCleanup(self.temp_dir.cleanup)
        self.manager = TaskManager()
        base_dir = Path(self.temp_dir.name) / "runtime_data"
        self.manager.set_storage_root(base_dir, runtime_archive_dir=Path(self.temp_dir.name) / "runtime_archive")
        self.manager._tasks.clear()
        self.manager._projects.clear()
        self.manager._index = {
            "next_project_id": 2,
            "latest_project_id": 1,
            "latest_project_by_user": {"1": 1},
        }

    def _persist(self, snapshot: dict[str, object]) -> None:
        self.manager._project_path(int(snapshot["project_id"])).write_text(
            json.dumps(snapshot, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def _read_docx_text(self, docx_path: Path) -> str:
        doc = Document(str(docx_path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text).strip()

    def _assert_no_forbidden_visible_content(self, text: str) -> None:
        forbidden_tokens = (
            "【待补全：补充人物定位】",
            "待补全",
            "补充人物定位",
            "待完善",
            "未提供",
            "暂无",
            "待填写",
            "待定",
            "TBD",
            "TODO",
            "None",
            "null",
            "[object Object]",
            "服装版本映射内容",
            "appearanceMapping",
            "scene_json",
            "character_setting",
            "outfit_variants",
            "FastGPT",
        )
        for token in forbidden_tokens:
            self.assertNotIn(token, text)
        self.assertNotIn("{", text)
        self.assertNotIn("}", text)

    def _assert_expected_natural_sections(self, text: str) -> None:
        self.assertIn("故事梗概", text)
        self.assertIn("世界观设定", text)
        self.assertIn("人物小传", text)
        self.assertIn("人物服饰说明", text)
        self.assertIn("核心场景", text)
        self.assertIn("剧本正文", text)

    def test_character_setting_export_block_parses_inline_natural_text(self) -> None:
        snapshot = _snapshot()
        snapshot["artifacts"]["character_summary"] = (
            "林夏：作为项目负责人，性格冷静克制，核心目标是保住团队。\n"
            "顾川（关键对手）：表面强势，实际被旧案束缚。"
        )
        snapshot["artifacts"].pop("characters", None)
        snapshot["artifacts"].pop("character_bios", None)

        block = self.manager._build_character_setting_export_block(snapshot)

        self.assertIsNotNone(block)
        characters = block["character_setting"]["characters"]  # type: ignore[index]
        self.assertEqual([item["character_name"] for item in characters], ["林夏", "顾川"])
        self.assertIn("项目负责人", characters[0]["dramatic_value"])
        self.assertIn("林夏：是故事中的主角", text)
        self.assertIn("顾川：是故事中的关键对手", text)
        self.assertIn("旧港调度塔：悬疑调查场", text)
        self.assertIn("服饰版本自然语言说明", text)
        self.assertIn("资源高度集中的近未来港城", text)
        self.assertIn("第1集：风起", text)

    def test_build_docx_export_source_prefers_natural_language_sections(self) -> None:
        snapshot = _snapshot_with_export_artifacts()

        content = self.manager._build_docx_export_source_text(snapshot)

        self.assertIn("故事梗概\n故事从主角重返故乡展开", content)
        self.assertIn("世界观设定\n故事发生在资源高度紧张的近未来港城", content)
        self.assertIn("人物小传\n林夏：作为项目负责人", content)
        self.assertIn("人物服饰说明\n【角色】林夏", content)
        self.assertIn("核心场景\n旧码头负责承载悬念与秘密交易", content)
        self.assertIn("剧本正文\n第1集：风起", content)
        self.assertNotIn("分集计划", content)
        self.assertNotIn("opening", content)
        self.assertNotIn("character_design_principle", content)
        self.assertNotIn("```json", content)

    def test_build_docx_export_source_falls_back_to_readable_text_without_raw_json(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"].pop("framework_natural_language", None)
        snapshot["artifacts"].pop("character_natural_language", None)
        snapshot["artifacts"].pop("character_summary", None)
        snapshot["artifacts"].pop(APPEARANCE_NATURAL_LANGUAGE_ARTIFACT, None)
        snapshot["artifacts"]["story_outline"] = {
            "opening": "主角被迫回到故乡",
            "inciting_incident": "旧案重启调查",
            "middle_escalation": "关系在追查中持续失衡",
            "final_climax": "码头对决揭开真正的幕后真相",
            "theme": "身份与选择",
        }
        snapshot["artifacts"]["characters"] = json.dumps(
            {
                "character_setting": {
                    "characters": [
                        {
                            "character_name": "林夏",
                            "story_role": "主角",
                            "personality": "冷静克制",
                            "core_motivation": "保住团队",
                            "growth_arc": "学会直面旧创伤",
                        }
                    ]
                }
            },
            ensure_ascii=False,
        )
        snapshot["artifacts"]["appearanceMapping"] = json.dumps(
            {
                "appearanceMapping": {
                    "characters": [
                        {
                            "character_name": "林夏",
                            "default_name": "林夏【日常】",
                            "same_person_anchor": "深色风衣与冷静眼神",
                            "forbidden_generic_names": ["女主", "配角"],
                            "outfit_variants": [
                                {
                                    "alias_name": "林夏【会议室交锋态】",
                                    "usage_rule": "高压谈判时使用",
                                    "episode_range_hint": "第1集后段",
                                }
                            ],
                        }
                    ]
                }
            },
            ensure_ascii=False,
        )

        content = self.manager._build_docx_export_source_text(snapshot)

        self.assertIn("故事从主角被迫回到故乡展开。", content)
        self.assertIn("主角因为旧案重启调查卷入冲突。", content)
        self.assertIn("林夏：是故事中的主角", content)
        self.assertIn("人物服饰说明\n【角色】林夏", content)
        self.assertNotIn("{", content)
        self.assertNotIn("opening", content)
        self.assertNotIn("outfit_variants", content)

    def test_framework_natural_language_trims_trailing_structured_dump_before_export(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"]["framework_natural_language"] = (
            "故事梗概：林夏重返港城调查沉船旧案，旧同盟与旧敌人同时逼近。\n"
            "主要人物小传：林夏与顾川围绕真相持续对抗。\n"
            '{"story_outline":{"opening":"这段 JSON 不应该进入导出","theme":"误导尾巴"}}'
        )

        content = self.manager._build_docx_export_source_text(snapshot)

        self.assertIn("故事梗概\n林夏重返港城调查沉船旧案", content)
        self.assertNotIn("这段 JSON 不应该进入导出", content)
        self.assertNotIn('{"story_outline"', content)
        self.assertNotIn("{", content)

    def test_worldview_natural_language_rejects_placeholder_heavy_candidate_and_falls_back(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"]["worldview_natural_language"] = (
            "世界观概述：待补全。\n"
            "时代背景：待补全。\n"
            "社会规则：待补全。"
        )
        snapshot["artifacts"]["worldview"] = {
            "worldview_summary": "港城资源高度集中，航运财团与城市治理深度绑定。",
            "social_rules": ["效率优先", "身份等级决定话语权"],
            "atmosphere": "冷硬压迫",
        }

        content = self.manager._build_docx_export_source_text(snapshot)

        self.assertIn("世界观设定\n港城资源高度集中", content)
        self.assertIn("这个世界的运行规则主要体现在：效率优先", content)
        self.assertIn("身份等级决定话语权。", content)
        self.assertNotIn("待补全", content)

    def test_character_export_rejects_single_character_placeholder_natural_language(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"]["character_natural_language"] = "林夏：待补全。"
        snapshot["artifacts"]["character_summary"] = "林夏：待补全。"
        snapshot["artifacts"]["characters"] = _structured_characters_payload()

        content = self.manager._build_docx_export_source_text(snapshot)

        self.assertIn("人物小传\n林夏：是故事中的主角", content)
        self.assertIn("顾川：是故事中的关键对手", content)
        self.assertIn("周临：是故事中的关键配角", content)
        self.assertNotIn("待补全", content)

    def test_character_export_rejects_placeholder_heavy_natural_language(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"]["character_natural_language"] = (
            "林夏：待补全。\n顾川：待补全。\n周临：待补全。"
        )
        snapshot["artifacts"]["character_summary"] = snapshot["artifacts"]["character_natural_language"]
        snapshot["artifacts"]["characters"] = _structured_characters_payload()

        content = self.manager._build_docx_export_source_text(snapshot)

        self.assertIn("林夏：是故事中的主角", content)
        self.assertIn("顾川：是故事中的关键对手", content)
        self.assertIn("周临：是故事中的关键配角", content)
        self.assertNotIn("待补全", content)

    def test_character_export_placeholder_warning_is_logged_once_per_signature(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"]["character_natural_language"] = (
            "林夏：待补全。\n顾川：待补全。\n周临：待补全。"
        )
        snapshot["artifacts"]["character_summary"] = snapshot["artifacts"]["character_natural_language"]
        snapshot["artifacts"]["characters"] = _structured_characters_payload()

        with self.assertLogs("task_manager", level="WARNING") as logs:
            first = self.manager._build_docx_export_source_text(snapshot)
            second = self.manager._build_docx_export_source_text(snapshot)

        joined = "\n".join(logs.output)
        self.assertEqual(joined.count("character_export_text_rejected"), 1)
        self.assertIn("林夏：是故事中的主角", first)
        self.assertEqual(first, second)

    def test_character_fallback_skips_placeholder_fragments(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"].pop("character_natural_language", None)
        snapshot["artifacts"].pop("character_summary", None)
        snapshot["artifacts"]["characters"] = json.dumps(
            {
                "character_setting": {
                    "characters": [
                        {
                            "character_name": "林夏",
                            "story_role": "主角",
                            "personality": "待补全",
                            "core_motivation": "待补全",
                            "growth_arc": "学会直面旧创伤",
                            "appearance_anchor": "待补全",
                        }
                    ]
                }
            },
            ensure_ascii=False,
        )

        content = self.manager._build_docx_export_source_text(snapshot)

        self.assertIn("林夏：是故事中的主角", content)
        self.assertIn("人物成长会落在学会直面旧创伤", content)
        self.assertNotIn("待补全", content)

    def test_character_fallback_with_name_only_does_not_emit_placeholder_sentence(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"].pop("character_natural_language", None)
        snapshot["artifacts"].pop("character_summary", None)
        snapshot["artifacts"]["characters"] = json.dumps(
            {
                "character_setting": {
                    "characters": [
                        {
                            "character_name": "林夏",
                        }
                    ]
                }
            },
            ensure_ascii=False,
        )

        content = self.manager._build_docx_export_source_text(snapshot)

        self.assertIn("人物小传\n林夏", content)
        self.assertNotIn("人物设定暂未补充完整", content)
        self.assertNotIn("待补充", content)

    def test_completed_snapshot_does_not_keep_user_visible_appearance_json_artifacts(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"]["appearanceMapping"] = _appearanceMapping_payload()
        snapshot["artifacts"]["character_registry"] = {"林夏": {"default_name": "林夏【日常】"}}
        snapshot["artifacts"]["character_alias_registry"] = {"林夏【会议室交锋态】": "林夏"}
        snapshot["artifacts"]["episode_alias_plan"] = [{"episode": 1, "alias": "林夏【会议室交锋态】"}]

        compacted = self.manager._compact_completed_snapshot(snapshot)
        artifacts = compacted.get("artifacts") or {}

        self.assertIn(APPEARANCE_NATURAL_LANGUAGE_ARTIFACT, artifacts)
        self.assertNotIn("appearanceMapping", artifacts)
        self.assertNotIn("character_registry", artifacts)
        self.assertNotIn("character_alias_registry", artifacts)
        self.assertNotIn("episode_alias_plan", artifacts)

    def test_save_final_script_exports_docx_with_readable_preface_sections(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        self._persist(snapshot)

        docx_path = self.manager.save_final_script(1, user_id=1)

        self.assertTrue(docx_path.exists())
        doc = Document(str(docx_path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("长夜回潮", text)
        self.assertIn("故事梗概", text)
        self.assertIn("世界观设定", text)
        self.assertIn("人物服饰说明", text)
        self.assertIn("剧本正文", text)
        self.assertIn("旧码头负责承载悬念与秘密交易", text)
        self.assertIn("第1集：风起", text)
        self.assertNotIn("分集计划", text)
        self.assertNotIn("character_design_principle", text)
        self.assertNotIn("opening", text)

    def test_save_final_script_hydrates_character_natural_language_before_export(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"].pop("character_natural_language", None)
        snapshot["artifacts"].pop("character_summary", None)
        snapshot["artifacts"]["characters"] = _structured_characters_payload()
        self._persist(snapshot)

        with patch(
            "workflow_code_skeleton.app.orchestrators.fastgpt_hybrid_workflow.run_stage_with_contract_guard",
            return_value={
                CHARACTER_NATURAL_LANGUAGE_VAR: "林夏：作为项目负责人始终冷静克制，她的目标是保住团队并追出旧案真相。\n顾川：他表面强势克制，实际被旧案反噬，是主角的重要对手与镜像。\n周临：这位关键配角长期在夹缝中求生，最终会被迫站到真相一侧。"
            },
        ) as mocked_stage:
            docx_path = self.manager.save_final_script(1, user_id=1)

        self.assertTrue(docx_path.exists())
        mocked_stage.assert_called_once()
        saved_snapshot = self.manager.get_project_snapshot(1, user_id=1, public_view=False)
        self.assertIn("林夏：作为项目负责人始终冷静克制", saved_snapshot["artifacts"]["character_natural_language"])
        self.assertIn(
            "周临：这位关键配角长期在夹缝中求生",
            saved_snapshot["debug_state"]["variables"][CHARACTER_NATURAL_LANGUAGE_VAR],
        )

    def test_build_export_character_naturalize_stage_variables_uses_raw_character_bios_when_structured_missing(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"].pop("characters", None)
        snapshot["artifacts"].pop("character_bios", None)
        snapshot["debug_state"]["variables"].pop(CHARACTER_BIOS_VAR, None)
        snapshot["input_payload"]["character_bios"] = (
            "林夏：项目负责人，冷静克制，目标是保住团队并追出旧案真相。\n"
            "顾川：关键对手，强势隐忍，想守住自己在财团里的位置。"
        )

        variables = self.manager._build_export_character_naturalize_stage_variables(snapshot)

        self.assertEqual(variables[UNSTRUCTURED_KIND_VAR], "generic")
        self.assertIn("林夏：项目负责人", variables[UNSTRUCTURED_SOURCE_VAR])
        self.assertNotIn("待补全", variables[UNSTRUCTURED_SOURCE_VAR])

    def test_save_final_script_persists_structured_character_fallback_when_naturalize_fails(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"].pop("character_natural_language", None)
        snapshot["artifacts"].pop("character_summary", None)
        snapshot["artifacts"]["characters"] = _structured_characters_payload()
        self._persist(snapshot)

        with patch(
            "workflow_code_skeleton.app.orchestrators.fastgpt_hybrid_workflow.run_stage_with_contract_guard",
            side_effect=RuntimeError("upstream unavailable"),
        ):
            docx_path = self.manager.save_final_script(1, user_id=1)

        self.assertTrue(docx_path.exists())
        saved_snapshot = self.manager.get_project_snapshot(1, user_id=1, public_view=False)
        persisted = str(saved_snapshot["artifacts"]["character_natural_language"] or "")
        self.assertIn("林夏：是故事中的主角", persisted)
        self.assertIn("顾川：是故事中的关键对手", persisted)
        self.assertNotIn("待补全", persisted)

        txt_text = docx_path.with_suffix(".txt").read_text(encoding="utf-8")
        docx_text = self._read_docx_text(docx_path)
        for text in (txt_text, docx_text):
            self.assertIn("林夏：是故事中的主角", text)
            self.assertIn("顾川：是故事中的关键对手", text)
            self.assertNotIn("【待补全：补充人物定位】", text)

    def test_save_final_script_skips_character_hydration_when_existing_text_available(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        self._persist(snapshot)

        with patch(
            "workflow_code_skeleton.app.orchestrators.fastgpt_hybrid_workflow.run_stage_with_contract_guard",
        ) as mocked_stage:
            docx_path = self.manager.save_final_script(1, user_id=1)

        self.assertTrue(docx_path.exists())
        mocked_stage.assert_not_called()

    def test_save_final_script_hydrates_appearance_natural_language_before_export(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"].pop(APPEARANCE_NATURAL_LANGUAGE_ARTIFACT, None)
        snapshot["debug_state"]["variables"].pop(APPEARANCE_NATURAL_LANGUAGE_VAR, None)
        snapshot["artifacts"]["appearanceMapping"] = _appearanceMapping_payload()
        self._persist(snapshot)

        with patch(
            "workflow_code_skeleton.app.orchestrators.fastgpt_hybrid_workflow.run_stage_with_contract_guard",
            return_value={
                APPEARANCE_NATURAL_LANGUAGE_VAR: "【角色】林夏\n默认称呼：林夏【日常】\n固定识别锚点：深色风衣与冷静眼神"
            },
        ) as mocked_stage:
            docx_path = self.manager.save_final_script(1, user_id=1)

        self.assertTrue(docx_path.exists())
        mocked_stage.assert_called_once()
        saved_snapshot = self.manager.get_project_snapshot(1, user_id=1, public_view=False)
        appearance_text = saved_snapshot["artifacts"][APPEARANCE_NATURAL_LANGUAGE_ARTIFACT]
        self.assertIn("【角色】林夏", appearance_text)
        self.assertIn("默认称呼：林夏【日常】", appearance_text)
        self.assertIn("固定识别锚点：深色风衣与冷静眼神", appearance_text)
        self.assertIn(
            "固定识别锚点：深色风衣与冷静眼神",
            saved_snapshot["debug_state"]["variables"][APPEARANCE_NATURAL_LANGUAGE_VAR],
        )

    def test_save_final_script_falls_back_when_appearance_natural_language_hydration_fails(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        snapshot["artifacts"].pop(APPEARANCE_NATURAL_LANGUAGE_ARTIFACT, None)
        snapshot["debug_state"]["variables"].pop(APPEARANCE_NATURAL_LANGUAGE_VAR, None)
        snapshot["artifacts"]["appearanceMapping"] = _appearanceMapping_payload()
        self._persist(snapshot)

        with patch(
            "workflow_code_skeleton.app.orchestrators.fastgpt_hybrid_workflow.run_stage_with_contract_guard",
            side_effect=RuntimeError("upstream unavailable"),
        ):
            docx_path = self.manager.save_final_script(1, user_id=1)

        self.assertTrue(docx_path.exists())
        doc = Document(str(docx_path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("人物服饰说明", text)
        self.assertIn("默认称呼：林夏【日常】", text)
        saved_snapshot = self.manager.get_project_snapshot(1, user_id=1, public_view=False)
        self.assertNotIn(APPEARANCE_NATURAL_LANGUAGE_ARTIFACT, saved_snapshot["artifacts"])

    def test_save_final_script_skips_appearance_hydration_when_existing_text_available(self) -> None:
        snapshot = _snapshot_with_export_artifacts()
        self._persist(snapshot)

        with patch(
            "workflow_code_skeleton.app.orchestrators.fastgpt_hybrid_workflow.run_stage_with_contract_guard",
        ) as mocked_stage:
            docx_path = self.manager.save_final_script(1, user_id=1)

        self.assertTrue(docx_path.exists())
        mocked_stage.assert_not_called()

    def test_counterexample_fixture_round_trips_txt_docx_and_public_snapshot_without_structured_leaks(self) -> None:
        snapshot = _counterexample_snapshot()
        self._persist(snapshot)

        with patch(
            "workflow_code_skeleton.app.services.runtime_export_store.use_fastgpt_backend",
            return_value=False,
        ):
            docx_path = self.manager.save_final_script(1, user_id=1)

        self.assertTrue(docx_path.exists())
        txt_path = docx_path.with_suffix(".txt")
        self.assertTrue(txt_path.exists())

        txt_text = txt_path.read_text(encoding="utf-8")
        docx_text = self._read_docx_text(docx_path)
        public_snapshot = self.manager.get_project_snapshot(1, user_id=1, public_view=True) or {}
        public_visible_text = "\n".join(
            [
                str(public_snapshot.get("display_stage_output") or ""),
                str(public_snapshot.get("message") or ""),
                "\n".join(str(value or "") for value in (public_snapshot.get("artifacts") or {}).values()),
            ]
        )

        for text in (txt_text, docx_text, public_visible_text):
            self._assert_no_forbidden_visible_content(text)
        for text in (txt_text, docx_text):
            self._assert_expected_natural_sections(text)

        self.assertIn("第1集：风起", public_visible_text)
        self.assertNotIn("scene_json", public_visible_text)
        self.assertNotIn("appearanceMapping", public_visible_text)


if __name__ == "__main__":
    unittest.main()
