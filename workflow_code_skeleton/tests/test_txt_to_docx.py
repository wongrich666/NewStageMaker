from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from workflow_code_skeleton.app.utils.txt_to_docx import (
    convert,
    render_character,
    render_script,
)


class TxtToDocxCompatibilityTests(unittest.TestCase):
    def test_render_character_accepts_string_personality_and_misc_types(self) -> None:
        doc = Document()
        render_character(
            doc,
            {
                "character_name": "林夏",
                "story_role": "主角",
                "personality": "冷静敏锐，但极度克制。",
                "appearance": "总穿深色风衣",
                "relationships": [{"name": "顾川", "description": "旧友兼对手"}],
                "speech_profile": ["短句", "惜字如金"],
                "actable_evidence": {"value": "看人前会先观察出口位置"},
            },
        )

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("林夏", text)
        self.assertIn("冷静敏锐", text)
        self.assertIn("总穿深色风衣", text)
        self.assertIn("顾川", text)

    def test_render_character_accepts_non_dict_character(self) -> None:
        doc = Document()
        render_character(doc, "旧数据里直接存下来的人物纯文本描述")
        render_character(doc, None)
        render_character(doc, ["寡言", "行动力强"])

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("人物信息", text)
        self.assertIn("旧数据里直接存下来的人物纯文本描述", text)

    def test_render_script_accepts_dict_payload(self) -> None:
        doc = Document()
        render_script(
            doc,
            {
                "content": "剧本正文\n第1集：风起\n场景1：旧码头\n林夏：先查人，再查船。",
            },
        )

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("第1集：风起", text)
        self.assertIn("林夏：", text)

    def test_convert_succeeds_with_legacy_character_shapes(self) -> None:
        source = """测试短剧

故事梗概
一个旧项目的导出兼容测试。

人物小传
```json
{
  "character_setting": {
    "character_design_principle": "人人都要能落地导出。",
    "characters": [
      {
        "character_name": "林夏",
        "story_role": "主角",
        "personality": "冷静敏锐，但极度克制。",
        "appearance": "总穿深色风衣",
        "speech_profile": ["短句", "惜字如金"],
        "relationships": {
          "顾川": "旧友兼对手"
        }
      }
    ]
  }
}
```

剧本正文
第1集：风起
场景1：旧码头
林夏：先查人，再查船。
"""
        with tempfile.TemporaryDirectory() as tmpdir:
            txt_path = Path(tmpdir) / "legacy_export.txt"
            docx_path = Path(tmpdir) / "legacy_export.docx"
            txt_path.write_text(source, encoding="utf-8")

            result_path = convert(str(txt_path), str(docx_path))

            self.assertEqual(result_path, str(docx_path))
            self.assertTrue(docx_path.exists())
            doc = Document(str(docx_path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertIn("测试短剧", text)
            self.assertIn("林夏", text)
            self.assertIn("第1集：风起", text)

    def test_convert_supports_readable_preface_sections_without_json_blocks(self) -> None:
        source = """长夜回潮

故事梗概
故事从主角重返故乡展开，旧案迫使她再次入局。

世界观设定
故事发生在资源高度紧张的近未来港城，效率优先与身份等级并行。

人物小传
林夏：作为项目负责人，性格冷静克制，核心目标是保住团队。

人物服饰说明
【角色】林夏
默认称呼：林夏【日常】
固定识别锚点：深色风衣与冷静眼神

核心场景
旧码头负责承载悬念与秘密交易，会议室集中呈现角色对峙。

分集计划
第1集《风起》
林夏回城后被迫接手旧案调查。

剧本正文
第1集：风起
场景1：旧码头
林夏：先查人，再查船。
"""
        with tempfile.TemporaryDirectory() as tmpdir:
            txt_path = Path(tmpdir) / "readable_export.txt"
            docx_path = Path(tmpdir) / "readable_export.docx"
            txt_path.write_text(source, encoding="utf-8")

            result_path = convert(str(txt_path), str(docx_path))

            self.assertEqual(result_path, str(docx_path))
            doc = Document(str(docx_path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertIn("长夜回潮", text)
            self.assertIn("人物服饰说明", text)
            self.assertIn("默认称呼：林夏【日常】", text)
            self.assertIn("剧本正文", text)
            self.assertIn("第1集：风起", text)

    def test_convert_skips_legacy_appearance_and_registry_json_blocks(self) -> None:
        source = """长夜回潮

人物服饰说明
【角色】林夏
默认称呼：林夏【日常】

appearance_mapping
```json
{
  "appearance_mapping": {
    "characters": [
      {
        "character_name": "林夏",
        "default_name": "林夏【日常】"
      }
    ]
  }
}
```

character_registry
```json
{
  "character_registry": {
    "林夏": {
      "default_name": "林夏【日常】"
    }
  }
}
```

剧本正文
第1集：风起
林夏：先查人，再查船。
"""
        with tempfile.TemporaryDirectory() as tmpdir:
            txt_path = Path(tmpdir) / "skip_legacy_json.txt"
            docx_path = Path(tmpdir) / "skip_legacy_json.docx"
            txt_path.write_text(source, encoding="utf-8")

            convert(str(txt_path), str(docx_path))

            doc = Document(str(docx_path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertIn("人物服饰说明", text)
            self.assertIn("默认称呼：林夏【日常】", text)
            self.assertIn("剧本正文", text)
            self.assertNotIn("appearance_mapping", text)
            self.assertNotIn("character_registry", text)
            self.assertNotIn("附加数据", text)


if __name__ == "__main__":
    unittest.main()
